import sys
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, 
    QLabel, QFrame, QStackedWidget, QFileDialog, QComboBox, QLineEdit, QTextEdit, 
    QProgressBar, QMessageBox, QSizePolicy, QGridLayout, QScrollArea, QDialog, 
    QProgressDialog, QGroupBox, QCheckBox, QListWidget, QListWidgetItem, QInputDialog,
    QDialogButtonBox, QTabWidget, QTableWidget, QTableWidgetItem, QShortcut
)
from PyQt5.QtCore import Qt, QTimer, QSettings, QObject, pyqtSignal, QThread
from PyQt5.QtGui import QColor, QBrush, QFont, QPixmap, QImage, QKeySequence
from docx import Document
import fitz  # PyMuPDF
import os
import re
import openai
from openai import OpenAI
import subprocess
import networkx as nx
import json
import argparse
import requests
from bs4 import BeautifulSoup
import zipfile
import difflib
from urllib.parse import urljoin
from collections import defaultdict
from difflib import SequenceMatcher
from pathlib import Path
from datetime import datetime, timedelta
import sqlite3
from functools import wraps, partial
import logging
import pandas as pd  # For Excel support
import html
from typing import Optional
from database_functions import initialize_database, save_to_database, fetch_from_database
from evaluate_overall_similarity import evaluate_responses, EvaluationError, extract_dataset_ies, extract_response_ies, normalize_ie
import tempfile
import textwrap
from openai import AzureOpenAI
from dotenv import load_dotenv
import importlib.util
from log_layer_severity_impact_analyzer import LogLayerSeverityAnalyzer, PROTOCOL_DEPENDENCIES
from code_testing_engine import CodeTestingEngine
from spec_reference_evaluator import SpecReferenceEvaluator
from llm_judge_evaluator import LLMJudgeEvaluator

# Import the complete error fixing pipeline
try:
    from Error_fixing_pipelin import CompleteErrorFixingPipeline
except ImportError as e:
    # Fallback: try direct import with path manipulation
    try:
        import sys
        import os
        # Add the Error_fixing_pipelin directory to the Python path
        error_pipeline_path = os.path.join(os.path.dirname(__file__), 'Error_fixing_pipelin')
        if error_pipeline_path not in sys.path:
            sys.path.append(error_pipeline_path)
        
        from complete_error_fixing_pipeline import CompleteErrorFixingPipeline
    except ImportError as e2:
        # Final fallback if the module is not found
        print(f"Warning: Could not import CompleteErrorFixingPipeline: {e2}")
        CompleteErrorFixingPipeline = None

# Theme and Style Constants
THEME = {
    'colors': {
        'primary': '#0078D4',
        'secondary': '#106EBE',
        'background': '#F5F5F5',
        'surface': '#FFFFFF',
        'border': '#E1E5EA',
        'hover': '#F0F0F0',
        'text': {
            'primary': '#323130',
            'secondary': '#605E5C'
        }
    },
    'typography': {
        'title': """
            QLabel {
                color: #323130;
                font-size: 24px;
                font-weight: bold;
                margin-bottom: 8px;
            }
        """,
        'subtitle': """
            QLabel {
                color: #605E5C;
                font-size: 14px;
                margin-bottom: 24px;
            }
        """,
        'label': """
            QLabel {
                color: #323130;
                font-size: 13px;
                margin-bottom: 4px;
            }
        """,
        'nav_title': """
            QLabel {
                color: #323130;
                font-size: 20px;
                font-weight: 600;
                margin-bottom: 16px;
            }
        """,
        'nav_label': """
            QLabel {
                color: #605E5C;
                font-size: 12px;
                font-weight: 600;
                margin: 16px 0 8px 0;
            }
        """,
        'body': """
            QLabel {
                color: #323130;
                font-size: 14px;
                line-height: 1.4;
            }
        """,
        'caption': """
            QLabel {
                color: #605E5C;
                font-size: 12px;
            }
        """
    }
}

INPUT_STYLE = """
    QLineEdit, QComboBox {
        background-color: #FFFFFF;
        border: 1px solid #E1E5EA;
        border-radius: 4px;
        padding: 8px;
        font-size: 13px;
    }
    QLineEdit:focus, QComboBox:focus {
        border-color: #0078D4;
    }
"""

BUTTON_STYLE = """
    QPushButton {
        background-color: #FFFFFF;
        border: 1px solid #E1E5EA;
        border-radius: 4px;
        padding: 8px 16px;
        font-size: 13px;
        color: #323130;
    }
    QPushButton:hover {
        background-color: #F0F0F0;
        border-color: #0078D4;
    }
"""

GROUP_BOX_STYLE = """
    QGroupBox {
        background-color: transparent;
        border: 1px solid #E1E5EA;
        border-radius: 4px;
        margin-top: 12px;
        padding-top: 16px;
    }
    QGroupBox::title {
        subcontrol-origin: margin;
        subcontrol-position: top left;
        padding: 0 4px;
        color: #323130;
        font-weight: 600;
    }
"""

CHECKBOX_STYLE = """
    QCheckBox {
        color: #323130;
        font-size: 13px;
        padding: 4px;
    }
    QCheckBox::indicator {
        width: 18px;
        height: 18px;
    }
    QCheckBox::indicator:unchecked {
        border: 1px solid #E1E5EA;
        background: white;
        border-radius: 3px;
    }
    QCheckBox::indicator:checked {
        border: 1px solid #0078D4;
        background: #0078D4;
        border-radius: 3px;
    }
    QCheckBox::indicator:checked::mark {
        image: url(checkmark.png);
    }
"""
#Azure key start

# Check if .env file exists
if not os.path.exists('.env'):
    print("Warning: .env file not found in the current directory.")

# Load environment variables from .env file
load_dotenv()

# Get Azure OpenAI credentials from environment variables
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_MODEL_NAME = os.getenv("AZURE_OPENAI_MODEL_NAME")

missing_vars = []
if not AZURE_OPENAI_ENDPOINT:
    missing_vars.append("AZURE_OPENAI_ENDPOINT")
if not AZURE_OPENAI_API_KEY:
    missing_vars.append("AZURE_OPENAI_API_KEY")
if not AZURE_OPENAI_MODEL_NAME:
    missing_vars.append("AZURE_OPENAI_MODEL_NAME")

if missing_vars:
    print(f"Error: Missing environment variables: {', '.join(missing_vars)}. Please set them in your .env file.")
    sys.exit(1)

try:
    client = AzureOpenAI(
        api_key=AZURE_OPENAI_API_KEY,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_version="2024-02-01"  # Check Azure OpenAI documentation for the latest supported API version
    )
except Exception as e:
    print(f"Failed to initialize AzureOpenAI client: {e}")
    sys.exit(1)

# Azure key end

directory = r"C:\Users\ChanduVangala\Documents\GenAI\extract"

# Navigation Button Style
NAV_BUTTON_STYLE = """
    QPushButton {
        background-color: transparent;
        border: none;
        border-radius: 4px;
        color: #605E5C;
        font-size: 14px;
        padding: 8px 12px;
        text-align: left;
    }
    QPushButton:hover {
        background-color: #F5F7FA;
        color: #323130;
    }
    QPushButton:checked {
        background-color: #F0F6FF;
        color: #0078D4;
        font-weight: 600;
    }
"""

# Enterprise Button Styles
BUTTON_STYLE = """
    QPushButton {
        background-color: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 4px;
        color: #2d3748;
        font-size: 12px;
        padding: 6px 12px;
        text-align: left;
    }
    QPushButton:hover {
        background-color: #edf2f7;
    }
    QPushButton:pressed {
        background-color: #e2e8f0;
    }
    QPushButton:disabled {
        color: #a0aec0;
        background-color: #f8f9fa;
    }
"""

# Enterprise Input Styles
INPUT_STYLE = """
    QLineEdit, QTextEdit, QComboBox {
        background-color: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 4px;
        color: #2d3748;
        font-size: 12px;
        padding: 6px;
    }
    QLineEdit:focus, QTextEdit:focus, QComboBox:focus {
        border: 1px solid #2c5282;
    }
    QComboBox::drop-down {
        border: none;
        width: 20px;
    }
    QComboBox::down-arrow {
        image: url(down_arrow.png);
        width: 12px;
        height: 12px;
    }
"""

# Add style constants
GROUP_BOX_STYLE = """
    QGroupBox {
        background-color: transparent;
        border: 1px solid #E1E5EA;
        border-radius: 4px;
        margin-top: 12px;
        padding-top: 16px;
    }
    QGroupBox::title {
        subcontrol-origin: margin;
        subcontrol-position: top left;
        padding: 0 4px;
        color: #323130;
        font-weight: 600;
    }
"""

CHECKBOX_STYLE = """
    QCheckBox {
        color: #323130;
        font-size: 13px;
        padding: 4px;
    }
    QCheckBox::indicator {
        width: 18px;
        height: 18px;
    }
    QCheckBox::indicator:unchecked {
        border: 1px solid #E1E5EA;
        background: white;
        border-radius: 3px;
    }
    QCheckBox::indicator:checked {
        border: 1px solid #0078D4;
        background: #0078D4;
        border-radius: 3px;
    }
    QCheckBox::indicator:checked::mark {
        image: url(checkmark.png);
    }
"""

#For extraction progress bar start
class ExtractionThread(QThread):
    progress_signal = pyqtSignal(int)  # Signal to update progress bar
    finished = pyqtSignal()  # Signal for completion
    error_signal = pyqtSignal(str)  # Signal for errors
    
    def __init__(self, prompt, parent=None):
        super().__init__(parent)
        self.prompt = prompt
        self.text_content = None
        self.selected_prompt = None
        self.result = None
        
        
    def run(self):
        try:
            # Initialize progress
            self.progress_signal.emit(10)
            
            # Generate response using OpenAI
            if hasattr(self.parent(), 'generate_response_from_text'):
                self.result = self.parent().generate_response_from_text(
                    self.progress_signal.emit,
                    self.text_content,
                    self.selected_prompt
                )
                self.progress_signal.emit(90)
                
                if self.result and isinstance(self.result, str):
                    if "Error:" in self.result:
                        self.error_signal.emit(self.result)
                    else:
                        self.progress_signal.emit(100)
                        self.finished.emit()
                else:
                    self.error_signal.emit("Invalid response format received")
            else:
                self.error_signal.emit("Response generation method not found")
                
        except Exception as e:
            self.error_signal.emit(f"Error in response generation: {str(e)}")
        finally:
            if not self.result:
                self.progress_signal.emit(0)

    def update_progress(self, value):
        self.progress_signal.emit(value)
#For extraction progress bar end


class DeploymentContextDialog(QDialog):
    """Dialog for editing deployment context settings"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Deployment Context Settings")
        self.setMinimumWidth(600)
        self.setMinimumHeight(600)
        
        # Load default values from JSON
        print("\n" + "="*80)
        print("🔍 DEPLOYMENT CONTEXT DIALOG - INITIALIZATION")
        print("="*80)
        self.default_values = self.load_default_values()
        self.current_values = self.default_values.copy()
        
        print(f"\n📦 default_values dictionary contains {len(self.default_values)} items:")
        for key, value in self.default_values.items():
            print(f"   '{key}': '{value}'")
        
        print("\n🔎 Checking specific fields:")
        check_fields = ['nssai_sst', 'nssai_sd', 'dnn', 'nmc_size', 'Deploy_command_cu_gnb_conf', 'Deploy_command_du_gnb_conf']
        for field in check_fields:
            value = self.default_values.get(field, "NOT FOUND")
            print(f"   {field}: {value}")
        print("="*80 + "\n")
        
        self.init_ui()
    
    def load_default_values(self):
        """Load default deployment context from JSON file"""
        try:
            # Try multiple paths
            possible_paths = [
                'database/error_patterns_structured.json',
                'Error_fixing_pipelin/database/error_patterns_structured.json',
                os.path.join(os.path.dirname(__file__), 'database', 'error_patterns_structured.json'),
                os.path.join(os.path.dirname(__file__), 'Error_fixing_pipelin', 'database', 'error_patterns_structured.json')
            ]
            
            for json_path in possible_paths:
                if os.path.exists(json_path):
                    print(f"✅ Loading deployment context from: {json_path}")
                    with open(json_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        deployment_context = data.get('deployment_context', {})
                        print(f"📊 Loaded {len(deployment_context)} deployment context values:")
                        for key, value in deployment_context.items():
                            print(f"   - {key}: {value}")
                        return deployment_context
            
            print(f"⚠️ Could not find deployment context JSON in any of these paths:")
            for path in possible_paths:
                print(f"   - {path} (exists: {os.path.exists(path)})")
                
        except Exception as e:
            print(f"❌ Error loading deployment context: {e}")
            import traceback
            traceback.print_exc()
        
        # Return empty dict if loading fails
        return {}
    
    def init_ui(self):
        """Initialize the UI components"""
        layout = QVBoxLayout()
        
        # Title
        title_label = QLabel("Configure Deployment Context")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #0078D4;
                margin-bottom: 10px;
            }
        """)
        layout.addWidget(title_label)
        
        # Subtitle
        subtitle_label = QLabel("Edit the deployment context values that will be used in bug analysis. "
                               "These values will override the default JSON configuration.")
        subtitle_label.setWordWrap(True)
        subtitle_label.setStyleSheet("""
            QLabel {
                font-size: 12px;
                color: #605E5C;
                margin-bottom: 15px;
            }
        """)
        layout.addWidget(subtitle_label)
        
        # Scroll area for fields
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        
        # Create input fields
        self.fields = {}
        
        # Define field configurations
        field_configs = [
            ("cu_ip_address", "CU IP Address", "IP address of the Central Unit"),
            ("du_ip_address", "DU IP Address", "IP address of the Distributed Unit"),
            ("gnb_ip_address", "gNB IP Address", "IP address of the gNodeB"),
            ("amf_ip_address", "AMF IP Address", "IP address of the Access and Mobility Management Function"),
            ("core_network_machine_ip", "Core Network Machine IP", "IP address of the Core Network machine"),
            ("local_s_portc", "Local S Port C", "Local control plane port"),
            ("local_s_portd", "Local S Port D", "Local data plane port"),
            ("remote_s_portc", "Remote S Port C", "Remote control plane port"),
            ("remote_s_portd", "Remote S Port D", "Remote data plane port"),
            ("nssai_sst", "NSSAI SST", "Network Slice Selection Assistance Information - Slice/Service Type"),
            ("nssai_sd", "NSSAI SD", "Network Slice Selection Assistance Information - Slice Differentiator"),
            ("nmc_size", "NMC Size", "Network Management Component size"),
            ("dnn", "DNN", "Data Network Name"),
            ("Deploy_command_cu_gnb_conf", "CU Deployment Command", "Command used to deploy CU gNB configuration"),
            ("Deploy_command_du_gnb_conf", "DU Deployment Command", "Command used to deploy DU gNB configuration"),
        ]
        
        print("\n📋 CREATING INPUT FIELDS:")
        print("-" * 80)
        for field_key, field_label, field_tooltip in field_configs:
            print(f"\n🔹 Creating field: {field_label} (key: {field_key})")
            
            field_group = QGroupBox(field_label)
            field_group.setStyleSheet("""
                QGroupBox {
                    font-weight: bold;
                    border: 1px solid #E1E5EA;
                    border-radius: 5px;
                    margin-top: 10px;
                    padding-top: 10px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 10px;
                    padding: 0 5px;
                }
            """)
            field_layout = QVBoxLayout()
            
            # Input field
            input_field = QLineEdit()
            default_value = self.default_values.get(field_key, "")
            print(f"   📥 Retrieved value from default_values: '{default_value}' (type: {type(default_value).__name__})")
            
            # Convert to string, handling None and empty values
            display_value = str(default_value) if default_value is not None else ""
            print(f"   📤 Display value (after str conversion): '{display_value}'")
            
            input_field.setText(display_value)
            print(f"   ✏️ Set text in QLineEdit: '{input_field.text()}'")
            
            input_field.setPlaceholderText(f"Enter {field_label}")
            input_field.setToolTip(field_tooltip)
            
            # Debug logging for each field
            if default_value:
                print(f"   ✅ SUCCESS - Field '{field_key}' loaded with value: '{display_value}'")
            else:
                print(f"   ⚠️ WARNING - Field '{field_key}' has no value in JSON (will be empty)")
            input_field.setStyleSheet("""
                QLineEdit {
                    padding: 8px;
                    border: 1px solid #E1E5EA;
                    border-radius: 4px;
                    font-size: 13px;
                }
                QLineEdit:focus {
                    border: 2px solid #0078D4;
                }
            """)
            field_layout.addWidget(input_field)
            
            # Help text
            help_label = QLabel(field_tooltip)
            help_label.setWordWrap(True)
            help_label.setStyleSheet("""
                QLabel {
                    font-size: 11px;
                    color: #8A8886;
                    margin-top: 2px;
                }
            """)
            field_layout.addWidget(help_label)
            
            field_group.setLayout(field_layout)
            scroll_layout.addWidget(field_group)
            
            self.fields[field_key] = input_field
        
        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)
        
        print("\n" + "="*80)
        print("✅ DIALOG INITIALIZATION COMPLETE")
        print(f"📊 Total fields created: {len(self.fields)}")
        print("\n🔍 Field values that were set:")
        for field_key, input_field in self.fields.items():
            current_text = input_field.text()
            if current_text:
                print(f"   ✓ {field_key}: '{current_text}'")
            else:
                print(f"   ✗ {field_key}: EMPTY")
        print("="*80 + "\n")
        
        # Buttons
        button_layout = QHBoxLayout()
        
        # Clear custom settings button
        clear_btn = QPushButton("Clear All (Use JSON)")
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #D13438;
                color: white;
                font-weight: bold;
                padding: 10px 20px;
                border: none;
                border-radius: 4px;
                min-width: 140px;
            }
            QPushButton:hover {
                background-color: #A52A2D;
            }
        """)
        clear_btn.setToolTip("Clear all custom settings and use JSON defaults")
        clear_btn.clicked.connect(self.clear_all_settings)
        button_layout.addWidget(clear_btn)
        
        # Reset button
        reset_btn = QPushButton("Reset to JSON Defaults")
        reset_btn.setStyleSheet("""
            QPushButton {
                background-color: #8A8886;
                color: white;
                font-weight: bold;
                padding: 10px 20px;
                border: none;
                border-radius: 4px;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #797775;
            }
        """)
        reset_btn.clicked.connect(self.reset_to_defaults)
        button_layout.addWidget(reset_btn)
        
        button_layout.addStretch()
        
        # Cancel button
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #8A8886;
                color: white;
                font-weight: bold;
                padding: 10px 20px;
                border: none;
                border-radius: 4px;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #797775;
            }
        """)
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        
        # Save button
        save_btn = QPushButton("Save")
        save_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                font-weight: bold;
                padding: 10px 20px;
                border: none;
                border-radius: 4px;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
        """)
        save_btn.clicked.connect(self.save_values)
        button_layout.addWidget(save_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
    
    def clear_all_settings(self):
        """Clear all field values"""
        reply = QMessageBox.question(
            self,
            "Clear All Settings",
            "This will clear all custom settings and the system will use JSON defaults.\n\n"
            "Are you sure you want to continue?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            for field_key, input_field in self.fields.items():
                input_field.clear()
    
    def reset_to_defaults(self):
        """Reset all fields to default values from JSON"""
        for field_key, input_field in self.fields.items():
            input_field.setText(str(self.default_values.get(field_key, "")))
    
    def save_values(self):
        """Save the edited values and close dialog"""
        self.current_values = {}
        for field_key, input_field in self.fields.items():
            value = input_field.text().strip()
            if value:
                self.current_values[field_key] = value
        self.accept()
    
    def get_values(self):
        """Return the current deployment context values"""
        return self.current_values


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        initialize_database()
        
        # Initialize settings
        self.settings = QSettings('AgenticRAN', 'TestScriptGenerator')
        """
        # Initialize OpenAI client with API key from settings
        api_key = self.settings.value('openai_api_key')
        if api_key:
            os.environ['OPENAI_API_KEY'] = api_key
        """

        self.working_directory = self.settings.value('working_directory', '')
        self.output_directory = self.settings.value('output_directory', '')
        
        # Set window flags for custom title bar
        self.setWindowFlags(Qt.FramelessWindowHint)
        self.testcases_list = None
        self.testcases_name=None
        # Create directories if they don't exist
        if not self.working_directory:
            self.working_directory = os.path.join(os.path.expanduser('~'), 'AgenticRAN')
            os.makedirs(self.working_directory, exist_ok=True)
            self.settings.setValue('working_directory', self.working_directory)
            
        if not self.output_directory:
            self.output_directory = os.path.join(self.working_directory, 'output')
            os.makedirs(self.output_directory, exist_ok=True)
            self.settings.setValue('output_directory', self.output_directory)

        # Define role templates
        self.role_templates = {
            "tester": [
                ("Test Script", "Create automated test scripts", "📝"),
                ("Test Case", "Generate comprehensive test cases", "🧪"),
                ("Bug Analysis", "Analyze and fix issues", "🔍"),
                ("Performance Test", "Design performance test cases", "⚡")
            ],
            "developer": [
                ("Code Assistant", "Generate high-quality code", "💻"),
                ("Bug Analysis", "Fix identified issues", "🐛")
            ],
            "analyst": [
                ("Feature Design", "Analyze feature requirements", "📊"),
                ("Code Assistant", "Review code quality", "👀"),
                ("Bug Analysis", "Analyze issue patterns", "🔍")
            ]
        }
        
        # Initialize custom deployment context (None means use default from JSON)
        self.custom_deployment_context = None


        with open('C:\\Users\\ChanduVangala\\Documents\\AgenticRAN-V8_azure_key\\ue_attach_utils.py', 'r') as file:
            self.ue_attach_utils = file.read()
        
        # Define dummy variables for prompt templates
        self.VENDOR = "{VENDOR}"
        self.SYSTEM_TYPE = "{SYSTEM_TYPE}"
        self.DOMAIN = "{DOMAIN}"
        self.CONNECTION_METHOD = "{CONNECTION_METHOD}"
        self.LOGIN_CREDENTIALS = "{LOGIN_CREDENTIALS}"
        self.ACCESS_MODE = "{ACCESS_MODE}"
        self.PRIMARY_FEATURE = "{PRIMARY_FEATURE}"
        self.LANGUAGE = "{LANGUAGE}"
        self.commands_content = "{commands_content}"
        
        # Define prompts dictionary
        self.prompts = {
            "Test Script": '''You are an expert in 5G network testing and automation, with deep knowledge of RRC and NAS protocols. Given the dataset containing the complete 5G NSA attach procedure—including all signaling messages and their respective Information Elements (IEs)—generate comprehensive, production-ready {LANGUAGE} test automation scripts.

Requirements:
- First, trigger the UE attach procedure using the provided reference code.
- For each message in the attach sequence (RRC and NAS), generate a {LANGUAGE} function that:
    - Simulates the message exchange.
    - Extracts and validates all Information Elements (IEs) for that message.
    - Logs results and assertions for traceability.
- Ensure:
    - Every message in the attach procedure is covered, in correct sequence, with no skipped steps.
    - All IEs per message are explicitly validated.
    - Scripts are modular, readable, and maintainable.
    - No placeholder comments or "add your code here" lines—provide full logic.
    - Output is only {LANGUAGE} code (no explanations, no markdown, no extra text).
-Make sure the generated scripts covers the testcases which are in {self.testcases_name}
Dataset Summary:
[Attach procedure dataset including all messages in sequence and their respective Information Elements]

Expected Output:
- {LANGUAGE} automation scripts that:
    - Trigger the UE attach.
    - Validate each message and all IEs.
    - Log results for each step.
    - Ensure full message and IE validation coverage.

Must use this reference code for triggering the attach procedure and validating attach and every message in the attach procedure, use logic from this code:
{self.ue_attach_utils}

Instructions:

- Use the above reference code for triggering the attach procedure and validating attach status.
- For each message in the dataset, generate a function that validates all IEs, logs results, and asserts correctness.
- Do not output any explanations, markdown, or placeholder comments—only complete {LANGUAGE} code.

''',
            "Testing Strategy": '''You are an expert test strategist with deep knowledge of software testing methodologies and best practices. Based on the provided dataset and requirements, create a comprehensive testing strategy document.

Your task is to analyze the given dataset and generate a detailed testing strategy that includes:

1. **Test Scope and Objectives**
   - Define what needs to be tested
   - Identify key testing goals and success criteria
   - Determine testing boundaries and limitations

2. **Test Approach and Methodology**
   - Recommend appropriate testing methodologies (unit, integration, system, acceptance)
   - Define testing phases and their sequence
   - Specify testing techniques and tools to be used

3. **Test Planning and Organization**
   - Identify test deliverables and artifacts
   - Define roles and responsibilities
   - Create testing timeline and milestones

4. **Risk Assessment and Mitigation**
   - Identify potential testing risks
   - Propose mitigation strategies
   - Define contingency plans

5. **Test Environment and Infrastructure**
   - Specify test environment requirements
   - Define test data management strategy
   - Plan for test automation where applicable

6. **Quality Metrics and Reporting**
   - Define key performance indicators
   - Establish reporting mechanisms
   - Set quality gates and exit criteria

Provide a professional, actionable testing strategy that can be implemented by a testing team.''',
            "Bug Analysis": '''You are an expert in 5G/LTE telecommunications debugging. You should be generous with relevance scores for items that could plausibly help debug or fix the error.

Error: "{error_message}"

Evaluate these {candidate_type} for relevance to the error. You have access to the complete code/context for each candidate:

{candidate_details}

For each item, provide a relevance score (0.0-1.0) and brief reason based on the actual code/context.
Format: "1: 0.8 - reason, 2: 0.3 - reason, ..."

**CRITICAL FOR RRC ERRORS**: For RRC-related errors, be especially generous with:
- RRC rejection functions (rrc_gNB_generate_RRCReject, rrc_gNB_send_RRCReject)
- RRC error handling functions (handle_error, process_error)
- RRC setup failure functions (setup_failure, handle_failure)
- RRC context management functions (remove_ue, cleanup_context)
- Functions that handle RRCSetupRequest processing
- Functions that generate RRC messages (Reject, Failure, etc.)

Score generously - consider relevant:
- Any function containing keywords from the error (AMF, gNB, NGAP, RRC, etc.)
- Configuration parameters related to the error components
- Protocol handling functions even if indirectly related
- Error handling and setup functions
- Network interface and connection functions
- Functions that could be involved in the error scenario based on their actual implementation
- Config parameters that could affect the error based on their current values and context

**For RRC segmentation faults, ANY RRC error handling function should score 0.7+**

Score 0.5+ if there's ANY reasonable connection to the error scenario based on the actual code/context.''',
            "Feature Design": '''You are an expert telecommunications systems architect and feature design specialist with deep knowledge of 5G/LTE networks, protocol design, and system integration. Based on the provided requirements and context, design comprehensive feature specifications.

Your design should include:

1. **Feature Architecture and Design**
   - Define the feature's core functionality and scope
   - Design the system architecture and component interactions
   - Plan for protocol integration and compliance
   - Specify data flow and message sequences

2. **Technical Specifications**
   - Define detailed technical requirements
   - Specify protocol compliance (3GPP, RFC standards)
   - Design API interfaces and data structures
   - Plan for error handling and edge cases

3. **Implementation Strategy**
   - Break down the feature into implementable components
   - Define development phases and milestones
   - Specify testing and validation requirements
   - Plan for integration with existing systems

4. **Performance and Scalability**
   - Define performance requirements and metrics
   - Plan for scalability and load handling
   - Specify resource requirements and constraints
   - Design for high availability and reliability

5. **Security and Compliance**
   - Implement security best practices
   - Ensure regulatory compliance
   - Plan for data protection and privacy
   - Design for audit and monitoring

6. **Documentation and Standards**
   - Create comprehensive technical documentation
   - Define user interfaces and workflows
   - Plan for training and knowledge transfer
   - Ensure maintainability and extensibility

7. **Testing and Validation**
   - Define comprehensive test strategies
   - Plan for conformance and interoperability testing
   - Design for performance and stress testing
   - Plan for user acceptance testing

Provide a complete, production-ready feature design with detailed specifications, implementation guidance, and quality assurance plans.''',
            "Code Review": '''You are a senior software engineer and code review expert with extensive experience in code quality assessment, best practices, and technical leadership. Based on the provided code or dataset, conduct a comprehensive code review.

Your review should cover:

1. **Code Quality Assessment**
   - Evaluate code readability and maintainability
   - Check adherence to coding standards and conventions
   - Assess code organization and structure

2. **Technical Analysis**
   - Review algorithms and data structures
   - Evaluate performance considerations
   - Check for potential technical debt

3. **Security Review**
   - Identify potential security vulnerabilities
   - Check for secure coding practices
   - Assess data handling and validation

4. **Best Practices Compliance**
   - Verify design patterns usage
   - Check error handling and logging
   - Assess testing coverage and quality

5. **Documentation and Comments**
   - Evaluate code documentation
   - Check comment quality and relevance
   - Assess API documentation completeness

6. **Improvement Recommendations**
   - Provide specific refactoring suggestions
   - Recommend performance optimizations
   - Suggest architectural improvements

7. **Risk Assessment**
   - Identify potential production issues
   - Assess maintainability concerns
   - Evaluate scalability implications

Provide constructive feedback with specific examples and actionable recommendations for improvement.''',
            "Performance Test": '''You are a performance testing expert with deep knowledge of system performance analysis, load testing, and optimization strategies. Based on the provided dataset, create a comprehensive performance testing plan.

Your plan should include:

1. **Performance Requirements Analysis**
   - Define performance objectives and SLAs
   - Identify key performance indicators (KPIs)
   - Establish performance baselines and benchmarks

2. **Test Strategy and Approach**
   - Design load testing scenarios
   - Plan stress and volume testing
   - Define performance test types (load, stress, spike, endurance)

3. **Test Environment Setup**
   - Specify hardware and software requirements
   - Define test data requirements
   - Plan for monitoring and measurement tools

4. **Test Scenarios and Scripts**
   - Create realistic user scenarios
   - Design test data sets
   - Develop performance test scripts

5. **Performance Metrics and Monitoring**
   - Define key metrics to measure
   - Set up monitoring and alerting
   - Plan for performance data collection

6. **Analysis and Reporting**
   - Establish performance analysis procedures
   - Define reporting formats and frequency
   - Plan for performance optimization recommendations

7. **Risk Mitigation**
   - Identify performance risks
   - Plan for performance bottlenecks
   - Define performance degradation handling

Provide a detailed, implementable performance testing strategy that ensures system reliability and optimal performance.''',
            "API Design": '''You are a senior API architect and design expert with extensive experience in RESTful APIs, microservices, and system integration. Based on the provided requirements, design a comprehensive API solution.

Your design should include:

1. **API Architecture and Design**
   - Define API structure and organization
   - Design RESTful endpoints and resources
   - Plan for API versioning and evolution

2. **API Specifications**
   - Define request/response schemas
   - Specify authentication and authorization
   - Design error handling and status codes

3. **API Documentation**
   - Create comprehensive API documentation
   - Define usage examples and tutorials
   - Plan for interactive documentation (Swagger/OpenAPI)

4. **Security and Compliance**
   - Implement security best practices
   - Plan for data protection and privacy
   - Ensure compliance with standards

5. **Performance and Scalability**
   - Design for high performance
   - Plan for scalability and load handling
   - Implement caching strategies

6. **Testing Strategy**
   - Define API testing approach
   - Plan for automated testing
   - Design integration test scenarios

7. **Monitoring and Analytics**
   - Plan for API monitoring
   - Design analytics and metrics
   - Implement logging and debugging

Provide a complete, production-ready API design with detailed specifications and implementation guidance.''',
            "Code Assistant": '''You are an expert telecommunications software debugger and fix engineer specializing in 5G/LTE systems.  
You have deep knowledge of:

- NGAP (Next Generation Application Protocol)
- RRC (Radio Resource Control)
- NAS (Non-Access Stratum)
- AMF (Access and Mobility Management Function)
- gNB (Next Generation Node B) architecture
- SCTP, network configuration, and protocol stacks
- Network connectivity, subnet analysis, and routing

Your task is to analyze error contexts and propose **SURGICAL but MEANINGFUL** fixes that address the **ACTUAL** root cause.  
Fixes must intelligently **reconstruct missing or incorrect control-flow** when evidence shows the function's logic is incomplete.  
Do **NOT** rely solely on defensive null checks unless they are truly the cause.

🚨 **MANDATORY REQUIREMENT**: If you find any incomplete if-else if chains (missing final else clause), you MUST add the missing else branch to handle unrecognized cases. This is CRITICAL for preventing segmentation faults.

🔍 **STEP-BY-STEP SEARCH PROCESS**:
1. **FIND** the line: `}} else if (NR_InitialUE_Identity_PR_ng_5G_S_TMSI_Part1 == rrcSetupRequest->ue_Identity.present) {{`
2. **SCAN DOWN** to find the ACTUAL LAST line in this else if block (look for assignments like `UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1;`)
3. **FIND** the closing brace `}}` right after that last assignment
4. **USE** that closing section as your original_code

⚠️ **EXAMPLE OF MISSING ELSE**: If you see code like:
```
if (condition1) {{
    // some code
}} else if (condition2) {{
    // some code
    UE->some_field = value;  // <- FIND THE ACTUAL LAST ASSIGNMENT
}}
// <-- MISSING ELSE CLAUSE HERE!
NR_CellGroupConfig_t *cellGroupConfig = NULL;  // Code continues
```
You MUST find the LAST assignment (like `UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1;`) + closing brace and add the else clause there.

### ROOT CAUSE ANALYSIS REQUIREMENTS
- Always check **network connectivity first** for 5G issues.  
- Distinguish between **network vs. code logic** failures.  
- Trace error flow from **log messages** to code path.  
- **CRITICAL**: Look for **incomplete conditional logic patterns**:
  * Hardcoded constant comparisons (e.g., `== 3`) instead of runtime field checks
  * Missing `else` branches in multi-case scenarios (enum handling, protocol states)
  * Unvalidated assumptions about input data structure/presence
- If branches or identity handling are clearly missing or malformed, **reconstruct the expected flow** using:
  * 3GPP specifications
  * Variable names and surrounding code context
  * Typical OpenAirInterface gNB patterns  
- Do **NOT** require the original code to be supplied—derive expected behavior from specs and context.

### CODE PATCH RULES
- **CRITICAL**: Preserve existing correct logic; **insert or adjust only the minimal lines** necessary to restore intended behavior.  
- Use **exact surrounding code lines or variable names** for placement (`after line containing "..."`).  
- **CRITICAL**: Look for **hardcoded constant comparisons** (like `== 3`, `== 1`) and replace with **proper dynamic checks** using available variables/fields.
- **CRITICAL**: If the function has incomplete conditional logic (missing `else` branches, unhandled enum cases), **reconstruct the missing branches** instead of adding trivial guards.  
- **CRITICAL**: For segmentation faults, trace the **actual cause** (uninitialized variables, missing validation, incomplete state handling) rather than adding generic null checks.
- Never invent functions or config parameters that don't exist.  
- Validate variable scope before referencing them.  
- **AVOID GENERIC NULL CHECKS**: If a segmentation fault is due to skipped handling of input cases, **add the missing handling**, not just a pointer check.

### COMMON SEGMENTATION FAULT PATTERNS TO LOOK FOR
- **Hardcoded enum comparisons**: `if (ENUM_CONSTANT == hardcoded_value)` → should be `if (data->field == ENUM_CONSTANT)`
- **CRITICAL: Missing else/default cases**: Incomplete switch/if-else chains for enum/state handling
- **MANDATORY: Always add else branch to any if-else if chain that lacks final else clause**
- **Uninitialized pointers**: Variables that remain NULL when certain conditions aren't met
- **Protocol state violations**: Functions proceeding without validating required protocol fields
- **Array/structure access**: Accessing fields without checking structure validity first

### 🚨 CRITICAL REQUIREMENT: MISSING ELSE BRANCHES 🚨
**EXAMINE THE PROVIDED CODE CAREFULLY**: Look for incomplete if-else if chains and add the missing else clause:
- In the provided code, if you see: `if (condition1) {{ ... }} else if (condition2) {{ ... }}` WITHOUT a final `else`
- AND the code continues AFTER the closing brace without handling other cases
- You MUST create a SEPARATE patch with these EXACT specifications:
  * `"original_code": "FIND THE ACTUAL LAST ASSIGNMENT in the else if block (like UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1;) + closing brace }}"` 
  * `"patched_code": "SAME ENDING LINES + }} else {{\\n    LOG_E(NR_RRC, \\"Unhandled ue_Identity.present value: %d\\", rrcSetupRequest->ue_Identity.present);\\n    return;\\n}}"`
  * `"line_numbers": "replace the section ending with the actual last assignment + closing brace"`
  * **CRITICAL**: Look for the ACTUAL FINAL STATEMENT in the else if block (not function calls in the middle)
- **NEVER** use function calls like rrc_gNB_create_ue_context - find the REAL LAST assignment like UE->field = value;

### CONFIGURATION VALIDATION
- Use actual config names and values from candidate configs.  
- Verify subnet and port compatibility using deployment context when relevant.  
- Suggest meaningful corrections only (never no-ops).

### RESPONSE FORMAT
Respond with **ONLY valid JSON** in this EXACT structure:

{{
    "suspected_functions": ["function1", "function2"],
    "suspected_configs": ["config1", "config2"], 
    "reason": "Detailed root cause explanation",
    "config_fix": "Specific configuration corrections with exact values",
    "code_patches": [
        {{
            "function_name": "function_name",
            "file_path": "path/to/file.c",
            "patch_type": "targeted_insertion_or_adjustment",
            "original_code": "// Line(s) around the issue",
            "patched_code": "// The exact corrected code snippet",
            "line_numbers": "EXACT line numbers (e.g., '120-125') or specific context (e.g., 'after line containing \\"amf_desc_p = ngap_gNB_get_AMF\\"')",
            "description": "Why this correction resolves the error"
        }}
    ],
    "config_patches": [
        {{
            "config_name": "parameter_name",
            "file_path": "path/to/config.conf",
            "patch_type": "set_value",
            "current_value": "current_value",
            "new_value": "corrected_value",
            "line_number": "approximate_line_or_section",
            "relevance_score": "confidence_score_from_analysis",
            "description": "Why this config change resolves the error"
        }}
    ],
    "root_cause_analysis": "Deep technical analysis of why this error occurs",
    "investigation_steps": ["step1", "step2", "step3"]
}}

### ADDITIONAL HINTS
- **Protocol Compliance**: Use 3GPP specifications and RFCs to understand expected behavior and mandatory validations.
- **Network Connectivity**: For 5G/LTE issues, verify IP reachability, routing, and port accessibility between components.
- **State Validation**: Ensure all possible input states/cases are handled with appropriate error responses.
- **Container Networking**: If using Docker/containers, add static routes or network configuration as needed.

### GOOD PATCH EXAMPLES
- **CRITICAL**: For missing else clauses - EXAMPLE PATCH (shows ACTUAL CLOSING LINES):
  ```
  {{
    "original_code": "    UE->Initialue_identity_5g_s_TMSI.presence = true;\\n    UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1;\\n  }}",
    "patched_code": "    UE->Initialue_identity_5g_s_TMSI.presence = true;\\n    UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1;\\n  }} else {{\\n    LOG_E(NR_RRC, \\"Unhandled ue_Identity.present value: %d\\", rrcSetupRequest->ue_Identity.present);\\n    return;\\n  }}",
    "line_numbers": "replace the section ending with UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1; }}"
  }}
  ```
  **KEY**: Find the ACTUAL LAST assignment in the else if block (like UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1;) + closing brace
- Replace hardcoded comparisons: `if (enum_constant == 3)` → `if (variable->field == enum_constant)`
- Adding **routing checks** or static route commands for Dockerized AMF  

### BAD PATCH EXAMPLES
- Only adding generic NULL checks like `if (ptr == NULL) {{ return; }}` without addressing the root cause.
- Rewriting entire functions when only specific lines need adjustment.
- Using invented APIs, functions, or configuration parameters.
- Keeping hardcoded constants instead of using proper runtime checks.
- Adding defensive code without understanding why the error occurs.''',
            "Test Case": {
                "User Prompt": f"""

AVAILABLE COMMANDS/INTERFACES:
{self.commands_content}

TASK:
Generate the **maximum possible number of unique and comprehensive test cases** for the above document content by intelligently integrating relevant commands/interfaces (if provided) or creating appropriate generic test steps.

REQUIREMENTS:
1. **Comprehensive Analysis**: Extract test scenarios from every relevant sentence, condition, and rule in the document
2. **Intelligent Command/Interface Selection**: Analyze the available commands/interfaces (if provided) and select appropriate ones for each test scenario, or create generic test steps
3. **Category Coverage**: Ensure test cases span all 8 testing categories (Positive, Negative, Edge, Performance, Security, Integration, Usability, Compatibility)
4. **No Duplication**: Avoid repetitive or near-identical test cases
5. **Exhaustive Coverage**: If one sentence implies multiple scenarios, create separate test cases for each
6. **Command/Interface Relevance**: Only use commands/interfaces that are actually relevant and applicable to each test scenario, or create appropriate generic steps

COMMAND/INTERFACE SELECTION STRATEGY:
- **Analyze Before Selection**: Review all available commands/interfaces (if provided) to understand their purpose and functionality
- **Context-Based Matching**: Select commands/interfaces based on what the test step is trying to accomplish, or create generic actions
- **Logical Sequencing**: Arrange commands/actions in a logical order within test steps (setup → execute → verify → cleanup)
- **Appropriate Integration**: Integrate commands naturally into descriptive test steps using backticks, or use generic action descriptions

INTEGRATION EXAMPLES:

**With CLI Commands Available:**
- "Access the {self.SYSTEM_TYPE} CLI interface using `{self.CONNECTION_METHOD} {self.LOGIN_CREDENTIALS}` and authenticate with appropriate credentials"
- "Configure the feature by executing appropriate commands if available in {self.ACCESS_MODE}"
- "Verify configuration status using appropriate commands if available and confirm expected parameter values"
- "Execute functionality test using appropriate commands if available to validate behavior"

**Without CLI Commands Available:**
- "Access the {self.SYSTEM_TYPE} management interface through the provided connection method and authenticate with appropriate credentials"
- "Configure the feature using the system's configuration interface and set the required parameters"
- "Verify configuration status by checking the system's status display and confirm expected parameter values"
- "Execute functionality test through the system's testing interface to validate behavior"
- "Enable diagnostic monitoring through the system's monitoring tools to capture detailed operational information"
- "Export configuration data using the system's export functionality for documentation purposes"

COMMAND/INTERFACE MATCHING PROCESS:
1. **Identify Test Step Purpose**: Understand what each test step is trying to achieve
2. **Check Available Resources**: Look through the provided commands/interfaces list (if available) for relevant options
3. **Select Best Match**: Choose the command/interface that best fits the step's objective, or create generic action
4. **Integrate Naturally**: Write the test step in natural language with the command inline (if available) or generic action description
5. **Verify Relevance**: Ensure the selected command/interface actually makes sense for the scenario, or that the generic action is appropriate

EXPECTED OUTPUT:
Return a valid JSON array of test cases where:
- Commands/interfaces are selected from the provided list only (if available), or generic actions are created
- Commands/actions are integrated inline within descriptive test steps
- Each test case includes commands Used and verificationMethods arrays
- Commands/actions are chosen based on test scenario requirements, not arbitrary categories

Do not skip or merge scenarios — extract comprehensive test cases from every relevant aspect of the documentation while making intelligent use of the available commands/interfaces or creating appropriate generic test steps.""",
                "System Prompt": f"""You are a test case generation expert for {self.SYSTEM_TYPE} systems with adaptive testing approach expertise.

🎯 ADAPTIVE ROLE:
1. **Test Case Generation**: Analyze technical documentation and generate comprehensive test cases
2. **Command Integration**: Intelligently select and integrate relevant commands/interfaces when available, or create generic test steps when not available

✅ TEST GENERATION RESPONSIBILITIES:
- Carefully analyze every sentence, condition, and rule in the document content
- Translate each into one or more meaningful test cases
- Ensure no major point from the document is missed
- Create unique test cases across all testing categories
- Intelligently match and integrate relevant commands/interfaces when provided, or create appropriate test steps based on system capabilities

✅ REQUIRED TEST CASE CATEGORIES:
1. **POSITIVE TESTING** - Expected behavior with valid inputs
2. **NEGATIVE TESTING** - Invalid input, misconfiguration, or failure recovery
3. **EDGE CASES** - Boundary conditions and unusual inputs
4. **PERFORMANCE TESTING** - Load, stress, and scalability scenarios
5. **SECURITY TESTING** - Authentication, authorization, access control
6. **INTEGRATION TESTING** - Interactions across components or features
7. **USABILITY TESTING** - Simplicity, clarity, and user experience
8. **COMPATIBILITY TESTING** - Version, platform, or environment variations

📋 COMMAND/INTERFACE INTEGRATION STRATEGY:
- **Analyze Available Resources**: Study the provided commands/interfaces list (if available) to understand available operations
- **Adaptive Integration**: If commands are available, intelligently select relevant ones. If not available, create appropriate test steps based on system capabilities
- **Command Categories (Auto-Detect - when commands are available)**:
  - Configuration commands (keywords: configure, set, enable, create, add, modify)
  - Verification commands (keywords: show, display, get, list, describe, status, info)
  - Test commands (keywords: test, ping, check, validate, verify, execute)
  - Debug commands (keywords: debug, trace, log, monitor, troubleshoot)
  - Data commands (keywords: copy, export, backup, save, dump, archive)
  - Reset commands (keywords: reset, clear, delete, remove, restore, reload)

- **Non-CLI Test Step Creation (when commands are not available)**:
  - Use generic action verbs: "Access", "Configure", "Verify", "Test", "Monitor", "Validate"
  - Include specific system interfaces: GUI, API endpoints, configuration files, logs
  - Reference system-specific tools and methods mentioned in documentation

🧾 OUTPUT FORMAT (Strict JSON):
Each test case must include:
1. "testCaseId" (e.g., "TC-001")
2. "testCaseTitle"
3. "testObjective"
4. "preConditions"
5. "testSteps" (as a numbered array of strings with commands/interfaces inline when available, or generic actions)
6. "expectedResult"
7. "testData" (optional; include if relevant)
8. "priority" ("High", "Medium", "Low")
9. "testType" (one of: Positive, Negative, Edge, Performance, Security, Integration, Usability, Compatibility)
10. "commandsUsed" (array of commands/interfaces used - empty if no commands available)
11. "verificationMethods" (array of verification methods used from available list or generic methods)

📋 EXAMPLE TEST CASE FORMAT:
```
{{
  "testCaseId": "TC-001",
  "testCaseTitle": "Verify {self.PRIMARY_FEATURE} Configuration and Status",
  "testObjective": "Validate that {self.PRIMARY_FEATURE} can be properly configured and verified using CLI commands",
  "preConditions": "1. {self.SYSTEM_TYPE} is powered on and accessible\n2. Administrator credentials are available\n3. Network connectivity is established",
  "testSteps": [
    "Connect to the {self.SYSTEM_TYPE} CLI interface using `{self.CONNECTION_METHOD} {self.LOGIN_CREDENTIALS}` and authenticate with administrative privileges",
    "Enter {self.ACCESS_MODE} by executing the appropriate mode command from the available CLI commands",
    "Configure {self.PRIMARY_FEATURE} using appropriate commands if available with required parameters",
    "Verify the configuration by running appropriate commands if available and analyze the output",
    "Test functionality using appropriate commands if available to validate operational behavior",
    "Document results and exit configuration mode using appropriate commands"
  ],
  "expectedResult": "1. CLI connection established successfully\n2. Configuration commands execute without errors\n3. Verification commands show expected {self.PRIMARY_FEATURE} status\n4. Test commands demonstrate proper functionality\n5. System responds as documented",
  "testData": "{self.PRIMARY_FEATURE} configuration parameters, expected CLI output patterns, test validation criteria",
  "priority": "High",
  "testType": "Positive",
  "cliCommands": appropriate commands if available,
  "verificationCommands": appropriate commands if available
}}
```

⚠️ CRITICAL REQUIREMENTS:
- Generate the **maximum possible number of test cases** from the content
- **ONLY use CLI commands that exist in the provided commands list**
- Test Steps must be:
  - Written in natural language with CLI commands inline using backticks
  - Descriptive and specific to the scenario
  - Include connection, configuration, verification, and cleanup where applicable
- **Smart Command Selection**: Choose commands based on test scenario requirements, not predefined categories

🚫 AVOID Generic Steps Like:
- "Step 1: {self.CONNECTION_METHOD} {self.LOGIN_CREDENTIALS}"
- "Step 2: Run command"

✅ INSTEAD Use Descriptive Steps Like:
- "Log in to the {self.SYSTEM_TYPE} via {self.CONNECTION_METHOD} using `{self.CONNECTION_METHOD} {self.LOGIN_CREDENTIALS}` to begin the configuration"
- "Verify {self.PRIMARY_FEATURE} status by running appropriate commands if available and confirm the expected configuration parameters"
- "Enable debugging with appropriate commands if available to monitor system behavior during testing"

🔧 INTELLIGENT CLI COMMAND SELECTION GUIDELINES:
1. **For Setup/Configuration Steps**: Look for commands containing keywords like "configure", "set", "enable", "create"
2. **For Verification Steps**: Look for commands containing keywords like "show", "display", "status", "list", "get"
3. **For Testing Steps**: Look for commands containing keywords like "test", "ping", "check", "validate"
4. **For Debugging Steps**: Look for commands containing keywords like "debug", "trace", "log", "monitor"
5. **For Data Collection**: Look for commands containing keywords like "copy", "export", "backup", "save"
6. **For Cleanup Steps**: Look for commands containing keywords like "reset", "clear", "delete", "restore"

📊 ENHANCED TEST STEP EXAMPLES:
- "Connect to the {self.SYSTEM_TYPE} management interface using `{self.CONNECTION_METHOD} {self.LOGIN_CREDENTIALS}` and authenticate with admin privileges"
- "Check current system status by executing appropriate commands if available and record baseline settings"
- "Configure the target feature using appropriate commands if available with appropriate parameters and verify command acceptance"
- "Monitor system behavior during configuration changes using appropriate commands if available to capture detailed operational data"
- "Validate configuration persistence by running appropriate commands if available after system changes"

⚡ COMMAND MATCHING STRATEGY:
- Analyze each test scenario to understand what type of operation is needed
- Search the available commands list for commands that match the required operation
- Select the most appropriate command based on context and functionality
- If multiple commands are suitable, choose the most specific or comprehensive one
- Include expected output patterns in expectedResult when using verification commands
- Ensure selected commands are logically sequenced within test steps

🎯 COVERAGE REQUIREMENTS:
- Cover **every major sentence and condition** in the content
- Avoid duplicate or near-duplicate test cases
- Be exhaustive: if one sentence implies multiple scenarios, create multiple test cases
- Ensure CLI commands are relevant and executable for each test scenario
- Extract maximum value from both the documentation and available CLI commands"""

            }
        
        }

        self.setWindowTitle("5G Radio Network Test Automation Framework")
        self.setGeometry(100, 100, 1200, 800)

        self.current_prompt_key = None  # Keeps track of the currently selected prompt


        # Main widget layout
        main_widget = QWidget()
        self.main_layout = QHBoxLayout(main_widget)  # Horizontal layout for left navigation and main content
        self.setCentralWidget(main_widget)

         # Uniform size policy for input fields
        self.input_size_policy = QSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        # Left-Side Navigation Layout
        self.create_left_navigation()

        # Main Content Area
        self.content_layout = QVBoxLayout()
        self.main_layout.addLayout(self.content_layout)

        # Top Bar with Title and Home Button
        self.create_top_bar()

        # Content Area (Center)
        self.content_widget = QStackedWidget()
        self.content_widget.setStyleSheet(f"""
            QStackedWidget {{
                background-color: {THEME['colors']['background']};
                border: none;
                padding: 0;
            }}
        """)
        self.content_layout.addWidget(self.content_widget)

        # Add pages to content area
        self.home_page = self.create_home_page()
        self.tasks_page = self.create_tasks_page()
        self.prompt_templates_page = self.create_prompt_templates_page()  # Page for prompt templates
        self.user_history_page = self.create_user_history_page()  # Page for user history
        self.explore_page = self.create_explore_page()  # Added Explore Page
        self.eval_page = self.create_eval_page()  # Added Evaluation Page

        self.content_widget.addWidget(self.home_page)  # Index 0
        self.content_widget.addWidget(self.tasks_page)  # Index 1
        self.content_widget.addWidget(self.prompt_templates_page)  # Index 2
        self.content_widget.addWidget(self.user_history_page)  # Index 3
        self.content_widget.addWidget(self.explore_page)  # Index 4 for Explore Page
        self.content_widget.addWidget(self.eval_page)  # Index 5 for Evaluation Page

        # Show the home page by default
        self.content_widget.setCurrentIndex(0)

        # Store the buttons for easy reference
        self.buttons = {
            "tasks": self.tasks_button,
            "prompt_templates": self.prompt_button,
            "user_history": self.history_button,
            "eval": self.eval_button,
        }

        # Template dialog reference
        self.template_dialog = None

        # Set the API key
        
        #os.environ['OPENAI_API_KEY'] = self.api_key
        #self.client = OpenAI(api_key=self.api_key)

        # Store latest loaded dataset and generated response
        self.latest_dataset_content = None  # Store latest loaded dataset
        self.latest_generated_response = None  # Store latest generated response
        # Store latest loaded dataset and generated response end
        self.previous_response = None  # <-- Add this line to ensure it's always initialized
        self.code_testing_engine = CodeTestingEngine()
        self.spec_reference_evaluator = SpecReferenceEvaluator()
        self.llm_judge_evaluator = LLMJudgeEvaluator()

    def create_left_navigation(self):
        """Create the vertical navigation bar with modern styling."""
        nav_frame = QFrame()
        nav_frame.setFixedWidth(240)
        nav_frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border-right: 1px solid {THEME['colors']['border']};
                padding: 16px 8px;
            }}
        """)
        nav_layout = QVBoxLayout(nav_frame)
        nav_layout.setSpacing(2)
        nav_layout.setContentsMargins(8, 8, 8, 8)
        nav_layout.setAlignment(Qt.AlignTop)

        # Title
        title = QLabel("5G Testing")
        title.setStyleSheet(THEME['typography']['nav_title'])
        nav_layout.addWidget(title)

        # Navigation Label
        nav_label = QLabel("NAVIGATION")
        nav_label.setStyleSheet(THEME['typography']['nav_label'])
        nav_layout.addWidget(nav_label)

        # Home Button with icon
        self.home_button = QPushButton("🏠 Home")
        self.home_button.setStyleSheet(NAV_BUTTON_STYLE)
        self.home_button.setCheckable(True)
        self.home_button.clicked.connect(lambda: self.display_page(0))
        nav_layout.addWidget(self.home_button)

        # Tasks Button with icon
        self.tasks_button = QPushButton("⚡ Tasks")
        self.tasks_button.setStyleSheet(NAV_BUTTON_STYLE)
        self.tasks_button.setCheckable(True)
        self.tasks_button.clicked.connect(lambda: self.display_page(1))
        nav_layout.addWidget(self.tasks_button)

        # Prompt Templates Button with icon
        self.prompt_button = QPushButton("📝 Prompt Templates")
        self.prompt_button.setStyleSheet(NAV_BUTTON_STYLE)
        self.prompt_button.setCheckable(True)
        self.prompt_button.clicked.connect(lambda: self.display_page(2))
        nav_layout.addWidget(self.prompt_button)

        # User History Button with icon
        self.history_button = QPushButton("📊 User History")
        self.history_button.setStyleSheet(NAV_BUTTON_STYLE)
        self.history_button.setCheckable(True)
        self.history_button.clicked.connect(lambda: self.display_page(3))
        nav_layout.addWidget(self.history_button)

        self.eval_button = QPushButton("🧮 Response Evaluation")
        self.eval_button.setStyleSheet(NAV_BUTTON_STYLE)
        self.eval_button.setCheckable(True)
        self.eval_button.clicked.connect(lambda: self.display_page(5))
        nav_layout.addWidget(self.eval_button)

        nav_layout.addStretch()
        self.main_layout.addWidget(nav_frame)

    def create_top_bar(self):
        """Create the top bar with enterprise styling and window controls."""
        top_bar = QHBoxLayout()
        top_bar.setContentsMargins(8, 0, 8, 0)
        
        # Title
        self.title_label = QLabel("5G RAN Testing Framework")
        self.title_label.setStyleSheet(THEME['typography']['title'])
        self.title_label.setAlignment(Qt.AlignLeft)
        top_bar.addWidget(self.title_label)
        
        # Add stretch to push window controls to the right
        top_bar.addStretch()
        
        # Window control buttons
        control_layout = QHBoxLayout()
        control_layout.setSpacing(8)
        
        # Minimize button
        minimize_btn = QPushButton("−")
        minimize_btn.setFixedSize(30, 30)
        minimize_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                color: #323130;
                font-size: 16px;
            }
            QPushButton:hover {
                background-color: #E6E6E6;
            }
        """)
        minimize_btn.clicked.connect(self.showMinimized)
        control_layout.addWidget(minimize_btn)
        
        # Maximize button
        maximize_btn = QPushButton("□")
        maximize_btn.setFixedSize(30, 30)
        maximize_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                color: #323130;
                font-size: 16px;
            }
            QPushButton:hover {
                background-color: #E6E6E6;
            }
        """)
        maximize_btn.clicked.connect(self.toggle_maximize)
        control_layout.addWidget(maximize_btn)
        
        # Close button
        close_btn = QPushButton("×")
        close_btn.setFixedSize(30, 30)
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                color: #323130;
                font-size: 16px;
            }
            QPushButton:hover {
                background-color: #FF4444;
                color: white;
            }
        """)
        close_btn.clicked.connect(self.close)
        control_layout.addWidget(close_btn)
        
        top_bar.addLayout(control_layout)
        self.content_layout.addLayout(top_bar)

    def toggle_maximize(self):
        """Toggle between maximized and normal window state."""
        if self.isMaximized():
            self.showNormal()
        else:
            self.showMaximized()

    def mousePressEvent(self, event):
        """Handle mouse press events for window dragging."""
        if event.button() == Qt.LeftButton:
            self.drag_position = event.globalPos() - self.frameGeometry().topLeft()
            event.accept()

    def mouseMoveEvent(self, event):
        """Handle mouse move events for window dragging."""
        if event.buttons() == Qt.LeftButton and hasattr(self, 'drag_position'):
            self.move(event.globalPos() - self.drag_position)
            event.accept()

    def display_page(self, page_index):
        """Switch the content of the main page and update navigation button states."""
        self.content_widget.setCurrentIndex(page_index)
        
        # Update button states
        self.home_button.setChecked(page_index == 0)
        self.tasks_button.setChecked(page_index == 1)
        self.prompt_button.setChecked(page_index == 2)
        self.history_button.setChecked(page_index == 3)
        self.eval_button.setChecked(page_index == 5)

    def create_home_page(self):
        """Create a simple and informative home page."""
        home_widget = QWidget()
        home_widget.setStyleSheet(f"background-color: {THEME['colors']['background']};")
        home_layout = QVBoxLayout()
        home_layout.setContentsMargins(24, 24, 24, 24)
        home_layout.setSpacing(24)

        # Welcome Section
        welcome_frame = QFrame()
        welcome_frame.setStyleSheet("""
            QFrame {
                background-color: white;
                border-radius: 8px;
                padding: 32px;
            }
        """)
        welcome_layout = QVBoxLayout(welcome_frame)
        welcome_layout.setSpacing(20)

        # Title
        title = QLabel("Gen AI Based 5G RAN Test Automation Framework")
        title.setStyleSheet("""
            QLabel {
                color: #323130;
                font-size: 28px;
                font-weight: bold;
            }
        """)
        welcome_layout.addWidget(title)

        # Description
        description = QLabel(
            "Streamline testing lifecycle with AI Driven Approach for End to End 5G Radio Access Network testing.\n"
            "Click below to start exploring the framework's features."
        )
        description.setWordWrap(True)
        description.setStyleSheet("""
            QLabel {
                color: #605E5C;
                font-size: 16px;
                line-height: 1.6;
                margin-bottom: 16px;
            }
        """)
        welcome_layout.addWidget(description)

        # Quick Start Buttons
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(16)

        # Dataset Generator Button
        dataset_btn = QPushButton("Dataset Generator")
        dataset_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 12px 24px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
        """)
        dataset_btn.clicked.connect(self.show_dataset_generator)
        buttons_layout.addWidget(dataset_btn)

        # Explore Framework Button
        explore_btn = QPushButton("Explore Framework")
        explore_btn.setStyleSheet("""
            QPushButton {
                background-color: #FFFFFF;
                color: #0078D4;
                border: 1px solid #0078D4;
                border-radius: 4px;
                padding: 12px 24px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #F0F8FF;
            }
        """)
        explore_btn.clicked.connect(lambda: self.display_page(1))
        buttons_layout.addWidget(explore_btn)

        buttons_layout.addStretch()
        welcome_layout.addLayout(buttons_layout)
        home_layout.addWidget(welcome_frame)

        # Add stretch to push everything to the top
        home_layout.addStretch()

        home_widget.setLayout(home_layout)
        return home_widget

    def update_system_stats(self):
        """Update the system statistics display."""
        try:
            # CPU Usage
            cpu_percent = psutil.cpu_percent()
            self.cpu_value.setText(f"{cpu_percent:.1f}%")

            # Memory Usage
            memory = psutil.virtual_memory()
            memory_used = memory.used / (1024 * 1024)  # Convert to MB
            memory_available = memory.available / (1024 * 1024)  # Convert to MB
            
            self.memory_value.setText(f"{memory_used:.0f} MB")
            self.available_value.setText(f"{memory_available:.0f} MB")
        except:
            pass  # Ignore any errors in getting system stats

    def generate_unit_tests(self):
        """Generate unit test scripts."""
        self.generate_test_scripts("Unit Tests")

    def generate_integration_tests(self):
        """Generate integration test scripts."""
        self.generate_test_scripts("Integration Tests")

    def generate_performance_tests(self):
        """Generate performance test scripts."""
        self.generate_test_scripts("Performance Tests")

    def generate_conformance_tests(self):
        """Generate conformance test scripts."""
        self.generate_test_scripts("Conformance Tests")

    def generate_test_scripts(self, test_type):
        """Generate test scripts based on the selected type."""
        progress = QProgressDialog(f"Generating {test_type}...", "Cancel", 0, 100, self)
        progress.setWindowModality(Qt.WindowModal)
        progress.setAutoClose(True)
        progress.setMinimumDuration(0)
        
        # Generate appropriate prompts based on test type
        prompts = {
            "Unit Tests": "Generate unit test cases for component validation",
            "Integration Tests": "Create end-to-end integration test scenarios",
            "Performance Tests": "Design performance benchmarking tests",
            "Conformance Tests": "Generate standard compliance test cases"
        }
        
        selected_prompt = prompts.get(test_type, "")
        
        # Update progress
        for i in range(101):
            if progress.wasCanceled():
                break
            progress.setValue(i)
            QThread.msleep(50)  # Simulate processing time
        
        if not progress.wasCanceled():
            QMessageBox.information(self, "Success", f"{test_type} generated successfully!")

    def handle_test_generation(self, test_type, input_text):
        """Handle the generation of test scripts based on type and input."""
        try:
            # Create a progress dialog
            progress = QProgressDialog("Generating test scripts...", None, 0, 100)
            progress.setWindowModality(Qt.WindowModal)
            progress.show()

            # Generate appropriate prompt based on test type
            prompt = self.create_test_generation_prompt(test_type, input_text)
            
            # Update progress
            progress.setValue(30)
            
            # Generate test script using OpenAI
            response = self.generate_response(prompt)
            
            progress.setValue(60)
            
            # Post-process the response
            processed_script = self.post_process_test_script(response)
            
            progress.setValue(90)
            
            # Save the generated test script
            self.save_generated_test(test_type, processed_script)
            
            progress.setValue(100)
            
            QMessageBox.information(self, "Success", f"{test_type.title()} test script generated successfully!")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate test script: {str(e)}")

    def create_test_generation_prompt(self, test_type, input_text):
        """Create appropriate prompt for test generation based on type."""
        prompts = {
            "unit": f"Generate a Python unit test script for the following function/description:\n{input_text}\n\nInclude:\n- Test class setup\n- Multiple test cases\n- Edge cases\n- Error handling\n- Mocking if needed",
            "integration": f"Generate a Python integration test script for the following scenario:\n{input_text}\n\nInclude:\n- Test environment setup\n- Dependencies handling\n- Multiple integration scenarios\n- Cleanup procedures",
            "performance": f"Generate a Python performance test script with these requirements:\n{input_text}\n\nInclude:\n- Performance metrics collection\n- Load simulation\n- Results analysis\n- Threshold validation",
            "conformance": f"Generate a Python test script for 3GPP conformance testing based on:\n{input_text}\n\nInclude:\n- 3GPP specification compliance\n- Protocol conformance checks\n- Message flow validation\n- Success/failure criteria"
        }
        return prompts.get(test_type, "")

    def create_getting_started_section(self):
        """Create the Getting Started section with step-by-step guidance."""
        frame = QFrame()
        frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 20px;
            }}
        """)
        layout = QVBoxLayout(frame)

        # Section title
        title = QLabel("Getting Started")
        title.setStyleSheet(THEME['typography']['title'])
        layout.addWidget(title)

        # Steps
        steps = [
            ("1. Configure Workspace", "Set up your working and output directories", "⚙️", self.show_dataset_generator),
            ("2. Generate Dataset", "Process and extract data from 5G specifications", "📊", self.show_dataset_generator),
            ("3. Create Test Scripts", "Generate automated test cases", "📝", self.show_test_script_generator),
            ("4. Run Tests", "Execute and monitor test cases", "▶️", self.show_test_execution)
        ]

        for step_title, desc, icon, handler in steps:
            step_btn = QPushButton(f"{icon} {step_title}")
            step_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {THEME['colors']['surface']};
                    border: 1px solid {THEME['colors']['border']};
                    border-radius: 4px;
                    padding: 12px;
                    text-align: left;
                    font-size: 14px;
                }}
                QPushButton:hover {{
                    background-color: {THEME['colors']['hover']};
                    border-color: {THEME['colors']['primary']};
                }}
            """)
            step_btn.clicked.connect(handler)
            layout.addWidget(step_btn)

            # Step description
            desc_label = QLabel(desc)
            desc_label.setStyleSheet(THEME['typography']['caption'])
            desc_label.setContentsMargins(32, 0, 0, 16)
            layout.addWidget(desc_label)

        return frame

    def create_recent_activity_section(self):
        """Create the Recent Activity section."""
        frame = QFrame()
        frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 20px;
                margin-bottom: 16px;
            }}
        """)
        layout = QVBoxLayout(frame)

        # Section title with view all button
        header_layout = QHBoxLayout()
        title = QLabel("Recent Activity")
        title.setStyleSheet(THEME['typography']['title'])
        header_layout.addWidget(title)

        view_all_btn = QPushButton("View All")
        view_all_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: transparent;
                border: none;
                color: {THEME['colors']['primary']};
                font-size: 13px;
            }}
            QPushButton:hover {{
                text-decoration: underline;
            }}
        """)
        view_all_btn.clicked.connect(lambda: self.display_page(3))  # Switch to history page
        header_layout.addWidget(view_all_btn)
        layout.addLayout(header_layout)

        # Activity list (placeholder)
        activity_list = QFrame()
        activity_list.setStyleSheet("background: transparent;")
        activity_layout = QVBoxLayout(activity_list)
        activity_layout.setSpacing(8)

        # Add some placeholder activities
        for _ in range(3):
            activity_item = QLabel("No recent activity")
            activity_item.setStyleSheet(THEME['typography']['caption'])
            activity_layout.addWidget(activity_item)

        layout.addWidget(activity_list)
        return frame

    def create_statistics_section(self):
        """Create the Statistics section."""
        frame = QFrame()
        frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 20px;
            }}
        """)
        layout = QVBoxLayout(frame)

        # Section title
        title = QLabel("Statistics")
        title.setStyleSheet(THEME['typography']['title'])
        layout.addWidget(title)

        # Stats grid
        stats_grid = QGridLayout()
        stats_grid.setSpacing(16)

        stats = [
            ("Tests Generated", "0", "📊"),
            ("Success Rate", "0%", "✅"),
            ("Active Tests", "0", "🔄"),
            ("Issues Found", "0", "🔍")
        ]

        for i, (label, value, icon) in enumerate(stats):
            stat_frame = QFrame()
            stat_frame.setStyleSheet(f"""
                QFrame {{
                    background-color: {THEME['colors']['background']};
                    border-radius: 4px;
                    padding: 12px;
                }}
            """)
            stat_layout = QVBoxLayout(stat_frame)

            # Stat label
            stat_label = QLabel(f"{icon} {label}")
            stat_label.setStyleSheet("""
                QLabel {
                    font-size: 13px;
                    color: #605E5C;
                }
            """)
            stat_layout.addWidget(stat_label)

            # Stat value
            stat_value = QLabel(value)
            stat_value.setStyleSheet("""
                QLabel {
                    font-size: 24px;
                    font-weight: bold;
                    color: #323130;
                }
            """)
            stat_layout.addWidget(stat_value)

            stats_grid.addWidget(stat_frame, i // 2, i % 2)

        layout.addLayout(stats_grid)
        return frame

    def create_quick_actions_section(self):
        """Create the Quick Actions section."""
        frame = QFrame()
        frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 20px;
            }}
        """)
        layout = QVBoxLayout(frame)

        # Section title
        title = QLabel("Quick Actions")
        title.setStyleSheet(THEME['typography']['title'])
        layout.addWidget(title)

        # Actions grid
        actions_grid = QGridLayout()
        actions_grid.setSpacing(16)

        actions = [
            ("Generate Dataset", "Process 5G specifications", "📊", self.show_dataset_generator),
            ("Create Test Scripts", "Generate automated test cases", "🔧", self.show_test_script_generator),
            ("View History", "Check previous test results", "📋", lambda: self.display_page(3)),
            ("Configure System", "Adjust framework settings", "⚙️", self.show_dataset_generator)
        ]

        for i, (title, desc, icon, handler) in enumerate(actions):
            action_btn = QPushButton(f"{icon} {title}")
            action_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {THEME['colors']['surface']};
                    border: 1px solid {THEME['colors']['border']};
                    border-radius: 4px;
                    padding: 16px;
                    text-align: left;
                    font-size: 14px;
                }}
                QPushButton:hover {{
                    background-color: {THEME['colors']['hover']};
                    border-color: {THEME['colors']['primary']};
                }}
            """)
            action_btn.clicked.connect(handler)
            
            # Create a widget to hold the button and description
            action_widget = QWidget()
            action_layout = QVBoxLayout(action_widget)
            action_layout.setSpacing(4)
            action_layout.addWidget(action_btn)
            
            desc_label = QLabel(desc)
            desc_label.setStyleSheet(THEME['typography']['caption'])
            action_layout.addWidget(desc_label)
            
            actions_grid.addWidget(action_widget, i // 2, i % 2)

        layout.addLayout(actions_grid)
        return frame

    def show_quick_start_guide(self):
        """Display the quick start guide dialog."""
        guide = QDialog(self)
        guide.setWindowTitle("Quick Start Guide")
        guide.setFixedSize(600, 400)
        
        layout = QVBoxLayout()
        
        title = QLabel("Getting Started with 5G RAN Testing Framework")
        title.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 16px;")
        layout.addWidget(title)
        
        steps = [
            "1. Configure your workspace and settings",
            "2. Generate or import test datasets",
            "3. Create test scripts using the generator",
            "4. Execute tests and view results",
            "5. Analyze and export test reports"
        ]
        
        for step in steps:
            step_label = QLabel(step)
            step_label.setStyleSheet("font-size: 14px; margin: 8px 0;")
            layout.addWidget(step_label)
        
        close_btn = QPushButton("Got it!")
        close_btn.clicked.connect(guide.close)
        layout.addWidget(close_btn)
        
        guide.setLayout(layout)
        guide.exec_()

    def create_tasks_page(self):
        """Create the tasks page with modern card-based layout."""
        tasks_widget = QWidget()
        tasks_widget.setStyleSheet(f"background-color: {THEME['colors']['background']};")
        tasks_layout = QVBoxLayout()
        tasks_layout.setContentsMargins(24, 24, 24, 24)
        tasks_layout.setSpacing(24)

        # Title section
        title = QLabel("Framework Components")
        title.setStyleSheet(THEME['typography']['title'])
        tasks_layout.addWidget(title)

        subtitle = QLabel("Explore and interact with the testing framework modules")
        subtitle.setStyleSheet(THEME['typography']['subtitle'])
        tasks_layout.addWidget(subtitle)

        # Main content area with cards
        content_frame = QFrame()
        content_frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 24px;
            }}
        """)
        content_layout = QVBoxLayout(content_frame)
        content_layout.setSpacing(16)

        # Framework Components Grid
        components_grid = QGridLayout()
        components_grid.setSpacing(16)
        components_grid.setContentsMargins(0, 0, 0, 0)

        # Define framework components with their handlers
        components = [
            {
                "title": "Dataset Generator",
                "description": "Process 5G specifications",
                "icon": "📊",
                "handler": self.show_dataset_generator
            },
            {
                "title": "Test Script Generator",
                "description": "Generate test scripts",
                "icon": "🔧",
                "handler": self.show_test_script_generator
            },
            {
                "title": "Test Deployment",
                "description": "Deploy test environment",
                "icon": "🚀",
                "handler": self.show_test_deployment
            },
            {
                "title": "Test Execution",
                "description": "Run test scenarios",
                "icon": "▶️",
                "handler": self.show_test_execution
            },
            {
                "title": "Bug Discovery",
                "description": "Analyze test results",
                "icon": "🔍",
                "handler": self.show_bug_discovery
            },
            {
                "title": "Code Evaluation",
                "description": "Evaluate code changes",
                "icon": "🧪",
                "handler": self.show_code_testing
            },
            {
                "title": "Code Assistant",
                "description": "Review and optimize",
                "icon": "💻",
                "handler": self.show_code_assistant
            }
        ]

        # Create cards for components
        for i, component in enumerate(components):
            card = QFrame()
            card.setStyleSheet(f"""
                QFrame {{
                    background-color: {THEME['colors']['surface']};
                    border: 1px solid {THEME['colors']['border']};
                    border-radius: 8px;
                    padding: 16px;
                    cursor: pointer;
                }}
                QFrame:hover {{
                    border-color: {THEME['colors']['primary']};
                    background-color: {THEME['colors']['hover']};
                }}
            """)
            card.setMinimumHeight(210)
            card.setMaximumHeight(210)
            card.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
            card_layout = QVBoxLayout(card)
            card_layout.setSpacing(8)

            # Icon and title in horizontal layout
            header_layout = QHBoxLayout()
            icon_label = QLabel(component["icon"])
            icon_label.setStyleSheet("font-size: 24px;")
            header_layout.addWidget(icon_label)
            
            title_label = QLabel(component["title"])
            title_label.setStyleSheet(THEME['typography']['body'] + "font-weight: 600;")
            header_layout.addWidget(title_label)
            header_layout.addStretch()
            
            card_layout.addLayout(header_layout)

            desc_label = QLabel(component["description"])
            desc_label.setStyleSheet(THEME['typography']['caption'])
            card_layout.addWidget(desc_label)

            # Add a button to launch the component
            launch_btn = QPushButton("Launch")
            launch_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {THEME['colors']['primary']};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 8px 16px;
                    font-size: 13px;
                }}
                QPushButton:hover {{
                    background-color: {THEME['colors']['secondary']};
                }}
            """)
            launch_btn.clicked.connect(component["handler"])
            card_layout.addWidget(launch_btn)

            row = i // 3
            components_grid.addWidget(card, row, i % 3)
            components_grid.setRowMinimumHeight(row, 230)

        components_widget = QWidget()
        components_widget.setLayout(components_grid)

        components_scroll = QScrollArea()
        components_scroll.setWidgetResizable(True)
        components_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        components_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        components_scroll.setWidget(components_widget)

        content_layout.addWidget(components_scroll)
        tasks_layout.addWidget(content_frame)
        tasks_widget.setLayout(tasks_layout)
        return tasks_widget

    def show_dataset_generator(self):
        """Show the dataset generator dialog with simplified UI."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Dataset Generator")
        dialog.setModal(True)
        dialog.setMinimumWidth(900)
        dialog.setMinimumHeight(1000)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)

        # Title
        title_layout = QHBoxLayout()
        title = QLabel("Dataset Generator")
        title.setStyleSheet("""
            font-size: 24px;
            font-weight: bold;
            color: #323130;
            margin-bottom: 12px;
        """)
        title_layout.addWidget(title)
        title_layout.addStretch()
        layout.addLayout(title_layout)

        # Input Configuration Section
        config_group = QFrame()
        config_group.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #E1E5EA;
                border-radius: 8px;
                padding: 16px;
            }
            QLabel {
                font-size: 14px;
                color: #323130;
                font-weight: 500;
            }
            QLineEdit {
                font-size: 14px;
                padding: 8px;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
            }
            QPushButton {
                font-size: 14px;
                padding: 8px 16px;
                background: #f3f3f3;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
            }
            QPushButton:hover {
                background: #e9e9e9;
            }
        """)
        config_layout = QVBoxLayout(config_group)

        # Working Directory
        work_dir_layout = QHBoxLayout()
        work_dir_label = QLabel("Working Directory:")
        work_dir_layout.addWidget(work_dir_label)
        self.work_dir_display = QLineEdit(self.working_directory)
        self.work_dir_display.setReadOnly(True)
        work_dir_layout.addWidget(self.work_dir_display)
        work_dir_btn = QPushButton("Browse")
        work_dir_btn.clicked.connect(self.select_working_directory)
        work_dir_layout.addWidget(work_dir_btn)
        config_layout.addLayout(work_dir_layout)

        # Output Directory
        output_dir_layout = QHBoxLayout()
        output_dir_label = QLabel("Output Directory:")
        output_dir_layout.addWidget(output_dir_label)
        self.output_dir_display = QLineEdit(self.output_directory)
        self.output_dir_display.setReadOnly(True)
        output_dir_layout.addWidget(self.output_dir_display)
        output_dir_btn = QPushButton("Browse")
        output_dir_btn.clicked.connect(self.select_output_directory)
        output_dir_layout.addWidget(output_dir_btn)
        config_layout.addLayout(output_dir_layout)

        layout.addWidget(config_group)

        # Document Selection Section
        doc_group = QFrame()
        doc_group.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #E1E5EA;
                border-radius: 8px;
                padding: 16px;
            }
            QLabel {
                font-size: 14px;
                color: #323130;
                font-weight: 500;
            }
            QPushButton {
                font-size: 14px;
                padding: 8px 16px;
                background: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
            }
            QPushButton:hover {
                background: #106EBE;
            }
            QComboBox {
                font-size: 14px;
                padding: 8px;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
        """)
        doc_layout = QVBoxLayout(doc_group)

        # Document Selection
        doc_select_layout = QHBoxLayout()
        self.document_label = QLabel("No document loaded")
        self.document_label.setStyleSheet("font-size: 14px; color: #666;")
        doc_select_layout.addWidget(self.document_label)
        load_doc_btn = QPushButton("Load Document")
        load_doc_btn.clicked.connect(self.load_document)
        doc_select_layout.addWidget(load_doc_btn)
        doc_layout.addLayout(doc_select_layout)

        # Section Selection
        section_layout = QHBoxLayout()
        section_label = QLabel("Section:")
        section_layout.addWidget(section_label)
        self.section_combo = QComboBox()
        self.section_combo.addItem("Select Section")
        self.section_combo.currentTextChanged.connect(self.populate_subsections)
        section_layout.addWidget(self.section_combo)
        doc_layout.addLayout(section_layout)

        # Subsection Selection
        subsection_layout = QHBoxLayout()
        subsection_label = QLabel("Subsection:")
        subsection_layout.addWidget(subsection_label)
        self.subsection_combo = QComboBox()
        self.subsection_combo.addItem("Select Subsection")
        subsection_layout.addWidget(self.subsection_combo)
        doc_layout.addLayout(subsection_layout)

        layout.addWidget(doc_group)

        # Results & Actions Section
        results_group = QFrame()
        results_group.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #E1E5EA;
                border-radius: 8px;
                padding: 16px;
            }
            QLabel {
                font-size: 14px;
                color: #323130;
            }
            QPushButton {
                font-size: 14px;
                padding: 8px 16px;
                background: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
            }
            QPushButton:hover {
                background: #106EBE;
            }
            QProgressBar {
                font-size: 12px;
                text-align: center;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
            }
            QProgressBar::chunk {
                background: #0078D4;
            }
            QTextEdit {
                font-size: 14px;
                font-family: 'Consolas', monospace;
                padding: 8px;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
            }
        """)
        results_layout = QVBoxLayout(results_group)

        # Status and Progress
        self.status_label = QLabel("Ready")
        results_layout.addWidget(self.status_label)

        self.progress_bar = QProgressBar()
        results_layout.addWidget(self.progress_bar)

        # Extract Button
        extract_btn = QPushButton("Extract Dataset")
        extract_btn.setStyleSheet("""
            QPushButton {
                font-size: 14px;
                padding: 12px 24px;
                background: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #106EBE;
            }
        """)
        extract_btn.clicked.connect(self.start_extraction)
        results_layout.addWidget(extract_btn)

        # Results Display
        self.results_text = QTextEdit()
        self.results_text.setReadOnly(True)
        self.results_text.setMinimumHeight(100)
        self.results_text.setStyleSheet("""
            QTextEdit {
                font-size: 14px;
                font-family: 'Consolas', monospace;
                padding: 12px;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
                color: #323130;
            }
        """)
        results_layout.addWidget(self.results_text)

        # Progress Bars
        progress_layout = QVBoxLayout()
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                font-size: 12px;
                text-align: center;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
                height: 24px;
            }
            QProgressBar::chunk {
                background: #0078D4;
            }
        """)
        progress_layout.addWidget(self.progress_bar)

        self.progress_bar1 = QProgressBar()
        self.progress_bar1.setStyleSheet("""
            QProgressBar {
                font-size: 12px;
                text-align: center;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
                height: 24px;
            }
            QProgressBar::chunk {
                background: #0078D4;
            }
        """)
        progress_layout.addWidget(self.progress_bar1)
        results_layout.addLayout(progress_layout)

        # Response Display
        self.response_display = QTextEdit()
        self.response_display.setReadOnly(True)
        self.response_display.setMinimumHeight(200)
        self.response_display.setStyleSheet("""
            QTextEdit {
                font-size: 14px;
                font-family: 'Consolas', monospace;
                padding: 12px;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
                color: #323130;
            }
        """)
        results_layout.addWidget(self.response_display)

        # Generate Response Button
        generate_response_btn = QPushButton("Generate Response")
        generate_response_btn.setStyleSheet("""
            QPushButton {
                font-size: 14px;
                padding: 12px 24px;
                background: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #106EBE;
            }
        """)
        generate_response_btn.clicked.connect(self.start_response)
        results_layout.addWidget(generate_response_btn)

        layout.addWidget(results_group)

        # Update description boxes and labels
        self.document_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                color: #323130;
                font-weight: normal;
            }
        """)

        self.section_combo.setStyleSheet("""
            QComboBox {
                font-size: 14px;
                padding: 8px;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
        """)

        self.subsection_combo.setStyleSheet("""
            QComboBox {
                font-size: 14px;
                padding: 8px;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
        """)

        self.status_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                color: #323130;
                font-weight: 500;
            }
        """)

        self.progress_bar.setStyleSheet("""
            QProgressBar {
                font-size: 12px;
                text-align: center;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                background: white;
                height: 24px;
            }
            QProgressBar::chunk {
                background: #0078D4;
            }
        """)

        dialog.exec_()
    
    def load_document(self):
        """Load a document and extract main sections."""
        file_name, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "Documents (*.docx *.pdf)")
        if file_name:
            self.doc_path = file_name
            self.document_label.setText(f"Loaded: {file_name}")
            main_sections = self.get_main_sections(file_name)
            self.populate_sections(main_sections)
    
    def populate_sections(self, main_sections):
        """Populate the section dropdown with main sections."""
        self.section_combo.clear()
        self.section_combo.addItem("Select Section")
        for section in main_sections:
            self.section_combo.addItem(section)
        self.sections = {section: [] for section in main_sections}

    

    def get_main_sections(self, doc_path):
        """Extract main sections (Heading 1) from a DOCX document."""
        if doc_path.endswith(".docx"):
            doc = Document(doc_path)
            main_sections = []
            for para in doc.paragraphs:
                if para.style.name == 'Heading 1':  # Check for Heading 1
                    main_sections.append(para.text.strip())
            return main_sections
        elif doc_path.endswith(".pdf"):
            return self.get_pdf_main_sections(doc_path)
        return []

    def get_pdf_main_sections(self, pdf_path):
        """Extract main sections from a PDF."""
        import fitz
        sections = []
        document = fitz.open(pdf_path)
        for page_num in range(document.page_count):
            page = document.load_page(page_num)
            text = page.get_text()
            lines = text.split("\n")
            for line in lines:
                if line.lower().startswith("section"):
                    sections.append(line.strip())
        return sections
    
    def populate_subsections(self):
        """Populate the subsection dropdown based on the selected main section."""
        selected_section = self.section_combo.currentText()
        if selected_section and selected_section != "Select Section":
            self.sections[selected_section] = self.get_subsections(selected_section)
            self.subsection_combo.clear()
            self.subsection_combo.addItem("Select Subsection")
            self.subsection_combo.addItems(self.sections[selected_section])

    def get_subsections(self, main_section):
        """Extract subsections (Heading 2) under the specified main section."""
        subsections = []
        if self.doc_path.endswith(".docx"):
            doc = Document(self.doc_path)
            in_main_section = False
            for para in doc.paragraphs:
                if para.style.name == 'Heading 1':
                    in_main_section = (para.text.strip() == main_section)
                elif in_main_section and para.style.name == 'Heading 2':
                    subsections.append(para.text.strip())
        return subsections

    def select_working_directory(self):
        """Allow user to select working directory."""
        dir_path = QFileDialog.getExistingDirectory(self, "Select Working Directory", self.working_directory)
        if dir_path:
            self.working_directory = dir_path
            self.settings.setValue('working_directory', dir_path)
            self.work_dir_display.setText(dir_path)
            # Create necessary subdirectories
            os.makedirs(os.path.join(dir_path, 'datasets'), exist_ok=True)
            os.makedirs(os.path.join(dir_path, 'test_scripts'), exist_ok=True)

    def select_output_directory(self):
        """Allow user to select output directory."""
        dir_path = QFileDialog.getExistingDirectory(self, "Select Output Directory", self.output_directory)
        if dir_path:
            self.output_directory = dir_path
            self.settings.setValue('output_directory', dir_path)
            self.output_dir_display.setText(dir_path)

    def get_working_directory(self):
        """Get the current working directory."""
        return self.working_directory

    def get_output_directory(self):
        """Get the current output directory."""
        return self.output_directory

    def show_test_script_generator(self):
        """Show the test script generator dialog."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Test Script Generator")
        dialog.setMinimumWidth(900)
        dialog.setMinimumHeight(750)
        dialog.setModal(False)  # Make dialog modeless to allow interaction with other windows
        
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # Dataset Section
        dataset_frame = QFrame()
        dataset_frame.setStyleSheet("QFrame { background-color: white; border: 1px solid #E0E0E0; border-radius: 4px; }")
        dataset_layout = QVBoxLayout(dataset_frame)
        dataset_layout.setContentsMargins(10, 10, 10, 10)
        
        # Dataset selection row
        dataset_row = QHBoxLayout()
        dataset_label = QLabel("Dataset:")
        dataset_label.setFixedWidth(60)
        dataset_row.addWidget(dataset_label)
        
        self.dataset_path = QLineEdit()
        self.dataset_path.setPlaceholderText("Select dataset files...")
        dataset_row.addWidget(self.dataset_path)
        
        # Progress bar for dataset loading (local to dialog)
        progress_bar = QProgressBar()
        progress_bar.setVisible(False)
        dataset_row.addWidget(progress_bar)
        
        load_dataset_btn = QPushButton("Load Dataset")
        load_dataset_btn.setStyleSheet(self.button_style())
        load_dataset_btn.setFixedWidth(100)
        load_dataset_btn.clicked.connect(lambda: self.load_dataset(dialog, progress_bar))
        dataset_row.addWidget(load_dataset_btn)
        
        dataset_layout.addLayout(dataset_row)
        layout.addWidget(dataset_frame)

        # Prompt Section
        prompt_frame = QFrame()
        prompt_frame.setStyleSheet("QFrame { background-color: white; border: 1px solid #E0E0E0; border-radius: 4px; }")
        prompt_layout = QVBoxLayout(prompt_frame)
        prompt_layout.setContentsMargins(10, 10, 10, 10)
        
        # Prompt selection
        prompt_header = QHBoxLayout()
        prompt_label = QLabel("Prompt Template:")
        prompt_label.setFixedWidth(100)
        prompt_header.addWidget(prompt_label)
        
        self.prompt_combo = QComboBox()
        self.prompt_combo.addItem("Select Template...")
        self.prompt_combo.addItem("Test Script")
        self.prompt_combo.addItem("Test Case")
        self.prompt_combo.addItem("Custom")
        self.prompt_combo.currentTextChanged.connect(self.handle_prompt)
        self.prompt_combo.setStyleSheet("""
            QComboBox {
                padding: 5px;
                border: 1px solid #E0E0E0;
                border-radius: 4px;
                background: white;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                image: url(down_arrow.png);
                width: 12px;
                height: 12px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #E0E0E0;
                selection-background-color: #E3F2FD;
                selection-color: #1976D2;
                background: white;
            }
        """)
        prompt_header.addWidget(self.prompt_combo)
        
        # View Templates button
        view_templates_btn = QPushButton("View Templates")
        view_templates_btn.setStyleSheet(self.button_style())
        view_templates_btn.setFixedWidth(120)
        view_templates_btn.clicked.connect(self.show_templates_dialog)
        prompt_header.addWidget(view_templates_btn)
        
        prompt_layout.addLayout(prompt_header)
        
        # Test Case variable dropdowns (6 dropdowns)
        self.test_case_variables_container = QWidget()
        self.test_case_variables_container.setVisible(False)  # Initially hidden
        test_case_variables_layout = QVBoxLayout(self.test_case_variables_container)
        test_case_variables_layout.setContentsMargins(0, 8, 0, 8)
        test_case_variables_layout.setSpacing(12)
        
        # Create two rows of dropdowns
        row1_layout = QHBoxLayout()
        row1_layout.setSpacing(10)
        row2_layout = QHBoxLayout()
        row2_layout.setSpacing(10)
        
        # Create compact dropdowns
        self.domain_combo = QComboBox()
        self.domain_combo.addItem("Select Domain...")
        self.domain_combo.addItems([
            "Network Infrastructure",
            "Wireless Communication", 
            "Cloud Platform"
        ])
        self.domain_combo.setStyleSheet(self.compact_combo_style())
        self.domain_combo.currentTextChanged.connect(self.update_prompt_variables)
        row1_layout.addWidget(self.domain_combo)
        
        self.system_type_combo = QComboBox()
        self.system_type_combo.addItem("Select System Type...")
        self.system_type_combo.addItems([
            "Network Element",
            "Access Point", 
            "5G",
            "4G"
        ])
        self.system_type_combo.setStyleSheet(self.compact_combo_style())
        self.system_type_combo.currentTextChanged.connect(self.update_prompt_variables)
        row1_layout.addWidget(self.system_type_combo)
        
        
        self.primary_feature_combo = QComboBox()
        self.primary_feature_combo.addItem("Select Primary Feature...")
        self.primary_feature_combo.addItems([
            "Attach",
            "Dettach",
            "Register",
            "Deregister",
            "Handover"
        ])
        self.primary_feature_combo.setStyleSheet(self.compact_combo_style())
        self.primary_feature_combo.currentTextChanged.connect(self.update_prompt_variables)
        row1_layout.addWidget(self.primary_feature_combo)
        
        # Row 2: Connection Variables
        self.connection_method_combo = QComboBox()
        self.connection_method_combo.addItem("Select Connection Method...")
        self.connection_method_combo.addItems([
            "SSH",
            "NETCONF", 
            "gRPC",
            "RESTCONF",
            "Console",
            "API",
            "CLI",
            "Telnet"
        ])
        self.connection_method_combo.setStyleSheet(self.compact_combo_style())
        self.connection_method_combo.currentTextChanged.connect(self.update_prompt_variables)
        row2_layout.addWidget(self.connection_method_combo)
        
        self.login_credentials_combo = QComboBox()
        self.login_credentials_combo.addItem("Select Login Credentials...")
        self.login_credentials_combo.addItems([
            "admin@<IP>",
            "root@<IP>"
        ])
        self.login_credentials_combo.setStyleSheet(self.compact_combo_style())
        self.login_credentials_combo.currentTextChanged.connect(self.update_prompt_variables)
        row2_layout.addWidget(self.login_credentials_combo)
        
        
        self.access_mode_combo = QComboBox()
        self.access_mode_combo.addItem("Select Access Mode...")
        self.access_mode_combo.addItems([
            "CLI mode",
            "config mode", 
            "configuration mode",
            "admin mode",
            "edit mode",
            "privileged mode",
            "exec mode"
        ])
        self.access_mode_combo.setStyleSheet(self.compact_combo_style())
        self.access_mode_combo.currentTextChanged.connect(self.update_prompt_variables)
        row2_layout.addWidget(self.access_mode_combo)
        
        # Add rows to the test case variables layout
        test_case_variables_layout.addLayout(row1_layout)
        test_case_variables_layout.addLayout(row2_layout)
        prompt_layout.addWidget(self.test_case_variables_container)
        
        # Test Script specific options (Language and File Upload)
        self.test_script_variables_container = QWidget()
        self.test_script_variables_container.setVisible(False)  # Initially hidden
        test_script_variables_layout = QVBoxLayout(self.test_script_variables_container)
        test_script_variables_layout.setContentsMargins(0, 8, 0, 8)
        test_script_variables_layout.setSpacing(12)
        
        test_script_options_layout = QHBoxLayout()
        test_script_options_layout.setSpacing(15)
        
        # Language dropdown
        language_label = QLabel("Language:")
        language_label.setStyleSheet("font-weight: bold; color: #1e3a5f;")
        test_script_options_layout.addWidget(language_label)
        
        self.language_combo = QComboBox()
        self.language_combo.addItem("Select Language...")
        self.language_combo.addItems(["Python", "C (Coming Soon)", "C++ (Coming Soon)"])
        self.language_combo.setStyleSheet(self.compact_combo_style())
        self.language_combo.currentTextChanged.connect(self.handle_language_selection)
        test_script_options_layout.addWidget(self.language_combo)
        
        # File upload button
        upload_label = QLabel("Reference Code:")
        upload_label.setStyleSheet("font-weight: bold; color: #1e3a5f;")
        test_script_options_layout.addWidget(upload_label)
        
        self.upload_btn = QPushButton("Upload File")
        self.upload_btn.setStyleSheet("""
            QPushButton {
                background-color: #1e3a5f;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-weight: bold;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #2c4a6b;
            }
            QPushButton:pressed {
                background-color: #153157;
            }
        """)
        self.upload_btn.clicked.connect(self.upload_reference_code)
        test_script_options_layout.addWidget(self.upload_btn)
        
        # File name label
        self.uploaded_file_label = QLabel("No file selected")
        self.uploaded_file_label.setStyleSheet("color: #666; font-style: italic;")
        test_script_options_layout.addWidget(self.uploaded_file_label)
        
        # Add stretch to push everything to the left
        test_script_options_layout.addStretch()
        
        test_script_variables_layout.addLayout(test_script_options_layout)
        prompt_layout.addWidget(self.test_script_variables_container)
        
        # Prompt text editor
        self.prompt_text_edit = QTextEdit()
        self.prompt_text_edit.setPlaceholderText("Select a template or enter a custom prompt...")
        self.prompt_text_edit.setReadOnly(True)  # Initially read-only
        self.prompt_text_edit.setStyleSheet("""
            QTextEdit {
                border: 1px solid #E0E0E0;
                border-radius: 4px;
                padding: 8px;
                background-color: #FAFAFA;
                font-family: Consolas, Monaco, monospace;
            }
        """)
        prompt_layout.addWidget(self.prompt_text_edit)
        
        # Connect textChanged to keep latest custom prompt updated
        self.prompt_text_edit.textChanged.connect(self.update_latest_custom_prompt)
        
        # Buttons for prompt management
        prompt_buttons = QHBoxLayout()
        
        self.modify_btn = QPushButton("Modify")
        self.modify_btn.setStyleSheet(self.button_style())
        self.modify_btn.setFixedWidth(80)
        self.modify_btn.clicked.connect(self.handle_modify_button)
        self.modify_btn.setEnabled(False)  # Initially disabled
        prompt_buttons.addWidget(self.modify_btn)
        
        self.save_template_btn = QPushButton("Save Template")
        self.save_template_btn.setStyleSheet(self.button_style())
        self.save_template_btn.setFixedWidth(100)
        self.save_template_btn.clicked.connect(self.handle_save_button)
        self.save_template_btn.setEnabled(False)  # Initially disabled
        prompt_buttons.addWidget(self.save_template_btn)
        
        # Generate button
        generate_btn = QPushButton("Generate Response")
        generate_btn.setStyleSheet(self.active_button_style())
        generate_btn.clicked.connect(self.start_response)
        prompt_buttons.addWidget(generate_btn)
        
        prompt_buttons.addStretch()
        prompt_layout.addLayout(prompt_buttons)
        
        layout.addWidget(prompt_frame)

        # Response Section
        response_frame = QFrame()
        response_frame.setStyleSheet("QFrame { background-color: white; border: 1px solid #E0E0E0; border-radius: 4px; }")
        response_layout = QVBoxLayout(response_frame)
        response_layout.setContentsMargins(10, 10, 10, 10)
        
        response_label = QLabel("Generated Test Script:")
        response_layout.addWidget(response_label)
        
        # Initialize response text widget
        self.response_text = QTextEdit()
        self.response_text.setReadOnly(True)
        self.response_text.setPlaceholderText("Generated test script will appear here...")
        self.response_text.setMinimumHeight(100)
        self.response_text.setMaximumHeight(200)
        self.response_text.setStyleSheet("""
            QTextEdit {
                background-color: #FAFAFA;
                border: 1px solid #E0E0E0;
                border-radius: 4px;
                padding: 10px;
                font-family: Consolas, Monaco, monospace;
                font-size: 12px;
            }
        """)
        response_layout.addWidget(self.response_text)
        response_layout.addSpacing(15)
        # --- Add New Prompt UI ---
        self.new_prompt_label = QLabel("New Prompt:")
        response_layout.addWidget(self.new_prompt_label)
        self.new_prompt_edit = QTextEdit()
        self.new_prompt_edit.setPlaceholderText("Enter a new prompt to refine the script...")
        self.new_prompt_edit.setFixedHeight(60)
        response_layout.addWidget(self.new_prompt_edit)
        self.new_prompt_btn = QPushButton("Generate with New Prompt")
        self.new_prompt_btn.setStyleSheet(self.active_button_style())
        self.new_prompt_btn.clicked.connect(self.generate_with_new_prompt)
        response_layout.addWidget(self.new_prompt_btn)
        # --- End New Prompt UI ---
        
        # Save Response button
        save_response_btn = QPushButton("Save Test Script")
        save_response_btn.setStyleSheet(self.button_style())
        save_response_btn.clicked.connect(self.save_response)
        response_layout.addWidget(save_response_btn)
        
        layout.addWidget(response_frame)
        
        dialog.setLayout(layout)
        dialog.show()

    def start_response(self):
        """Start the response generation process."""
        try:
            # Get the text content and selected prompt
            text_content = self.text_content
            if not text_content:
                self.show_warning("Please load a dataset first.")
                return

            selected_prompt = self.get_current_prompt()
            if not selected_prompt:
                self.show_warning("Please enter a prompt or select a template.")
                return

            # Create and configure the progress dialog
            self.progress_dialog = QProgressDialog("Generating response...", "Cancel", 0, 100, self)
            self.progress_dialog.setWindowModality(Qt.WindowModal)
            self.progress_dialog.setAutoClose(True)
            self.progress_dialog.setAutoReset(True)
            self.progress_dialog.show()

            # Create the extraction thread
            self.extraction_thread = ExtractionThread(selected_prompt, self)
            
            # Connect signals using instance variables to maintain references
            self.progress_connection = self.extraction_thread.progress_signal.connect(self.update_progress)
            self.finished_connection = self.extraction_thread.finished.connect(self.handle_response_complete)
            self.error_connection = self.extraction_thread.error_signal.connect(self.show_error)
            
            # Store text_content and selected_prompt in the thread instance
            self.extraction_thread.text_content = text_content
            self.extraction_thread.selected_prompt = selected_prompt
            
            # Start the thread
            self.extraction_thread.start()

        except Exception as e:
            self.show_error(f"Error starting response generation: {str(e)}")
            if hasattr(self, 'progress_dialog'):
                self.progress_dialog.close()

    def update_progress(self, value):
        """Update progress dialog with proper error handling."""
        try:
            if hasattr(self, 'progress_dialog') and self.progress_dialog:
                self.progress_dialog.setValue(value)
                QApplication.processEvents()  # Keep UI responsive
        except Exception as e:
            print(f"Error updating progress: {str(e)}")

    def handle_response_complete(self):
        """Handle completion of response generation."""
        try:
            if hasattr(self, 'progress_dialog'):
                self.progress_dialog.close()
            
            if hasattr(self, 'extraction_thread') and self.extraction_thread.result:
                self.update_response_display(self.extraction_thread.result)
                
        except Exception as e:
            self.show_error(f"Error handling response completion: {str(e)}")
        finally:
            if hasattr(self, 'extraction_thread'):
                self.extraction_thread.deleteLater()

    def update_response_display(self, response_text):
        """Update the response display with proper validation."""
        try:
            if not hasattr(self, 'response_text'):
                self.show_error("Response display not initialized.")
                return False
                
            if not response_text:
                self.response_text.setPlainText("No response generated.")
                return False
                
            self.response_text.setPlainText(response_text)
            self.latest_generated_response = response_text  # <-- Store latest response
            self.previous_response = response_text  # <-- Store as previous response for iterative refinement
            self.response_text.verticalScrollBar().setValue(0)  # Scroll to top
            return True
            
        except Exception as e:
            self.show_error(f"Error updating response display: {str(e)}")
            return False

    def start_extraction(self):
        """Starts the extraction process by calling recursive_test_graph_attach.py as a subprocess."""
        selected_section = self.section_combo.currentText()
        selected_subsection = self.subsection_combo.currentText()
        doc_path = getattr(self, 'doc_path', None)
        if not doc_path or selected_section == "Select Section" or selected_subsection == "Select Subsection":
            QMessageBox.warning(self, "Invalid Selection", "Please select a valid document, section, and subsection.")
            return

        # Prepare output file path
        output_file = os.path.join(self.output_directory, "output.json")
        script_path = os.path.join(os.path.dirname(__file__), "recursive_test_graph_attach.py")
        # If script is not in the same directory, adjust path accordingly
        if not os.path.exists(script_path):
            script_path = os.path.abspath("recursive_test_graph_attach.py")
        if not os.path.exists(script_path):
            QMessageBox.critical(self, "Error", f"Could not find recursive_test_graph_attach.py at {script_path}")
            return

        # Build the command
        cmd = [
            sys.executable, script_path,
            output_file,
            self.working_directory,
            f"--file_path={doc_path}",
            f"--section={selected_section}",
            f"--subsection={selected_subsection}"
        ]

        # Show progress dialog
        progress = QProgressDialog("Extracting dataset...", None, 0, 0, self)
        progress.setWindowModality(Qt.WindowModal)
        progress.setAutoClose(True)
        progress.setMinimumDuration(0)
        progress.show()
        QApplication.processEvents()

        try:
            import subprocess
            result = subprocess.run(cmd, capture_output=True, text=True, check=False)
            progress.close()
            if result.returncode == 0:
                self.results_text.setPlainText(f"Extraction completed successfully.\nOutput saved to: {output_file}\n\n{result.stdout}")
                QMessageBox.information(self, "Extraction Status", "Extraction Completed Successfully!")
            else:
                self.results_text.setPlainText(f"Extraction failed.\n\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
                QMessageBox.critical(self, "Extraction Failed", f"Extraction failed.\n\n{result.stderr}")
        except Exception as e:
            progress.close()
            QMessageBox.critical(self, "Error", f"Failed to run extraction script: {str(e)}")

    def update_progress1(self, value):
        """Updates the progress bar and ensures UI refresh."""
        self.progress_bar.setValue(value)
        QApplication.processEvents()  # Forces UI update
    
    def show_message(self, message):
        """Shows a message box with the result of the extraction."""
        QMessageBox.information(self, "Extraction Status", message)
    
    def increment_progress(self):
        """Increment progress bar gradually."""
        if self.progress < 100:
            self.progress += 10  # Increase progress
            self.progress_bar.setValue(self.progress)


    #Jaysvi Code integration
    def extract_data(self, update_progress):
        selected_section = self.section_combo.currentText()
        selected_subsection = self.subsection_combo.currentText()
        #selected_prompt = self.prompt_combo.currentText()  # for prompt
        if selected_section and selected_subsection and selected_section != "Select Section" and selected_subsection != "Select Subsection":
            section_text = self.extract_text_from_subsection(selected_subsection)
            self.section_text = section_text  # Save extracted text to instance variable
            output_file= r"C:\Users\ChanduVangala\Documents\GenAI\extract\output.json"
            directory= r"C:\Users\ChanduVangala\Documents\GenAI\extract"
            #print("Extracted Text:\n", section_text)  # Print the extracted text for debugging
            update_progress(10)
            # Extract references and clauses using OpenAI API
            references, clauses, text = self.extract_references_and_clauses(section_text)

            # Filter out non-reference entries from references
            references = [ref for ref in references if re.match(r'^3GPP TS \d+\.\d+', ref)]

            # Debug prints for references and clauses
            print("References:", references)
            print("Clauses:", clauses)
            update_progress(30)
            download_directory = r"C:\Users\ChanduVangala\Documents\GenAI\extract"  # Your specified download directory
            if not os.path.exists(download_directory):
                os.makedirs(download_directory)
            self.download_specifications_from_references(references, download_directory)
            # Check for reference documents in the directory using OpenAI API
            update_progress(50)
            present_references, missing_references = self.check_references_in_directory(references, directory)

            if missing_references:
                print("Missing References:", missing_references)
            else: 
                print("All references are present in the directory.")
            update_progress(70)
            # Create a graph from the document, references, and clauses
            #graph = build_graph(text)
            graph=self.build_graph_from_input(text)
            #graph = create_graph_with_references_and_clauses(file_path, references, clauses)
            self.save_graph_to_json(graph, output_file)
            print("jayasvi before calling extract")
            section_text = self.extract_references_and_clauses1(graph, references, clauses, present_references, directory)
            self.save_graph_to_json(graph, output_file)
            update_progress(90)
            print(f"Graph saved to {output_file}")   
            update_progress(100)
        return "Extraction Completed Successfully!"
    #Jayasvi code integration end

    def extract_text_from_subsection(self, selected_subsection):
        """Extract text from the specified subsection."""
        if self.doc_path.endswith(".docx"):
            return self.extract_text_from_docx_subsection(self.doc_path, selected_subsection)
        elif self.doc_path.endswith(".pdf"):
            return self.extract_text_from_pdf_subsection(self.doc_path, selected_subsection)
        return ""
    
    def extract_text_from_docx_subsection(self, doc_path, selected_subsection):
        """Extract text from the specified subsection in DOCX file."""
        doc = Document(doc_path)
        section_text = ""
        extract = False
        for para in doc.paragraphs:
            if para.style.name == 'Heading 2':
                if para.text == selected_subsection:
                    extract = True
                    section_text += para.text + "\n"
                elif extract:
                    break  # Stop when reaching the next Heading 2
            elif extract:
                section_text += para.text + "\n"
        return section_text

    def extract_text_from_pdf_subsection(self, pdf_path, selected_subsection):
        """Extract text from the specified subsection in PDF file."""
        import fitz
        section_text = ""
        extract = False
        document = fitz.open(pdf_path)
        for page_num in range(document.page_count):
            page = document.load_page(page_num)
            text = page.get_text()
            lines = text.split("\n")
            for line in lines:
                if selected_subsection in line:
                    extract = True
                if extract:
                    section_text += line + "\n"
                    if line.startswith("Clause "):  # Assuming next clause starts with 'Clause '
                        break
        return section_text
    
    def extract_selected_subsection_text(self, status_label):
        """Extract text from the selected subsection."""
        selected_subsection = self.subsection_combo.currentText()
        if selected_subsection and selected_subsection != "Select Subsection":
            self.text = self.extract_text_from_subsection(selected_subsection)
            #self.result_label.setText(f"Extracted Text:\n{text}")
            status_label.setText("Status: Extracting...")
            QTimer.singleShot(1000, lambda: status_label.setText("Status: Dataset Created Successfully"))
        else:
            QMessageBox.warning(self, "Warning", "Please select a subsection to extract.")

    #Jayasvi code definations start
    def extract_references_and_clauses(self,text):
        """Extracts 3GPP references and associated clauses from the given text using OpenAI API."""
        messages = [
            {"role": "system", "content": "Extract 3GPP references and clauses seperately from the following text,dont add any special characters as cosmetic to the output, keep the clauses under respective references"},
            {"role": "user", "content": text}
        ]
        response = client.chat.completions.create(
            #model="gpt-3.5-turbo",
            #model="gpt-4o-mini",
            model=AZURE_OPENAI_MODEL_NAME,
            messages=messages,
            max_tokens=500,  # Increased token limit for potentially longer responses
            temperature=0.5
        )
        output = response.choices[0].message.content.strip()
        print("OpenAI Response for ref:", output)  # Debug print to see the response
        references, clauses = self.parse_openai_output(output)
        print("before returning to 1")
        return references, clauses, output

    def parse_openai_output(self,output):
        """Parses the OpenAI API output to extract references and clauses."""
        references = []
        clauses = []
        #print(f"jayasvi {output}")
        # Prompt to send to the OpenAI API
        prompt = f"""
        Extract the references and clauses from the following text. Separate the references and clauses clearly under headings "References" and "Clauses", dont assign any serial numbers:

        {output}
        """

        # Get the response from the OpenAI API
        api_output = self.get_openai_output(prompt)
        print("OpenAI Response for api_output:", api_output)
        #Split the output by lines and process
        lines = api_output.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if line.startswith("**References**") or line.startswith("References"):
                current_section = "references"
                continue
            elif line.startswith("**Clauses**:") or line.startswith("Clauses"):
                current_section = "clauses"
                continue

            if current_section == "references":
                if line:
                    references.append(line)
            elif current_section == "clauses":
                if line:
                    clauses.append(line)
        print("OpenAI Response for api_output after")
        return references, clauses
    

    def save_graph_to_json(self,graph, output_file):
        """Saves the graph structure to a JSON file."""
        print("chandu inside save_graph_to_json function")
        #data = nx.readwrite.json_graph.node_link_data(graph)
        data = nx.readwrite.json_graph.node_link_data(graph, edges="edges")
        with open(output_file, 'w') as f:
            json.dump(data, f, indent=4)
    # Function to find the latest ZIP file within the release
    def get_latest_zip_file(self,url):
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        zip_files = [a['href'] for a in soup.find_all('a', href=True) if a['href'].endswith('.zip')]

        if not zip_files:
            raise ValueError(f"No ZIP files found in the directory: {url}")
    
        latest_zip = sorted(zip_files, reverse=True)[0]
        return latest_zip
    
    def extract_references_and_clauses1(self,graph, references, clauses, present_references, directory):
        """Processes the graph to extract text from specified clauses of reference files.
        
        
        # Add nodes and edges to the graph
        for node in present_references:
            G.add_node(node)
            print("jayasvi node:", node)
        for i in range(len(clauses)):
            if i < len(present_references):
                G.add_edge(present_references[i], clauses[i])
                print("jayasvi node1:", present_references[i], clauses[i])    
        # Process the graph
        for node in G.nodes:
            for clause in nx.neighbors(G, node):
                file_path = os.path.join(directory, node)
        """
        text=''
        phrase=''
        G = nx.DiGraph()
        start_node = "start"
        G.add_node(start_node)  # Add the start node

    # Add reference nodes and connect to the start node
        for node in present_references:
            G.add_node(node)  # Add each reference node
            G.add_edge(start_node, node)  # Connect start node to reference node

    # Add edges from reference nodes to corresponding clauses
        for i in range(len(clauses)):
            if i < len(present_references):
                print("jayasvi1")
                G.add_edge(present_references[i], clauses[i]) # Reference to clause
        if len(present_references) == 1 or present_references[0] == present_references[1]:
            for i in range(len(clauses)):
                if i < len(clauses):
                    print("jayasvi2", present_references[0], clauses[i], type(present_references[0]), type(clauses[i]))
                    G.add_edge(present_references[0], clauses[i])
                    #print("jayasvi3")
                    print("Edge added:", present_references[0], "->", clauses[i])
                    print("Current edges in graph:", list(G.edges()))
        graph = G
    # Process the graph
        print("All nodes in the graph:", graph.nodes)
        for node in graph.nodes:
            if node != start_node:
                print("jayasvi node ",  node)
        # Check neighbors (connected clauses) of the current reference node
                for clause in nx.neighbors(graph, node):
                    print(f"jayasvi start node: {node}")
                    if(node):
                        print(f"jayasvi start1 node: {node}")
                        node1 = self.get_filename(node)
                        print (f"jayasvi in extract2nd: {node1} , clause: {clause}")
                        if(node1):
                            file_path = os.path.join(directory, node1)
                            clause = re.search(r'\b\d+(\.\d+)+\b', clause).group()
                            print(f"jayasvi before calling read from clause {clause} from file {file_path} for node {node1}")
                            text = self.read_clause_from_file(file_path, clause)
                            if text:
                                #print(f"Extracted text from {node} for clause {clause}:", text)
                                references, clauses, dummy_text = self.extract_references_and_clauses(text)
                                print("Jayasvi References:", references)
                                print("Jayasvi Clauses:", clauses)
                                ref_dict = self.extract_keywords_with_lines(text)

                                phrase = self.match_extract_from_ref(ref_dict)
                                print(f"jayasvi phrases {ref_dict}")
                                #if (phrase):
                                #    references,dummy_clauses, dummy_text = extract_references_and_clauses(phrase)
                                #present, missing = find_files_for_extracted_phrases(phrase)
                                matched_section=self.process_reference_extract_sections(phrase)
                                print("Jayasvi after level5 References:", references)
                                text=self.append_clause_sections(text,file_path,clauses)
                                print("Jayasvi before keywords:", clause)
                                base_dir = Path("C:\\Users\\JayasviBattu\\Downloads\\dataset")
                                file_name = f"{clause}_file.txt"
                                file_path = os.path.join(base_dir, file_name)
                                with open(file_path, "w") as file:
                                    file.write(text)
                                    file.write(matched_section)
                                test_scenarios=self.generate_test_scenarios(text)
                                #print("Generated Test Scenarios:\n")
                                #print(test_scenarios)
                                print(f"jayasvi end node: {node1}")
                            else:
                                print(f"No text found in {node} for clause {clause}.\n")
        return text
    



    def get_filename(self,main_part):
        # Extract the main part of the reference for comparison (e.g., "23.401" from "3GPP TS 23.401 [29]")
        files_in_directory = os.listdir(directory)
        print("main_part0",main_part)
        #main_part = re.search(r'\d+\.\d+', main_part)
        main_part=re.search(r'\d+-[a-zA-Z0-9]+', main_part)
        print("main_part",main_part)
        if main_part:
            main_part = main_part.group().replace('.', '')
            print("main_part1",main_part)
            """for file_name in files_in_directory:
                    # Extract numeric part of the file name for comparison (e.g., "23401" from "23401-i60.docx")
                file_numeric_part = re.search(r'\d+', file_name)
                print("file_numeric_part",file_numeric_part)
                if file_numeric_part and main_part in file_numeric_part.group():
                    print("file_name",file_name)
                    return file_name  # Store the exact file name with extension"""
            for file_name in files_in_directory:
              match = re.search(r'\d+-[a-zA-Z0-9]+', file_name)
              if match and match.group() == main_part:
                return file_name

    def read_clause_from_file(self,file_path, clause):
        """Reads the specified clause (section) from the given DOCX file by searching through sections/subsections."""
        try:
            print(f"Attempting to read DOCX file: {file_path}")  # Debug: file path

            # Load the DOCX file
            doc = Document(file_path)

            #Identify sections and subsections in the document
            sections = self.find_sections(doc)
            #print("jayasvi sections:", sections)
            print("Clauses:", clause)
            # Search for the clause within the identified sections
            clause_text = self.search_clause_in_sections(sections, clause)
            #print(f"jayasvi return {clause_text}")
            if clause_text:
                print(f"Extracted text from {file_path} for clause '{clause}':")
            else:
                clause_text = f"Clause '{clause}' not found in the document."

        except Exception as e:
            clause_text = f"Error reading {file_path}: {e}"
    
        return clause_text
    
    def find_sections(self,doc):
        sections = {}
        current_section = None
        current_content = []

        for paragraph in doc.paragraphs:
            # Check the paragraph style
            style = paragraph.style.name

            # Determine if the paragraph is a section heading
            if style in ['Heading 0', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
                if current_section is not None:
                    # Store the content of the previous section
                    sections[current_section] = '\n'.join(current_content)
            
                # Start a new section
                current_section = paragraph.text
                current_content = []
                #print(f"Jayasvi in Section: {current_section}")
            else:
                # Append content to the current section
                if current_section is not None:
                    current_content.append(paragraph.text)
    
        # Don't forget to store the last section's content
        if current_section is not None:
            sections[current_section] = '\n'.join(current_content)
            if current_section == '5.3.4B.6':
                print(f"Jayasvi in Section: {current_section}")
                print(f"Jayasvi in Section1: {sections[current_section]}")
        return sections


    def search_clause_in_sections(self,sections, clause):
        """Searches for the specified clause in the section titles and contents."""
        clause = clause.replace("Clause", "").strip()  # Adjust clause text for matching
        print(f"jayasvi {clause}")
        section_title_number=''
        #clause_number = clause.split(" ", 1)[0]  # Extract the number part of the clause
        #clause_number = clause.split("Clause")[1].strip().split(" ")[0]
        #match = re.search(r"\b\d+\.\d+\.\d+\.\d+\b", clause) or re.search(r"\b\d+(\.\d+)*[A-Za-z]?\b", clause)
        match = re.search(r"\b\d+(\.\d+)*[A-Za-z]?(\.\d+)?\b", clause) or re.search(r"\b\d+\.\d+\.\d+\.\d+\b", clause)
        clause_number = match.group(0)
        print(f"jayasvi clause number: {clause_number}")
        #print(f"jayasvi {sections}")
    
        for section in sections:
            #print(f"jayasvi inside loop")
            #section_title_number = data['title'].split("\t")[0].strip()  # Extract the section number from title
            match = re.search(r'(\d+(\.\d+)*)', section)
            if (match):
                section_title_number = match.group(1)
                #print(f"jayasvi section_titile_number {section_title_number}")
            #print(f"Checking Section {section} with title: {data['title']}")  # Debug: Print section title
            #print(f"jayasvi: clause_number {clause_number} and matched in Section is {section_title_number} ")
            # Check if clause number matches the section title number
            if clause_number == section_title_number:
                print(f"Clause '{clause}' matched in Section {section}.")
                #result = query_openai_for_clause(data['content'], clause)
                result = sections[section]
                return result
            #else:
                #print(f"Clause '{clause}' not found in any section.")


    def extract_keywords_with_lines(self,text):
 
        keyword_pattern = r"TS\s*\d{2}\.\d{3}"
        # Initialize an empty dictionary to store the results
        result_dict = defaultdict(list)
        sentences = re.split(r'(?<=\.)\s+',text)
        # Loop through each line in the text
        for sentence in sentences:
            # Search for the keyword pattern in the line
            match = re.search(keyword_pattern, sentence,re.IGNORECASE)
            #print("jayasvi keyword match",match, len(text))
            if match:
                ts_code = match.group(0)
                # Use the matched keyword as the key and store the whole line as the value
                #if ts_code in result_dict:
                #result_dict[ts_code].append(sentence.strip())
                #else:
                result_dict[ts_code] = sentence.strip()
                #print(f"Keyword: {match.group()}\nLine: {sentence.strip()}\n")
        return result_dict
    

    def match_extract_from_ref(self, ref_dict):
        """Extract technical phrases from references using OpenAI API."""
        for key, value in ref_dict.items():
            print(f"Key: {key}")
            print(f"Value: {value}")
            print("-" * 30)
        technical_phrase='' 
        extracted_phrases = []
        for key, value in ref_dict.items():
        # Prepare the prompt for OpenAI API
            prompt = f"Extract a meaningful 3GPP technical phrase from the following text:\n\n{value}"
            
            # Make an API call to ChatGPT for completion
            response = client.chat.completions.create(
                #model="gpt-4o-mini",
                model=AZURE_OPENAI_MODEL_NAME,
                temperature=0.7,
                messages=[
                    {"role": "system", "content": "You are an assistant specialized in 3GPP technical documentation."},
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Extract the content from the API response
            technical_phrase = response.choices[0].message.content.strip()
            formatted_string = f"Key: {key}, Extracted Phrase: {technical_phrase}"
            extracted_phrases.append(formatted_string)
            print(formatted_string)
                # Print the extracted phrase along with the key
            #print(f"Key: {key}, Extracted Phrase: {technical_phrase}")
        return extracted_phrases


    def process_reference_extract_sections(self, phrases):
        """Process reference sections using the working directory."""
        matched_sections = ''
        files_in_directory = os.listdir(self.get_working_directory())
        
        for phrase in phrases:
            match = re.search(r"Key: TS\s*(\d+\.\d+)", phrase)
            if match:
                main_part = match.group(1).replace('.', '')
                matched = False
                doc_path = None
                
                for file_name in files_in_directory:
                    file_numeric_part = re.search(r'\d+', file_name)
                    if file_numeric_part and main_part in file_numeric_part.group():
                        doc_path = os.path.join(self.get_working_directory(), file_name)
                        
                if doc_path and os.path.exists(doc_path):
                    doc = Document(doc_path)
                    sections = self.find_sections(doc)
                    
                    for section in sections:
                        extracted_phrase = re.search(r"Extracted Phrase:\s*(.*)", phrase).group(1)
                        if self.is_phrase_in_section(extracted_phrase, section, threshold=0.5):
                            matched = True
                            matched_sections += "\n" + "Section:" + "\t" + section + "\t"
                            matched_sections += sections[section]
                            if matched:
                                break
                                
        return matched_sections
    
    def append_clause_sections(self,text, filepath, clauses):
        doc = Document(filepath)
        print ("jayasvi in append clause")
        for clause in clauses:
            print ("jayasvi in append clause for:", clause )
            return_text = self.read_clause_from_file(filepath, clause)
            #print("jayasvi after return text")
            #print(return_text)
            text += "\n" + "Clause:" + "\t" + clause + "\t"
            text += return_text
        with open("total_content.txt", "w") as file:
            file.write(text)
        #print(text)
        return text
    
    def generate_test_scenarios(self,input_text):
        """
        Function to generate test scenarios according to 3GPP standards.
    
        Args:
        - input_text (str): The input text that describes the scenario or specification.
    
        Returns:
        - str: The generated test scenarios.
        """
        messages = [
            {"role": "system", "content": "You are an expert in 3GPP standards."},
            {"role": "user", "content": f"Based on the following description/text/input, generate all possible test scenarios python scripts for every section or case or scenario including edge cases as per 3GPP standards, take the redirected /referred clauses from the same text:\n\n{input_text}"}
        ]   

        try:
            response = client.chat.completions.create(
                #model="gpt-4o-mini",  # or use "gpt-4-32k" for larger inputs
                model=AZURE_OPENAI_MODEL_NAME,
                messages=messages,
                max_tokens=1000,  # Adjust based on the desired length of output
                temperature=0.7,  # Controls randomness, 0.7 is good for balanced output
            )

        # Extract the generated scenarios from the API response
            test_scenarios = response.choices[0].message.content
            return test_scenarios

        except openai.error.OpenAIError as e:
            print(f"OpenAI API error: {e}")
            return None


    def is_phrase_in_section(self,phrase, section, threshold=0.5):
        """
        Check if a phrase is in a section with partial matching.
    
        Args:
            phrase (str): The phrase to search for.
            section (str): The section of text in which to search.
            threshold (float): Similarity threshold for a match (0 to 1).
        
        Returns:
            bool: True if the phrase is found in the section with a similarity above the threshold, False otherwise.
        """
        # Tokenize the section into sentences for more granular matching
        sentences = section.split('. ')
    
        for sentence in sentences:
            similarity = SequenceMatcher(None, phrase.lower(), sentence.lower()).ratio()
            if similarity >= threshold:
                return True  # Found a similar enough match

        return False  # No sufficient match found
    
    # Main function to download and extract the latest specification
    def download_and_extract_specification(self, series, spec_number, download_directory):
        base_url = f"https://www.3gpp.org/ftp/Specs/archive/{series}_series/{spec_number}/"
        print(f"Base URL: {base_url}")
    
        latest_zip_file = self.get_latest_zip_file(base_url)
    
        # Use urljoin to correctly handle URL construction
        latest_zip_url = urljoin(base_url, latest_zip_file)
    
        print(f"Final URL with ZIP file: {latest_zip_url}")
        print(f"Downloading {latest_zip_file} from {latest_zip_url}")
    
        # Download the ZIP file
        response = requests.get(latest_zip_url)
    
        # Extract just the filename from the latest_zip_file
        zip_file_name = os.path.basename(latest_zip_url)
    
        # Save the file
        file_path = os.path.join(download_directory, zip_file_name)
        with open(file_path, 'wb') as file:
            file.write(response.content)
    
        print(f"Downloaded: {file_path}")
    
        # Unzip the downloaded file
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(download_directory)
                print(f"Extracted contents to: {download_directory}")
        except zipfile.BadZipFile:
            print(f"Error: The file downloaded is not a valid ZIP file: {file_path}")
    
        # Remove the ZIP file after extraction
        os.remove(file_path)
        print(f"Removed ZIP file: {file_path}")

    # Function to parse the input references and download the corresponding specifications
    def download_specifications_from_references(self,references, download_directory):
        for reference in references:
            match = re.search(r"3GPP TS (\d+)\.(\d+)", reference)
            if match:
                series = match.group(1)
                spec_number = f"{match.group(1)}.{match.group(2)}"
                print(f"Processing: Series {series}, Specification {spec_number}")
                try:
                    self.download_and_extract_specification(series, spec_number, download_directory)
                except ValueError as e:
                    print(f"Error: {e}")
            else:
                print(f"Invalid reference format: {reference}")

    def check_references_in_directory(self, references, directory):
        """Checks if the reference documents are present in the specified directory using regex for pattern matching."""
        present_references = []
        missing_references = []
        
        try:
            directory_path = Path(directory)
            if not directory_path.exists():
                directory_path.mkdir(parents=True, exist_ok=True)
            
            files_in_directory = list(directory_path.glob("*.docx")) + list(directory_path.glob("*.pdf"))
            
            for reference in references:
                matched = False
                # Extract the main part of the reference for comparison (e.g., "23.401" from "3GPP TS 23.401 [29]")
                main_part = re.search(r'\d+\.\d+', reference)
                if main_part:
                    main_part = main_part.group().replace('.', '')
                    for file_path in files_in_directory:
                        # Extract numeric part of the file name for comparison
                        file_numeric_part = re.search(r'\d+', file_path.name)
                        if file_numeric_part and main_part in file_numeric_part.group():
                            present_references.append(file_path.name)
                            matched = True
                            break
                if not matched:
                    missing_references.append(reference)
            
            print("Present References Files:", present_references)
            return present_references, missing_references
            
        except Exception as e:
            print(f"Error checking references: {str(e)}")
            return [], references
    
    def build_graph_from_input(self,input_text):
        """
        Function to build a graph from the given input text, where references are nodes and clauses are edges.
    
        Args:
        - input_text (str): The input text containing references and clauses.

        Returns:
        - G (networkx.DiGraph): The directed graph created from the input.
        """
        # Initialize a directed graph
        G = nx.DiGraph()

        # Add a start node
        G.add_node("start")

        current_ref = None
    
    # Process each line in the input text
        for line in input_text.splitlines():
            line = line.strip()
            if line.startswith("- 3GPP TS"):  # This identifies a reference node
                current_ref = line[2:].strip()
                G.add_node(current_ref)
                G.add_edge("start", current_ref)
            elif line.startswith("- Clause") and current_ref:  # This identifies a clause (edge)
                clause = line[2:].strip()
                G.add_node(clause)
                G.add_edge(current_ref, clause)

        return G
    
    def get_openai_output(self, prompt):
        """Get output from OpenAI API."""
        response = client.chat.completions.create(
            #model="gpt-4o",
            model=AZURE_OPENAI_MODEL_NAME,
            messages=[
                {"role": "system", "content": "You are an expert assistant specialized in analyzing and generating responses based on technical documentation and specifications. Provide detailed, accurate responses that directly address the user's requirements."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    

    def generate_response_from_text(self, update_progress, text_content, selected_prompt):
        """Generate test scripts with improved structure and error handling."""
        try:
            update_progress(30)
            
            # Create the prompt based on whether it's a template or custom prompt
            if self.current_prompt_key == "Custom":
                prompt = f'''
You are an expert test engineer.

ONLY use the information provided in the following dataset to answer the user's request.

DATASET:
{text_content}

USER REQUEST:
{selected_prompt}

INSTRUCTIONS:
- Do NOT use any generic knowledge.
- All test cases must be directly based on the dataset content above.
- Reference specific data, values, or steps from the dataset.
- If the dataset is not sufficient, say so.
'''         
            elif self.current_prompt_key == "Test Case":
                # For Test Case, use the merged prompt directly with dataset
                prompt = f'''
DATASET:
{text_content}

{selected_prompt}
'''
                self.testcases_list = "Test Case"
            elif self.current_prompt_key == "Test Script":
                prompt = f'''
You are an expert in 5G network test automation.

Based ONLY on the following dataset, generate a complete Python test script.

DATASET:
{text_content}

INSTRUCTIONS:
{selected_prompt}
Note: Make Sure the generated scripts covers the testcases which are in {self.testcases_name}
'''
            else:
                prompt = f"""
Based on the following specification and template, generate a complete Python test script.
Include proper test class, setup/teardown methods, and individual test cases.
Use unittest framework and follow best practices for test automation.

Specification:
{text_content}

Template:
{selected_prompt}

Requirements:
1. Follow the template structure exactly
2. Include proper exception handling
3. Add detailed comments and docstrings
4. Include setup and teardown methods
5. Add logging for test execution
6. Include test result validation
7. Follow 3GPP test specifications if applicable
"""
            
            update_progress(50)  # Update progress before API call
            
            # Generate response using OpenAI
            response = self.generate_response(prompt)
            if not response or "Error:" in response:
                return f"Error generating response: {response}"
                
            update_progress(70)  # Update progress after API call
            
            # Post-process the response
            processed_response = self.post_process_test_script(response)
            
            update_progress(90)  # Update progress after processing
            
            return processed_response
            
        except Exception as e:
            return f"An error occurred while generating the test script: {str(e)}"

    def post_process_test_script(self, script_content):
        """Return the script as generated by the LLM, no extra prepending."""
        if not script_content:
            return "Error: No script content to process"
        return script_content

    def generate_response(self, prompt):
        """Generate response using OpenAI API with proper error handling."""
        try:
            print("=== PROMPT SENT TO OPENAI ===")
            print(prompt)
            print("=============================")
            # Make the API call with proper parameters
            response = client.chat.completions.create(
                #model="gpt-3.5-turbo",
                #model="gpt-4o-mini",
                model=AZURE_OPENAI_MODEL_NAME,
                messages=[
                    {"role": "system", "content": "You are an expert test automation engineer specializing in 5G network testing."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000,
                top_p=1.0,
                frequency_penalty=0.0,
                presence_penalty=0.0
            )

            # Extract and validate the response
            if response and hasattr(response, 'choices') and response.choices:
                if len(response.choices) > 0 and hasattr(response.choices[0], 'message'):
                    content = response.choices[0].message.content
                    if content and isinstance(content, str):
                        return content.strip()
            
            return "Error: Invalid response structure"

        except Exception as e:
            error_msg = f"Error in response generation: {str(e)}"
            print(error_msg)  # Log the error
            return error_msg

    def save_response_to_file(self, response, folder_path):
        """Save generated test scripts in a structured test suite format."""
        try:
            # Create base test suite directory if it doesn't exist
            test_suite_dir = Path(folder_path) / "test_suite"
            test_suite_dir.mkdir(parents=True, exist_ok=True)

            # Create subdirectories based on test type
            if self.current_prompt_key:
                test_type = self.current_prompt_key.lower().replace(" ", "_")
                test_type_dir = test_suite_dir / test_type
                test_type_dir.mkdir(exist_ok=True)

                # Generate timestamp for unique file naming
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

                # Save as JSON if test type is 'test_case', else as .py
                if test_type == "test_case":
                    test_script_file = test_type_dir / f"{test_type}_{timestamp}.json"
                    # Try to parse response as JSON, else save as raw string
                    import json
                    try:
                        parsed = json.loads(response)
                        with open(test_script_file, "w", encoding="utf-8") as file:
                            json.dump(parsed, file, indent=2)
                    except Exception:
                        with open(test_script_file, "w", encoding="utf-8") as file:
                            file.write(response)
                else:
                    test_script_file = test_type_dir / f"{test_type}_{timestamp}.py"
                    with open(test_script_file, "w", encoding="utf-8") as file:
                        file.write("#!/usr/bin/env python\n")
                        file.write("# -*- coding: utf-8 -*-\n\n")
                        file.write(f"# Test Type: {self.current_prompt_key}\n")
                        file.write(f"# Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                        file.write("import unittest\n")
                        file.write("import sys\n")
                        file.write("import os\n\n")
                        file.write(response)

                # Create or update test suite configuration
                self.update_test_suite_config(test_type_dir, test_script_file.name)

                return True, test_script_file
            else:
                return False, "No test type selected"
        except Exception as e:
            return False, str(e)

    def update_test_suite_config(self, test_type_dir, script_name):
        """Update the test suite configuration file with new test script."""
        config_file = test_type_dir / "test_suite_config.json"
        
        # Load existing configuration or create new
        if config_file.exists():
            with open(config_file, "r") as f:
                config = json.load(f)
        else:
            config = {
                "suite_name": test_type_dir.name,
                "created_date": datetime.now().strftime("%Y-%m-%d"),
                "test_scripts": []
            }

        # Add new test script to configuration
        script_entry = {
            "name": script_name,
            "created_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "status": "new"
        }
        config["test_scripts"].append(script_entry)

        # Save updated configuration
        with open(config_file, "w") as f:
            json.dump(config, f, indent=4)

    def load_dataset(self, parent_dialog, progress_bar=None):
        """Load dataset from files with multiple file selection support."""
        try:
            file_dialog = QFileDialog(parent_dialog)
            file_dialog.setFileMode(QFileDialog.ExistingFiles)  # Enable multiple file selection
            file_dialog.setNameFilter(
                "Supported Files (*.txt *.pdf *.docx *.xlsx *.xls);;Text Files (*.txt);;"
                "PDF Files (*.pdf);;Word Files (*.docx);;Excel Files (*.xlsx *.xls);;All Files (*.*)"
            )
            
            if file_dialog.exec_():
                selected_files = file_dialog.selectedFiles()
                if not selected_files:
                    return
                    
                # Update dataset path display
                self.dataset_path.setText(f"{len(selected_files)} file(s) selected")
                
                # Show progress bar
                if progress_bar is not None:
                    progress_bar.setVisible(True)
                    progress_bar.setValue(0)
                
                # Process each file
                combined_content = []
                total_files = len(selected_files)
                
                for i, file_path in enumerate(selected_files):
                    try:
                        # Update progress
                        if progress_bar is not None:
                            progress = int((i + 1) / total_files * 100)
                            progress_bar.setValue(progress)
                        
                        # Process file based on type
                        if file_path.lower().endswith((".xlsx", ".xls")):
                            content = self.load_excel_file(file_path)
                        elif file_path.lower().endswith('.pdf'):
                            content = self.load_pdf_file(file_path)
                        elif file_path.lower().endswith('.docx'):
                            content = self.load_docx_file(file_path)
                        else:
                            content = None
                            encodings = ['utf-8', 'latin-1', 'cp1252']
                            for encoding in encodings:
                                try:
                                    with open(file_path, 'r', encoding=encoding) as f:
                                        content = f.read()
                                    break
                                except UnicodeDecodeError:
                                    continue
                            if content is None:
                                raise UnicodeDecodeError(
                                    f"Could not decode file {os.path.basename(file_path)} with any of the attempted encodings: {encodings}")
                        
                        if content:
                            combined_content.append(content)
                        
                    except Exception as e:
                        QMessageBox.warning(
                            parent_dialog,
                            "Warning",
                            f"Error processing file {os.path.basename(file_path)}: {str(e)}"
                        )
                
                # Hide progress bar
                if progress_bar is not None:
                    progress_bar.setVisible(False)
                
                if combined_content:
                    # Store the combined content
                    self.text_content = '\n\n'.join(combined_content)
                    self.latest_dataset_content = self.text_content  # <-- Store latest dataset
                    # Save to a temporary file
                    temp_file = os.path.join(self.get_working_directory(), 'combined_dataset.txt')
                    with open(temp_file, 'w', encoding='utf-8') as f:
                        f.write(self.text_content)
                    
                    # Update dataset path to point to the combined file
                    self.dataset_path.setText(temp_file)
                    
                    QMessageBox.information(
                        parent_dialog,
                        "Success",
                        f"Successfully loaded {len(selected_files)} file(s)"
                    )
                else:
                    QMessageBox.warning(
                        parent_dialog,
                        "Warning",
                        "No content was loaded from the selected files"
                    )
                
        except Exception as e:
            QMessageBox.critical(
                parent_dialog,
                "Error",
                f"Error loading dataset: {str(e)}"
            )
            # Hide progress bar in case of error
            if progress_bar is not None:
                progress_bar.setVisible(False)

    def load_excel_file(self, file_path):
        """Load and process Excel file."""
        try:
            df = pd.read_excel(file_path)
            return df.to_string()
        except Exception as e:
            raise Exception(f"Error reading Excel file: {str(e)}")

    def load_pdf_file(self, file_path):
        """Load and process PDF file."""
        try:
            pdf_document = fitz.open(file_path)
            text_content = []
            for page in pdf_document:
                text_content.append(page.get_text())
            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")

    def load_docx_file(self, file_path):
        """Load and process Word file."""
        try:
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise Exception(f"Error reading Word file: {str(e)}")

    def show_templates_dialog(self):
        """Show the templates dialog."""
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("Template Library")
            dialog.setMinimumWidth(600)
            dialog.setMinimumHeight(400)
            
            layout = QVBoxLayout(dialog)
            
            # Search bar
            search_layout = QHBoxLayout()
            search_label = QLabel("Search:")
            search_input = QLineEdit()
            search_input.setPlaceholderText("Search templates...")
            search_layout.addWidget(search_label)
            search_layout.addWidget(search_input)
            layout.addLayout(search_layout)
            
            # Template list
            template_list = QListWidget()
            template_list.setStyleSheet("""
                QListWidget {
                    border: 1px solid #E0E0E0;
                    border-radius: 4px;
                    padding: 4px;
                }
                QListWidget::item {
                    padding: 8px;
                    border-bottom: 1px solid #E0E0E0;
                }
                QListWidget::item:selected {
                    background-color: #E3F2FD;
                    color: #1976D2;
                }
            """)
            
            # Add templates to list
            for key in self.prompts.keys():
                template_list.addItem(key)
            
            layout.addWidget(template_list)
            
            # Preview area
            preview_label = QLabel("Preview:")
            layout.addWidget(preview_label)
            
            preview_text = QTextEdit()
            preview_text.setReadOnly(True)
            preview_text.setMinimumHeight(150)
            preview_text.setStyleSheet("""
                QTextEdit {
                    border: 1px solid #E0E0E0;
                    border-radius: 4px;
                    padding: 8px;
                    background-color: #FAFAFA;
                    font-family: Consolas, Monaco, monospace;
                }
            """)
            layout.addWidget(preview_text)
            
            # Button row
            button_layout = QHBoxLayout()
            select_btn = QPushButton("Select")
            select_btn.setStyleSheet(self.button_style())
            cancel_btn = QPushButton("Cancel")
            cancel_btn.setStyleSheet(self.button_style())
            button_layout.addWidget(select_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            # Connect signals
            def update_preview(item):
                if item and item.text() in self.prompts:
                    preview_text.setText(self.prompts[item.text()])
            
            def handle_selection():
                if template_list.currentItem():
                    selected = template_list.currentItem().text()
                    self.prompt_combo.setCurrentText(selected)
                    self.handle_prompt(selected)
                    dialog.accept()
            
            def filter_templates(text):
                template_list.clear()
                for key in self.prompts.keys():
                    if text.lower() in key.lower():
                        template_list.addItem(key)
            
            template_list.itemClicked.connect(update_preview)
            select_btn.clicked.connect(handle_selection)
            cancel_btn.clicked.connect(dialog.reject)
            search_input.textChanged.connect(filter_templates)
            
            dialog.exec_()
            
        except Exception as e:
            self.show_error(f"Error showing templates dialog: {str(e)}")

    def apply_selected_template(self):
        """Apply the selected template to the main window."""
        try:
            if not hasattr(self, 'template_list') or not hasattr(self, 'template_dialog'):
                return
                
            current_item = self.template_list.currentItem()
            if current_item and current_item.data(Qt.UserRole):
                template_name = current_item.data(Qt.UserRole)
                if template_name in self.prompts:
                    # Update the combo box and text edit in the main window
                    if hasattr(self, 'prompt_combo'):
                        self.prompt_combo.setCurrentText(template_name)
                    if hasattr(self, 'prompt_text_edit'):
                        self.prompt_text_edit.setPlainText(self.prompts[template_name])
                        self.prompt_text_edit.setReadOnly(True)
                    if hasattr(self, 'modify_btn'):
                        self.modify_btn.setEnabled(True)
                    if hasattr(self, 'save_template_btn'):
                        self.save_template_btn.setEnabled(False)
                    
                    # Close the template dialog
                    if self.template_dialog:
                        self.template_dialog.close()
                        
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error applying template: {str(e)}")

    def handle_template_dialog_close(self, event):
        """Handle template dialog close event."""
        try:
            if hasattr(self, 'template_dialog'):
                self.template_dialog = None
            event.accept()
        except Exception as e:
            print(f"Error in handle_template_dialog_close: {str(e)}")
            event.accept()

    def populate_template_list(self):
        """Populate the template list with categorized items."""
        try:
            if not hasattr(self, 'template_list'):
                return
                
            self.template_list.clear()
            
            # Define categories
            categories = {
                'Test Scripts': [],
                'Analysis': [],
                'Documentation': [],
                'Other': []
            }
            
            # Categorize templates
            for template_name in self.prompts.keys():
                if 'test' in template_name.lower():
                    categories['Test Scripts'].append(template_name)
                elif 'analysis' in template_name.lower():
                    categories['Analysis'].append(template_name)
                elif 'doc' in template_name.lower():
                    categories['Documentation'].append(template_name)
                else:
                    categories['Other'].append(template_name)
            
            # Add items by category
            for category, templates in categories.items():
                if templates:  # Only add categories that have templates
                    category_item = QListWidgetItem(category)
                    category_item.setFlags(category_item.flags() & ~Qt.ItemIsSelectable)
                    category_item.setBackground(QBrush(QColor('#F5F5F5')))
                    font = category_item.font()
                    font.setBold(True)
                    category_item.setFont(font)
                    self.template_list.addItem(category_item)
                    
                    # Add templates in this category
                    for template_name in sorted(templates):
                        template_item = QListWidgetItem(f"  {template_name}")  # Indent for visual hierarchy
                        template_item.setData(Qt.UserRole, template_name)  # Store actual template name
                        self.template_list.addItem(template_item)
                        
        except Exception as e:
            print(f"Error in populate_template_list: {str(e)}")

    def filter_templates(self, search_text):
        """Filter templates based on search text."""
        try:
            if not hasattr(self, 'template_list'):
                return
                
            search_text = search_text.lower()
            
            # Show all items if search is empty
            if not search_text:
                for i in range(self.template_list.count()):
                    self.template_list.item(i).setHidden(False)
                return
            
            # Track if we need to show category headers
            show_category = {
                'Test Scripts': False,
                'Analysis': False,
                'Documentation': False,
                'Other': False
            }
            
            # First pass: check which templates match and mark categories
            current_category = None
            for i in range(self.template_list.count()):
                item = self.template_list.item(i)
                text = item.text().strip()
                
                # Check if this is a category header
                if not item.data(Qt.UserRole):
                    current_category = text
                    continue
                
                # This is a template item
                if search_text in text.lower():
                    show_category[current_category] = True
            
            # Second pass: show/hide items based on search
            current_category = None
            for i in range(self.template_list.count()):
                item = self.template_list.item(i)
                text = item.text().strip()
                
                # Handle category headers
                if not item.data(Qt.UserRole):
                    current_category = text
                    item.setHidden(not show_category[current_category])
                    continue
                
                # Handle template items
                item.setHidden(not (search_text in text.lower()))
                
        except Exception as e:
            print(f"Error in filter_templates: {str(e)}")

    def handle_template_selection(self, item):
        """Handle template selection from the list."""
        try:
            if not item or not hasattr(self, 'preview_text'):
                return
                
            template_name = item.data(Qt.UserRole)
            if template_name and template_name in self.prompts:
                self.preview_text.setPlainText(self.prompts[template_name])
            else:
                self.preview_text.clear()
                
        except Exception as e:
            print(f"Error in handle_template_selection: {str(e)}")

    def show_test_deployment(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("🚀 Test Deployment")
        dialog.setMinimumWidth(700)
        dialog.setMinimumHeight(400)
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(18)

        # Title and subtitle
        title = QLabel("🚀 Test Deployment")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #0078D4;")
        layout.addWidget(title)
        subtitle = QLabel("Deploy your latest test scripts to the target environment.")
        subtitle.setStyleSheet("font-size: 14px; color: #555; margin-bottom: 8px;")
        layout.addWidget(subtitle)

        # Config selection
        config_layout = QHBoxLayout()
        self.config_file_path = QLineEdit()
        self.config_file_path.setPlaceholderText("Select a configuration file...")
        self.config_file_path.setReadOnly(True)
        config_layout.addWidget(self.config_file_path)
        browse_btn = QPushButton("📄 Change Config")
        browse_btn.setToolTip("Select a different deployment configuration file")
        browse_btn.clicked.connect(self.select_config_file)
        config_layout.addWidget(browse_btn)
        layout.addLayout(config_layout)

        # Deploy button
        self.deploy_btn = QPushButton("🚀 Deploy")
        self.deploy_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                border-radius: 6px;
                padding: 14px 32px;
                font-size: 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #005A9E;
            }
        """)
        self.deploy_btn.setToolTip("Start deployment with the selected configuration")
        self.deploy_btn.clicked.connect(self.run_test_deployment)
        layout.addWidget(self.deploy_btn)

        # Status label
        self.status_label = QLabel("Ready to deploy.")
        self.status_label.setStyleSheet("font-size: 13px; color: #0078D4; font-weight: 600; margin-top: 8px;")
        layout.addWidget(self.status_label)

        # Progress bar (indeterminate, just for UI feedback)
        self.deployment_progress = QProgressBar()
        self.deployment_progress.setRange(0, 0)  # Indeterminate
        self.deployment_progress.setVisible(False)
        layout.addWidget(self.deployment_progress)

        # Deployment log
        self.deployment_status = QTextEdit()
        self.deployment_status.setReadOnly(True)
        self.deployment_status.setPlaceholderText("Deployment output will be displayed here...")
        self.deployment_status.setMinimumHeight(150)
        self.deployment_status.setStyleSheet("""
            QTextEdit {
                background-color: #FAFAFA;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                padding: 10px;
                font-family: Consolas, Monaco, monospace;
                font-size: 12px;
            }
        """)
        layout.addWidget(self.deployment_status)

        dialog.setLayout(layout)
        dialog.exec_()

    def select_config_file(self):
        """Allow the user to select a new configuration file."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Configuration File", "", "Config Files (*.cfg *.json *.yaml *.ini)")
        if file_path:
            self.config_file_path.setText(file_path)
            # Only update ready_message if it exists (for dialog context)
            if hasattr(self, 'ready_message') and self.ready_message is not None:
                self.ready_message.setText("Configuration file selected: Ready for deployment")
            # Optionally, update a status label if present
            if hasattr(self, 'status_label') and self.status_label is not None:
                self.status_label.setText("Configuration file selected: Ready for deployment")
            self.deploy_btn.setEnabled(True)
            self.deploy_btn.setStyleSheet("""
                QPushButton {
                    background-color: #0078D4;
                    color: white;
                    border-radius: 4px;
                    padding: 12px 24px;
                    font-size: 14px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #106EBE;
                }
            """)
        #new_end
    def run_test_deployment(self):
        """Execute the test deployment script."""
        script_path = r"C:\Users\ChanduVangala\Documents\AgenticRAN-V3\test_deployment.py"
        config_path = self.config_file_path.text() if self.config_file_path.text() else "Default Config"

        self.deployment_status.setText(f"Starting deployment...\nUsing configuration: {config_path}")

        try:
            # Run test_deployment.py with the selected config
            process = subprocess.Popen(
                ["python", script_path, config_path], 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE, 
                shell=True, 
                text=True
            )

            stdout, stderr = process.communicate()

            if process.returncode == 0 and "Error" not in stdout and "Error" not in stderr:
                self.deployment_status.setText(f"Test Deployment completed successfully!\n\nTest environment is ready for execution.\n\nOutput:\n{stdout}")
                #QMessageBox.information(self, "Success", "Test Deployment completed successfully!\n\nTest environment is ready for execution.")
            else:
                error_message = stderr if stderr else stdout  # Prioritize stderr, but use stdout if needed
                self.deployment_status.setText(f"Test Deployment failed.\n\nError:\n{error_message}")
                #QMessageBox.critical(self, "Error", "Test Deployment failed!")

        except Exception as e:
            self.deployment_status.setText(f"Failed to launch Test Deployment:\n{str(e)}")
            #QMessageBox.critical(self, "Error", f"Failed to launch Test Deployment:\n{str(e)}")


    def show_test_execution(self):
        """Show the test execution dialog."""
        #QMessageBox.information(self, "Coming Soon", "Test Execution functionality will be available soon!")
            #chandu_deploy_code_start
        dialog = QDialog(self)
        dialog.setWindowTitle("Test Execution")
        dialog.setMinimumSize(600, 500)
        dialog.setStyleSheet("""
            QDialog {
                background: #F5F5F5;
            }
        """)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(32, 32, 32, 32)
        layout.setSpacing(18)

        # Section Header
        section_header = QLabel("Integrate CI/CD")
        section_header.setStyleSheet("font-size: 20px; font-weight: bold; color: #323130; margin-bottom: 8px;")
        layout.addWidget(section_header)

        # Start Execution Button
        exec_btn_layout = QHBoxLayout()
        self.start_exec_button = QPushButton("▶ Start Execution")
        self.start_exec_button.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                border-radius: 6px;
                padding: 14px 32px;
                font-size: 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #005A9E;
            }
        """)
        self.start_exec_button.setToolTip("Start the deployment and test execution process")
        self.start_exec_button.clicked.connect(self.start_deployment)
        exec_btn_layout.addWidget(self.start_exec_button)
        exec_btn_layout.addStretch()
        layout.addLayout(exec_btn_layout)

        # Progress Label
        progress_label = QLabel("Execution Progress")
        progress_label.setStyleSheet("font-size: 14px; font-weight: 600; color: #605E5C; margin-top: 18px;")
        layout.addWidget(progress_label)

        # Progress Bar
        self.deploy_progress_bar = QProgressBar()
        self.deploy_progress_bar.setValue(0)
        self.deploy_progress_bar.setFixedHeight(28)
        self.deploy_progress_bar.setStyleSheet("""
            QProgressBar {
                font-size: 14px;
                text-align: center;
                border: 1px solid #E1E5EA;
                border-radius: 6px;
                background: white;
                height: 28px;
            }
            QProgressBar::chunk {
                background: #0078D4;
            }
        """)
        self.deploy_progress_bar.setFormat("%p%")
        self.deploy_progress_bar.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.deploy_progress_bar)

        # Status Label
        from datetime import datetime
        now = datetime.now().strftime("%H:%M:%S")
        self.deploy_status_label = QLabel(f"Status: Not Started ({now})")
        self.deploy_status_label.setContentsMargins(10, 10, 10, 10)
        self.deploy_status_label.setStyleSheet("font-size: 15px; font-weight: bold; color: #605E5C;")
        layout.addWidget(self.deploy_status_label)

        # Results Box
        self.results_box = QTextEdit()
        self.results_box.setReadOnly(True)
        self.results_box.setPlaceholderText("Awaiting test results...")
        self.results_box.setMinimumHeight(180)
        self.results_box.setStyleSheet("""
            QTextEdit {
                background-color: #FAFAFA;
                border: 1px solid #E1E5EA;
                border-radius: 6px;
                padding: 14px;
                font-family: Consolas, Monaco, monospace;
                font-size: 14px;
                color: #323130;
            }
        """)
        layout.addWidget(self.results_box)

        dialog.setLayout(layout)
        dialog.exec_()

    def start_deployment(self):
        """Starts the deployment process by running jenkins_trigger.py."""
        self.deploy_progress_bar.setValue(0)
        self.deploy_status_label.setText("Execution Status: In Progress...")

        try:
            self.process = subprocess.Popen(
                ['python', r'/mnt/data/jenkins_trigger.py'],  # Check if this file exists!
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
            )
        except FileNotFoundError as e:
            self.deploy_status_label.setText(f"Error: {e}")
            return
        except Exception as e:
            self.deploy_status_label.setText(f"Unexpected Error: {e}")
            return

        self.timer = QTimer()
        self.timer.timeout.connect(self.update_deployment_progress)
        self.progress_value = 0
        self.timer.start(1000)

    def update_deployment_progress(self):
        """Updates the deployment progress bar and checks for errors."""
        if self.process.poll() is None:
            if self.progress_value < 90:
                self.progress_value += 10
            self.deploy_progress_bar.setValue(self.progress_value)
        else:
            self.timer.stop()
            output, error = self.process.communicate()  # Get both stdout and stderr
            result = "Fail"

            # Check if Jenkins job completed successfully
            for line in output.split("\n"):
                if "Job completed with result:" in line:
                    result = line.split(":")[-1].strip()
                    break

            self.deploy_progress_bar.setValue(100)
            self.deploy_status_label.setText(f"Execution Status: {result}")

            # Show both stdout and stderr in the results box
            if error:
                self.results_box.setPlainText(f"STDOUT:\n{output}\n\nSTDERR:\n{error}")
            else:
                self.results_box.setPlainText(f"STDOUT:\n{output}")


    #chandu_deploy_code_end

    def show_bug_discovery(self):
        """Show the bug discovery dialog - Analysis Phase."""
        self.bug_discovery_dialog = QDialog(self)
        dialog = self.bug_discovery_dialog
        dialog.setWindowTitle("Bug Discovery - Error Analysis")
        
        # Make dialog resizable and set a better default size
        dialog.setMinimumSize(1000, 750)
        dialog.resize(1200, 850)
        dialog.setSizeGripEnabled(True)  # Enable resize grip
        
        # Main layout with scroll area
        main_layout = QVBoxLayout(dialog)
        main_layout.setContentsMargins(10, 10, 10, 10)
        
        # Create scroll area for content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        
        # Container widget for scroll area
        container_widget = QWidget()
        layout = QVBoxLayout(container_widget)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(15)
        
        # Header
        header = QLabel("Bug Discovery - Error Analysis")
        header.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: bold;
                color: #323130;
                margin-bottom: 8px;
            }
        """)
        layout.addWidget(header)
        
        description = QLabel("Analyze errors from log files and generate fix suggestions")
        description.setStyleSheet("""
            QLabel {
                font-size: 14px;
                color: #605E5C;
                margin-bottom: 16px;
            }
        """)
        layout.addWidget(description)
        
        # Load Previous Analysis Section
        load_group = QGroupBox("📂 Bug Analysis History")
        load_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #E1E5EA;
                border-radius: 6px;
                margin-top: 12px;
                padding: 15px;
                background-color: #F0F8FF;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
                color: #0078D4;
            }
        """)
        load_layout = QVBoxLayout()
        load_layout.setSpacing(10)
        
        # Info label
        load_info_label = QLabel("Load a previous analysis to view results and artifacts without re-running RCA")
        load_info_label.setStyleSheet("font-size: 11px; color: #605E5C; font-style: italic;")
        load_layout.addWidget(load_info_label)
        
        # History combo and buttons layout
        history_control_layout = QHBoxLayout()
        history_control_layout.setSpacing(10)
        
        # History dropdown
        self.bug_discovery_history_combo = QComboBox()
        self.bug_discovery_history_combo.setMinimumWidth(500)
        self.bug_discovery_history_combo.setMinimumHeight(35)
        self.bug_discovery_history_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 1px solid #C8C6C4;
                border-radius: 4px;
                background-color: white;
                font-size: 12px;
            }
            QComboBox:focus {
                border: 2px solid #0078D4;
            }
            QComboBox::drop-down {
                border: none;
                width: 30px;
            }
        """)
        history_control_layout.addWidget(self.bug_discovery_history_combo, 1)
        
        # Load button
        load_prev_btn = QPushButton("📂 Load Previous Run")
        load_prev_btn.setMinimumHeight(35)
        load_prev_btn.setMinimumWidth(150)
        load_prev_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
            QPushButton:pressed {
                background-color: #005A9E;
            }
        """)
        load_prev_btn.setToolTip("Load all results and artifacts from a previous analysis")
        load_prev_btn.clicked.connect(self.load_previous_bug_analysis)
        history_control_layout.addWidget(load_prev_btn)
        
        # Refresh button
        refresh_prev_btn = QPushButton("🔄")
        refresh_prev_btn.setMinimumHeight(35)
        refresh_prev_btn.setMaximumWidth(40)
        refresh_prev_btn.setStyleSheet("""
            QPushButton {
                background-color: #8A8886;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #797775;
            }
        """)
        refresh_prev_btn.setToolTip("Refresh history list")
        refresh_prev_btn.clicked.connect(self.load_bug_discovery_history)
        history_control_layout.addWidget(refresh_prev_btn)
        
        load_layout.addLayout(history_control_layout)
        load_group.setLayout(load_layout)
        layout.addWidget(load_group)
        
        # Source Code Directory Selection
        dir_group = QGroupBox("Source Code Directory")
        dir_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #E1E5EA;
                border-radius: 6px;
                margin-top: 12px;
                padding: 15px;
                background-color: #FAFAFA;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }
        """)
        dir_layout = QHBoxLayout()
        dir_layout.setSpacing(10)
        self.bug_code_dir_edit = QLineEdit()
        self.bug_code_dir_edit.setPlaceholderText("Select OAI gNodeB source code directory")
        self.bug_code_dir_edit.setMinimumHeight(35)
        self.bug_code_dir_edit.setStyleSheet("""
            QLineEdit {
                padding: 8px;
                border: 1px solid #C8C6C4;
                border-radius: 4px;
                background-color: white;
                font-size: 12px;
            }
            QLineEdit:focus {
                border: 2px solid #0078D4;
            }
        """)
        browse_code_btn = QPushButton("Browse")
        browse_code_btn.setMinimumHeight(35)
        browse_code_btn.setMinimumWidth(100)
        browse_code_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
            QPushButton:pressed {
                background-color: #005A9E;
            }
        """)
        browse_code_btn.clicked.connect(self.select_bug_code_directory)
        dir_layout.addWidget(self.bug_code_dir_edit)
        dir_layout.addWidget(browse_code_btn)
        dir_group.setLayout(dir_layout)
        layout.addWidget(dir_group)
        
        # Log Files Selection
        log_group = QGroupBox("Log Files")
        log_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #E1E5EA;
                border-radius: 6px;
                margin-top: 12px;
                padding: 15px;
                background-color: #FAFAFA;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }
        """)
        log_layout = QVBoxLayout()
        log_layout.setSpacing(12)
        
        # Log directory selection
        log_dir_layout = QHBoxLayout()
        log_dir_layout.setSpacing(10)
        self.bug_log_dir_edit = QLineEdit()
        self.bug_log_dir_edit.setPlaceholderText("Select log directory")
        self.bug_log_dir_edit.setMinimumHeight(35)
        self.bug_log_dir_edit.setStyleSheet("""
            QLineEdit {
                padding: 8px;
                border: 1px solid #C8C6C4;
                border-radius: 4px;
                background-color: white;
                font-size: 12px;
            }
            QLineEdit:focus {
                border: 2px solid #0078D4;
            }
        """)
        browse_log_btn = QPushButton("Browse")
        browse_log_btn.setMinimumHeight(35)
        browse_log_btn.setMinimumWidth(100)
        browse_log_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
            QPushButton:pressed {
                background-color: #005A9E;
            }
        """)
        browse_log_btn.clicked.connect(lambda: self.select_log_directory(self.bug_log_dir_edit))
        log_dir_layout.addWidget(self.bug_log_dir_edit)
        log_dir_layout.addWidget(browse_log_btn)
        log_layout.addLayout(log_dir_layout)
        
        # Log file dropdown
        file_dropdown_layout = QHBoxLayout()
        file_dropdown_layout.setSpacing(10)
        file_dropdown_label = QLabel("Log File:")
        file_dropdown_label.setStyleSheet("font-size: 12px; font-weight: bold; min-width: 70px;")
        self.bug_log_file_combo = QComboBox()
        self.bug_log_file_combo.setMinimumWidth(350)
        self.bug_log_file_combo.setMinimumHeight(35)
        self.bug_log_file_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 1px solid #C8C6C4;
                border-radius: 4px;
                background-color: white;
                font-size: 12px;
            }
            QComboBox:focus {
                border: 2px solid #0078D4;
            }
            QComboBox::drop-down {
                border: none;
                width: 30px;
            }
            QComboBox::down-arrow {
                image: url(down_arrow.png);
                width: 12px;
                height: 12px;
            }
        """)
        file_dropdown_layout.addWidget(file_dropdown_label)
        file_dropdown_layout.addWidget(self.bug_log_file_combo, 1)
        
        # Add spacer
        file_dropdown_layout.addSpacing(20)
        
        # Crash Analysis Toggle Slider (on same line as log file selector)
        crash_toggle_label = QLabel("Core Dump Analysis:")
        crash_toggle_label.setStyleSheet("font-size: 11px; font-weight: bold; color: #605E5C; min-width: 100px;")
        file_dropdown_layout.addWidget(crash_toggle_label)
        
        # Use QPushButton as checkable toggle with text inside
        self.crash_analysis_checkbox = QPushButton("OFF")
        self.crash_analysis_checkbox.setCheckable(True)
        self.crash_analysis_checkbox.setMinimumWidth(65)
        self.crash_analysis_checkbox.setMaximumWidth(65)
        self.crash_analysis_checkbox.setMinimumHeight(28)
        self.crash_analysis_checkbox.setStyleSheet("""
            QPushButton {
                background-color: #AAAAAA;  /* Dark grey when OFF */
                color: #FFFFFF;  /* White text */
                border: 2px solid #999999;
                border-radius: 14px;
                font-size: 10px;
                font-weight: bold;
                padding: 4px;
            }
            QPushButton:checked {
                background-color: #28A745;  /* Green when ON */
                border: 2px solid #1E7E34;
                color: #FFFFFF;  /* White text */
            }
            QPushButton:hover {
                border: 2px solid #666666;
            }
            QPushButton:checked:hover {
                border: 2px solid #1E7E34;
                background-color: #218838;  /* Darker green on hover */
            }
        """)
        self.crash_analysis_checkbox.setToolTip("Toggle ON for segmentation faults and crash analysis")
        
        # Update button text when toggled
        def update_crash_toggle(checked):
            if checked:
                self.crash_analysis_checkbox.setText("ON")
            else:
                self.crash_analysis_checkbox.setText("OFF")
        
        self.crash_analysis_checkbox.toggled.connect(update_crash_toggle)
        file_dropdown_layout.addWidget(self.crash_analysis_checkbox)
        
        log_layout.addLayout(file_dropdown_layout)

        # Update dropdown when directory changes
        def update_bug_log_file_combo():
            self.bug_log_file_combo.clear()
            self.bug_log_file_combo.addItem("Select a log file")
            dir_path = self.bug_log_dir_edit.text()
            if os.path.isdir(dir_path):
                for fname in sorted(os.listdir(dir_path)):
                    if fname.lower().endswith((".log", ".txt")):
                        self.bug_log_file_combo.addItem(fname)
        
        self.bug_log_dir_edit.textChanged.connect(update_bug_log_file_combo)
        browse_log_btn.clicked.connect(update_bug_log_file_combo)
        update_bug_log_file_combo()
        
        # Start RCA Analysis button
        rca_btn_layout = QHBoxLayout()
        rca_btn_layout.setSpacing(12)
        self.start_bug_rca_btn = QPushButton("🔍 Start RCA Analysis")
        self.start_bug_rca_btn.setMinimumHeight(42)
        self.start_bug_rca_btn.setMinimumWidth(180)
        self.start_bug_rca_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                font-weight: bold;
                padding: 12px 24px;
                border: none;
                border-radius: 5px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
            QPushButton:pressed {
                background-color: #005A9E;
            }
        """)
        self.start_bug_rca_btn.clicked.connect(self.handle_bug_rca_analysis)
        rca_btn_layout.addWidget(self.start_bug_rca_btn)
        
        # Deployment Context Settings button
        self.deployment_settings_btn = QPushButton("⚙️ Deployment Settings")
        self.deployment_settings_btn.setMinimumHeight(42)
        self.deployment_settings_btn.setMinimumWidth(180)
        self.deployment_settings_btn.setStyleSheet("""
            QPushButton {
                background-color: #8A8886;
                color: white;
                font-weight: bold;
                padding: 12px 24px;
                border: none;
                border-radius: 5px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #797775;
            }
            QPushButton:pressed {
                background-color: #605E5C;
            }
        """)
        self.deployment_settings_btn.clicked.connect(self.open_deployment_context_settings)
        self.deployment_settings_btn.setToolTip("Configure deployment context values for RCA analysis")
        rca_btn_layout.addWidget(self.deployment_settings_btn)
        
        # Fix Already Present button (initially hidden)
        self.bug_fix_present_btn = QPushButton("✅ Fix Already Present - Click to View")
        self.bug_fix_present_btn.setMinimumHeight(42)
        self.bug_fix_present_btn.setMinimumWidth(250)
        self.bug_fix_present_btn.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                font-weight: bold;
                padding: 12px 24px;
                border: none;
                border-radius: 5px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1e7e34;
            }
        """)
        self.bug_fix_present_btn.setVisible(False)  # Hidden by default
        self.bug_fix_present_btn.clicked.connect(self.show_bug_existing_fix_details)
        rca_btn_layout.addWidget(self.bug_fix_present_btn)
        
        rca_btn_layout.addStretch()
        log_layout.addLayout(rca_btn_layout)
        
        # Connect log file selection to check for existing fixes
        self.bug_log_file_combo.currentTextChanged.connect(self.check_bug_for_existing_fix)
        
        log_group.setLayout(log_layout)
        layout.addWidget(log_group)
        
        # Results Display Area
        results_group = QGroupBox("Analysis Results")
        results_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #E1E5EA;
                border-radius: 6px;
                margin-top: 12px;
                padding: 15px;
                background-color: #FAFAFA;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }
        """)
        results_layout = QVBoxLayout()
        results_layout.setSpacing(12)
        
        self.bug_results_display = QTextEdit()
        self.bug_results_display.setReadOnly(True)
        self.bug_results_display.setMinimumHeight(250)
        self.bug_results_display.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.bug_results_display.setPlaceholderText("Analysis results will appear here after running RCA analysis...")
        self.bug_results_display.setStyleSheet("""
            QTextEdit {
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                padding: 10px;
                background-color: #FFFFFF;
                font-family: 'Consolas', 'Monaco', monospace;
                font-size: 12px;
                line-height: 1.5;
            }
        """)
        results_layout.addWidget(self.bug_results_display)
        
        # Buttons layout (View Artifacts and Go to Code Evaluation) - Placed at bottom
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(10)
        buttons_layout.addStretch()  # Push buttons to the right
        
        # View Artifacts button
        self.view_artifacts_bug_btn = QPushButton("📦 View Artifacts")
        self.view_artifacts_bug_btn.setMinimumHeight(38)
        self.view_artifacts_bug_btn.setStyleSheet("""
            QPushButton {
                background-color: #6B46C1;
                color: white;
                font-weight: bold;
                padding: 10px 20px;
                border: none;
                border-radius: 4px;
                font-size: 13px;
            }
            QPushButton:hover {
                background-color: #553C9A;
            }
            QPushButton:disabled {
                background-color: #C8C6C4;
                color: #8A8886;
            }
        """)
        self.view_artifacts_bug_btn.clicked.connect(self.show_bug_artifacts)
        self.view_artifacts_bug_btn.setEnabled(False)
        buttons_layout.addWidget(self.view_artifacts_bug_btn)
        
        # Go to Code Evaluation button
        self.goto_code_evaluation_btn = QPushButton("🧪 Go to Code Evaluation")
        self.goto_code_evaluation_btn.setMinimumHeight(38)
        self.goto_code_evaluation_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                font-weight: bold;
                padding: 10px 20px;
                border: none;
                border-radius: 4px;
                font-size: 13px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
            QPushButton:disabled {
                background-color: #C8C6C4;
                color: #8A8886;
            }
        """)
        self.goto_code_evaluation_btn.clicked.connect(self.go_to_code_evaluation_from_bug_discovery)
        self.goto_code_evaluation_btn.setEnabled(False)
        buttons_layout.addWidget(self.goto_code_evaluation_btn)
        
        results_layout.addLayout(buttons_layout)
        
        results_group.setLayout(results_layout)
        layout.addWidget(results_group)
        
        # Progress Bar
        self.bug_analysis_progress = QProgressBar()
        self.bug_analysis_progress.hide()
        layout.addWidget(self.bug_analysis_progress)
        
        # Set the scroll area widget
        scroll_area.setWidget(container_widget)
        main_layout.addWidget(scroll_area)
        
        # Button bar at the bottom (not scrolled)
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        close_btn = QPushButton("Close")
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #8A8886;
                color: white;
                font-weight: bold;
                padding: 8px 20px;
                border: none;
                border-radius: 4px;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #797775;
            }
        """)
        close_btn.clicked.connect(dialog.accept)
        button_layout.addWidget(close_btn)
        
        main_layout.addLayout(button_layout)
        
        # Load bug discovery history to populate dropdown
        self.load_bug_discovery_history()
        
        dialog.exec_()

    def open_deployment_context_settings(self):
        """Open dialog to configure deployment context settings"""
        dialog = DeploymentContextDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            # Get the edited values
            self.custom_deployment_context = dialog.get_values()
            
            # Check if any custom values were set
            if self.custom_deployment_context and len(self.custom_deployment_context) > 0:
                # Update button appearance to show custom settings are active
                self.deployment_settings_btn.setText("⚙️ Deployment Settings ✓")
                self.deployment_settings_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #107C10;
                        color: white;
                        font-weight: bold;
                        padding: 12px 24px;
                        border: none;
                        border-radius: 5px;
                        font-size: 14px;
                    }
                    QPushButton:hover {
                        background-color: #0E6B0E;
                    }
                    QPushButton:pressed {
                        background-color: #0C5A0C;
                    }
                """)
                self.deployment_settings_btn.setToolTip(
                    f"Custom deployment context active ({len(self.custom_deployment_context)} values configured)\n"
                    "Click to edit settings"
                )
                
                QMessageBox.information(
                    self,
                    "Settings Saved",
                    f"Deployment context settings have been saved ({len(self.custom_deployment_context)} values).\n\n"
                    "These values will be used in the next RCA analysis.\n\n"
                    "The Deployment Settings button now shows a checkmark to indicate custom settings are active."
                )
                print(f"✅ Custom deployment context saved with {len(self.custom_deployment_context)} values:")
                for key, value in self.custom_deployment_context.items():
                    print(f"   - {key}: {value}")
            else:
                # No custom values - reset to JSON defaults
                self.custom_deployment_context = None
                self.deployment_settings_btn.setText("⚙️ Deployment Settings")
                self.deployment_settings_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #8A8886;
                        color: white;
                        font-weight: bold;
                        padding: 12px 24px;
                        border: none;
                        border-radius: 5px;
                        font-size: 14px;
                    }
                    QPushButton:hover {
                        background-color: #797775;
                    }
                    QPushButton:pressed {
                        background-color: #605E5C;
                    }
                """)
                self.deployment_settings_btn.setToolTip("Configure deployment context values for RCA analysis")
                
                QMessageBox.information(
                    self,
                    "Settings Cleared",
                    "All custom deployment settings have been cleared.\n\n"
                    "The system will now use default values from JSON file."
                )
                print("🔄 Custom deployment context cleared, using JSON defaults")

    def load_bug_discovery_history(self):
        """Load bug analysis history from files to populate the dropdown."""
        self.bug_discovery_history_combo.clear()
        self.bug_discovery_history_combo.addItem("Or select a previous analysis...")
        
        history_dir = "backend/resources/bug_history"
        if not os.path.exists(history_dir):
            return
        
        # Get all history files and sort by timestamp (newest first)
        history_files = [f for f in os.listdir(history_dir) if f.endswith('.json')]
        history_files.sort(reverse=True)  # Sort by filename (timestamp) - newest first
        
        # Store mapping of display text to file path
        self.bug_discovery_history_map = {}
        
        for filename in history_files:
            file_path = os.path.join(history_dir, filename)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    error_message = data.get('error_message', 'Unknown error')
                    timestamp = data.get('timestamp', '')
                    
                    # Handle None error message
                    if error_message is None:
                        error_message = 'Unknown error'
                    
                    # Format display text: "Error message (truncated)"
                    display_text = error_message[:80] + "..." if len(error_message) > 80 else error_message
                    
                    # Add timestamp to display
                    if timestamp:
                        try:
                            dt = datetime.fromisoformat(timestamp)
                            time_str = dt.strftime("%Y-%m-%d %H:%M")
                            display_text = f"[{time_str}] {display_text}"
                        except:
                            pass
                    
                    # Show patch counts in this history file
                    results = data.get('results', {})
                    phase3_fixes = results.get('phase3_fixes', {})
                    fix_suggestion = phase3_fixes.get('fix_suggestion', {})
                    code_count = len(fix_suggestion.get('code_patches', []))
                    config_count = len(fix_suggestion.get('config_patches', []))
                    display_text += f" [Code:{code_count}, Config:{config_count}]"
                    
                    self.bug_discovery_history_combo.addItem(display_text)
                    self.bug_discovery_history_map[display_text] = file_path
            except Exception as e:
                print(f"Error loading history file {filename}: {e}")
                continue

    def load_previous_bug_analysis(self):
        """Load selected previous analysis with complete results and artifacts."""
        selected_text = self.bug_discovery_history_combo.currentText()
        
        if selected_text == "Or select a previous analysis...":
            QMessageBox.warning(self, "No Selection", "Please select a bug analysis from the dropdown.")
            return
        
        if selected_text not in self.bug_discovery_history_map:
            QMessageBox.warning(self, "Error", "Could not find the selected analysis file.")
            return
        
        file_path = self.bug_discovery_history_map[selected_text]
        
        try:
            # Show loading indicator
            QApplication.setOverrideCursor(Qt.WaitCursor)
            
            # Load JSON file
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Populate Input Fields
            code_dir = data.get('code_dir', '')
            log_path = data.get('log_path', '')
            log_file = data.get('log_file', '')
            
            self.bug_code_dir_edit.setText(code_dir)
            
            # Populate log directory
            if log_path and os.path.exists(log_path):
                log_dir = os.path.dirname(log_path)
                self.bug_log_dir_edit.setText(log_dir)
                
                # Update log file combo
                self.bug_log_file_combo.clear()
                self.bug_log_file_combo.addItem("Select a log file")
                
                # Populate log files from directory
                if os.path.exists(log_dir):
                    log_files = [f for f in os.listdir(log_dir) if f.endswith('.log')]
                    for lf in log_files:
                        self.bug_log_file_combo.addItem(lf)
                    
                    # Select the original log file
                    if log_file in log_files:
                        self.bug_log_file_combo.setCurrentText(log_file)
            else:
                # Log file path doesn't exist anymore
                self.bug_log_dir_edit.setText(log_dir if log_dir else "")
                print(f"⚠️  Original log file not found: {log_path}")
            
            # CRITICAL: Restore Complete Analysis State
            self.current_bug_analysis = {
                'error_message': data.get('error_message', ''),
                'log_file': log_file,
                'log_path': log_path,
                'code_dir': code_dir,
                'results': data.get('results', {}),  # Complete results for artifacts
                'timestamp': data.get('timestamp', '')
            }
            
            # Display Complete Formatted Results (exact same as after analysis)
            self.display_bug_analysis_results(
                data.get('results', {}), 
                data.get('error_message', 'Unknown error')
            )
            
            # Enable Action Buttons - Artifacts and Code Evaluation now accessible!
            self.view_artifacts_bug_btn.setEnabled(True)
            self.goto_code_evaluation_btn.setEnabled(True)
            
            # Clear loading cursor
            QApplication.restoreOverrideCursor()
            
            # Show success message with details
            timestamp_str = data.get('timestamp', 'N/A')
            try:
                dt = datetime.fromisoformat(timestamp_str)
                timestamp_str = dt.strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass
            
            results = data.get('results', {})
            phase3_fixes = results.get('phase3_fixes', {})
            fix_suggestion = phase3_fixes.get('fix_suggestion', {})
            code_count = len(fix_suggestion.get('code_patches', []))
            config_count = len(fix_suggestion.get('config_patches', []))
            
            error_msg_preview = data.get('error_message', 'Unknown')[:100]
            if len(data.get('error_message', '')) > 100:
                error_msg_preview += "..."
            
            QMessageBox.information(
                self,
                "✅ Analysis Loaded Successfully",
                f"Previous analysis loaded successfully!\n\n"
                f"📝 Error: {error_msg_preview}\n"
                f"📁 Log File: {log_file}\n"
                f"📅 Timestamp: {timestamp_str}\n"
                f"🔧 Code Patches: {code_count}\n"
                f"⚙️  Config Patches: {config_count}\n\n"
                f"✅ Results displayed in window\n"
                f"✅ Artifacts available (click 'View Artifacts')\n"
                f"✅ Ready for Code Evaluation"
            )
            
            print(f"✅ Loaded previous analysis from: {file_path}")
            print(f"   Error: {data.get('error_message', 'Unknown')}")
            print(f"   Code Patches: {code_count}, Config Patches: {config_count}")
            
        except Exception as e:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(
                self, 
                "Error Loading Analysis", 
                f"Failed to load previous analysis:\n\n{str(e)}"
            )
            print(f"❌ Error loading previous analysis: {e}")
            import traceback
            traceback.print_exc()

    def select_bug_code_directory(self):
        """Select the OAI gNodeB source code directory for bug discovery."""
        folder_path = QFileDialog.getExistingDirectory(self, "Select OAI gNodeB Source Directory")
        if folder_path:
            # Validate if it's an OAI gNodeB directory
            if self.validate_oai_directory(folder_path):
                self.bug_code_dir_edit.setText(folder_path)
            else:
                QMessageBox.warning(
                    self,
                    "Invalid Directory",
                    "The selected directory does not appear to be an OAI gNodeB source directory."
                )

    def handle_bug_rca_analysis(self):
        """Handle RCA analysis in Bug Discovery module."""
        # Validate inputs
        code_dir = self.bug_code_dir_edit.text()
        log_dir = self.bug_log_dir_edit.text()
        log_file = self.bug_log_file_combo.currentText()
        
        if not code_dir:
            QMessageBox.warning(self, "Missing Input", "Please select the source code directory.")
            return
        
        if not log_dir or log_file == "Select a log file":
            QMessageBox.warning(self, "Missing Input", "Please select a log file.")
            return
        
        # Check if CompleteErrorFixingPipeline is available
        if CompleteErrorFixingPipeline is None:
            QMessageBox.warning(
                self,
                "Module Not Available",
                "Complete Error Fixing Pipeline module is not available.\n\nPlease check if Error_fixing_pipelin directory exists."
            )
            return
        
        # Store original directory for cleanup
        original_cwd = os.getcwd()
        
        try:
            # Show progress
            self.bug_analysis_progress.show()
            self.bug_analysis_progress.setValue(0)
            self.bug_results_display.clear()
            self.bug_results_display.setText("Starting RCA Analysis...\n")
            
            # Disable View Artifacts button until analysis completes
            self.view_artifacts_bug_btn.setEnabled(False)
            
            QApplication.processEvents()
            
            # Construct log file path
            log_file_path = os.path.join(log_dir, log_file)
            
            # Extract error message from log
            error_message = self.extract_error_from_log(log_file_path)
            
            self.bug_results_display.append(f"Log File: {log_file}\n")
            self.bug_results_display.append(f"Analyzing error...\n")
            QApplication.processEvents()
            
            # Run the complete error fixing pipeline
            # original_cwd already stored above
            error_pipeline_dir = os.path.join(os.path.dirname(__file__), 'Error_fixing_pipelin')
            os.chdir(error_pipeline_dir)
            
            relative_log_path = os.path.relpath(log_file_path, error_pipeline_dir)
            openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            
            self.bug_analysis_progress.setValue(30)
            QApplication.processEvents()
            
            pipeline = CompleteErrorFixingPipeline(openair_codebase_file_name=openair_codebase_file_name)
            
            # Check if crash analysis is enabled
            crash_analysis_enabled = self.crash_analysis_checkbox.isChecked()
            
            if crash_analysis_enabled:
                # Use crash analysis flow for segmentation faults - Run complete flow (Phase 1, 2, 2.5, 3)
                self.bug_results_display.append("🔬 Crash Analysis Mode Enabled\n")
                self.bug_results_display.append("Running complete crash analysis pipeline...\n\n")
                self.bug_results_display.append("Phase 1: Extracting error and backtrace from segmentation fault log...\n")
                QApplication.processEvents()
                
                # Run complete crash analysis (Phase 1, 2, 2.5, 3)
                results = pipeline.process_crash_analysis(relative_log_path, phase="full")
                
                # Extract error message for display
                error_message = results.get('error_message', 'Segmentation Fault Detected')
                
                # Display crash analysis results and skip terminal commands
                self.bug_results_display.append("\n✅ Crash Analysis Complete!\n\n")
                QApplication.processEvents()
                
                # Display the results
                self.display_bug_analysis_results(results, error_message)
                
                # Save crash-specific results
                crash_output_file = os.path.join(error_pipeline_dir, "output/crash_phase3_fixes.json")
                if os.path.exists(crash_output_file):
                    try:
                        with open(crash_output_file, 'r', encoding='utf-8') as f:
                            crash_detailed = json.load(f)
                            results['crash_detailed_fixes'] = crash_detailed
                    except Exception as e:
                        print(f"Could not load detailed crash fixes: {e}")
                
                # Save complete crash results
                complete_crash_file = os.path.join(error_pipeline_dir, "output/complete_crash_analysis.json")
                with open(complete_crash_file, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)
                
                print(f"💾 Complete crash results saved to: {complete_crash_file}")
                
                # Store crash analysis results for artifacts view
                self.current_bug_analysis = {
                    'error_message': error_message,
                    'log_file': log_file,
                    'log_path': log_file_path,
                    'code_dir': code_dir,
                    'results': results,
                    'timestamp': datetime.now().isoformat(),
                    'crash_analysis': True  # Flag for special handling
                }
                
                # Save to history
                self.save_bug_analysis_to_history()
                
                # Enable buttons
                self.view_artifacts_bug_btn.setEnabled(True)
                self.goto_code_evaluation_btn.setEnabled(True)  # Enable code evaluation for crash too

                # Show completion message
                QMessageBox.information(
                    self,
                    "Crash Analysis Complete",
                    "Crash analysis completed successfully!\n\n"
                    "You can:\n• View Artifacts to see detailed analysis\n• Go to Code Evaluation to continue investigation"
                )

                # Restore working directory and return early (skip normal flow)
                os.chdir(original_cwd)
                return
            else:
                # Pass custom deployment context if user has configured it
                custom_context = self.custom_deployment_context if hasattr(self, 'custom_deployment_context') and self.custom_deployment_context else None
                if custom_context:
                    self.bug_results_display.append("Using custom deployment context settings...\n")
                    QApplication.processEvents()
                
                results = pipeline.process_error_with_context(error_message, relative_log_path, custom_deployment_context=custom_context)
            
            # Generate Phase 4 terminal commands (like the old version does)
            phase3_fixes = results.get('phase3_fixes', {})
            fix_suggestion = phase3_fixes.get('fix_suggestion', {})
            investigation_steps = fix_suggestion.get('investigation_steps', [])
            deployment_context = results.get('deployment_context')
            
            # Get troubleshooting hints
            troubleshooting_hints = []
            try:
                patterns_file = os.path.join(error_pipeline_dir, 'database', 'error_patterns_structured.json')
                with open(patterns_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    patterns = data.get('patterns', {})
                    error_lower = error_message.lower()
                    
                    # Find matching pattern
                    pattern_found = False
                    for pattern_name, pattern_data in patterns.items():
                        keywords = pattern_data.get('keywords', [])
                        if any(keyword in error_lower for keyword in keywords):
                            suggested_fixes = pattern_data.get('suggested_fixes', [])
                            troubleshooting_hints.extend(suggested_fixes)
                            pattern_found = True
                            break
                    
                    # If no pattern matches, generate dynamic pattern (like old version)
                    if not pattern_found:
                        print(f"🔄 No pattern found for Phase 4, generating dynamic pattern for: {error_message}")
                        from fix_suggestion_pipeline import FixSuggestionPipeline
                        fix_pipeline = FixSuggestionPipeline(openair_codebase_file_name=openair_codebase_file_name)
                        dynamic_pattern = fix_pipeline._generate_dynamic_error_pattern(error_message)
                        fix_pipeline._add_pattern_to_json(error_message, dynamic_pattern)
                        
                        # Use the generated pattern
                        suggested_fixes = dynamic_pattern.get('suggested_fixes', [])
                        troubleshooting_hints.extend(suggested_fixes)
                        
            except Exception as e:
                print(f"Could not load troubleshooting hints: {e}")
                troubleshooting_hints = [
                    "Validate network configuration and parameters in config files",
                    "Check network reachability between endpoints",
                    "Verify protocol-specific configuration settings",
                    "Review error logs for additional context"
                ]
            
            # Generate terminal commands
            terminal_commands = self.generate_terminal_commands(
                error_message=error_message,
                investigation_steps=investigation_steps,
                deployment_context=deployment_context,
                troubleshooting_hints=troubleshooting_hints,
                openair_codebase_file_name=openair_codebase_file_name
            )
            
            # Add commands to results
            results['phase4_commands'] = {
                "terminal_commands": terminal_commands,
                "command_count": len(terminal_commands)
            }
            
            print(f"DEBUG [Bug Discovery]: Generated {len(terminal_commands)} terminal commands")
            
            # Display commands (like old version)
            if terminal_commands:
                print(f"\n💻 Generated Terminal Commands ({len(terminal_commands)} commands):")
                for i, cmd in enumerate(terminal_commands, 1):
                    print(f"\n   {i}. {cmd['command']}")
                    print(f"      💡 {cmd['explanation']}")
            else:
                print(f"\n⚠️  No terminal commands generated")
            
            # Save complete results to output directory (like old version)
            output_file = "output/complete_error_analysis.json"
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            
            print(f"💾 Complete results saved to: {output_file}")
            
            # Save fix suggestions separately (including terminal commands)
            fix_suggestions_file = "output/fix_suggestions.json"
            fix_suggestions_data = results.get('phase3_fixes', {}).copy()
            fix_suggestions_data['terminal_commands'] = results.get('phase4_commands', {})
            
            with open(fix_suggestions_file, 'w', encoding='utf-8') as f:
                json.dump(fix_suggestions_data, f, indent=2, ensure_ascii=False)
            
            print(f"🔧 Fix suggestions saved to: {fix_suggestions_file}")
            
            # Save deployment context if available
            if results.get('deployment_context'):
                context_file = "output/deployment_context.json"
                with open(context_file, 'w', encoding='utf-8') as f:
                    json.dump(results['deployment_context'], f, indent=2, ensure_ascii=False)
                print(f"🌐 Deployment context saved to: {context_file}")
            
            # Save summary report (like old version)
            summary_file = "output/error_fix_summary.txt"
            with open(summary_file, 'w', encoding='utf-8') as f:
                f.write(f"Error Fix Summary Report\n")
                f.write(f"=" * 60 + "\n\n")
                f.write(f"Error: {error_message}\n")
                f.write(f"Log File: {log_file_path or 'None'}\n")
                f.write(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                # Deployment context
                if results.get('deployment_context'):
                    ctx = results['deployment_context']
                    f.write(f"Deployment Context:\n")
                    f.write(f"- Role: {ctx.get('role', 'Unknown')}\n")
                    f.write(f"- Active Configs: {len(ctx.get('active_configs', []))}\n")
                    f.write(f"- Network: gNB={ctx.get('network_params', {}).get('gnb_ipv4')}, AMF={ctx.get('network_params', {}).get('amf_ipv4')}\n\n")
                
                # Phase 2
                phase2 = results.get('phase2_analysis', {})
                f.write(f"Phase 2 Results:\n")
                f.write(f"- Retrieval Method: {phase2.get('retrieval_method', 'standard')}\n")
                f.write(f"- Functions: {len(phase2.get('suspected_functions', []))}\n")
                f.write(f"- Configs: {len(phase2.get('suspected_configs', []))}\n\n")
                
                # Phase 3
                phase3 = results.get('phase3_fixes', {})
                fix_suggestion = phase3.get('fix_suggestion', {})
                f.write(f"Phase 3 Results:\n")
                f.write(f"- Root Cause: {fix_suggestion.get('reason', 'Not provided')[:200]}...\n")
                f.write(f"- Fix Available: {'Yes' if fix_suggestion.get('config_fix') or fix_suggestion.get('code_patch') else 'No'}\n")
            
            print(f"📄 Summary report saved to: {summary_file}")
            
            os.chdir(original_cwd)
            
            self.bug_analysis_progress.setValue(70)
            QApplication.processEvents()
            
            # Display results - use error message from pipeline results if available
            pipeline_error_message = results.get('error_message', error_message)
            
            # Store results for later use (use pipeline's error message)
            self.current_bug_analysis = {
                'error_message': pipeline_error_message,  # Use pipeline's extracted error
                'log_file': log_file,
                'log_path': log_file_path,
                'code_dir': code_dir,
                'results': results,
                'timestamp': datetime.now().isoformat()
            }
            
            self.display_bug_analysis_results(results, pipeline_error_message)
            
            # Save to history
            self.save_bug_analysis_to_history()
            
            self.bug_analysis_progress.setValue(100)
            self.view_artifacts_bug_btn.setEnabled(True)
            self.goto_code_evaluation_btn.setEnabled(True)
            
            QMessageBox.information(self, "Analysis Complete", "RCA Analysis completed successfully!")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to run RCA analysis:\n{str(e)}")
            print(f"Error in handle_bug_rca_analysis: {e}")
            import traceback
            traceback.print_exc()
        finally:
            self.bug_analysis_progress.hide()
            if os.getcwd() != original_cwd:
                os.chdir(original_cwd)

    def display_bug_analysis_results(self, results, error_message):
        """Display the bug analysis results in the text area."""
        self.bug_results_display.clear()
        
        # Check if this is crash analysis
        is_crash_analysis = results.get('crash_analysis', False)
        
        # Format the results with enhanced styling (matching old Code Assistant format exactly)
        if is_crash_analysis:
            formatted_text = "🔬 CRASH ANALYSIS RESULTS\n"
        else:
            formatted_text = "🔧 RCA ANALYSIS RESULTS\n"
        formatted_text += "=" * 80 + "\n\n"
        
        # Error Information
        formatted_text += f"🔥 ERROR DETECTED:\n"
        formatted_text += "-" * 40 + "\n"
        formatted_text += f"{error_message}\n\n"
        
        # Crash-specific information
        if is_crash_analysis and 'extraction_summary' in results:
            extraction = results['extraction_summary']
            formatted_text += "💥 CRASH DETAILS:\n"
            formatted_text += "-" * 40 + "\n"
            formatted_text += f"  Signal: {extraction.get('signal', 'Unknown')}\n"
            formatted_text += f"  Crash Type: {extraction.get('crash_type', 'Unknown')}\n"
            formatted_text += f"  Fault Location: {extraction.get('fault_location', 'Unknown')}\n"
            formatted_text += f"  Faulting Function: {extraction.get('faulting_function', 'Unknown')}\n"
            formatted_text += f"  Thread: {extraction.get('crash_thread', {}).get('name', 'Unknown')}\n"
            formatted_text += f"  Backtrace Frames: {extraction.get('backtrace_frames', 0)}\n"
            formatted_text += f"  Scenario Steps Before Crash: {extraction.get('scenario_steps_before_crash', 0)}\n\n"
            
            # Show backtrace
            if results.get('backtrace'):
                formatted_text += "📚 BACKTRACE:\n"
                formatted_text += "-" * 40 + "\n"
                for frame in results['backtrace'][:5]:  # Top 5 frames
                    formatted_text += f"  #{frame.get('frame_number', '?')}: {frame.get('function', 'Unknown')}\n"
                    formatted_text += f"      at {frame.get('file', 'Unknown')}:{frame.get('line', '?')}\n"
                formatted_text += "\n"
            
            # Show scenario flow
            if results.get('scenario_flow'):
                formatted_text += "🔄 SCENARIO FLOW (Steps before crash):\n"
                formatted_text += "-" * 40 + "\n"
                for i, step in enumerate(results['scenario_flow'][:10], 1):  # Top 10 steps
                    formatted_text += f"  {i}. {step}\n"
                formatted_text += "\n"
        
        # Phase 3 Analysis Results
        phase3_fixes = results.get('phase3_fixes', {})
        fix_suggestion = phase3_fixes.get('fix_suggestion', {})
        
        # Extract counts for summary
        code_patches = fix_suggestion.get('code_patches', [])
        config_patches = fix_suggestion.get('config_patches', [])
        suspected_functions = fix_suggestion.get('suspected_functions', [])
        suspected_configs = fix_suggestion.get('suspected_configs', [])
        
        # Suspected Functions (for investigation - not actual patches)
        if suspected_functions:
            formatted_text += f"🔍 SUSPECTED FUNCTIONS:\n"
            formatted_text += "-" * 40 + "\n"
            for i, func in enumerate(suspected_functions, 1):
                formatted_text += f"  {i}. {func}\n"
            formatted_text += "\n"
        
        # Suspected Configs (for investigation - not actual patches)
        if suspected_configs:
            formatted_text += f"⚙️ SUSPECTED CONFIGURATIONS:\n"
            formatted_text += "-" * 40 + "\n"
            for i, config in enumerate(suspected_configs, 1):
                formatted_text += f"  {i}. {config}\n"
            formatted_text += "\n"
        
        # Code Patches (Enhanced) - ACTUAL PATCHES TO APPLY
        if code_patches:
            formatted_text += f"\n💻 CODE PATCHES:\n"
            formatted_text += "─" * 50 + "\n"
            for i, patch in enumerate(code_patches, 1):
                formatted_text += f"  {i}. Function: {patch.get('function_name', 'Unknown')}\n"
                formatted_text += f"     File: {patch.get('file_path', 'Unknown')}\n"
                formatted_text += f"     Type: {patch.get('patch_type', 'Unknown')}\n"
                formatted_text += f"     Lines: {patch.get('line_numbers', 'Unknown')}\n"
                formatted_text += f"     Description: {patch.get('description', 'No description')}\n"
                
                if patch.get('original_code') and patch.get('patched_code'):
                    # Format as git-style diff
                    formatted_text += f"     📝 Code Changes:\n"
                    formatted_text += f"     {'─' * 50}\n"
                    
                    # Parse original and patched code for diff-style display
                    original_lines = patch['original_code'].strip().split('\n')
                    patched_lines = patch['patched_code'].strip().split('\n')
                    
                    # Show original code in red (deletion style)
                    formatted_text += f"     🔴 Original Code:\n"
                    for line in original_lines:
                        if line.strip():
                            formatted_text += f"     - {line}\n"
                    
                    # Show patched code in green (addition style)
                    formatted_text += f"\n     🟢 Patched Code:\n"
                    for line in patched_lines:
                        if line.strip():
                            formatted_text += f"     + {line}\n"
                    
                    formatted_text += f"     {'─' * 50}\n"
                elif patch.get('original_code'):
                    formatted_text += f"     Original Code:\n"
                    formatted_text += f"     {patch['original_code']}\n"
                elif patch.get('patched_code'):
                    formatted_text += f"     Patched Code:\n"
                    formatted_text += f"     {patch['patched_code']}\n"
                formatted_text += "\n"
        
        # Config Patches - ACTUAL PATCHES TO APPLY
        if config_patches:
            formatted_text += f"\n⚙️ CONFIG PATCHES:\n"
            formatted_text += "─" * 50 + "\n"
            for i, patch in enumerate(config_patches, 1):
                formatted_text += f"  {i}. Config: {patch.get('config_name', 'Unknown')}\n"
                formatted_text += f"     File: {patch.get('file_path', 'Unknown')}\n"
                formatted_text += f"     Line: {patch.get('line_number', 'Unknown')}\n"
                formatted_text += f"     Description: {patch.get('description', 'No description')}\n"
                
                # Format config changes as git-style diff
                if patch.get('current_value') and patch.get('new_value'):
                    formatted_text += f"     📝 Configuration Changes:\n"
                    formatted_text += f"     {'─' * 50}\n"
                    formatted_text += f"     🔴 Current Value: {patch.get('current_value', 'Unknown')}\n"
                    formatted_text += f"     🟢 New Value:     {patch.get('new_value', 'Unknown')}\n"
                    formatted_text += f"     {'─' * 50}\n"
                else:
                    formatted_text += f"     Current: {patch.get('current_value', 'Unknown')}\n"
                    formatted_text += f"     New: {patch.get('new_value', 'Unknown')}\n"
                
                formatted_text += "\n"
        
        # Investigation Steps
        investigation_steps = fix_suggestion.get('investigation_steps', [])
        if investigation_steps:
            formatted_text += "\n📋 INVESTIGATION STEPS:\n"
            formatted_text += "─" * 50 + "\n"
            for i, step in enumerate(investigation_steps, 1):
                formatted_text += f"  {i}. {step}\n"
            formatted_text += "\n"
        
        # Detailed Root Cause (if available - show AFTER investigation steps)
        if fix_suggestion.get('root_cause_analysis'):
            formatted_text += "\n🔬 DETAILED ROOT CAUSE:\n"
            formatted_text += "─" * 50 + "\n"
            formatted_text += f"{fix_suggestion['root_cause_analysis']}\n\n"
        
        # Terminal Commands (Investigation Commands) - Skip for crash analysis
        # Based on the old version, commands are stored in phase4_commands
        is_crash_analysis = results.get('crash_analysis', False)
        terminal_commands = None
        
        # Skip terminal commands for crash analysis
        if is_crash_analysis:
            formatted_text += "\n💡 TIP: Use the existing patch applicator to apply the generated code patches.\n"
            formatted_text += "   Command: python unified_patch_applicator.py --input output/crash_phase3_fixes.json\n\n"
        else:
            # Normal flow: Show terminal commands
            # Option 1: In phase4_commands (this is the correct location based on old code)
            if 'phase4_commands' in results:
                phase4 = results['phase4_commands']
                if isinstance(phase4, dict):
                    terminal_commands = phase4.get('terminal_commands', None)
            
            # Option 2: Direct in results['terminal_commands'] (fallback)
            if not terminal_commands and 'terminal_commands' in results:
                if isinstance(results['terminal_commands'], dict):
                    terminal_commands = results['terminal_commands'].get('terminal_commands', None)
                elif isinstance(results['terminal_commands'], list):
                    terminal_commands = results['terminal_commands']
            
            if terminal_commands and len(terminal_commands) > 0:
                formatted_text += "\n💻 INVESTIGATION COMMANDS:\n"
                formatted_text += "─" * 50 + "\n"
                for i, cmd in enumerate(terminal_commands, 1):
                    if isinstance(cmd, dict):
                        formatted_text += f"  {i}. {cmd.get('command', 'Unknown command')}\n"
                        if cmd.get('explanation'):
                            formatted_text += f"     💡 {cmd['explanation']}\n"
                    elif isinstance(cmd, str):
                        formatted_text += f"  {i}. {cmd}\n"
                    formatted_text += "\n"
        
        # Analysis Context
        context_summary = results.get('summary', {})
        if context_summary:
            formatted_text += "📊 ANALYSIS CONTEXT:\n"
            formatted_text += "-" * 40 + "\n"
            formatted_text += f"  Functions analyzed: {context_summary.get('total_functions_analyzed', 0)}\n"
            formatted_text += f"  Configs analyzed: {context_summary.get('total_configs_analyzed', 0)}\n"
            formatted_text += f"  Call graph entries: {context_summary.get('call_graph_entries', 0)}\n"
            formatted_text += f"  Pattern matched: {context_summary.get('pattern_matched', False)}\n\n"
        
        # Analysis Summary
        formatted_text += "📊 ANALYSIS SUMMARY:\n"
        formatted_text += "-" * 40 + "\n"
        
        # Generate timestamp for results file
        timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        results_filename = f"bug_analysis_{timestamp_str}.json"
        
        formatted_text += f"Analysis completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        formatted_text += f"Results file: {results_filename}\n"
        
        # Count total sections (deployment_context, phase2, phase3, phase4)
        total_sections = 0
        if 'deployment_context' in results and results.get('deployment_context'): total_sections += 1
        if 'phase2_analysis' in results and results.get('phase2_analysis'): total_sections += 1
        if 'phase3_fixes' in results and results.get('phase3_fixes'): total_sections += 1
        if 'phase4_commands' in results and results.get('phase4_commands'): total_sections += 1
        
        formatted_text += f"Total sections: {total_sections}\n"
        
        # Debug output to verify patch counts
        print(f"DEBUG [Bug Discovery Display]: Suspected Functions: {len(suspected_functions)}, Code Patches: {len(code_patches)}, Config Patches: {len(config_patches)}")
        
        self.bug_results_display.setText(formatted_text)

    def show_bug_artifacts(self):
        """Show detailed error artifacts from Bug Discovery analysis."""
        if not hasattr(self, 'current_bug_analysis'):
            QMessageBox.warning(self, "No Analysis", "Please run RCA analysis first.")
            return
        
        try:
            artifacts_data = self.current_bug_analysis.get('results', {})
            
            # Create dialog for artifacts display
            dialog = QDialog(self)
            dialog.setWindowTitle("Error Analysis Artifacts")
            dialog.setMinimumSize(1200, 750)
            
            # Apply professional blue theme styling to the dialog
            dialog.setStyleSheet("""
                QDialog {
                    background-color: #f0f4f8;
                    color: #1e3a5f;
                }
                QTabWidget::pane {
                    border: 1px solid #b3d9ff;
                    background-color: #ffffff;
                    border-radius: 4px;
                }
                QTabBar::tab {
                    background-color: #e6f2ff;
                    color: #1e3a5f;
                    padding: 8px 16px;
                    margin-right: 2px;
                    border-top-left-radius: 4px;
                    border-top-right-radius: 4px;
                }
                QTabBar::tab:selected {
                    background-color: #ffffff;
                    color: #1e3a5f;
                    border-bottom: 2px solid #0066cc;
                }
                QTabBar::tab:hover {
                    background-color: #cce7ff;
                }
            """)
            
            layout = QVBoxLayout()
            
            # Create tab widget for different artifact types
            tab_widget = QTabWidget()
            
            # Add bug analysis specific data to artifacts_data for compatibility
            artifacts_data['error_message'] = self.current_bug_analysis.get('error_message', 'Unknown error')
            artifacts_data['timestamp'] = self.current_bug_analysis.get('timestamp', 'Unknown')
            # Use log_path which contains the full path to the log file
            artifacts_data['log_file'] = self.current_bug_analysis.get('log_path', self.current_bug_analysis.get('log_file', 'None'))
            
            # Error Information Tab
            error_tab = self.create_error_info_tab(artifacts_data)
            tab_widget.addTab(error_tab, "🔍 Error Information")
            
            # Crash-specific tabs
            is_crash = artifacts_data.get('crash_analysis', False)
            if is_crash:
                # Backtrace Tab (crash-specific)
                backtrace_tab = self.create_crash_backtrace_tab(artifacts_data)
                tab_widget.addTab(backtrace_tab, "📚 Backtrace")
                
                # Scenario Flow Tab (crash-specific)
                scenario_tab = self.create_crash_scenario_tab(artifacts_data)
                tab_widget.addTab(scenario_tab, "🔄 Scenario Flow")
            
            # Suspected Functions Tab
            functions_tab = self.create_functions_artifacts_tab(artifacts_data)
            tab_widget.addTab(functions_tab, "⚙️ Suspected Functions")
            
            # Suspected Configs Tab (may be empty for crashes)
            configs_tab = self.create_configs_artifacts_tab(artifacts_data)
            tab_widget.addTab(configs_tab, "⚙️ Suspected Configs")
            
            # Context Information Tab
            context_tab = self.create_context_artifacts_tab(artifacts_data)
            tab_widget.addTab(context_tab, "📋 Context Information")
            
            # Analysis Summary Tab
            summary_tab = self.create_analysis_summary_tab(artifacts_data)
            tab_widget.addTab(summary_tab, "📊 Analysis Summary")
            
            # 3GPP Spec Reference Tab
            spec_tab = self.create_3gpp_spec_tab(artifacts_data)
            tab_widget.addTab(spec_tab, "📚 3GPP Spec Reference")
            
            # Impact Analysis Tab
            impact_tab = self.create_impact_analysis_tab(artifacts_data)
            tab_widget.addTab(impact_tab, "⚠️ Impact Analysis")
            
            layout.addWidget(tab_widget)
            
            # Export button
            export_btn = QPushButton("Export Artifacts")
            export_btn.clicked.connect(lambda: self.export_artifacts(artifacts_data))
            layout.addWidget(export_btn)
            
            dialog.setLayout(layout)
            dialog.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to show error artifacts: {str(e)}")
            print(f"Error in show_bug_artifacts: {e}")
            import traceback
            traceback.print_exc()

    def save_bug_analysis_to_history(self):
        """Save the bug analysis results to history file."""
        if not hasattr(self, 'current_bug_analysis'):
            return
        
        # Save current directory
        current_dir = os.getcwd()
        
        # Get the main workspace root (where backend/ folder is)
        workspace_root = os.path.dirname(os.path.abspath(__file__))
        
        # Create history directory if it doesn't exist (use absolute path)
        history_dir = os.path.join(workspace_root, "backend", "resources", "bug_history")
        os.makedirs(history_dir, exist_ok=True)
        
        # Create unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        history_file = os.path.join(history_dir, f"bug_analysis_{timestamp}.json")
        
        # Prepare data to save
        history_data = {
            'error_message': self.current_bug_analysis.get('error_message', ''),
            'log_file': self.current_bug_analysis.get('log_file', ''),
            'log_path': self.current_bug_analysis.get('log_path', ''),
            'code_dir': self.current_bug_analysis.get('code_dir', ''),
            'timestamp': self.current_bug_analysis.get('timestamp', ''),
            'results': self.current_bug_analysis.get('results', {}),
            'history_file': history_file
        }
        
        # Save to file (history_file already has full path from os.path.join above)
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history_data, f, indent=2)
        
        # Restore original directory if changed
        if os.getcwd() != current_dir:
            os.chdir(current_dir)
        
        # Debug: Verify patch counts before saving
        results = history_data.get('results', {})
        phase3_fixes = results.get('phase3_fixes', {})
        fix_suggestion = phase3_fixes.get('fix_suggestion', {})
        code_patches = fix_suggestion.get('code_patches', [])
        config_patches = fix_suggestion.get('config_patches', [])
        suspected_functions = fix_suggestion.get('suspected_functions', [])
        
        print(f"\n{'='*60}")
        print(f"DEBUG [Save to History]: Saving analysis to history")
        print(f"{'='*60}")
        print(f"DEBUG: Suspected Functions: {len(suspected_functions)} (investigation only)")
        if suspected_functions:
            print(f"       Names: {suspected_functions}")
        print(f"DEBUG: Code Patches: {len(code_patches)} (will be available in Code Assistant)")
        if code_patches:
            print(f"       Functions: {[p.get('function_name', 'N/A') for p in code_patches]}")
        print(f"DEBUG: Config Patches: {len(config_patches)} (will be available in Code Assistant)")
        if config_patches:
            print(f"       Configs: {[p.get('config_name', p.get('parameter_name', 'N/A')) for p in config_patches]}")
        print(f"DEBUG: File: {history_file}")
        print(f"{'='*60}\n")

    def go_to_code_evaluation_from_bug_discovery(self):
        """Navigate from Bug Discovery to Code Evaluation with the latest analysis selected."""
        try:
            if hasattr(self, 'bug_discovery_dialog'):
                self.bug_discovery_dialog.accept()

            self.show_code_testing()

            def select_latest_analysis():
                if hasattr(self, 'testing_history_combo'):
                    self.load_testing_bug_history()
                    if self.testing_history_combo.count() > 1:
                        self.testing_history_combo.setCurrentIndex(1)
                        QTimer.singleShot(50, self.load_selected_testing_analysis)
                    else:
                        QMessageBox.information(
                            self,
                            "No Saved Analyses",
                            "No bug analyses found in history. Run an analysis or save results first."
                        )

            QTimer.singleShot(300, select_latest_analysis)

        except Exception as exc:
            QMessageBox.warning(
                self,
                "Navigation Error",
                f"Could not open Code Evaluation: {exc}"
            )

    def go_to_code_assistant_from_bug_discovery(self):
        """Navigate from Bug Discovery to Code Assistant with latest bug selected."""
        try:
            # Close the Bug Discovery dialog
            if hasattr(self, 'bug_discovery_dialog'):
                self.bug_discovery_dialog.accept()
            
            # Open Code Assistant dialog (this will call load_bug_history automatically)
            self.show_code_assistant()
            
            # Wait a bit for the dialog to fully initialize, then select latest bug
            def select_and_load_latest():
                if hasattr(self, 'bug_history_combo'):
                    # Reload to ensure we have the latest saved bug
                    self.load_bug_history()
                    
                    print(f"DEBUG: Bug history count: {self.bug_history_combo.count()}")
                    
                    # Select the latest bug (index 1, since 0 is "Select a bug analysis...")
                    if self.bug_history_combo.count() > 1:
                        self.bug_history_combo.setCurrentIndex(1)
                        print(f"DEBUG: Selected bug at index 1: {self.bug_history_combo.currentText()}")
                        
                        # Automatically load the selected bug
                        QTimer.singleShot(50, self.load_selected_bug_analysis)
                        
                        print("✅ Navigated to Code Assistant with latest bug selected")
                    else:
                        print(f"⚠️  No bug history found. Count: {self.bug_history_combo.count()}")
                        QMessageBox.information(
                            self,
                            "No Bug History",
                            "No bug analyses found in history. The analysis may not have been saved yet."
                        )
                else:
                    print("⚠️  bug_history_combo not found")
            
            # Use longer delay to ensure dialog is fully initialized
            QTimer.singleShot(300, select_and_load_latest)
            
        except Exception as e:
            print(f"❌ Error navigating to Code Assistant: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.warning(
                self,
                "Navigation Error",
                f"Could not navigate to Code Assistant: {str(e)}"
            )

    def go_to_code_assistant_from_testing(self):
        """Open Code Assistant from the Code Evaluation dialog."""
        try:
            if hasattr(self, 'code_testing_dialog'):
                self.code_testing_dialog.accept()

            self.show_code_assistant()

            def select_latest_bug():
                if hasattr(self, 'bug_history_combo'):
                    self.load_bug_history()
                    if self.bug_history_combo.count() > 1:
                        self.bug_history_combo.setCurrentIndex(1)
                        QTimer.singleShot(50, self.load_selected_bug_analysis)

            QTimer.singleShot(300, select_latest_bug)

        except Exception as exc:
            QMessageBox.warning(
                self,
                "Navigation Error",
                f"Could not open Code Assistant: {exc}"
            )

    def check_bug_for_existing_fix(self):
        """Check if a fix already exists in Git commits for the selected log file in Bug Discovery."""
        try:
            if not hasattr(self, 'bug_fix_present_btn'):
                return
            
            # Hide button by default
            self.bug_fix_present_btn.setVisible(False)
            self.bug_existing_fix_result = None
            
            # Get selected log file
            log_file = self.bug_log_file_combo.currentText()
            log_dir = self.bug_log_dir_edit.text()
            
            # Skip if placeholder is selected or no file/dir
            if not log_file or not log_dir or log_file == "Select a log file":
                return
            
            log_file_path = os.path.join(log_dir, log_file)
            
            if not os.path.exists(log_file_path):
                return
            
            # Create progress dialog
            progress = QProgressDialog("Checking for existing fixes in Git history...", "Cancel", 0, 100, self)
            progress.setWindowTitle("Analyzing Log File")
            progress.setWindowModality(Qt.WindowModal)
            progress.setMinimumDuration(0)  # Show immediately
            progress.setValue(0)
            QApplication.processEvents()
            
            print(f"\n🔍 Checking for existing fixes for: {log_file}")
            
            # Step 1: Extract error from log file
            progress.setLabelText("Extracting error message from log file...")
            progress.setValue(20)
            QApplication.processEvents()
            
            if progress.wasCanceled():
                return
            
            from Error_fixing_pipelin.parse_log_context import LogContextParser
            
            # Get the source code directory
            code_dir = self.bug_code_dir_edit.text() if hasattr(self, 'bug_code_dir_edit') else None
            if code_dir:
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            else:
                openair_codebase_file_name = "openairinterface5g-develop"  # Default fallback
            
            log_parser = LogContextParser(openair_codebase_file_name=openair_codebase_file_name)
            error_message = log_parser.extract_error_message(log_file_path)
            
            if not error_message:
                print("   No clear error extracted from log")
                progress.close()
                return
            
            print(f"   Error extracted: {error_message[:100]}...")
            
            # Step 2: Search embeddings
            progress.setLabelText("Searching 30,000+ commits for similar fixes...")
            progress.setValue(50)
            QApplication.processEvents()
            
            if progress.wasCanceled():
                return
            
            from Error_fixing_pipelin.smart_commit_selector import CommitSearcher, SmartSelector
            
            searcher = CommitSearcher(embeddings_dir='Error_fixing_pipelin/resources/embeddings', 
                                     validate_commits=False,
                                     openair_codebase_file_name=openair_codebase_file_name)
            search_results = searcher.search(error_message, top_k=10)
            
            if not search_results:
                print("   No similar commits found")
                progress.close()
                return
            
            # Step 3: Intelligent selection
            progress.setLabelText("Analyzing matches with smart selector...")
            progress.setValue(75)
            QApplication.processEvents()
            
            if progress.wasCanceled():
                return
            
            selector = SmartSelector(use_llm=False)
            selection_result = selector.select_best_fix(error_message, search_results)
            
            # Step 4: Check results
            progress.setLabelText("Finalizing results...")
            progress.setValue(90)
            QApplication.processEvents()
            
            # Check if a good fix was found
            if selection_result['status'] in ['auto_selected', 'suggested', 'llm_verified']:
                if selection_result['commit'] and selection_result['commit']['is_rca_commit']:
                    # RCA fix found!
                    confidence = selection_result['confidence']
                    score = selection_result['commit'].get('boosted_score', 
                                                           selection_result['commit']['similarity'])
                    
                    print(f"   ✅ Found existing RCA fix! Confidence: {confidence}, Score: {score:.2%}")
                    
                    # Store the result
                    self.bug_existing_fix_result = selection_result
                    
                    # Show the button
                    self.bug_fix_present_btn.setVisible(True)
                    self.bug_fix_present_btn.setText(f"✅ Fix Already Present ({confidence}) - Click to View")
                else:
                    print("   Found similar commits but no RCA fixes")
            else:
                print("   No suitable fix found")
            
            # Complete
            progress.setValue(100)
            progress.close()
                
        except Exception as e:
            if 'progress' in locals():
                progress.close()
            print(f"⚠️ Error checking for existing fix: {e}")
            import traceback
            traceback.print_exc()

    def show_bug_existing_fix_details(self):
        """Show details of the existing fix found in Git history for Bug Discovery."""
        if not hasattr(self, 'bug_existing_fix_result') or not self.bug_existing_fix_result:
            QMessageBox.information(self, "No Fix", "No existing fix data available.")
            return
        
        result = self.bug_existing_fix_result
        commit = result.get('commit')
        
        if not commit:
            QMessageBox.information(self, "No Fix", "No fix commit available.")
            return
        
        try:
            # Get the source code directory
            code_dir = self.bug_code_dir_edit.text() if hasattr(self, 'bug_code_dir_edit') else None
            if code_dir:
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            else:
                openair_codebase_file_name = "openairinterface5g-develop"
            
            # Get full commit details
            commit_hash = commit.get('commit_hash', commit.get('commit_hash_short', ''))
            if not commit_hash:
                self.bug_results_display.setText("Error: No commit hash available.")
                return
            
            # Show loading message
            self.bug_results_display.setText("Loading commit details from Git repository...\n")
            QApplication.processEvents()
            
            # Try to get commit details from Git
            import subprocess
            git_show_cmd = ['git', 'show', '--stat', '--format=fuller', '--unified=3', commit_hash]
            
            possible_paths = [
                os.path.join(os.getcwd(), openair_codebase_file_name),
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', openair_codebase_file_name),
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', 'openairinterface5g-develop'),
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', 'openairinterface5g-test'),
            ]
            
            git_dir = None
            for path in possible_paths:
                if os.path.exists(path) and os.path.exists(os.path.join(path, '.git')):
                    try:
                        check_result = subprocess.run(
                            ['git', 'cat-file', '-e', commit_hash],
                            cwd=path,
                            capture_output=True,
                            timeout=2
                        )
                        if check_result.returncode == 0:
                            git_dir = path
                            break
                    except:
                        continue
            
            if not git_dir:
                self.bug_results_display.setText("Error: Could not find Git repository with this commit.")
                return
            
            git_result = subprocess.run(
                git_show_cmd,
                cwd=git_dir,
                capture_output=True,
                text=True,
                timeout=10
            )
            
            if git_result.returncode != 0:
                # Fallback to basic display
                self._display_bug_basic_fix_details(commit, result)
                # Still prepare the fix data for Code Assistant
                self._prepare_existing_fix_for_code_assistant(commit, result)
                return
            
            # Parse git show output
            git_output = git_result.stdout
            lines = git_output.split('\n')
            
            # Extract commit info
            commit_info = {}
            in_message = False
            message_lines = []
            
            for line in lines:
                if line.startswith('commit '):
                    commit_info['full_hash'] = line.split()[1]
                elif line.startswith('Author:'):
                    commit_info['author'] = line.replace('Author:', '').strip()
                elif line.startswith('AuthorDate:'):
                    commit_info['date'] = line.replace('AuthorDate:', '').strip()
                elif line.startswith('Commit:'):
                    commit_info['committer'] = line.replace('Commit:', '').strip()
                elif line.startswith('CommitDate:'):
                    commit_info['commit_date'] = line.replace('CommitDate:', '').strip()
                elif line.startswith('    ') and not in_message:
                    in_message = True
                    message_lines.append(line.strip())
                elif in_message and line.startswith('    '):
                    message_lines.append(line.strip())
                elif in_message and not line.startswith('    '):
                    break
            
            # Format the detailed display
            details = self._format_bug_fix_details(commit, result, commit_info, message_lines, git_output)
            
            # Display in the results text area
            self.bug_results_display.setText(details)
            
            # Prepare the fix data for Code Assistant
            self._prepare_existing_fix_for_code_assistant(commit, result)
            
        except Exception as e:
            print(f"Error getting commit details: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to basic display
            self._display_bug_basic_fix_details(commit, result)
            # Still prepare the fix data for Code Assistant
            self._prepare_existing_fix_for_code_assistant(commit, result)
    
    def _display_bug_basic_fix_details(self, commit, result):
        """Display basic fix details when git show fails in Bug Discovery"""
        msg = f"🔍 EXISTING FIX FOUND\n"
        msg += f"{'='*60}\n\n"
        
        # Basic commit info
        msg += f"📋 COMMIT DETAILS\n"
        msg += f"{'-'*30}\n"
        msg += f"Hash: {commit.get('commit_hash_short', 'N/A')}\n"
        msg += f"Author: {commit.get('author_name', 'N/A')}\n"
        msg += f"Date: {commit.get('date_iso', 'N/A')}\n"
        msg += f"Subject: {commit.get('subject', 'N/A')}\n\n"
        
        # Branch information
        msg += f"🌿 BRANCH INFORMATION\n"
        msg += f"{'-'*30}\n"
        msg += f"Fix Source: origin/develop\n"
        
        # Get current branch of target repository
        current_branch = self._get_current_branch()
        msg += f"Target Branch: {current_branch}\n\n"
        
        # Match details
        msg += f"🎯 MATCH DETAILS\n"
        msg += f"{'-'*30}\n"
        msg += f"Confidence: {result.get('confidence', 'N/A')}\n"
        msg += f"Similarity Score: {commit.get('boosted_score', commit.get('similarity', 0)):.2%}\n"
        msg += f"Reasoning: {result.get('reasoning', 'N/A')}\n\n"
        
        # Commit message from embeddings
        if commit.get('body'):
            msg += f"📝 COMMIT MESSAGE\n"
            msg += f"{'-'*30}\n"
            msg += f"{commit.get('body', 'N/A')}\n\n"
        
        # Patch information
        msg += f"🔧 PATCH INFORMATION\n"
        msg += f"{'-'*30}\n"
        
        if commit.get('code_patches'):
            msg += f"Code Patches: {commit.get('code_patch_count', 0)}\n"
            for i, patch in enumerate(commit.get('code_patches', [])[:5], 1):
                if isinstance(patch, dict):
                    msg += f"  {i}. {patch.get('function', 'N/A')} ({patch.get('file', 'N/A')})\n"
                else:
                    msg += f"  {i}. {patch}\n"
            if len(commit.get('code_patches', [])) > 5:
                msg += f"  ... and {len(commit.get('code_patches', [])) - 5} more\n"
        
        if commit.get('config_patches'):
            msg += f"Config Patches: {commit.get('config_patch_count', 0)}\n"
            for i, patch in enumerate(commit.get('config_patches', [])[:5], 1):
                if isinstance(patch, dict):
                    msg += f"  {i}. {patch.get('parameter', 'N/A')} ({patch.get('file', 'N/A')})\n"
                else:
                    msg += f"  {i}. {patch}\n"
            if len(commit.get('config_patches', [])) > 5:
                msg += f"  ... and {len(commit.get('config_patches', [])) - 5} more\n"
        
        # Files changed
        if commit.get('files_changed'):
            msg += f"\n📁 FILES CHANGED\n"
            msg += f"{'-'*30}\n"
            for file_info in commit.get('files_changed', [])[:10]:
                if isinstance(file_info, dict):
                    msg += f"  • {file_info.get('file', 'N/A')} ({file_info.get('changes', 'N/A')})\n"
                else:
                    msg += f"  • {file_info}\n"
            if len(commit.get('files_changed', [])) > 10:
                msg += f"  ... and {len(commit.get('files_changed', [])) - 10} more files\n"
        
        # Recommendation
        msg += f"\n💡 RECOMMENDATION\n"
        msg += f"{'-'*30}\n"
        msg += f"This fix was previously applied to resolve a similar issue.\n"
        msg += f"Review the patches above and apply similar changes to your codebase.\n"
        msg += f"\nNote: Full git details not available. The commit may have been\n"
        msg += f"deleted or the repository path is incorrect.\n"
        
        self.bug_results_display.setText(msg)
    
    def _format_bug_fix_details(self, commit, result, commit_info, message_lines, git_output):
        """Format comprehensive fix details for Bug Discovery"""
        details = f"🔍 EXISTING FIX FOUND\n"
        details += f"{'='*60}\n\n"
        
        # Basic commit info
        details += f"📋 COMMIT DETAILS\n"
        details += f"{'-'*30}\n"
        details += f"Hash: {commit_info.get('full_hash', commit.get('commit_hash_short', 'N/A'))}\n"
        details += f"Author: {commit_info.get('author', commit.get('author_name', 'N/A'))}\n"
        details += f"Date: {commit_info.get('date', commit.get('date_iso', 'N/A'))}\n"
        details += f"Committer: {commit_info.get('committer', 'N/A')}\n"
        details += f"Commit Date: {commit_info.get('commit_date', 'N/A')}\n\n"
        
        # Branch information
        details += f"🌿 BRANCH INFORMATION\n"
        details += f"{'-'*30}\n"
        details += f"Fix Source: origin/develop\n"
        
        # Get current branch of target repository
        current_branch = self._get_current_branch()
        details += f"Target Branch: {current_branch}\n\n"
        
        # Match details
        details += f"🎯 MATCH DETAILS\n"
        details += f"{'-'*30}\n"
        details += f"Confidence: {result.get('confidence', 'N/A')}\n"
        details += f"Similarity Score: {commit.get('boosted_score', commit.get('similarity', 0)):.2%}\n"
        details += f"Reasoning: {result.get('reasoning', 'N/A')}\n\n"
        
        # Commit message
        if message_lines:
            details += f"📝 COMMIT MESSAGE\n"
            details += f"{'-'*30}\n"
            details += '\n'.join(message_lines) + "\n\n"
        
        # Patch information
        details += f"🔧 PATCH INFORMATION\n"
        details += f"{'-'*30}\n"
        if commit.get('code_patches'):
            details += f"Code Patches: {commit.get('code_patch_count', 0)}\n"
            for i, patch in enumerate(commit.get('code_patches', [])[:3], 1):
                if isinstance(patch, dict):
                    details += f"  {i}. {patch.get('function', 'N/A')} ({patch.get('file', 'N/A')})\n"
                else:
                    details += f"  {i}. {patch}\n"
            if len(commit.get('code_patches', [])) > 3:
                details += f"  ... and {len(commit.get('code_patches', [])) - 3} more\n"
        
        if commit.get('config_patches'):
            details += f"Config Patches: {commit.get('config_patch_count', 0)}\n"
            for i, patch in enumerate(commit.get('config_patches', [])[:3], 1):
                if isinstance(patch, dict):
                    details += f"  {i}. {patch.get('parameter', 'N/A')} ({patch.get('file', 'N/A')})\n"
                else:
                    details += f"  {i}. {patch}\n"
            if len(commit.get('config_patches', [])) > 3:
                details += f"  ... and {len(commit.get('config_patches', [])) - 3} more\n"
        
        # Files changed (from git show)
        if 'files changed' in git_output.lower():
            details += f"\n📁 FILES CHANGED\n"
            details += f"{'-'*30}\n"
            file_section = False
            for line in git_output.split('\n'):
                if 'files changed' in line.lower():
                    file_section = True
                    details += f"{line}\n"
                elif file_section and line.strip() and not line.startswith(' '):
                    break
                elif file_section and line.strip():
                    details += f"{line}\n"
        
        # Show actual code changes (diff)
        details += f"\n🔍 CODE CHANGES (DIFF)\n"
        details += f"{'-'*30}\n"
        
        # Parse and format the diff output
        diff_lines = []
        in_diff = False
        current_file = None
        
        for line in git_output.split('\n'):
            if line.startswith('diff --git'):
                in_diff = True
                current_file = line.split()[-1] if len(line.split()) > 3 else "Unknown file"
                diff_lines.append(f"\n📄 File: {current_file}")
                diff_lines.append("─" * 50)
            elif in_diff and line.startswith('@@'):
                continue
            elif in_diff and (line.startswith('+') or line.startswith('-') or line.startswith(' ')):
                if line.startswith('+'):
                    diff_lines.append(f"➕ {line[1:]}")
                elif line.startswith('-'):
                    diff_lines.append(f"➖ {line[1:]}")
                else:
                    diff_lines.append(f"  {line[1:]}")
            elif in_diff and line.strip() == '':
                diff_lines.append("")
            elif in_diff and not line.startswith(('diff', 'index', '+++', '---', '@@', '+', '-', ' ')):
                in_diff = False
        
        if diff_lines:
            details += '\n'.join(diff_lines[:100])
            if len(diff_lines) > 100:
                details += f"\n\n... and {len(diff_lines) - 100} more lines of changes"
        else:
            details += "No code changes visible in this commit."
        
        # Recommendation
        details += f"\n\n💡 RECOMMENDATION\n"
        details += f"{'-'*30}\n"
        details += f"This fix was previously applied to resolve a similar issue.\n"
        details += f"You can review the changes above and apply similar patches\n"
        details += f"to your current codebase.\n"
        
        return details

    def _prepare_existing_fix_for_code_assistant(self, commit, result):
        """Prepare existing fix data in the format expected by Code Assistant and enable buttons."""
        try:
            # Get log file and directory information
            log_file = self.bug_log_file_combo.currentText() if hasattr(self, 'bug_log_file_combo') else 'N/A'
            log_dir = self.bug_log_dir_edit.text() if hasattr(self, 'bug_log_dir_edit') else ''
            log_file_path = os.path.join(log_dir, log_file) if log_dir and log_file != 'N/A' else ''
            code_dir = self.bug_code_dir_edit.text() if hasattr(self, 'bug_code_dir_edit') else ''
            
            # Extract error message from the search result
            # Try to get it from the log file if available
            error_message = "Existing fix from Git history"
            if log_file_path and os.path.exists(log_file_path):
                try:
                    from Error_fixing_pipelin.parse_log_context import LogContextParser
                    openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep)) if code_dir else "openairinterface5g-develop"
                    log_parser = LogContextParser(openair_codebase_file_name=openair_codebase_file_name)
                    extracted_error = log_parser.extract_error_message(log_file_path)
                    if extracted_error:
                        error_message = extracted_error
                except Exception as e:
                    print(f"Could not extract error message: {e}")
            
            # Get the git diff output to extract actual code changes
            commit_hash = commit.get('commit_hash', commit.get('commit_hash_short', ''))
            git_diff_data = self._get_git_diff_for_commit(commit_hash, code_dir)
            
            # Convert commit patches to the format expected by Code Assistant
            code_patches = []
            config_patches = []
            
            # Extract patches from commit data with actual code from git diff
            if commit.get('code_patches'):
                for patch in commit.get('code_patches', []):
                    if isinstance(patch, dict):
                        function_name = patch.get('function', 'Unknown')
                        file_path_from_commit = patch.get('file', 'Unknown')
                        
                        # Extract actual code changes from git diff and get the full path
                        original_code, patched_code, full_file_path = self._extract_code_from_git_diff(
                            git_diff_data, file_path_from_commit, function_name
                        )
                        
                        # Construct the full path for the file
                        # The git diff gives us the path relative to repo root
                        if code_dir:
                            openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
                        else:
                            openair_codebase_file_name = "openairinterface5g-develop"
                        
                        # Use the full path from git diff if available, otherwise construct it
                        if full_file_path:
                            final_file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/{full_file_path}"
                            print(f"DEBUG: Constructed path from git: {final_file_path}")
                        else:
                            # Fallback: try to construct path based on common patterns
                            file_basename = os.path.basename(file_path_from_commit)
                            if file_basename == 'rrc_gNB.c':
                                final_file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/openair2/RRC/NR/{file_basename}"
                            else:
                                final_file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/{file_path_from_commit}"
                            print(f"DEBUG: Constructed path from fallback: {final_file_path}")
                        
                        # Verify the file exists
                        if os.path.exists(final_file_path):
                            print(f"✅ File verified: {final_file_path}")
                        else:
                            print(f"⚠️  File not found: {final_file_path}")
                        
                        # Debug: Show sample of extracted code
                        print(f"DEBUG: Code patch for {function_name}:")
                        print(f"  Original code (first 100 chars): {original_code[:100]}...")
                        print(f"  Patched code (first 100 chars): {patched_code[:100]}...")
                        
                        code_patches.append({
                            'function_name': function_name,
                            'file_path': final_file_path,
                            'original_code': original_code,
                            'patched_code': patched_code,
                            'suggested_code': patched_code,  # Alias for compatibility
                            'description': f"From commit {commit.get('commit_hash_short', 'N/A')}: {commit.get('subject', 'N/A')}",
                            'patch_type': 'targeted_insertion_or_adjustment',  # Required for replacement mode
                            'commit_hash': commit_hash,
                            'confidence': result.get('confidence', 'N/A'),
                            'similarity_score': commit.get('boosted_score', commit.get('similarity', 0))
                        })
            
            if commit.get('config_patches'):
                for patch in commit.get('config_patches', []):
                    if isinstance(patch, dict):
                        param_name = patch.get('parameter', 'Unknown')
                        file_path_from_commit = patch.get('file', 'Unknown')
                        
                        # Extract config values from git diff
                        current_value, new_value = self._extract_config_from_git_diff(
                            git_diff_data, param_name
                        )
                        
                        # Construct the full path for config files
                        if code_dir:
                            openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
                        else:
                            openair_codebase_file_name = "openairinterface5g-develop"
                        
                        # Config files are typically in specific locations
                        file_basename = os.path.basename(file_path_from_commit)
                        if file_basename.endswith('.conf') or file_basename.endswith('.cfg'):
                            if 'gnb' in file_basename.lower():
                                final_file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/targets/PROJECTS/GENERIC-NR-5GC/CONF/{file_basename}"
                            elif 'ue' in file_basename.lower():
                                final_file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/openair3/NAS/TOOLS/{file_basename}"
                            else:
                                final_file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/{file_path_from_commit}"
                        else:
                            final_file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/{file_path_from_commit}"
                        
                        config_patches.append({
                            'parameter_name': param_name,
                            'config_name': param_name,
                            'file_path': final_file_path,
                            'current_value': current_value,
                            'new_value': new_value,
                            'suggested_value': new_value,  # Alias for compatibility
                            'description': f"From commit {commit.get('commit_hash_short', 'N/A')}: {commit.get('subject', 'N/A')}",
                            'commit_hash': commit_hash,
                            'confidence': result.get('confidence', 'N/A'),
                            'similarity_score': commit.get('boosted_score', commit.get('similarity', 0))
                        })
            
            # Create the results structure in the format expected by Code Assistant
            results = {
                'phase3_fixes': {
                    'fix_suggestion': {
                        'code_patches': code_patches,
                        'config_patches': config_patches,
                        'fix_strategy': result.get('reasoning', 'Apply fix from similar Git commit'),
                        'confidence_level': result.get('confidence', 'N/A'),
                        'commit_info': {
                            'hash': commit_hash,
                            'subject': commit.get('subject', 'N/A'),
                            'author': commit.get('author_name', 'N/A'),
                            'date': commit.get('date_iso', 'N/A'),
                            'similarity_score': commit.get('boosted_score', commit.get('similarity', 0))
                        }
                    }
                }
            }
            
            # Store in current_bug_analysis (same format as RCA analysis)
            self.current_bug_analysis = {
                'error_message': error_message,
                'log_file': log_file,
                'log_path': log_file_path,
                'code_dir': code_dir,
                'results': results,
                'timestamp': datetime.now().isoformat(),
                'source': 'existing_fix'  # Mark this as from existing fix
            }
            
            # Enable the buttons
            if hasattr(self, 'view_artifacts_bug_btn'):
                self.view_artifacts_bug_btn.setEnabled(True)
            if hasattr(self, 'goto_code_evaluation_btn'):
                self.goto_code_evaluation_btn.setEnabled(True)
            
            # Save to history so it can be accessed from Code Assistant
            self.save_bug_analysis_to_history()
            
            print(f"✅ Existing fix prepared for Code Assistant: {len(code_patches)} code patches, {len(config_patches)} config patches")
            
        except Exception as e:
            print(f"Error preparing existing fix for Code Assistant: {e}")
            import traceback
            traceback.print_exc()

    def _get_git_diff_for_commit(self, commit_hash, code_dir):
        """Get the full git diff output for a commit."""
        try:
            import subprocess
            
            if not code_dir:
                openair_codebase_file_name = "openairinterface5g-develop"
            else:
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            
            # Find the git repository
            possible_paths = [
                os.path.join(os.getcwd(), openair_codebase_file_name),
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', openair_codebase_file_name),
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', 'openairinterface5g-develop'),
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', 'openairinterface5g-test'),
            ]
            
            git_dir = None
            for path in possible_paths:
                if os.path.exists(path) and os.path.exists(os.path.join(path, '.git')):
                    try:
                        check_result = subprocess.run(
                            ['git', 'cat-file', '-e', commit_hash],
                            cwd=path,
                            capture_output=True,
                            timeout=2
                        )
                        if check_result.returncode == 0:
                            git_dir = path
                            break
                    except:
                        continue
            
            if not git_dir:
                print(f"Could not find Git repository with commit {commit_hash}")
                return None
            
            # Get the full git show output
            git_result = subprocess.run(
                ['git', 'show', '--unified=10', commit_hash],
                cwd=git_dir,
                capture_output=True,
                text=True,
                timeout=10
            )
            
            if git_result.returncode == 0:
                return git_result.stdout
            else:
                print(f"Git show failed: {git_result.stderr}")
                return None
                
        except Exception as e:
            print(f"Error getting git diff: {e}")
            return None

    def _extract_code_from_git_diff(self, git_diff_data, file_path, function_name):
        """Extract original_code, patched_code, and full file path from git diff output."""
        import re
        
        if not git_diff_data:
            return "", "", None
        
        try:
            # Normalize the file path for matching
            file_basename = os.path.basename(file_path)
            
            # Find the diff section for this file
            file_pattern = rf'diff --git a/[^\s]+ b/([^\s]+)'
            diff_sections = re.split(file_pattern, git_diff_data)
            
            target_diff = None
            full_file_path = None
            for i in range(1, len(diff_sections), 2):
                if i + 1 < len(diff_sections):
                    current_file = diff_sections[i]
                    if file_basename in current_file:
                        target_diff = diff_sections[i + 1]
                        full_file_path = current_file  # This is the path from git (e.g., openair2/RRC/NR/rrc_gNB.c)
                        print(f"DEBUG: Found file in git diff: {full_file_path}")
                        break
            
            if not target_diff:
                print(f"Could not find diff for file: {file_basename}")
                return "", "", None
            
            # Always use the comprehensive extraction method for better results
            original_code, patched_code = self._extract_code_with_more_context(target_diff)
            
            print(f"DEBUG: Extracted code - Original: {len(original_code)} chars, Patched: {len(patched_code)} chars")
            return original_code, patched_code, full_file_path
            
        except Exception as e:
            print(f"Error extracting code from git diff: {e}")
            import traceback
            traceback.print_exc()
            return "", "", None

    def _extract_code_with_more_context(self, diff_section):
        """Extract code changes with more surrounding context.
        
        This extracts code properly for replacement patches by:
        1. Including sufficient context before and after changes
        2. Building separate original and patched versions
        3. Ensuring the original code can be uniquely found in the file
        """
        import re
        
        lines = diff_section.split('\n')
        
        # Separate the lines into context, removed, and added
        original_lines = []  # Lines that exist in original (context + removed)
        patched_lines = []   # Lines that exist in patched (context + added)
        
        # Track if we've seen any actual changes
        has_changes = False
        
        # Track unique identifiers in the context
        context_before = []  # Context lines before the change
        context_after = []   # Context lines after the change
        change_started = False
        change_ended = False
        
        for line in lines:
            if line.startswith('@@'):
                continue
            elif line.startswith('---') or line.startswith('+++'):
                continue
            elif line.startswith('-') and not line.startswith('---'):
                # Line removed in patch - only in original
                original_lines.append(line[1:])
                has_changes = True
                change_started = True
            elif line.startswith('+') and not line.startswith('+++'):
                # Line added in patch - only in patched
                patched_lines.append(line[1:])
                has_changes = True
                change_started = True
            elif line.startswith(' '):
                # Context line - appears in both
                context_line = line[1:]
                original_lines.append(context_line)
                patched_lines.append(context_line)
                
                # Track context for better matching
                if not change_started:
                    context_before.append(context_line)
                elif change_started and has_changes:
                    context_after.append(context_line)
        
        # Build the result strings
        original_code = '\n'.join(original_lines)
        patched_code = '\n'.join(patched_lines)
        
        # For better matching, ensure we have enough unique context
        # Look for unique identifiers like function names, variable declarations
        unique_markers = []
        for line in context_before[-5:]:  # Last 5 lines before change
            if any(marker in line for marker in ['uint64_t', 'LOG_', 'AssertFatal', 'UE->', '==']):
                unique_markers.append(line.strip())
        
        print(f"DEBUG [_extract_code_with_more_context]:")
        print(f"  Has changes: {has_changes}")
        print(f"  Original lines: {len(original_lines)}")
        print(f"  Patched lines: {len(patched_lines)}")
        print(f"  Context before: {len(context_before)}")
        print(f"  Context after: {len(context_after)}")
        print(f"  Unique markers found: {len(unique_markers)}")
        print(f"  Original code length: {len(original_code)}")
        print(f"  Patched code length: {len(patched_code)}")
        if unique_markers:
            print(f"  Sample marker: {unique_markers[0][:50]}...")
        
        return original_code, patched_code

    def _extract_config_from_git_diff(self, git_diff_data, param_name):
        """Extract old and new config values from git diff output."""
        import re
        
        if not git_diff_data:
            return "N/A", "N/A"
        
        try:
            # Look for lines containing the parameter name
            lines = git_diff_data.split('\n')
            old_value = None
            new_value = None
            
            # Need to match pairs of - and + lines for the same parameter
            i = 0
            while i < len(lines):
                line = lines[i]
                if param_name in line:
                    if line.startswith('-') and not line.startswith('---'):
                        # Old value (removed line)
                        old_line = line[1:].strip()
                        old_value = self._extract_value_from_config_line(old_line)
                        
                        # Check if the next non-empty line is the corresponding + line
                        j = i + 1
                        while j < len(lines):
                            next_line = lines[j]
                            if next_line.strip() == '':
                                j += 1
                                continue
                            if param_name in next_line and next_line.startswith('+') and not next_line.startswith('+++'):
                                # Found the corresponding new value
                                new_line = next_line[1:].strip()
                                new_value = self._extract_value_from_config_line(new_line)
                                i = j  # Skip to the + line
                                break
                            else:
                                # Not a matching + line, stop looking
                                break
                            j += 1
                    elif line.startswith('+') and not line.startswith('+++'):
                        # Only process + line if we haven't already matched it with a - line
                        if new_value is None:
                            new_line = line[1:].strip()
                            new_value = self._extract_value_from_config_line(new_line)
                i += 1
            
            return old_value or "N/A", new_value or "N/A"
            
        except Exception as e:
            print(f"Error extracting config from git diff: {e}")
            return "N/A", "N/A"

    def _extract_value_from_config_line(self, line):
        """Extract the value from a configuration line."""
        import re
        
        # Handle different config formats
        # Format 1: param = "value";
        match = re.search(r'=\s*"([^"]+)"', line)
        if match:
            return match.group(1)
        
        # Format 2: param = value;
        match = re.search(r'=\s*([^;]+);', line)
        if match:
            return match.group(1).strip()
        
        # Format 3: ipv4 = "value"
        match = re.search(r'ipv4\s*=\s*"([^"]+)"', line)
        if match:
            return match.group(1)
        
        # Return the whole line if we can't parse it
        return line.strip()

    def handle_rca_test_completion(self):
        """Handle RCA test completion and update analysis options."""
        # First run the RCA test
        self.handle_start_rca_test()
        
        # Note: The update_analysis_options_after_test() will be called 
        # automatically after the test completes in handle_start_rca_test()

    def update_analysis_options_after_test(self):
        """Update analysis options after RCA test completion."""
        print("RCA test completed. Analysis options are now available.")
        
        # The analysis options are now fixed and already created in show_bug_discovery
        # This function is kept for compatibility but doesn't need to do anything
        # since the options are already created as checkboxes in the UI
        
        # Show a message that RCA test is complete
        QMessageBox.information(
            self,
            "RCA Test Complete",
            "RCA test has been completed. You can now select analysis options and start the analysis."
        )

    def start_bug_analysis(self):
        """Run RCA test script and then extract data from JSON file based on selected options."""
        if not self.validate_bug_analysis_inputs():
            return
            
        self.analysis_progress.show()
        
        # Clear UI and show initial message
        self.bug_results_text.clear()
        self.bug_results_text.setText("Starting analysis...\n")
        QApplication.processEvents()
        
        # Get log file path
        log_dir = self.log_dir_edit.text()
        log_file = self.log_file_combo.currentText()
        log_file_path = os.path.join(log_dir, log_file)
        
        # Create and start the RCA script worker
        self.rca_worker = RCAScriptWorker(log_file_path, self)
        self.rca_worker.finished.connect(self.on_rca_script_finished)
        self.rca_worker.error.connect(self.on_rca_script_error)
        self.rca_worker.start()
        
        # Show progress message
        self.bug_results_text.clear()
        self.bug_results_text.setText("Running RCA analysis...\n")
        QApplication.processEvents()
    




    def validate_bug_analysis_inputs(self):
        """Validate inputs for bug analysis."""
        if not self.log_dir_edit.text():
            QMessageBox.warning(
                self,
                "Invalid Input",
                "Please select a log directory."
            )
            return False
            
        if not os.path.exists(self.log_dir_edit.text()):
            QMessageBox.warning(
                self,
                "Invalid Directory",
                "The selected log directory does not exist."
            )
            return False
            
        # Check if any analysis options are selected
        has_analysis_options = False
        if hasattr(self, 'analysis_checkboxes'):
            has_analysis_options = any(checkbox.isChecked() for checkbox in self.analysis_checkboxes.values())
        
        if not has_analysis_options:
            QMessageBox.warning(
                self,
                "No Options Selected",
                "Please select at least one analysis option."
            )
            return False
            
        return True

    def extract_rca_data_and_display(self):
        """Display the raw content of the generated RCA JSON file."""
        try:
            # Get the log file name
            log_file = self.log_file_combo.currentText()
            if not log_file:
                self.bug_results_text.setText("No log file selected.")
                return
            
            # Construct the RCA JSON file path based on the selected log file
            log_name = os.path.splitext(log_file)[0]
            possible_paths = [
                os.path.join("RCA-Code-Updated/RCA-Code-Updated", f"{log_name}.rca.json"),
                os.path.join("RCA-Code-Updated/RCA-Code-Updated/test_logs", f"{log_name}.rca.json"),
                os.path.join("RCA-Code-Updated/RCA-Code-Updated", f"{log_name}.log.rca.json"),
                os.path.join("RCA-Code-Updated/RCA-Code-Updated/test_logs", f"{log_name}.log.rca.json"),
                f"{log_name}.rca.json",
                f"{log_name}.log.rca.json"
            ]
            
            # Also check the absolute paths
            rca_code_dir = r"C:\Users\ChanduVangala\Documents\AgenticRAN-V8_azure_key_with_RCA_v2_working\RCA-Code-Updated\RCA-Code-Updated"
            absolute_paths = [
                os.path.join(rca_code_dir, f"{log_name}.rca.json"),
                os.path.join(rca_code_dir, "test_logs", f"{log_name}.rca.json"),
                os.path.join(rca_code_dir, f"{log_name}.log.rca.json"),
                os.path.join(rca_code_dir, "test_logs", f"{log_name}.log.rca.json"),
            ]
            
            all_paths = possible_paths + absolute_paths
            
            json_file_path = None
            for path in all_paths:
                if os.path.exists(path):
                    json_file_path = path
                    break
            
            if not json_file_path:
                self.bug_results_text.setText(f"RCA analysis file not found for {log_name}.\nTried paths:\n" + 
                                            "\n".join(all_paths) + 
                                            "\n\nPlease ensure the RCA test script has completed successfully.")
                return
            
            # Print which JSON file is being used
            print(f"Found RCA JSON file: {json_file_path}")
            
            # Read and display the raw JSON content
            with open(json_file_path, 'r') as f:
                json_content = f.read()
            
            # Display the raw JSON content
            self.bug_results_text.setText(json_content)
            
        except Exception as e:
            error_msg = f"Error reading JSON file: {str(e)}"
            self.bug_results_text.setText(error_msg)
            print(f"Error in extract_rca_data_and_display: {e}")

    def show_analysis_results(self):
        """Show analysis results based on selected options in a table format."""
        try:
            # Get the log file name
            log_file = self.log_file_combo.currentText()
            if not log_file:
                self.bug_results_text.setText("No log file selected.")
                return
            
            # Construct the RCA JSON file path based on the selected log file
            log_name = os.path.splitext(log_file)[0]
            possible_paths = [
                os.path.join("RCA-Code-Updated/RCA-Code-Updated", f"{log_name}.rca.json"),
                os.path.join("RCA-Code-Updated/RCA-Code-Updated/test_logs", f"{log_name}.rca.json"),
                os.path.join("RCA-Code-Updated/RCA-Code-Updated", f"{log_name}.log.rca.json"),
                os.path.join("RCA-Code-Updated/RCA-Code-Updated/test_logs", f"{log_name}.log.rca.json"),
                f"{log_name}.rca.json",
                f"{log_name}.log.rca.json"
            ]
            
            # Also check the absolute paths
            rca_code_dir = r"C:\Users\ChanduVangala\Documents\AgenticRAN-V8_azure_key_with_RCA_v2_working\RCA-Code-Updated\RCA-Code-Updated"
            absolute_paths = [
                os.path.join(rca_code_dir, f"{log_name}.rca.json"),
                os.path.join(rca_code_dir, "test_logs", f"{log_name}.rca.json"),
                os.path.join(rca_code_dir, f"{log_name}.log.rca.json"),
                os.path.join(rca_code_dir, "test_logs", f"{log_name}.log.rca.json"),
            ]
            
            all_paths = possible_paths + absolute_paths
            
            json_file_path = None
            for path in all_paths:
                if os.path.exists(path):
                    json_file_path = path
                    break
            
            if not json_file_path:
                self.bug_results_text.setText(f"RCA analysis file not found for {log_name}.\nTried paths:\n" + 
                                            "\n".join(all_paths) + 
                                            "\n\nPlease ensure the RCA test script has completed successfully.")
                return
            
            # Print which JSON file is being used
            print(f"Found RCA JSON file: {json_file_path}")
            
            # Load the JSON data
            with open(json_file_path, 'r') as f:
                rca_data = json.load(f)
            
            # Get selected analysis options
            analysis_options = {}
            for key, checkbox in self.analysis_checkboxes.items():
                analysis_options[key] = checkbox.isChecked()
            
            # Create a new dialog for table display
            table_dialog = QDialog(self)
            table_dialog.setWindowTitle(f"Analysis Results - {log_name}")
            table_dialog.setMinimumSize(1200, 800)

            layout = QVBoxLayout()
            
            # Create tab widget for different analysis types
            tab_widget = QTabWidget()
            
            # Error Analysis Tab
            if analysis_options.get('error_analysis', False) and 'error_analysis' in rca_data:
                error_tab = self.create_error_analysis_tab(rca_data['error_analysis'])
                tab_widget.addTab(error_tab, "🔍 Error Analysis")
            
            # Function Analysis Tab
            if analysis_options.get('function_analysis', False) and 'function_analysis' in rca_data:
                function_tab = self.create_function_analysis_tab(rca_data['function_analysis'])
                tab_widget.addTab(function_tab, "⚙️ Function Analysis")
            
            # Impact Analysis Tab
            if analysis_options.get('impact_analysis', False) and 'impact_analysis' in rca_data:
                impact_tab = self.create_impact_analysis_tab(rca_data['impact_analysis'])
                tab_widget.addTab(impact_tab, "💥 Impact Analysis")
            
            # KB Analysis Tab
            if analysis_options.get('kb_analysis', False) and 'kb_analysis' in rca_data:
                kb_tab = self.create_kb_analysis_tab(rca_data['kb_analysis'])
                tab_widget.addTab(kb_tab, "📚 KB Analysis")
            
            # AI Analysis Tab
            if analysis_options.get('ai_analysis', False) and 'ai_analysis' in rca_data:
                ai_tab = self.create_ai_analysis_tab(rca_data['ai_analysis'])
                tab_widget.addTab(ai_tab, "🤖 AI Analysis")
            
            # Meta Information Tab - Removed as requested by user
            
            # If no options selected, show message
            if tab_widget.count() == 0:
                no_data_label = QLabel("❌ Please select at least one analysis option to view results.")
                no_data_label.setAlignment(Qt.AlignCenter)
                no_data_label.setStyleSheet("font-size: 16px; color: #666; padding: 20px;")
                layout.addWidget(no_data_label)
            else:
                layout.addWidget(tab_widget)
            
            # Add export button
            export_btn = QPushButton("Export Results")
            export_btn.clicked.connect(lambda: self.export_analysis_results(rca_data, analysis_options))
            layout.addWidget(export_btn)
            
            table_dialog.setLayout(layout)
            table_dialog.exec_()
            
        except Exception as e:
            error_msg = f"❌ Error showing analysis results: {str(e)}"
            self.bug_results_text.setText(error_msg)
            print(f"Error in show_analysis_results: {e}")
    
    def create_error_analysis_tab(self, error_data):
        """Create tab for error analysis with tables."""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Summary section
        summary_group = QGroupBox("📊 Error Summary")
        summary_layout = QVBoxLayout()
        
        # Create summary table
        summary_table = QTableWidget()
        summary_table.setColumnCount(5)
        summary_table.setHorizontalHeaderLabels(["Error Type", "Count", "Severity", "Layer", "Total"])
        
        # Populate summary table
        if 'identified_errors' in error_data:
            row = 0
            total_errors = 0
            for error_type, errors in error_data['identified_errors'].items():
                if errors:
                    summary_table.insertRow(row)
                    summary_table.setItem(row, 0, QTableWidgetItem(error_type.replace('_', ' ').title()))
                    summary_table.setItem(row, 1, QTableWidgetItem(str(len(errors))))
                    summary_table.setItem(row, 2, QTableWidgetItem(errors[0].get('severity', 'N/A')))
                    summary_table.setItem(row, 3, QTableWidgetItem(errors[0].get('layer', 'N/A')))
                    summary_table.setItem(row, 4, QTableWidgetItem(str(total_errors + len(errors))))
                    total_errors += len(errors)
                    row += 1
        
            summary_table.resizeColumnsToContents()
            # Enable scrolling
            summary_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            summary_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            summary_layout.addWidget(summary_table)
            summary_group.setLayout(summary_layout)
            layout.addWidget(summary_group)
        
        # Detailed errors section
        if 'identified_errors' in error_data:
            details_group = QGroupBox("📋 Detailed Errors")
            details_layout = QVBoxLayout()
            
            # Create detailed errors table
            details_table = QTableWidget()
            details_table.setColumnCount(6)
            details_table.setHorizontalHeaderLabels(["Error Type", "Line", "Pattern", "Severity", "Layer", "Confidence"])
            
            row = 0
            for error_type, errors in error_data['identified_errors'].items():
                for error in errors:
                    details_table.insertRow(row)
                    details_table.setItem(row, 0, QTableWidgetItem(error_type.replace('_', ' ').title()))
                    details_table.setItem(row, 1, QTableWidgetItem(str(error.get('line_number', 'N/A'))))
                    details_table.setItem(row, 2, QTableWidgetItem(str(error.get('pattern', 'N/A'))))
                    details_table.setItem(row, 3, QTableWidgetItem(str(error.get('severity', 'N/A'))))
                    details_table.setItem(row, 4, QTableWidgetItem(str(error.get('layer', 'N/A'))))
                    details_table.setItem(row, 5, QTableWidgetItem(str(error.get('confidence', 'N/A'))))
                    row += 1
            
            details_table.resizeColumnsToContents()
            # Enable scrolling
            details_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            details_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            details_layout.addWidget(details_table)
            details_group.setLayout(details_layout)
            layout.addWidget(details_group)
        
        # Statistics section
        if 'error_counts' in error_data or 'severity_counts' in error_data or 'layer_counts' in error_data:
            stats_group = QGroupBox("📈 Statistics")
            stats_layout = QVBoxLayout()
            
            # Create unified statistics table
            stats_table = QTableWidget()
            stats_table.setColumnCount(3)
            stats_table.setHorizontalHeaderLabels(["Category", "Type", "Count"])
                
            row = 0
            
            # Add Error Type statistics
            if 'error_counts' in error_data:
                for error_type, count in error_data['error_counts'].items():
                    stats_table.insertRow(row)
                    stats_table.setItem(row, 0, QTableWidgetItem("Error Type"))
                    stats_table.setItem(row, 1, QTableWidgetItem(error_type.replace('_', ' ').title()))
                    stats_table.setItem(row, 2, QTableWidgetItem(str(count)))
                    row += 1
                
            # Add Severity statistics
            if 'severity_counts' in error_data:
                for severity, count in error_data['severity_counts'].items():
                    stats_table.insertRow(row)
                    stats_table.setItem(row, 0, QTableWidgetItem("Severity"))
                    stats_table.setItem(row, 1, QTableWidgetItem(severity.title()))
                    stats_table.setItem(row, 2, QTableWidgetItem(str(count)))
                    row += 1
                
            # Add Layer statistics
            if 'layer_counts' in error_data:
                for layer, count in error_data['layer_counts'].items():
                    stats_table.insertRow(row)
                    stats_table.setItem(row, 0, QTableWidgetItem("Layer"))
                    stats_table.setItem(row, 1, QTableWidgetItem(layer))
                    stats_table.setItem(row, 2, QTableWidgetItem(str(count)))
                    row += 1
                
            stats_table.resizeColumnsToContents()
                # Enable scrolling
            stats_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            stats_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            stats_layout.addWidget(stats_table)
            stats_group.setLayout(stats_layout)
            layout.addWidget(stats_group)
        
        layout.addStretch()
        tab.setLayout(layout)
        return tab
    
    def create_function_analysis_tab(self, func_data):
        """Create tab for function analysis with tables."""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Affected Functions section
        if 'affected_functions' in func_data:
            affected_group = QGroupBox("🎯 Affected Functions")
            affected_layout = QVBoxLayout()
            
            affected_table = QTableWidget()
            affected_table.setColumnCount(4)
            affected_table.setHorizontalHeaderLabels(["Function Name", "File", "Relevance", "Score"])
            
            row = 0
            for func in func_data['affected_functions']:
                affected_table.insertRow(row)
                affected_table.setItem(row, 0, QTableWidgetItem(str(func.get('name', 'N/A'))))
                affected_table.setItem(row, 1, QTableWidgetItem(str(func.get('file', 'N/A'))))
                affected_table.setItem(row, 2, QTableWidgetItem(str(func.get('relevance', 'N/A'))))
                affected_table.setItem(row, 3, QTableWidgetItem(str(func.get('score', 'N/A'))))
                row += 1
            
            affected_table.resizeColumnsToContents()
            # Enable scrolling
            affected_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            affected_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            affected_layout.addWidget(affected_table)
            affected_group.setLayout(affected_layout)
            layout.addWidget(affected_group)
        
        # Affected Configuration Files section
        if 'affected_configs' in func_data and func_data['affected_configs']:
            config_group = QGroupBox("⚙️ Affected Configuration Files")
            config_layout = QVBoxLayout()
            
            config_table = QTableWidget()
            config_table.setColumnCount(5)
            config_table.setHorizontalHeaderLabels([
                "Config Name", "File Path", "Line", "Recommended Value", "Relevance"
            ])
            
            row = 0
            for config in func_data['affected_configs']:
                config_table.insertRow(row)
                
                # Config Name (same font as affected functions table)
                config_table.setItem(row, 0, QTableWidgetItem(str(config.get('config_name', 'N/A'))))
                
                # File Path
                file_path = config.get('file_path', 'N/A')
                if '/' in file_path:
                    file_path = file_path.split('/')[-1]  # Show only filename
                config_table.setItem(row, 1, QTableWidgetItem(file_path))
                
                # Line Number
                config_table.setItem(row, 2, QTableWidgetItem(str(config.get('line_number', 'N/A'))))
                
                # Recommended Value
                rec_val_item = QTableWidgetItem(str(config.get('recommended_value', 'N/A')))
                rec_val_item.setForeground(QColor(50, 150, 50))  # Green for recommended value
                config_table.setItem(row, 3, rec_val_item)
                
                # Relevance
                relevance = str(config.get('relevance', 'N/A'))
                relevance_item = QTableWidgetItem(relevance)
                
                # Color code by relevance
                if 'critical' in relevance.lower():
                    relevance_item.setBackground(QColor(255, 200, 200))  # Light red
                    relevance_item.setForeground(QColor(150, 0, 0))  # Dark red text
                elif 'high' in relevance.lower():
                    relevance_item.setBackground(QColor(255, 220, 180))  # Light orange
                    relevance_item.setForeground(QColor(150, 80, 0))  # Dark orange text
                elif 'medium' in relevance.lower():
                    relevance_item.setBackground(QColor(255, 255, 200))  # Light yellow
                    relevance_item.setForeground(QColor(120, 120, 0))  # Dark yellow text
                
                config_table.setItem(row, 4, relevance_item)
                
                # Add tooltip with description
                description = config.get('description', 'No description available')
                for col in range(5):
                    if config_table.item(row, col):
                        config_table.item(row, col).setToolTip(description)
                
                row += 1
            
            config_table.resizeColumnsToContents()
            # Enable scrolling
            config_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            config_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            
            # Add a label with instructions
            info_label = QLabel("💡 Tip: Hover over cells to see detailed descriptions")
            info_label.setStyleSheet("color: #666; font-size: 10px; font-style: italic;")
            
            config_layout.addWidget(config_table)
            config_layout.addWidget(info_label)
            config_group.setLayout(config_layout)
            layout.addWidget(config_group)
        
        # Critical Functions section - Removed as requested by user
        
        layout.addStretch()
        tab.setLayout(layout)
        return tab
    
    def create_impact_analysis_tab(self, impact_data):
        """Create tab for impact analysis with tables - supports both RCA and Bug Discovery."""
        tab = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)  # Remove margins for full-height display
        layout.setSpacing(0)  # Remove spacing
        
        # Check if this is Bug Discovery (has log_file) or RCA (has error_impacts)
        is_bug_discovery = 'log_file' in impact_data and 'error_impacts' not in impact_data
        
        if is_bug_discovery:
            # Run the log analyzer for Bug Discovery
            try:
                log_file = impact_data.get('log_file', '')
                error_patterns_file = "Error_fixing_pipelin/database/error_patterns_enhanced.json"
                
                if not log_file or log_file == 'None' or not os.path.exists(log_file):
                    no_data_label = QLabel("⚠️ No log file available for impact analysis.\n\nPlease ensure a valid log file is provided for Bug Discovery.")
                    no_data_label.setAlignment(Qt.AlignCenter)
                    no_data_label.setStyleSheet("color: #888; font-style: italic; padding: 40px; font-size: 14px;")
                    layout.addWidget(no_data_label)
                    tab.setLayout(layout)
                    return tab
                
                if not os.path.exists(error_patterns_file):
                    no_data_label = QLabel("⚠️ Error patterns file not found.\n\nPlease ensure error_patterns_enhanced.json exists.")
                    no_data_label.setAlignment(Qt.AlignCenter)
                    no_data_label.setStyleSheet("color: #888; font-style: italic; padding: 40px; font-size: 14px;")
                    layout.addWidget(no_data_label)
                    tab.setLayout(layout)
                    return tab
                
                # Initialize analyzer and run analysis
                analyzer = LogLayerSeverityAnalyzer(error_patterns_file)
                results = analyzer.analyze_log_file(log_file)
                
                if not results or not results.get("identified_errors"):
                    no_data_label = QLabel("✓ No errors found in log file.\n\nThe log analysis completed successfully with no issues detected.")
                    no_data_label.setAlignment(Qt.AlignCenter)
                    no_data_label.setStyleSheet("color: #2d862d; font-style: italic; padding: 40px; font-size: 14px;")
                    layout.addWidget(no_data_label)
                    tab.setLayout(layout)
                    return tab
                
                # Create Bug Discovery Impact Analysis UI
                self._create_bug_discovery_impact_ui(layout, results, analyzer)
                
            except Exception as e:
                error_label = QLabel(f"❌ Error running impact analysis:\n\n{str(e)}")
                error_label.setAlignment(Qt.AlignCenter)
                error_label.setStyleSheet("color: #d32f2f; padding: 40px; font-size: 13px;")
                layout.addWidget(error_label)
                print(f"Impact analysis error: {e}")
                import traceback
                traceback.print_exc()
        
        else:
            # Original RCA Impact Analysis
            # Error Impacts section
            if 'error_impacts' in impact_data:
                impacts_group = QGroupBox("📈 Error Impacts")
                impacts_layout = QVBoxLayout()
                
                impacts_table = QTableWidget()
                impacts_table.setColumnCount(4)
                impacts_table.setHorizontalHeaderLabels(["Error Type", "Primary Layer", "Severity", "Description"])
                
                row = 0
                for error_type, impact in impact_data['error_impacts'].items():
                    impacts_table.insertRow(row)
                    impacts_table.setItem(row, 0, QTableWidgetItem(error_type.replace('_', ' ').title()))
                    impacts_table.setItem(row, 1, QTableWidgetItem(str(impact.get('primary_layer', 'N/A'))))
                    impacts_table.setItem(row, 2, QTableWidgetItem(str(impact.get('severity', 'N/A'))))
                    impacts_table.setItem(row, 3, QTableWidgetItem(str(impact.get('impact_description', 'N/A'))))
                    row += 1
                
                impacts_table.resizeColumnsToContents()
                impacts_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
                impacts_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
                impacts_layout.addWidget(impacts_table)
                impacts_group.setLayout(impacts_layout)
                layout.addWidget(impacts_group)
            
            # Cascade Effects section
            if 'cascade_effects' in impact_data:
                cascade_group = QGroupBox("🌊 Cascade Effects")
                cascade_layout = QVBoxLayout()
                
                cascade_table = QTableWidget()
                cascade_table.setColumnCount(3)
                cascade_table.setHorizontalHeaderLabels(["Error Type", "Layer", "Impact"])
                
                row = 0
                for error_type, cascade in impact_data['cascade_effects'].items():
                    for stage in cascade.get('cascade', []):
                        cascade_table.insertRow(row)
                        cascade_table.setItem(row, 0, QTableWidgetItem(error_type.replace('_', ' ').title()))
                        cascade_table.setItem(row, 1, QTableWidgetItem(str(stage.get('layer', 'N/A'))))
                        cascade_table.setItem(row, 2, QTableWidgetItem(str(stage.get('impact', 'N/A'))))
                        row += 1
                
                cascade_table.resizeColumnsToContents()
                cascade_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
                cascade_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
                cascade_layout.addWidget(cascade_table)
                cascade_group.setLayout(cascade_layout)
                layout.addWidget(cascade_group)
        
        # Don't add stretch - let content fill naturally
        tab.setLayout(layout)
        return tab
    
    def _create_bug_discovery_impact_ui(self, layout, results, analyzer):
        """Create organized UI for Bug Discovery impact analysis results"""
        
        # Create scroll area for content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setMinimumHeight(600)  # Ensure minimum height
        scroll.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)  # Expand to fill available space
        scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background-color: #f8f9fa;
            }
        """)
        
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        content_layout.setSpacing(15)
        content_layout.setContentsMargins(10, 10, 10, 10)
        
        total_errors = sum(results["error_counts"].values())
        
        # ===== SUMMARY HEADER =====
        summary_header = QGroupBox("📊 Analysis Summary")
        summary_header.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: #1e3a5f;
                border: 1px solid #b3d9ff;
                border-radius: 4px;
                margin-top: 10px;
                padding: 15px;
                background-color: #ffffff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 4px 8px;
                color: #1e3a5f;
            }
        """)
        summary_layout = QVBoxLayout()
        
        summary_info = QLabel(f"""
            <div style='font-size: 13px; line-height: 1.8;'>
                <b style='color: #d32f2f; font-size: 15px;'>Total Errors Found:</b> <span style='font-size: 16px; color: #d32f2f;'>{total_errors}</span><br>
                <b style='color: #1e3a5f;'>Unique Error Types:</b> <span style='color: #1e3a5f;'>{len(results['identified_errors'])}</span><br>
                <b style='color: #1e3a5f;'>Layers Affected:</b> <span style='color: #1e3a5f;'>{len(results['layer_counts'])}</span>
            </div>
        """)
        summary_layout.addWidget(summary_info)
        summary_header.setLayout(summary_layout)
        content_layout.addWidget(summary_header)
        
        # ===== DETAILED ERROR INFORMATION =====
        details_group = QGroupBox("📋 Detailed Error Information")
        details_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: #1e3a5f;
                border: 1px solid #b3d9ff;
                border-radius: 4px;
                margin-top: 10px;
                padding: 10px;
                background-color: #ffffff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                padding: 4px 8px;
                color: #1e3a5f;
            }
        """)
        details_layout = QVBoxLayout()
        
        # Define severity colors for error cards - professional blue and red only
        severity_colors = {
            "critical": "#d32f2f",  # Red for critical
            "high": "#d32f2f",      # Red for high
            "medium": "#0066cc",    # Blue for medium
            "low": "#0066cc"        # Blue for low
        }
        
        # Sort errors by severity and count
        severity_priority = {"critical": 0, "high": 1, "medium": 2, "low": 3}
        sorted_errors = sorted(
            results["identified_errors"].items(),
            key=lambda x: (
                severity_priority.get(x[1][0]["severity"].lower(), 4),
                -len(x[1])
            )
        )
        
        for error_type, occurrences in sorted_errors:
            if not occurrences:
                continue
            
            first_occurrence = occurrences[0]
            primary_layer = first_occurrence.get("layer", "unknown")
            severity = first_occurrence.get("severity", "medium")
            impact = first_occurrence.get("impact", "Unknown impact")
            
            # Calculate cascading layers with detailed impacts
            cascade_chain = analyzer.calculate_cascading_layers(primary_layer, impact)
            impact_scores = analyzer.calculate_impact_score(error_type, severity)
            max_impact = max(impact_scores.values())
            impact_level = "High" if max_impact >= 0.7 else "Medium" if max_impact >= 0.4 else "Low"
            
            # Severity color coding
            severity_color = severity_colors.get(severity.lower(), "#666")
            
            # Create error card - use emoji indicator based on severity
            severity_emoji = "🔴" if severity.lower() in ["critical", "high"] else "🔵"
            error_card = QGroupBox(f"{severity_emoji} {error_type.replace('_', ' ').title()}")
            error_card.setStyleSheet(f"""
                QGroupBox {{
                    font-weight: bold;
                    font-size: 13px;
                    color: #1e3a5f;
                    border: 1px solid #b3d9ff;
                    border-left: 4px solid {severity_color};
                    border-radius: 4px;
                    margin-top: 8px;
                    padding: 10px;
                    background-color: white;
                }}
                QGroupBox::title {{
                    subcontrol-origin: margin;
                    padding: 4px 8px;
                    color: #1e3a5f;
                }}
            """)
            error_card_layout = QVBoxLayout()
            
            # Error details
            error_info = QLabel(f"""
                <div style='font-size: 12px; line-height: 1.8;'>
                    <table style='width: 100%; border-collapse: collapse;'>
                        <tr>
                            <td style='padding: 5px; width: 150px;'><b>Occurrences:</b></td>
                            <td style='padding: 5px; color: #d32f2f; font-weight: bold;'>{len(occurrences)}</td>
                        </tr>
                        <tr style='background-color: #f5f5f5;'>
                            <td style='padding: 5px;'><b>Primary Layer:</b></td>
                            <td style='padding: 5px; color: #0066cc; font-weight: bold;'>{primary_layer}</td>
                        </tr>
                        <tr>
                            <td style='padding: 5px;'><b>Severity:</b></td>
                            <td style='padding: 5px; color: {severity_color}; font-weight: bold; text-transform: uppercase;'>{severity}</td>
                        </tr>
                        <tr style='background-color: #f5f5f5;'>
                            <td style='padding: 5px;'><b>Impact Level:</b></td>
                            <td style='padding: 5px; font-weight: bold;'>{impact_level}</td>
                        </tr>
                        <tr>
                            <td style='padding: 5px;'><b>Description:</b></td>
                            <td style='padding: 5px; color: #424242;'>{impact}</td>
                        </tr>
                    </table>
                </div>
            """)
            error_card_layout.addWidget(error_info)
            
            # Cascade Effects with detailed impact descriptions
            cascade_html = """
                <div style='margin-top: 10px; padding: 10px; background-color: #f5f7fa; border: 1px solid #e1e5ea; border-radius: 4px;'>
                    <b style='color: #1e3a5f; font-size: 13px;'>🌊 Cascade Effects:</b><br>
                    <div style='margin-top: 8px;'>
            """
            
            stage_colors = {
                1: "#0066cc",  # Primary - professional blue
                2: "#4d88cc",  # First level - lighter blue
                3: "#7fa7d9"   # Second level - subtle blue
            }
            
            stage_labels = {
                1: "PRIMARY",
                2: "1ST LEVEL",
                3: "2ND LEVEL"
            }
            
            for cascade_item in cascade_chain:
                layer = cascade_item["layer"]
                cascade_impact = cascade_item["impact"]
                stage = cascade_item["stage"]
                stage_color = stage_colors.get(stage, "#666")
                stage_label = stage_labels.get(stage, f"STAGE {stage}")
                
                # Escape any HTML characters in impact text
                cascade_impact_escaped = cascade_impact.replace('<', '&lt;').replace('>', '&gt;')
                
                cascade_html += f"""
                    <div style='margin: 6px 0; padding: 8px; background-color: white; border-left: 4px solid {stage_color}; border-radius: 3px;'>
                        <div style='margin-bottom: 3px;'>
                            <span style='background-color: {stage_color}; color: white; padding: 2px 6px; border-radius: 3px; font-size: 9px; font-weight: bold;'>{stage_label}</span>
                            <b style='color: {stage_color}; margin-left: 8px; font-size: 12px;'>{layer}</b>
                        </div>
                        <div style='color: #424242; font-size: 11px; margin-left: 8px;'>
                            {cascade_impact_escaped}
                        </div>
                    </div>
                """
            
            cascade_html += "</div></div>"
            cascade_label = QLabel(cascade_html)
            cascade_label.setWordWrap(True)
            error_card_layout.addWidget(cascade_label)
            
            # Example occurrence
            example_label = QLabel(f"""
                <div style='margin-top: 10px; padding: 8px; background-color: #f5f7fa; border: 1px solid #e1e5ea; border-radius: 4px;'>
                    <b style='color: #1e3a5f;'>📝 Example Occurrence:</b><br>
                    <span style='color: #666; font-size: 11px;'>Line {first_occurrence.get('line_number', 'N/A')}:</span><br>
                    <code style='font-size: 10px; color: #1e3a5f; background-color: white; padding: 4px; display: block; margin-top: 5px; border: 1px solid #e1e5ea; border-radius: 3px; overflow: hidden;'>{first_occurrence.get('line', 'N/A')[:150]}</code>
                </div>
            """)
            error_card_layout.addWidget(example_label)
            
            error_card.setLayout(error_card_layout)
            details_layout.addWidget(error_card)
        
        details_group.setLayout(details_layout)
        content_layout.addWidget(details_group)
        
        content_layout.addStretch()
        scroll.setWidget(content_widget)
        
        # Make sure the scroll area expands to fill the entire tab
        layout.addWidget(scroll, 1)  # Stretch factor of 1 to expand
        layout.setContentsMargins(0, 0, 0, 0)  # Remove margins to use full space
    
    def create_kb_analysis_tab(self, kb_data):
        """Create tab for KB analysis with tables."""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Helper function to clean up numbering
        def clean_numbering(text):
            """Remove duplicate numbering like 1.1, 2.2 and replace with sequential numbers."""
            import re
            # Remove patterns like "1.1", "2.2", "3.3" etc.
            cleaned = re.sub(r'\d+\.\d+\.?\s*', '', str(text))
            return cleaned.strip()
        
        # Fix Recommendations section - Removed as requested by user
        
        # Comprehensive RCA section
        if 'comprehensive_rca' in kb_data:
            comp_group = QGroupBox("📋 Comprehensive RCA")
            comp_layout = QVBoxLayout()
            
            comp_rca = kb_data['comprehensive_rca']
            
            # Summary text with scrolling
            summary_text = QTextEdit()
            summary_text.setPlainText(f"Summary: {comp_rca.get('summary', 'N/A')}")
            summary_text.setMaximumHeight(100)
            summary_text.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            summary_text.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            comp_layout.addWidget(summary_text)
            
            # RCA Fix Recommendations
            if 'fix_recommendations' in comp_rca:
                rca_fixes_table = QTableWidget()
                rca_fixes_table.setColumnCount(1)
                rca_fixes_table.setHorizontalHeaderLabels(["RCA Recommendation"])
                
                row = 0
                for i, rec in enumerate(comp_rca['fix_recommendations'], 1):
                    rca_fixes_table.insertRow(row)
                    cleaned_rec = clean_numbering(rec)
                    rca_fixes_table.setItem(row, 0, QTableWidgetItem(cleaned_rec))
                    row += 1
                
                rca_fixes_table.resizeColumnsToContents()
                # Enable scrolling
                rca_fixes_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
                rca_fixes_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
                comp_layout.addWidget(rca_fixes_table)
            
            comp_group.setLayout(comp_layout)
            layout.addWidget(comp_group)
        
        # Protocol References section
        if 'protocol_references' in kb_data:
            refs_group = QGroupBox("📖 Protocol References")
            refs_layout = QVBoxLayout()
            
            refs_table = QTableWidget()
            refs_table.setColumnCount(3)
            refs_table.setHorizontalHeaderLabels(["Protocol", "Title", "Spec ID"])
            
            row = 0
            for protocol, ref in kb_data['protocol_references'].items():
                refs_table.insertRow(row)
                refs_table.setItem(row, 0, QTableWidgetItem(protocol))
                refs_table.setItem(row, 1, QTableWidgetItem(str(ref.get('title', 'N/A'))))
                refs_table.setItem(row, 2, QTableWidgetItem(str(ref.get('spec_id', 'N/A'))))
                row += 1
            
            refs_table.resizeColumnsToContents()
            # Enable scrolling
            refs_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            refs_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            refs_layout.addWidget(refs_table)
            refs_group.setLayout(refs_layout)
            layout.addWidget(refs_group)
        
        layout.addStretch()
        tab.setLayout(layout)
        return tab
    
    def create_ai_analysis_tab(self, ai_data):
        """Create tab for AI analysis with tables."""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Helper function to clean up numbering
        def clean_numbering(text):
            """Remove duplicate numbering like 1.1, 2.2 and replace with sequential numbers."""
            import re
            # Remove patterns like "1.1", "2.2", "3.3" etc.
            cleaned = re.sub(r'\d+\.\d+\.?\s*', '', str(text))
            return cleaned.strip()
        
        # Root Cause section
        if 'root_cause' in ai_data:
            root_group = QGroupBox("🎯 Root Cause")
            root_layout = QVBoxLayout()
            
            root_text = QTextEdit()
            root_text.setPlainText(ai_data['root_cause'])
            root_text.setMaximumHeight(150)
            root_text.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            root_text.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            root_layout.addWidget(root_text)
            root_group.setLayout(root_layout)
            layout.addWidget(root_group)
        
        # Suggested Fixes section
        if 'suggested_fixes' in ai_data:
            fixes_group = QGroupBox("🔧 Suggested Fixes")
            fixes_layout = QVBoxLayout()
            
            # Create a more detailed table with verification buttons
            fixes_table = QTableWidget()
            fixes_table.setColumnCount(3)
            fixes_table.setHorizontalHeaderLabels(["Fix", "Status", "Verify"])
            
            row = 0
            for i, fix in enumerate(ai_data['suggested_fixes'], 1):
                fixes_table.insertRow(row)
                cleaned_fix = clean_numbering(fix)
                fixes_table.setItem(row, 0, QTableWidgetItem(cleaned_fix))
                
                # Status column
                status_item = QTableWidgetItem("⏳ Pending")
                status_item.setTextAlignment(Qt.AlignCenter)
                fixes_table.setItem(row, 1, status_item)
                
                # Verify button column
                verify_button = QPushButton("🔍 Verify")
                verify_button.setStyleSheet("""
                    QPushButton {
                        background-color: #0078D4;
                        color: white;
                        border: none;
                        padding: 5px 10px;
                        border-radius: 3px;
                        font-size: 11px;
                    }
                    QPushButton:hover {
                        background-color: #106EBE;
                    }
                    QPushButton:pressed {
                        background-color: #005A9E;
                    }
                """)
                
                # Store the fix text and row index for verification
                verify_button.clicked.connect(partial(self.verify_configuration_fix, cleaned_fix, row, fixes_table))
                
                fixes_table.setCellWidget(row, 2, verify_button)
                row += 1
            
            fixes_table.resizeColumnsToContents()
            # Enable scrolling
            fixes_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            fixes_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            
            # Set column widths
            fixes_table.setColumnWidth(0, 400)  # Fix column
            fixes_table.setColumnWidth(1, 100)  # Status column
            fixes_table.setColumnWidth(2, 80)   # Verify column
            
            fixes_layout.addWidget(fixes_table)
            fixes_group.setLayout(fixes_layout)
            layout.addWidget(fixes_group)
        
        # Specifications section
        if 'specifications' in ai_data:
            specs_group = QGroupBox("📖 Specifications")
            specs_layout = QVBoxLayout()
            
            specs_table = QTableWidget()
            specs_table.setColumnCount(1)
            specs_table.setHorizontalHeaderLabels(["Specification"])
            
            row = 0
            for i, spec in enumerate(ai_data['specifications'], 1):
                specs_table.insertRow(row)
                specs_table.setItem(row, 0, QTableWidgetItem(str(spec)))
                row += 1
            
            specs_table.resizeColumnsToContents()
            # Enable scrolling
            specs_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            specs_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            specs_layout.addWidget(specs_table)
            specs_group.setLayout(specs_layout)
            layout.addWidget(specs_group)
        
        # Function Analysis section
        if 'function_analysis' in ai_data:
            func_group = QGroupBox("⚙️ Function Analysis")
            func_layout = QVBoxLayout()
            
            func_table = QTableWidget()
            func_table.setColumnCount(1)
            func_table.setHorizontalHeaderLabels(["Analysis"])
            
            row = 0
            for i, analysis in enumerate(ai_data['function_analysis'], 1):
                func_table.insertRow(row)
                func_table.setItem(row, 0, QTableWidgetItem(str(analysis)))
                row += 1
            
            func_table.resizeColumnsToContents()
            # Enable scrolling
            func_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            func_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            func_layout.addWidget(func_table)
            func_group.setLayout(func_layout)
            layout.addWidget(func_group)
        
        layout.addStretch()
        tab.setLayout(layout)
        return tab
    
    def create_meta_analysis_tab(self, meta_data):
        """Create tab for meta information with tables."""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Meta Information table
        meta_group = QGroupBox("📊 Meta Information")
        meta_layout = QVBoxLayout()
        
        meta_table = QTableWidget()
        meta_table.setColumnCount(2)
        meta_table.setHorizontalHeaderLabels(["Property", "Value"])
        
        row = 0
        if 'file' in meta_data:
            meta_table.insertRow(row)
            meta_table.setItem(row, 0, QTableWidgetItem("Source File"))
            meta_table.setItem(row, 1, QTableWidgetItem(str(meta_data['file'])))
            row += 1
            
        if 'lines_analyzed' in meta_data:
            meta_table.insertRow(row)
            meta_table.setItem(row, 0, QTableWidgetItem("Lines Analyzed"))
            meta_table.setItem(row, 1, QTableWidgetItem(str(meta_data['lines_analyzed'])))
            row += 1
            
        if 'timestamp' in meta_data:
            meta_table.insertRow(row)
            meta_table.setItem(row, 0, QTableWidgetItem("Analysis Time"))
            meta_table.setItem(row, 1, QTableWidgetItem(str(meta_data['timestamp'])))
            row += 1
            
        if 'protocol_filter' in meta_data and meta_data['protocol_filter']:
            meta_table.insertRow(row)
            meta_table.setItem(row, 0, QTableWidgetItem("Protocol Filter"))
            meta_table.setItem(row, 1, QTableWidgetItem(str(meta_data['protocol_filter'])))
            row += 1
        
        meta_table.resizeColumnsToContents()
        meta_layout.addWidget(meta_table)
        meta_group.setLayout(meta_layout)
        layout.addWidget(meta_group)
        
        layout.addStretch()
        tab.setLayout(layout)
        return tab
    
    def export_analysis_results(self, rca_data, analysis_options):
        """Export analysis results to CSV files."""
        try:
            # Get save directory
            save_dir = QFileDialog.getExistingDirectory(self, "Select Directory to Save Results")
            if not save_dir:
                return
            
            # Export each analysis type
            for analysis_type, is_selected in analysis_options.items():
                if is_selected and analysis_type in rca_data:
                    self.export_analysis_to_csv(rca_data[analysis_type], analysis_type, save_dir)
            
            QMessageBox.information(self, "Export Complete", f"Analysis results exported to {save_dir}")
            
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Error exporting results: {str(e)}")
    
    def export_analysis_to_csv(self, data, analysis_type, save_dir):
        """Export specific analysis data to CSV."""
        import csv
        
        filename = f"{analysis_type}_results.csv"
        filepath = os.path.join(save_dir, filename)
        
        with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            if analysis_type == 'error_analysis':
                # Export error details
                if 'identified_errors' in data:
                    writer.writerow(['Error Type', 'Line', 'Pattern', 'Severity', 'Layer', 'Confidence'])
                    for error_type, errors in data['identified_errors'].items():
                        for error in errors:
                            writer.writerow([
                                error_type.replace('_', ' ').title(),
                                error.get('line_number', 'N/A'),
                                error.get('pattern', 'N/A'),
                                error.get('severity', 'N/A'),
                                error.get('layer', 'N/A'),
                                error.get('confidence', 'N/A')
                            ])
            
            elif analysis_type == 'function_analysis':
                # Export function details
                if 'affected_functions' in data:
                    writer.writerow(['Function Name', 'File', 'Relevance', 'Score', 'Centrality'])
                    for func in data['affected_functions']:
                        writer.writerow([
                            func.get('name', 'N/A'),
                            func.get('file', 'N/A'),
                            func.get('relevance', 'N/A'),
                            func.get('score', 'N/A'),
                            func.get('centrality', 'N/A')
                        ])
            
            elif analysis_type == 'impact_analysis':
                # Export impact details
                if 'error_impacts' in data:
                    writer.writerow(['Error Type', 'Primary Layer', 'Severity', 'Description'])
                    for error_type, impact in data['error_impacts'].items():
                        writer.writerow([
                            error_type.replace('_', ' ').title(),
                            impact.get('primary_layer', 'N/A'),
                            impact.get('severity', 'N/A'),
                            impact.get('impact_description', 'N/A')
                        ])
            
            elif analysis_type == 'kb_analysis':
                # Export KB details
                if 'fix_recommendations' in data:
                    writer.writerow(['#', 'Recommendation'])
                    for i, rec in enumerate(data['fix_recommendations'], 1):
                        writer.writerow([i, rec])
            
            elif analysis_type == 'ai_analysis':
                # Export AI details
                if 'suggested_fixes' in data:
                    writer.writerow(['#', 'Fix'])
                    for i, fix in enumerate(data['suggested_fixes'], 1):
                        writer.writerow([i, fix])
            


    def handle_analysis_result(self, result):
        """Handle the bug analysis results and initiate RCA if needed."""
        try:
            data = json.loads(result)
            if 'error' in data:
                QMessageBox.warning(
                    self,
                    "Analysis Error",
                    f"Bug analysis failed: {data['error']}"
                )
                return

            # Format and display results
            formatted_result = ""
            for category, items in data.items():
                formatted_result += f"\n{category.upper()}\n"
                formatted_result += "=" * len(category) + "\n"
                if isinstance(items, list):
                    for item in items:
                        formatted_result += f"- {item}\n"
                elif isinstance(items, dict):
                    for key, value in items.items():
                        formatted_result += f"{key}: {value}\n"
                else:
                    formatted_result += f"{items}\n"

            self.bug_results_text.setText(formatted_result)

            # Check for critical bugs that need RCA
            critical_bugs = self.identify_critical_bugs(data)
            if critical_bugs:
                reply = QMessageBox.question(
                    self,
                    "Critical Issues Detected",
                    "Critical issues have been detected. Would you like to perform Root Cause Analysis?",
                    QMessageBox.Yes | QMessageBox.No
                )
                
                if reply == QMessageBox.Yes:
                    self.show_rca_dialog(critical_bugs)

        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Error processing analysis results: {str(e)}"
            )

    def identify_critical_bugs(self, analysis_data):
        """Identify critical bugs that need RCA."""
        critical_bugs = []
        
        # Define criteria for critical bugs
        critical_patterns = {
            'crash': r'(segmentation fault|core dump|stack overflow)',
            'deadlock': r'(deadlock|stuck|frozen)',
            'memory': r'(memory leak|buffer overflow|null pointer)',
            'performance': r'(timeout|high cpu|high memory|slow response)',
            'data_corruption': r'(corrupt|invalid|inconsistent) data'
        }
        
        try:
            # Check each error/warning in the analysis data
            for category, items in analysis_data.items():
                if isinstance(items, list):
                    for item in items:
                        for issue_type, pattern in critical_patterns.items():
                            if re.search(pattern, str(item), re.IGNORECASE):
                                critical_bugs.append({
                                    'bug_id': f"BUG_{len(critical_bugs) + 1}",
                                    'type': issue_type,
                                    'description': item,
                                    'severity': 'HIGH',
                                    'error_signatures': [pattern],
                                    'affected_files': self.extract_affected_files(item)
                                })
                                break
                
        except Exception as e:
            logging.error(f"Error identifying critical bugs: {str(e)}")
        
        return critical_bugs

    def extract_affected_files(self, error_message):
        """Extract affected files from error message."""
        files = []
        
        # Look for file paths in the error message
        file_patterns = [
            r'/[\w/\-\.]+\.[ch](?:pp)?',  # C/C++ source files
            r'[\w\-\.]+\.[ch](?:pp)?:\d+',  # Files with line numbers
        ]
        
        for pattern in file_patterns:
            matches = re.finditer(pattern, error_message)
            for match in matches:
                file_path = match.group(0).split(':')[0]
                if os.path.exists(file_path):
                    files.append(file_path)
        
        return files

    def save_analysis_results(self):
        """Save the analysis results to a file and database."""
        try:
            # Get file save location
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"bug_analysis_{timestamp}.txt"
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Analysis Results",
                os.path.join(os.path.expanduser("~"), default_filename),
                "Text Files (*.txt);;All Files (*.*)"
            )

            if not file_path:
                return

            # Save to file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(self.results_text.toPlainText())

            # Save to database
            analysis_data = {
                'action_type': 'ANALYZE_BUG',
                'description': f"Log analysis from {self.log_dir_edit.text()}",
                'details': json.dumps({
                    'timestamp': datetime.now().isoformat(),
                    'log_directory': self.log_dir_edit.text(),
                    'log_type': self.log_file_combo.currentText(),
                    'analysis_options': {
                        'error_pattern': self.error_pattern_check.isChecked(),
                        'performance': self.performance_check.isChecked(),
                        'resource': self.resource_check.isChecked(),
                        'timing': self.timing_check.isChecked()
                    },
                    'results': self.results_text.toPlainText(),
                    'file_path': file_path
                })
            }
            
            # Save to database
            save_to_database('user_history', analysis_data)

            QMessageBox.information(
                self,
                "Success",
                f"Analysis results have been saved to:\n{file_path}\nand recorded in the history database."
            )

        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to save analysis results: {str(e)}"
            )

    def show_code_testing(self):
        """Show the code testing dialog - pre-validation UI scaffold."""
        self.code_testing_dialog = QDialog(self)
        dialog = self.code_testing_dialog
        dialog.setWindowTitle("Code Evaluation - Evaluate Code Changes")

        screen = QApplication.primaryScreen()
        if screen:
            screen_geometry = screen.availableGeometry()
            dialog_width = min(1250, int(screen_geometry.width() * 0.85))
            dialog_height = min(900, int(screen_geometry.height() * 0.85))
            dialog.resize(dialog_width, dialog_height)

        dialog.setMinimumSize(1000, 650)
        dialog.setSizeGripEnabled(True)

        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setStyleSheet("QScrollArea { border: none; }")

        container_widget = QWidget()
        layout = QVBoxLayout(container_widget)
        layout.setSpacing(14)
        layout.setContentsMargins(16, 16, 16, 16)

        # Bug analysis selection
        history_group = QGroupBox("Bug Analysis Selection")
        history_layout = QHBoxLayout()
        history_layout.setSpacing(8)

        history_label = QLabel("Analysis:")
        history_label.setStyleSheet("font-weight: 600;")

        self.testing_history_combo = QComboBox()
        self.testing_history_combo.setMinimumWidth(420)

        load_btn = QPushButton("Load Analysis")
        load_btn.setStyleSheet(
            """
            QPushButton {
                background-color: #0078D4;
                color: white;
                font-weight: bold;
                padding: 6px 16px;
                border: none;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
            """
        )
        load_btn.clicked.connect(self.load_selected_testing_analysis)

        refresh_btn = QPushButton("↻")
        refresh_btn.setMaximumWidth(36)
        refresh_btn.setToolTip("Refresh analysis list")
        refresh_btn.clicked.connect(self.load_testing_bug_history)

        history_layout.addWidget(history_label)
        history_layout.addWidget(self.testing_history_combo, 1)
        history_layout.addWidget(load_btn)
        history_layout.addWidget(refresh_btn)
        history_group.setLayout(history_layout)
        history_group.setMaximumHeight(70)

        layout.addWidget(history_group)

        # Initialize containers for patches and preview
        self.testing_code_patch_checkboxes = []
        self.testing_config_patch_checkboxes = []
        self.testing_code_patch_data = []
        self.testing_config_patch_data = []
        self.testing_patch_value_editors = {}
        self.testing_loaded_bug_analysis = None
        self.testing_code_dir = None

        # Error details and preview side-by-side
        info_layout = QHBoxLayout()
        info_layout.setSpacing(10)

        error_group = QGroupBox("Error Details")
        error_layout = QVBoxLayout()
        error_layout.setContentsMargins(6, 6, 6, 6)

        self.testing_error_display = QTextEdit()
        self.testing_error_display.setReadOnly(True)
        self.testing_error_display.setMinimumHeight(200)
        self.testing_error_display.setMaximumHeight(260)
        self.testing_error_display.setPlaceholderText("Load an analysis to view error context.")

        error_layout.addWidget(self.testing_error_display)
        error_group.setLayout(error_layout)
        info_layout.addWidget(error_group, 2)

        preview_group = QGroupBox("Selected Patch Preview")
        preview_layout = QVBoxLayout()
        preview_layout.setContentsMargins(6, 6, 6, 6)

        self.testing_patch_preview = QTextEdit()
        self.testing_patch_preview.setReadOnly(True)
        self.testing_patch_preview.setMinimumHeight(200)
        self.testing_patch_preview.setMaximumHeight(260)
        self.testing_patch_preview.setPlaceholderText("Selected patch details will appear here.")
        self.testing_patch_preview.setStyleSheet(
            "QTextEdit {"
            "    font-family: 'Consolas', monospace;"
            "    font-size: 10pt;"
            "    background-color: #f9fafc;"
            "    border: 1px solid #d0d7de;"
            "    border-radius: 4px;"
            "}"
        )

        preview_layout.addWidget(self.testing_patch_preview)
        preview_group.setLayout(preview_layout)
        info_layout.addWidget(preview_group, 3)

        layout.addLayout(info_layout)

        # Patch selection area
        patches_group = QGroupBox("Patch Selection")
        patches_layout = QHBoxLayout()
        patches_layout.setSpacing(12)

        # Code patches column
        code_group = QGroupBox("Code Patches")
        code_group_layout = QVBoxLayout()
        code_group_layout.setContentsMargins(6, 6, 6, 6)

        code_scroll = QScrollArea()
        code_scroll.setWidgetResizable(True)
        code_scroll.setMinimumHeight(260)

        self.testing_code_patches_widget = QWidget()
        self.testing_code_patches_layout = QVBoxLayout()
        self.testing_code_patches_layout.setSpacing(6)
        self.testing_code_patches_layout.setContentsMargins(0, 0, 0, 0)

        self.testing_code_patches_placeholder = QLabel("No code patches available. Load an analysis to begin.")
        self.testing_code_patches_placeholder.setStyleSheet("color: #888; font-style: italic;")
        self.testing_code_patches_layout.addWidget(self.testing_code_patches_placeholder)

        self.testing_code_patches_widget.setLayout(self.testing_code_patches_layout)
        code_scroll.setWidget(self.testing_code_patches_widget)
        code_group_layout.addWidget(code_scroll)
        code_group.setLayout(code_group_layout)

        # Config patches column
        config_group = QGroupBox("Config Patches")
        config_group_layout = QVBoxLayout()
        config_group_layout.setContentsMargins(6, 6, 6, 6)

        config_scroll = QScrollArea()
        config_scroll.setWidgetResizable(True)
        config_scroll.setMinimumHeight(260)

        self.testing_config_patches_widget = QWidget()
        self.testing_config_patches_layout = QVBoxLayout()
        self.testing_config_patches_layout.setSpacing(6)
        self.testing_config_patches_layout.setContentsMargins(0, 0, 0, 0)

        self.testing_config_patches_placeholder = QLabel("No config patches available. Load an analysis to begin.")
        self.testing_config_patches_placeholder.setStyleSheet("color: #888; font-style: italic;")
        self.testing_config_patches_layout.addWidget(self.testing_config_patches_placeholder)

        self.testing_config_patches_widget.setLayout(self.testing_config_patches_layout)
        config_scroll.setWidget(self.testing_config_patches_widget)
        config_group_layout.addWidget(config_scroll)
        config_group.setLayout(config_group_layout)

        patches_layout.addWidget(code_group)
        patches_layout.addWidget(config_group)
        patches_group.setLayout(patches_layout)

        layout.addWidget(patches_group)

        # Test layer selection
        test_group = QGroupBox("Test Options")
        test_layout = QVBoxLayout()
        test_layout.setSpacing(6)
        test_layout.setContentsMargins(8, 8, 8, 8)

        test_description = QLabel("Select which validation layers to execute during code evaluation.")
        test_description.setStyleSheet("color: #605E5C; font-size: 12px;")
        test_layout.addWidget(test_description)

        self.testing_layer_checkboxes = {}
        layer_definitions = [
            ("layer1", "Layer 1: Syntax & Structural Validation"),
            ("layer2", "Layer 2: 3GPP Spec Reference Analysis"),
            ("layer3", "Layer 3: LLM as Judge"),
            ("layer4", "Layer 4: Variable Impact Analysis"),
        ]

        for key, label_text in layer_definitions:
            checkbox = QCheckBox(label_text)
            checkbox.setChecked(key in {"layer1", "layer2"})
            checkbox.setStyleSheet("font-size: 13px; padding: 4px;")
            self.testing_layer_checkboxes[key] = checkbox
            test_layout.addWidget(checkbox)

        test_group.setLayout(test_layout)
        layout.addWidget(test_group)

        # Control buttons
        control_layout = QHBoxLayout()
        control_layout.setSpacing(8)

        run_btn = QPushButton("Run Selected Tests")
        run_btn.setStyleSheet(
            """
            QPushButton {
                background-color: #107C10;
                color: white;
                font-weight: bold;
                padding: 6px 20px;
                border: none;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #0F5F0F;
            }
            """
        )
        run_btn.clicked.connect(self.run_code_testing)

        assistant_btn = QPushButton("Open Code Assistant")
        assistant_btn.setStyleSheet(
            """
            QPushButton {
                background-color: #ffffff;
                border: 1px solid #d0d7de;
                padding: 6px 16px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #f3f4f6;
            }
            """
        )
        assistant_btn.clicked.connect(self.go_to_code_assistant_from_testing)

        self.testing_results_button = QPushButton("View Results")
        self.testing_results_button.setStyleSheet(
            """
            QPushButton {
                background-color: #ffffff;
                border: 1px solid #d0d7de;
                padding: 6px 16px;
                border-radius: 3px;
            }
            QPushButton:disabled {
                color: #9aa0a6;
                background-color: #f3f4f6;
            }
            QPushButton:hover:!disabled {
                background-color: #f3f4f6;
            }
            """
        )
        self.testing_results_button.setEnabled(False)
        self.testing_results_button.clicked.connect(self.show_testing_results_dialog)

        clear_output_btn = QPushButton("Clear Output")
        clear_output_btn.setStyleSheet(
            """
            QPushButton {
                background-color: #ffffff;
                border: 1px solid #d0d7de;
                padding: 6px 16px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #f3f4f6;
            }
            """
        )
        clear_output_btn.clicked.connect(self.clear_testing_output)

        control_layout.addStretch()
        control_layout.addWidget(run_btn)
        control_layout.addWidget(assistant_btn)
        control_layout.addWidget(self.testing_results_button)
        control_layout.addWidget(clear_output_btn)

        layout.addLayout(control_layout)

        # Output display
        output_group = QGroupBox("Test Output")
        output_layout = QVBoxLayout()
        output_layout.setContentsMargins(6, 6, 6, 6)

        self.testing_output_display = QTextEdit()
        self.testing_output_display.setReadOnly(True)
        self.testing_output_display.setMinimumHeight(220)
        self.testing_output_display.setPlaceholderText("Execution summaries and pass/fail results will appear here.")
        self.testing_output_display.setStyleSheet(
            "QTextEdit {"
            "    background-color: #0b1623;"
            "    color: #f8f9fb;"
            "    font-family: 'Consolas', monospace;"
            "    font-size: 10pt;"
            "    border-radius: 4px;"
            "    padding: 8px;"
            "}"
        )

        output_layout.addWidget(self.testing_output_display)
        output_group.setLayout(output_layout)
        layout.addWidget(output_group)

        main_scroll.setWidget(container_widget)
        dialog_layout = QVBoxLayout()
        dialog_layout.addWidget(main_scroll)
        dialog.setLayout(dialog_layout)

        # Populate analysis list when dialog opens
        self.load_testing_bug_history()

        dialog.exec_()

    def load_testing_bug_history(self):
        """Load bug analysis history for the code testing dialog."""
        if not hasattr(self, 'testing_history_combo'):
            return

        self.testing_history_combo.clear()
        self.testing_history_combo.addItem("Select a bug analysis...")

        workspace_root = os.path.dirname(os.path.abspath(__file__))
        history_dir = os.path.join(workspace_root, "backend", "resources", "bug_history")

        self.testing_bug_history_map = {}

        if not os.path.exists(history_dir):
            print(f"DEBUG [Code Testing]: Bug history directory not found: {history_dir}")
            return

        history_files = [f for f in os.listdir(history_dir) if f.endswith('.json')]
        history_files.sort(reverse=True)

        for filename in history_files:
            file_path = os.path.join(history_dir, filename)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                error_message = data.get('error_message', 'Unknown error') or 'Unknown error'
                timestamp = data.get('timestamp', '')
                display_text = error_message[:100] + "..." if len(error_message) > 100 else error_message

                if timestamp:
                    try:
                        dt = datetime.fromisoformat(timestamp)
                        display_text = f"[{dt.strftime('%Y-%m-%d %H:%M')}] {display_text}"
                    except Exception:
                        pass

                results = data.get('results', {})
                phase3_fixes = results.get('phase3_fixes', {})
                fix_suggestion = phase3_fixes.get('fix_suggestion', {})
                code_count = len(fix_suggestion.get('code_patches', []))
                config_count = len(fix_suggestion.get('config_patches', []))

                display_text += f" [Code:{code_count}, Config:{config_count}]"

                self.testing_history_combo.addItem(display_text)
                self.testing_bug_history_map[display_text] = file_path
            except Exception as exc:
                print(f"DEBUG [Code Testing]: Failed to load {filename}: {exc}")

    def load_selected_testing_analysis(self):
        """Load analysis chosen in the code testing dialog and populate patches."""
        if not hasattr(self, 'testing_history_combo'):
            return

        selected_text = self.testing_history_combo.currentText()

        if selected_text == "Select a bug analysis...":
            QMessageBox.warning(self, "No Selection", "Please select a bug analysis from the dropdown.")
            return

        if selected_text not in getattr(self, 'testing_bug_history_map', {}):
            QMessageBox.warning(self, "Error", "Could not find the selected analysis file.")
            return

        file_path = self.testing_bug_history_map[selected_text]

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            self.testing_loaded_bug_analysis = data
            self.testing_code_dir = data.get('code_dir') or data.get('code_directory')

            error_message = data.get('error_message', 'Unknown error') or 'Unknown error'
            log_file = data.get('log_file', 'N/A')
            timestamp = data.get('timestamp', 'N/A')

            error_details = [
                f"Error Message:\n{error_message}",
                f"Log File: {log_file}",
                f"Timestamp: {timestamp}"
            ]
            self.testing_error_display.setText("\n\n".join(error_details))

            results = data.get('results', {})
            phase3_fixes = results.get('phase3_fixes', {})
            fix_suggestion = phase3_fixes.get('fix_suggestion', {})

            code_patches = fix_suggestion.get('code_patches', [])
            config_patches = fix_suggestion.get('config_patches', [])

            self.populate_testing_patches(code_patches, config_patches)
            self.update_testing_patch_preview()

            QMessageBox.information(self, "Analysis Loaded", (
                f"Loaded {len(code_patches)} code patches and {len(config_patches)} config patches."))

        except Exception as exc:
            QMessageBox.critical(self, "Error", f"Failed to load analysis:\n{exc}")
            print(f"DEBUG [Code Testing]: Error loading analysis: {exc}")

    def populate_testing_patches(self, code_patches, config_patches):
        """Populate patch selection lists for the code testing dialog."""
        self.clear_testing_patch_checkboxes()

        self.testing_code_patch_data = code_patches or []
        self.testing_config_patch_data = config_patches or []

        # Code patches
        if self.testing_code_patch_data:
            self.testing_code_patches_placeholder.hide()
            for patch in self.testing_code_patch_data:
                function_name = patch.get('function_name') or patch.get('function') or 'Unknown'
                file_path = patch.get('file_path', 'Unknown')
                file_name = os.path.basename(file_path) if file_path != 'Unknown' else 'Unknown file'
                checkbox_text = f"{function_name} ({file_name})"

                checkbox = QCheckBox(checkbox_text)
                checkbox.setChecked(True)
                checkbox.setToolTip(f"File: {file_path}\nDescription: {patch.get('description', 'N/A')}")
                checkbox.stateChanged.connect(self.update_testing_patch_preview)
                self.testing_code_patches_layout.addWidget(checkbox)
                self.testing_code_patch_checkboxes.append(checkbox)
            self.testing_code_patches_layout.addStretch()
        else:
            self.testing_code_patches_placeholder.show()

        # Config patches
        if self.testing_config_patch_data:
            self.testing_config_patches_placeholder.hide()
            for patch in self.testing_config_patch_data:
                param_name = patch.get('config_name', patch.get('parameter_name', 'Unknown'))
                file_path = patch.get('file_path', 'Unknown')
                file_name = os.path.basename(file_path) if file_path != 'Unknown' else 'Unknown file'
                checkbox_text = f"{param_name} ({file_name})"

                checkbox = QCheckBox(checkbox_text)
                checkbox.setChecked(True)
                checkbox.setToolTip(f"File: {file_path}\nDescription: {patch.get('description', 'N/A')}")
                checkbox.stateChanged.connect(self.update_testing_patch_preview)
                self.testing_config_patches_layout.addWidget(checkbox)
                self.testing_config_patch_checkboxes.append(checkbox)
            self.testing_config_patches_layout.addStretch()
        else:
            self.testing_config_patches_placeholder.show()

    def clear_testing_patch_checkboxes(self):
        """Clear existing patch selections in the code testing dialog."""
        if hasattr(self, 'testing_code_patch_checkboxes'):
            for checkbox in self.testing_code_patch_checkboxes:
                self.testing_code_patches_layout.removeWidget(checkbox)
                checkbox.deleteLater()
            self.testing_code_patch_checkboxes = []
            self._remove_layout_spacers(self.testing_code_patches_layout)

        if hasattr(self, 'testing_config_patch_checkboxes'):
            for checkbox in self.testing_config_patch_checkboxes:
                self.testing_config_patches_layout.removeWidget(checkbox)
                checkbox.deleteLater()
            self.testing_config_patch_checkboxes = []
            self._remove_layout_spacers(self.testing_config_patches_layout)

        if hasattr(self, 'testing_code_patches_placeholder'):
            self.testing_code_patches_placeholder.show()
            if self.testing_code_patches_layout.indexOf(self.testing_code_patches_placeholder) == -1:
                self.testing_code_patches_layout.insertWidget(0, self.testing_code_patches_placeholder)

        if hasattr(self, 'testing_config_patches_placeholder'):
            self.testing_config_patches_placeholder.show()
            if self.testing_config_patches_layout.indexOf(self.testing_config_patches_placeholder) == -1:
                self.testing_config_patches_layout.insertWidget(0, self.testing_config_patches_placeholder)

    def _remove_layout_spacers(self, layout):
        """Remove spacer items from a layout to prevent content from being pushed to the bottom."""
        if layout is None:
            return

        for i in reversed(range(layout.count())):
            item = layout.itemAt(i)
            if item is not None and item.spacerItem() is not None:
                layout.takeAt(i)

    def update_testing_patch_preview(self):
        """Refresh the patch preview based on current selections."""
        if not hasattr(self, 'testing_patch_preview'):
            return

        if not self.testing_loaded_bug_analysis:
            self.testing_patch_preview.setPlainText("Load a bug analysis to view patch details.")
            return

        lines = []

        selected_code = [
            (checkbox, idx)
            for idx, checkbox in enumerate(self.testing_code_patch_checkboxes)
            if checkbox.isChecked() and idx < len(self.testing_code_patch_data)
        ]

        selected_config = [
            (checkbox, idx)
            for idx, checkbox in enumerate(self.testing_config_patch_checkboxes)
            if checkbox.isChecked() and idx < len(self.testing_config_patch_data)
        ]

        if not selected_code and not selected_config:
            self.testing_patch_preview.setPlainText("Select at least one code or config patch to preview details.")
            return

        for pos, (checkbox, idx) in enumerate(selected_code, start=1):
            patch = self.testing_code_patch_data[idx]
            function_name = patch.get('function_name', 'Unknown Function')
            file_path = patch.get('file_path', 'Unknown File')
            description = patch.get('description', 'No description provided.')
            patched_code = patch.get('suggested_code') or patch.get('patched_code') or patch.get('new_code')

            lines.append(f"CODE PATCH {pos}")
            lines.append(f"Function : {function_name}")
            lines.append(f"File     : {file_path}")
            lines.append(f"Details  : {description}")
            if patched_code:
                lines.append("Suggested Code:")
                lines.append(patched_code)
            lines.append("\n")

        for pos, (checkbox, idx) in enumerate(selected_config, start=1):
            patch = self.testing_config_patch_data[idx]
            param_name = patch.get('config_name', patch.get('parameter_name', 'Unknown Parameter'))
            file_path = patch.get('file_path', 'Unknown File')
            description = patch.get('description', 'No description provided.')
            suggested_value = patch.get('suggested_value') or patch.get('patched_value') or patch.get('new_value')

            lines.append(f"CONFIG PATCH {pos}")
            lines.append(f"Parameter: {param_name}")
            lines.append(f"File     : {file_path}")
            lines.append(f"Details  : {description}")
            if suggested_value:
                lines.append(f"Suggested Value: {suggested_value}")
            lines.append("\n")

        preview_text = "\n".join(lines).strip()
        self.testing_patch_preview.setPlainText(preview_text)

    def gather_selected_testing_layers(self):
        """Return list of selected validation layers for code testing."""
        if not hasattr(self, 'testing_layer_checkboxes'):
            return []

        selected_layers = []
        for key, checkbox in self.testing_layer_checkboxes.items():
            if checkbox.isChecked():
                selected_layers.append((key, checkbox.text()))
        return selected_layers

    def run_code_testing(self):
        """Run selected code testing validation layers."""
        if not hasattr(self, 'testing_output_display'):
            return

        selected_layers = self.gather_selected_testing_layers()

        if not self.testing_loaded_bug_analysis:
            QMessageBox.warning(self, "No Analysis Loaded", "Load a bug analysis before running tests.")
            return

        if not selected_layers:
            QMessageBox.warning(self, "No Tests Selected", "Select at least one validation layer to execute.")
            return

        if not hasattr(self, 'testing_last_results'):
            self.testing_last_results = []

        if hasattr(self, 'testing_results_button'):
            self.testing_results_button.setEnabled(False)
        self.testing_last_results = []

        if hasattr(self, 'testing_output_display'):
            self.testing_output_display.setPlainText("Running selected tests...")
            QApplication.processEvents()

        progress_dialog = QProgressDialog("Running selected validation layers...", None, 0, 0, self)
        progress_dialog.setWindowModality(Qt.ApplicationModal)
        progress_dialog.setCancelButton(None)
        progress_dialog.setMinimumDuration(0)
        progress_dialog.show()
        QApplication.processEvents()

        try:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            layer_labels = ", ".join(label for _, label in selected_layers)
            log_lines = [
                f"[{timestamp}] Code testing initiated.",
                f"Selected layers: {layer_labels if layer_labels else 'None'}",
                ""
            ]

            selected_layer_keys = {key for key, _ in selected_layers}
            selected_code_patches = self.get_selected_testing_code_patches()
            selected_config_patches = self.get_selected_testing_config_patches()
            layer_results = []

            if 'layer1' in selected_layer_keys:
                layer1_lines = self.code_testing_engine.run_layer1_syntax_validation(
                    selected_code_patches,
                    self.testing_code_dir,
                    config_patches=selected_config_patches,
                )
                layer1_plain = "\n".join(layer1_lines)
                layer1_html = self._format_layer1_result_html(layer1_lines)
                log_lines.extend(layer1_lines)
                log_lines.append("")
                layer_results.append({
                    "label": "Layer 1: Syntax & Structural Validation",
                    "content": layer1_html or self._preformatted_html(layer1_plain),
                    "is_html": True,
                })
            else:
                message = "Layer 1: Syntax & Structural Validation not selected."
                log_lines.append(message)
                log_lines.append("")
                layer_results.append({
                    "label": "Layer 1: Syntax & Structural Validation",
                    "content": self._preformatted_html(message),
                    "is_html": True,
                })

            if 'layer2' in selected_layer_keys:
                layer2_lines, layer2_html = self.run_spec_reference_analysis(
                    selected_code_patches,
                    selected_config_patches,
                )
                log_lines.extend(layer2_lines)
                log_lines.append("")
                layer_results.append({
                    "label": "Layer 2: 3GPP Spec Reference Analysis",
                    "content": layer2_html,
                    "is_html": True,
                })
            else:
                message = "Layer 2: 3GPP Spec Reference Analysis not selected."
                log_lines.append(message)
                log_lines.append("")
                layer_results.append({
                    "label": "Layer 2: 3GPP Spec Reference Analysis",
                    "content": self._preformatted_html(message),
                    "is_html": True,
                })

            if 'layer3' in selected_layer_keys:
                layer3_lines, layer3_html = self.run_llm_judge_analysis(
                    selected_code_patches,
                    selected_config_patches,
                )
                log_lines.extend(layer3_lines)
                log_lines.append("")
                layer_results.append({
                    "label": "Layer 3: LLM as Judge",
                    "content": layer3_html,
                    "is_html": True,
                })
            else:
                message = "Layer 3: LLM as Judge not selected."
                log_lines.append(message)
                log_lines.append("")
                layer_results.append({
                    "label": "Layer 3: LLM as Judge",
                    "content": self._preformatted_html(message),
                    "is_html": True,
                })

            if 'layer4' in selected_layer_keys:
                analysis = getattr(self, 'testing_loaded_bug_analysis', {}) or {}
                variable_error_summary = (
                    analysis.get('error_message')
                    or analysis.get('results', {}).get('error_message')
                    or None
                )
                layer4_lines, layer4_results = self.code_testing_engine.run_variable_impact_analysis(
                    selected_code_patches,
                    self.testing_code_dir,
                    error_summary=variable_error_summary,
                )
                log_lines.extend(layer4_lines)
                log_lines.append("")
                layer_results.append({
                    "label": "Layer 4: Variable Impact Analysis",
                    "content": self._format_variable_impact_html(layer4_results, layer4_lines),
                    "is_html": True,
                })
            else:
                message = "Layer 4: Variable Impact Analysis not selected."
                log_lines.append(message)
                log_lines.append("")
                layer_results.append({
                    "label": "Layer 4: Variable Impact Analysis",
                    "content": self._preformatted_html(message),
                    "is_html": True,
                })

            summary = "\n".join(line for line in log_lines if line is not None)
            self.testing_last_results = layer_results
            self.testing_run_history_log = summary

            if hasattr(self, 'testing_output_display'):
                self.testing_output_display.setPlainText(
                    "Test run completed.\nClick 'View Results' to inspect detailed output for each layer."
                )

            if hasattr(self, 'testing_results_button'):
                self.testing_results_button.setEnabled(True)
        finally:
            try:
                progress_dialog.close()
            except Exception:
                pass

    def show_testing_results_dialog(self):
        """Display the latest test results inside a tabbed dialog."""
        results = getattr(self, "testing_last_results", None)
        if not results:
            QMessageBox.information(self, "No Results", "Run tests before viewing results.")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("Code Testing Results")
        dialog.resize(900, 600)

        layout = QVBoxLayout(dialog)
        tabs = QTabWidget(dialog)
        layout.addWidget(tabs)

        for entry in results:
            label = entry.get("label", "Layer")
            content = entry.get("content", "")
            is_html = entry.get("is_html", False)
            tab = QWidget()
            tab_layout = QVBoxLayout(tab)
            text_edit = QTextEdit()
            text_edit.setReadOnly(True)
            if is_html:
                text_edit.setHtml(content)
            else:
                text_edit.setPlainText(content)
            tab_layout.addWidget(text_edit)
            tabs.addTab(tab, label)

        button_box = QDialogButtonBox(QDialogButtonBox.Close)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)

        dialog.exec_()

    def clear_testing_output(self):
        """Clear output display and reset stored results."""
        if hasattr(self, "testing_output_display"):
            self.testing_output_display.clear()
        self.testing_last_results = []
        if hasattr(self, "testing_results_button"):
            self.testing_results_button.setEnabled(False)


    def get_selected_testing_code_patches(self):
        """Return code patch data for currently selected checkboxes."""
        patches = []
        if not hasattr(self, 'testing_code_patch_checkboxes'):
            return patches

        for idx, checkbox in enumerate(self.testing_code_patch_checkboxes):
            if checkbox.isChecked() and idx < len(self.testing_code_patch_data):
                patches.append(self.testing_code_patch_data[idx])
        return patches

    def get_selected_testing_config_patches(self):
        """Return config patch data for currently selected checkboxes."""
        patches = []
        if not hasattr(self, 'testing_config_patch_checkboxes'):
            return patches

        for idx, checkbox in enumerate(self.testing_config_patch_checkboxes):
            if checkbox.isChecked() and idx < len(self.testing_config_patch_data):
                patches.append(self.testing_config_patch_data[idx])
        return patches

    def run_spec_reference_analysis(self, selected_code_patches, selected_config_patches):
        """Execute the 3GPP spec reference review layer."""
        lines = ["Layer 2: 3GPP Spec Reference Analysis"]

        if not selected_code_patches and not selected_config_patches:
            lines.append("  ⚠️ No code or config patches selected – skipping spec reference analysis.")
            return lines, self._preformatted_html("\n".join(lines))

        if not getattr(self, 'spec_reference_evaluator', None):
            lines.append("  ❌ Spec reference evaluator is not configured.")
            return lines, self._preformatted_html("\n".join(lines))

        spec_context, fix_suggestion = self._extract_spec_context()
        if not spec_context:
            lines.append("  ⚠️ No specification context found in the analysis data.")
            return lines, self._preformatted_html("\n".join(lines))

        error_summary = (
            self.testing_loaded_bug_analysis.get('error_message') or
            self.testing_loaded_bug_analysis.get('results', {}).get('error_message', 'Unknown error')
        )

        additional_notes = self._extract_additional_notes(fix_suggestion)
        code_changes = self._build_spec_change_entries(selected_code_patches)
        config_changes = self._build_config_change_entries(selected_config_patches)
        combined_changes = code_changes + config_changes

        if not combined_changes:
            lines.append("  ⚠️ No reviewable changes were generated from the selected patches.")
            return lines, self._preformatted_html("\n".join(lines))

        result = self.spec_reference_evaluator.evaluate(
            error_summary=error_summary,
            code_changes=combined_changes,
            spec_context=spec_context,
            additional_notes=additional_notes,
        )

        if result.status != "success":
            for message in result.messages:
                lines.extend(self._wrap_and_indent(message))
            return lines, self._preformatted_html("\n".join(lines))

        verdict = (result.verdict or "UNKNOWN").upper()
        verdict_icon = {
            "APPROVE": "✅",
            "APPROVE_WITH_WARNINGS": "⚠️",
            "REJECT": "❌",
        }.get(verdict, "ℹ️")
        verdict_line = f"  {verdict_icon} Verdict: {verdict.replace('_', ' ').title()}"
        if result.confidence is not None:
            verdict_line += f" (Confidence: {result.confidence:.2f})"
        lines.append(verdict_line)

        if result.will_fix_issue is not None:
            fix_text = "Yes" if result.will_fix_issue else "No"
            lines.append(f"  Fix Effectiveness: {fix_text}")

        if result.summary:
            lines.append("  Summary:")
            lines.extend(self._wrap_and_indent(result.summary, indent="    "))

        if result.issues:
            lines.append("  Spec Findings:")
            for issue in result.issues:
                severity = issue.get('severity', 'UNKNOWN')
                reference = issue.get('spec_reference', 'Unknown reference')
                description = issue.get('description', '')
                lines.append(f"    - [{severity}] {reference}: {description}")
                action = issue.get('suggested_action')
                if action:
                    lines.extend(self._wrap_and_indent(f"Suggested action: {action}", indent="       "))

        if result.reasoning:
            lines.append("  Reasoning:")
            lines.extend(self._wrap_and_indent(result.reasoning, indent="    "))

        if result.recommendations:
            lines.append("  Recommendations:")
            for idx, rec in enumerate(result.recommendations, start=1):
                lines.extend(self._wrap_and_indent(f"{idx}. {rec}", indent="    "))

        for message in result.messages:
            lines.extend(self._wrap_and_indent(message))

        html_content = self._format_spec_result_html(result)
        return lines, html_content

    def run_llm_judge_analysis(self, selected_code_patches, selected_config_patches):
        """Execute the LLM-as-judge review layer."""
        lines = ["Layer 3: LLM as Judge"]

        if not selected_code_patches and not selected_config_patches:
            lines.append("  ⚠️ No code or config patches selected – skipping judge evaluation.")
            return lines, self._preformatted_html("\n".join(lines))

        if not getattr(self, 'llm_judge_evaluator', None):
            lines.append("  ❌ LLM judge evaluator is not configured.")
            return lines, self._preformatted_html("\n".join(lines))

        analysis = getattr(self, 'testing_loaded_bug_analysis', {}) or {}
        error_summary = (
            analysis.get('error_message')
            or analysis.get('results', {}).get('error_message')
            or 'Unknown error'
        )

        code_changes = self._build_spec_change_entries(selected_code_patches)

        config_changes = self._build_config_change_entries(selected_config_patches)

        result = self.llm_judge_evaluator.evaluate(
            error_summary=error_summary,
            code_changes=code_changes,
            config_changes=config_changes,
        )

        if result.status != "success":
            for message in result.messages:
                lines.extend(self._wrap_and_indent(message))
            return lines, self._preformatted_html("\n".join(lines))

        verdict = (result.verdict or "UNKNOWN").upper()
        icon = {
            "PASS": "✅",
            "FAIL": "❌",
            "NEEDS_REVIEW": "⚠️",
        }.get(verdict, "ℹ️")
        verdict_line = f"  {icon} Verdict: {verdict.replace('_', ' ').title()}"
        if result.confidence is not None:
            verdict_line += f" (Confidence: {result.confidence:.2f})"
        lines.append(verdict_line)

        if result.summary:
            lines.append("  Summary:")
            lines.extend(self._wrap_and_indent(result.summary, indent="    "))

        if result.reasoning:
            lines.append("  Reasoning:")
            lines.extend(self._wrap_and_indent(result.reasoning, indent="    "))

        if result.criteria:
            lines.append("  Criteria Evaluation:")
            for key, value in result.criteria.items():
                lines.append(f"    - {key.replace('_', ' ').title()}: {value}")

        if result.messages:
            for message in result.messages:
                lines.extend(self._wrap_and_indent(message))

        html_content = self._format_judge_result_html(result)
        return lines, html_content

    def _extract_spec_context(self):
        """Retrieve spec context and fix suggestion from the loaded analysis."""
        analysis = getattr(self, 'testing_loaded_bug_analysis', {}) or {}
        results = analysis.get('results', {})
        phase3 = results.get('phase3_fixes', {})
        fix_suggestion = phase3.get('fix_suggestion', {})

        spec_context = (
            fix_suggestion.get('specification_context') or
            phase3.get('specification_context') or
            results.get('specification_context')
        )

        return spec_context, fix_suggestion

    def _extract_additional_notes(self, fix_suggestion):
        """Compose additional context for the LLM review."""
        notes = []
        root_cause = fix_suggestion.get('root_cause_analysis')
        if root_cause:
            notes.append(f"Root Cause Analysis: {root_cause}")

        investigation_steps = fix_suggestion.get('investigation_steps')
        if investigation_steps:
            formatted_steps = "\n".join(
                f"- {step}" for step in investigation_steps if step
            )
            if formatted_steps:
                notes.append("Investigation Steps:\n" + formatted_steps)

        return "\n\n".join(notes) if notes else None

    def _build_spec_change_entries(self, selected_code_patches):
        """Prepare code change entries for the spec evaluator prompt."""
        entries = []
        for patch in selected_code_patches:
            original = patch.get('original_code') or patch.get('existing_code') or ""
            patched = patch.get('patched_code') or patch.get('suggested_code') or patch.get('new_code') or ""
            diff = self._generate_unified_diff(
                original,
                patched,
                from_label=patch.get('file_path') or 'original',
                to_label=(patch.get('file_path') or 'patched') + ' (new)',
            )

            function_original = None
            function_patched = None
            function_name = patch.get('function_name')
            if function_name:
                function_original, function_patched = self._extract_function_context(
                    patch.get('file_path'),
                    function_name,
                    original,
                    patched,
                )

            entries.append({
                "change_type": "code",
                "function_name": patch.get('function_name'),
                "file_path": patch.get('file_path'),
                "description": patch.get('description'),
                "diff": diff,
                "original_code": original,
                "patched_code": patched,
                "function_original": function_original,
                "function_patched": function_patched,
            })
        return entries

    def _build_config_change_entries(self, selected_config_patches):
        """Prepare configuration change entries for the spec evaluator prompt."""
        entries = []
        for patch in selected_config_patches:
            param_name = patch.get('config_name') or patch.get('parameter_name') or patch.get('setting')
            file_path = patch.get('file_path')
            description = patch.get('description')
            current_value = patch.get('current_value')
            new_value = (
                patch.get('new_value')
                or patch.get('suggested_value')
                or patch.get('patched_value')
            )

            diff_lines = []
            if param_name:
                diff_lines.append(f"Parameter: {param_name}")
            if current_value is not None:
                diff_lines.append(f"Current Value: {current_value}")
            if new_value is not None:
                diff_lines.append(f"Proposed Value: {new_value}")
            else:
                diff_lines.append("Proposed Value: <unspecified>")

            diff_text = "\n".join(diff_lines)

            entries.append(
                {
                    "change_type": "configuration",
                    "function_name": None,
                    "file_path": file_path,
                    "description": description or f"Update configuration parameter {param_name}",
                    "diff": diff_text,
                    "parameter_name": param_name,
                    "current_value": current_value,
                    "proposed_value": new_value,
                }
            )
        return entries

    def _extract_function_context(self, file_path, function_name, original_snippet, patched_snippet):
        """Return the full original and patched function bodies for a change."""
        if not function_name:
            return None, None

        cache = getattr(self, "_function_source_cache", None)
        if cache is None:
            cache = {}
            setattr(self, "_function_source_cache", cache)

        cache_key = (file_path or "", function_name)
        cached_original = cache.get(cache_key)
        if cached_original is None:
            resolved_path = self._resolve_patch_file_path(file_path)
            if not resolved_path:
                cache[cache_key] = None
                return None, None
            try:
                with open(resolved_path, "r", encoding="utf-8", errors="ignore") as handle:
                    content = handle.read()
            except OSError:
                cache[cache_key] = None
                return None, None

            function_text = self._locate_function_text(content, function_name, original_snippet)
            cache[cache_key] = function_text

        function_original = cache.get(cache_key)
        if not function_original:
            return None, None

        function_patched = None
        if original_snippet and patched_snippet and original_snippet in function_original:
            function_patched = function_original.replace(original_snippet, patched_snippet, 1)
            if function_patched == function_original:
                function_patched = None

        return function_original, function_patched

    def _resolve_patch_file_path(self, file_path):
        """Resolve a patch file path to an absolute path on disk."""
        if not file_path:
            return None
        if os.path.isabs(file_path):
            return file_path if os.path.exists(file_path) else None

        code_dir = getattr(self, "testing_code_dir", None)
        if code_dir:
            candidate = os.path.join(code_dir, file_path)
            if os.path.exists(candidate):
                return candidate

        workspace_root = os.path.dirname(os.path.abspath(__file__))
        candidate = os.path.join(workspace_root, file_path)
        if os.path.exists(candidate):
            return candidate

        candidate = os.path.join(workspace_root, "Error_fixing_pipelin", file_path)
        if os.path.exists(candidate):
            return candidate

        return None

    @staticmethod
    def _locate_function_text(content, function_name, original_snippet):
        """Extract the full text of a function by name from source content."""

        matches = []
        token = f"{function_name}("
        search_pos = 0
        while True:
            idx = content.find(token, search_pos)
            if idx == -1:
                break

            line_start = content.rfind("\n", 0, idx)
            signature_start = 0 if line_start == -1 else line_start + 1

            # Include preceding signature lines with qualifiers/attributes.
            while signature_start > 0:
                prev_newline = content.rfind("\n", 0, signature_start - 1)
                if prev_newline == -1:
                    signature_start = 0
                    break
                candidate_line = content[prev_newline + 1:signature_start].strip()
                if not candidate_line or candidate_line.startswith("#"):
                    break
                signature_start = prev_newline + 1

            brace_idx = content.find("{", idx)
            if brace_idx == -1:
                search_pos = idx + len(token)
                continue

            brace_depth = 0
            end = None
            pos = brace_idx
            while pos < len(content):
                char = content[pos]
                if char == "{":
                    brace_depth += 1
                elif char == "}":
                    brace_depth -= 1
                    if brace_depth == 0:
                        end = pos + 1
                        break
                pos += 1

            if end is None:
                search_pos = idx + len(token)
                continue

            func_text = content[signature_start:end]
            matches.append(func_text)
            search_pos = end

        if not matches:
            return None

        if original_snippet:
            trimmed = original_snippet.strip()
            if trimmed:
                for match in matches:
                    if trimmed in match:
                        return match

        return matches[0]

    @staticmethod
    def _preformatted_html(text: str) -> str:
        """Wrap plain text in a styled preformatted HTML block."""
        escaped = html.escape(text)
        return (
            "<div style='background-color:#0b1623;color:#f8f9fb;"
            "font-family:Consolas, monospace;font-size:10pt;border-radius:6px;"
            "padding:12px;border:1px solid #1f2a35;'>"
            f"<pre style='margin:0;white-space:pre-wrap;'>{escaped}</pre>"
            "</div>"
        )

    def _format_variable_impact_html(self, structured_results, plain_lines):
        """Render variable impact analysis results as HTML."""

        if not structured_results:
            return self._preformatted_html("\n".join(plain_lines))

        parts = [
            "<div style='font-family:\"Segoe UI\",sans-serif;font-size:13px;color:#202124;'>",
            "<h3 style='margin-top:0;'>Layer 4: Variable Impact Analysis</h3>",
        ]

        for result in structured_results:
            patch_label = result.get("patch_label") or "Patch"
            verdict = result.get("verdict") or result.get("status") or "UNKNOWN"
            confidence = result.get("confidence")
            summary = result.get("summary")
            variables = result.get("per_variable") or result.get("analysis_variables") or []
            issues = result.get("issues") or []
            recommendations = result.get("recommendations") or []
            messages = result.get("messages") or []

            parts.append(
                "<div style='margin:12px 0;padding:12px;border:1px solid #dfe1e5;border-radius:6px;background:#f8f9fa;'>"
            )
            parts.append(f"<h4 style='margin:0 0 6px;'>{html.escape(str(patch_label))}</h4>")

            verdict_text = html.escape(str(verdict))
            if confidence is not None:
                verdict_text += f" (confidence {confidence:.2f})"
            parts.append(f"<p style='margin:4px 0;'><strong>Verdict:</strong> {verdict_text}</p>")

            if summary:
                parts.append(f"<p style='margin:4px 0;'><strong>Summary:</strong> {html.escape(summary)}</p>")

            if variables:
                parts.append("<div style='margin-top:8px;'><strong>Variables:</strong><ul>")
                for variable in variables:
                    name = variable.get("name") or variable.get("variable") or "<unknown>"
                    assessment = variable.get("assessment") or variable.get("change_summary")
                    notes = variable.get("notes") or variable.get("rationale")
                    risk = variable.get("risk_level") or variable.get("scope")
                    parts.append(f"<li><b>{html.escape(str(name))}</b> ({html.escape(str(risk or 'N/A'))})")
                    if assessment:
                        parts.append(f"<div style='margin-left:12px;'>{html.escape(str(assessment))}</div>")
                    if notes:
                        parts.append(f"<div style='margin-left:12px;color:#5f6368;'>{html.escape(str(notes))}</div>")
                    related = variable.get("related_functions")
                    if related:
                        parts.append("<div style='margin-left:12px;margin-top:4px;'><em>Related functions:</em><ul>")
                        for rel in related[:5]:
                            fname = rel.get("function_name") or "<unknown>"
                            fpath = rel.get("file_path") or ""
                            usage = rel.get("usage_snippet") or ""
                            parts.append(
                                f"<li>{html.escape(str(fname))} <span style='color:#5f6368;'>({html.escape(str(fpath))})</span>"
                            )
                            if usage:
                                parts.append(
                                    "<pre style='margin:4px 0 8px 0;padding:8px;background:#fff;border:1px solid #dfe1e5;border-radius:4px;'>"
                                    f"{html.escape(usage)}"
                                    "</pre>"
                                )
                            parts.append("</li>")
                        parts.append("</ul></div>")
                    parts.append("</li>")
                parts.append("</ul></div>")

            if issues:
                parts.append("<div style='margin-top:8px;'><strong>Issues:</strong><ul>")
                for issue in issues:
                    desc = issue.get("description") or issue.get("message") or ""
                    severity = issue.get("severity") or ""
                    parts.append(f"<li><b>{html.escape(str(severity))}</b>: {html.escape(str(desc))}</li>")
                parts.append("</ul></div>")

            if recommendations:
                parts.append("<div style='margin-top:8px;'><strong>Recommendations:</strong><ul>")
                for rec in recommendations:
                    parts.append(f"<li>{html.escape(str(rec))}</li>")
                parts.append("</ul></div>")

            if messages:
                parts.append("<div style='margin-top:8px;color:#5f6368;'><strong>Notes:</strong><ul>")
                for message in messages:
                    parts.append(f"<li>{html.escape(str(message))}</li>")
                parts.append("</ul></div>")

            parts.append("</div>")

        parts.append("</div>")
        return "".join(parts)
    def _format_layer1_result_html(self, lines):
        """Convert Layer 1 output into a richer HTML layout."""

        if not lines:
            return self._preformatted_html("No results.")

        header = lines[0]
        patches = []
        summary_lines = []
        general = []
        current = None

        for line in lines[1:]:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("Summary"):
                summary_lines.append(stripped)
                continue

            if stripped.startswith("Overall Summary"):
                summary_lines.append(stripped)
                continue

            lstripped = line.lstrip()
            if lstripped.startswith("→"):
                if current:
                    patches.append(current)
                title = lstripped.replace("→", "", 1).strip()
                current = {"title": title, "messages": []}
                continue

            if current:
                current["messages"].append(lstripped)
            else:
                general.append(lstripped)

        if current:
            patches.append(current)

        parts = [
            "<div style='font-family:\"Segoe UI\",sans-serif;font-size:13px;color:#202124;'>",
            f"<h3 style='margin-top:0;'>{html.escape(header)}</h3>",
        ]

        if general:
            parts.append("<ul>")
            for msg in general:
                parts.append(f"<li>{html.escape(msg)}</li>")
            parts.append("</ul>")

        for patch in patches:
            parts.append("<div style='margin:12px 0;padding:12px;border:1px solid #dfe1e5;border-radius:6px;background:#f8f9fa;'>")
            parts.append(f"<h4 style='margin:0 0 8px;'>{html.escape(patch['title'])}</h4>")
            if patch["messages"]:
                parts.append("<ul style='margin:0;padding-left:18px;'>")
                for msg in patch["messages"]:
                    parts.append(f"<li>{html.escape(msg)}</li>")
                parts.append("</ul>")
            parts.append("</div>")

        if summary_lines:
            parts.append("<h4 style='margin:12px 0 6px;'>Summary</h4><ul>")
            for summary in summary_lines:
                summary_lower = summary.lower()
                highlight = ""
                match = re.search(r'(\d+)\s+failed', summary_lower)
                if match:
                    failed_count = int(match.group(1))
                    if failed_count == 0:
                        highlight = "color:#1a7f37;font-weight:bold;"
                    else:
                        highlight = "color:#d93025;font-weight:bold;"

                parts.append(f"<li><span style='{highlight}'>{html.escape(summary)}</span></li>")
            parts.append("</ul>")

        parts.append("</div>")
        return "".join(parts)

    def _format_judge_result_html(self, result):
        """Render the LLM judge evaluation in rich HTML."""

        verdict = (result.verdict or "UNKNOWN").upper()
        icon = {
            "PASS": "✅",
            "FAIL": "❌",
            "NEEDS_REVIEW": "⚠️",
        }.get(verdict, "ℹ️")

        def paragraph(text: Optional[str]) -> str:
            if not text:
                return ""
            escaped = html.escape(text)
            escaped = escaped.replace("\n", "<br>")
            return f"<p style='margin:6px 0;'>{escaped}</p>"

        parts = [
            "<div style='font-family:\"Segoe UI\",sans-serif;font-size:13px;color:#202124;'>",
            f"<h3 style='margin-top:0;'>{icon} Verdict: {verdict.replace('_', ' ').title()}</h3>",
        ]

        if result.confidence is not None:
            parts.append(
                f"<p style='margin:4px 0;color:#5f6368;'>Confidence: {result.confidence:.2f}</p>"
            )

        if result.summary:
            parts.append("<h4 style='margin:12px 0 6px;'>Summary</h4>")
            parts.append(paragraph(result.summary))

        if result.reasoning:
            parts.append("<h4 style='margin:12px 0 6px;'>Reasoning</h4>")
            parts.append(paragraph(result.reasoning))

        if result.criteria:
            parts.append(
                "<h4 style='margin:12px 0 6px;'>Criteria Evaluation</h4>"
                "<table style='border-collapse:collapse;width:100%;'>"
                "<thead><tr style='background:#f1f3f4;'>"
                "<th style='text-align:left;padding:6px;border:1px solid #dfe1e5;'>Criterion</th>"
                "<th style='text-align:left;padding:6px;border:1px solid #dfe1e5;'>Result</th>"
                "</tr></thead><tbody>"
            )
            for criterion, value in result.criteria.items():
                parts.append(
                    "<tr>"
                    f"<td style='padding:6px;border:1px solid #dfe1e5;'>{html.escape(criterion.replace('_', ' ').title())}</td>"
                    f"<td style='padding:6px;border:1px solid #dfe1e5;'>{html.escape(str(value))}</td>"
                    "</tr>"
                )
            parts.append("</tbody></table>")

        if result.issues:
            parts.append("<h4 style='margin:12px 0 6px;'>Identified Issues</h4><ul>")
            for issue in result.issues:
                category = html.escape(issue.get("category", "General"))
                severity = html.escape(issue.get("severity", "UNKNOWN"))
                description = html.escape(issue.get("description", ""))
                parts.append(
                    f"<li><b>[{severity}] {category}</b><br><span>{description}</span></li>"
                )
            parts.append("</ul>")

        if result.recommendations:
            parts.append("<h4 style='margin:12px 0 6px;'>Recommendations</h4><ol>")
            for rec in result.recommendations:
                parts.append(f"<li>{html.escape(rec)}</li>")
            parts.append("</ol>")

        if result.messages:
            parts.append("<h4 style='margin:12px 0 6px;'>Notes</h4><ul>")
            for message in result.messages:
                parts.append(f"<li>{html.escape(message)}</li>")
            parts.append("</ul>")

        parts.append("</div>")
        return "".join(parts)

    def _format_spec_result_html(self, result):
        """Render the spec reference evaluation in rich HTML."""

        verdict = (result.verdict or "UNKNOWN").upper()
        icon = {
            "APPROVE": "✅",
            "APPROVE_WITH_WARNINGS": "⚠️",
            "REJECT": "❌",
        }.get(verdict, "ℹ️")

        def paragraph(text: Optional[str]) -> str:
            if not text:
                return ""
            escaped = html.escape(text)
            escaped = escaped.replace("\n", "<br>")
            return f"<p style='margin:6px 0;'>{escaped}</p>"

        parts = [
            "<div style='font-family:\"Segoe UI\",sans-serif;font-size:13px;color:#202124;'>",
            f"<h3 style='margin-top:0;'>{icon} Verdict: {verdict.replace('_', ' ').title()}</h3>",
        ]

        if result.confidence is not None:
            parts.append(
                f"<p style='margin:4px 0;color:#5f6368;'>Confidence: {result.confidence:.2f}</p>"
            )

        if result.will_fix_issue is not None:
            fix_text = "Yes" if result.will_fix_issue else "No"
            parts.append(
                f"<p style='margin:4px 0;color:#5f6368;'>Fix Effectiveness: <b>{fix_text}</b></p>"
            )

        if result.summary:
            parts.append("<h4 style='margin:12px 0 6px;'>Summary</h4>")
            parts.append(paragraph(result.summary))

        if result.reasoning:
            parts.append("<h4 style='margin:12px 0 6px;'>Reasoning</h4>")
            parts.append(paragraph(result.reasoning))

        if result.issues:
            parts.append("<h4 style='margin:12px 0 6px;'>Spec Findings</h4><ul>")
            for issue in result.issues:
                severity = html.escape(issue.get("severity", "UNKNOWN"))
                reference = html.escape(issue.get("spec_reference", "Unknown reference"))
                description = html.escape(issue.get("description", ""))
                parts.append(
                    f"<li><b>[{severity}] {reference}</b><br><span>{description}</span></li>"
                )
                action = issue.get("suggested_action")
                if action:
                    action_esc = html.escape(action)
                    parts.append(
                        f"<div style='margin:4px 0 8px 12px;color:#5f6368;'>Suggested action: {action_esc}</div>"
                    )
            parts.append("</ul>")

        if result.recommendations:
            parts.append("<h4 style='margin:12px 0 6px;'>Recommendations</h4><ol>")
            for rec in result.recommendations:
                parts.append(f"<li>{html.escape(rec)}</li>")
            parts.append("</ol>")

        if result.messages:
            parts.append("<h4 style='margin:12px 0 6px;'>Notes</h4><ul>")
            for message in result.messages:
                parts.append(f"<li>{html.escape(message)}</li>")
            parts.append("</ul>")

        parts.append("</div>")
        return "".join(parts)

    def _generate_unified_diff(self, original_code, patched_code, *, from_label='original', to_label='patched'):
        """Create a unified diff between original and patched code."""
        original_lines = (original_code or '').rstrip().splitlines()
        patched_lines = (patched_code or '').rstrip().splitlines()
        diff = difflib.unified_diff(
            original_lines,
            patched_lines,
            fromfile=from_label,
            tofile=to_label,
            lineterm=''
        )
        diff_text = "\n".join(diff)
        return diff_text or patched_code

    def _wrap_and_indent(self, text, indent="  ", width=110):
        """Wrap text for display in the output pane."""
        if not text:
            return []
        wrapper = textwrap.TextWrapper(width=width, initial_indent=indent, subsequent_indent=indent)
        lines = []
        for paragraph in str(text).splitlines():
            if not paragraph.strip():
                lines.append("")
            else:
                lines.append(wrapper.fill(paragraph.strip()))
        return lines



    def show_code_assistant(self):
        """Show the code assistant dialog - Patch Application Phase."""
        # Store dialog reference to prevent garbage collection
        self.code_assistant_dialog = QDialog(self)
        dialog = self.code_assistant_dialog
        dialog.setWindowTitle("Code Assistant - Apply Patches")
        
        # Get screen size and set dialog to fit within it
        screen = QApplication.primaryScreen()
        screen_geometry = screen.availableGeometry()
        screen_height = screen_geometry.height()
        screen_width = screen_geometry.width()
        
        # Set dialog size to 90% of screen height to ensure buttons are visible
        dialog_width = min(1300, int(screen_width * 0.9))
        dialog_height = min(900, int(screen_height * 0.9))
        
        dialog.setMinimumSize(1000, 600)
        dialog.resize(dialog_width, dialog_height)
        dialog.setSizeGripEnabled(True)
        
        # Connect dialog close event to cleanup
        dialog.finished.connect(self.cleanup_code_assistant_widgets)
        
        # Main layout with tight spacing
        layout = QVBoxLayout()
        layout.setSpacing(8)  # Reduce spacing between widgets
        layout.setContentsMargins(10, 10, 10, 10)  # Reduce margins
        
        # Bug History Selection - Compact layout
        history_group = QGroupBox("Bug Analysis Selection")
        history_layout = QHBoxLayout()
        history_layout.setSpacing(8)
        
        history_label = QLabel("Analysis:")
        self.bug_history_combo = QComboBox()
        self.bug_history_combo.setMinimumWidth(400)
        
        load_btn = QPushButton("Load Analysis")
        load_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                font-weight: bold;
                padding: 6px 16px;
                border: none;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
        """)
        load_btn.clicked.connect(self.load_selected_bug_analysis)
        
        refresh_btn = QPushButton("↻")
        refresh_btn.setMaximumWidth(35)
        refresh_btn.setToolTip("Refresh analysis list")
        refresh_btn.clicked.connect(self.load_bug_history)
        
        history_layout.addWidget(history_label)
        history_layout.addWidget(self.bug_history_combo, 1)
        history_layout.addWidget(load_btn)
        history_layout.addWidget(refresh_btn)
        
        history_group.setLayout(history_layout)
        history_group.setMaximumHeight(70)
        layout.addWidget(history_group)
        
        # Load bug history on dialog open
        self.load_bug_history()
        
        # Create a split container for Error Details and Patch Preview (side by side)
        top_split = QHBoxLayout()
        top_split.setSpacing(8)
        
        # Error Details Display (Left - 40%)
        error_group = QGroupBox("Error Details")
        error_layout = QVBoxLayout()
        error_layout.setContentsMargins(5, 5, 5, 5)
        
        self.code_error_display = QTextEdit()
        self.code_error_display.setReadOnly(True)
        self.code_error_display.setMinimumHeight(150)
        self.code_error_display.setMaximumHeight(200)
        self.code_error_display.setPlaceholderText("Error details will appear here after loading an analysis...")
        error_layout.addWidget(self.code_error_display)
        
        error_group.setLayout(error_layout)
        top_split.addWidget(error_group, 2)  # 2 parts
        
        # Patch Preview Section (Right - 60%)
        preview_group = QGroupBox("Selected Patches Preview (Editable)")
        preview_layout = QVBoxLayout()
        preview_layout.setContentsMargins(5, 5, 5, 5)
        
        # Create scroll area for preview
        preview_scroll = QScrollArea()
        preview_scroll.setWidgetResizable(True)
        preview_scroll.setMinimumHeight(150)
        preview_scroll.setMaximumHeight(200)
        preview_scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid #d1d1d1;
                background-color: #f5f5f5;
            }
        """)
        
        # Create widget to hold preview content
        self.preview_content_widget = QWidget()
        self.preview_content_layout = QVBoxLayout()
        self.preview_content_layout.setAlignment(Qt.AlignTop)
        self.preview_content_layout.setSpacing(5)
        self.preview_content_layout.setContentsMargins(8, 8, 8, 8)
        
        # Placeholder label
        self.preview_placeholder = QLabel("Selected patches will be displayed here...")
        self.preview_placeholder.setStyleSheet("color: #888; font-style: italic; padding: 10px;")
        self.preview_content_layout.addWidget(self.preview_placeholder)
        
        self.preview_content_widget.setLayout(self.preview_content_layout)
        preview_scroll.setWidget(self.preview_content_widget)
        preview_layout.addWidget(preview_scroll)
        
        # Store references for editable fields
        self.patch_value_editors = {}  # {patch_index: {'type': 'code'/'config', 'editor': QLineEdit}}
        
        preview_group.setLayout(preview_layout)
        top_split.addWidget(preview_group, 3)  # 3 parts
        
        layout.addLayout(top_split)
        
        # Patches selection area (side by side)
        patches_container = QHBoxLayout()
        patches_container.setSpacing(8)
        
        # Code Patches Section (Left side)
        code_patches_group = QGroupBox("Code Patches")
        code_patches_outer_layout = QVBoxLayout()
        code_patches_outer_layout.setContentsMargins(5, 5, 5, 5)
        code_patches_outer_layout.setSpacing(4)
        
        # Create scroll area for code patches
        code_patches_scroll = QScrollArea()
        code_patches_scroll.setWidgetResizable(True)
        code_patches_scroll.setMinimumHeight(120)
        code_patches_scroll.setMaximumHeight(180)
        code_patches_scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid #d1d1d1;
                background-color: white;
            }
        """)
        
        # Create widget to hold patches
        code_patches_widget = QWidget()
        code_patches_layout = QVBoxLayout()
        code_patches_layout.setAlignment(Qt.AlignTop)
        code_patches_layout.setSpacing(3)
        code_patches_layout.setContentsMargins(5, 5, 5, 5)
        
        # Placeholder for code patches checkboxes
        self.code_patches_checkboxes = []
        self.code_patches_layout = code_patches_layout
        
        # Initially show placeholder message
        code_patches_placeholder = QLabel("No code patches available.")
        code_patches_placeholder.setStyleSheet("color: #888; font-style: italic; padding: 5px;")
        code_patches_layout.addWidget(code_patches_placeholder)
        self.code_patches_placeholder = code_patches_placeholder
        
        code_patches_widget.setLayout(code_patches_layout)
        code_patches_scroll.setWidget(code_patches_widget)
        code_patches_outer_layout.addWidget(code_patches_scroll)
        code_patches_group.setLayout(code_patches_outer_layout)
        
        # Config Patches Section (Right side)  
        config_patches_group = QGroupBox("Config Patches")
        config_patches_outer_layout = QVBoxLayout()
        config_patches_outer_layout.setContentsMargins(5, 5, 5, 5)
        config_patches_outer_layout.setSpacing(4)
        
        # Create scroll area for config patches
        config_patches_scroll = QScrollArea()
        config_patches_scroll.setWidgetResizable(True)
        config_patches_scroll.setMinimumHeight(120)
        config_patches_scroll.setMaximumHeight(180)
        config_patches_scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid #d1d1d1;
                background-color: white;
            }
        """)
        
        # Create widget to hold patches
        config_patches_widget = QWidget()
        config_patches_layout = QVBoxLayout()
        config_patches_layout.setAlignment(Qt.AlignTop)
        config_patches_layout.setSpacing(3)
        config_patches_layout.setContentsMargins(5, 5, 5, 5)
        
        # Placeholder for config patches checkboxes
        self.config_patches_checkboxes = []
        self.config_patches_layout = config_patches_layout
        
        # Initially show placeholder message
        config_patches_placeholder = QLabel("No config patches available.")
        config_patches_placeholder.setStyleSheet("color: #888; font-style: italic; padding: 5px;")
        config_patches_layout.addWidget(config_patches_placeholder)
        self.config_patches_placeholder = config_patches_placeholder
        
        config_patches_widget.setLayout(config_patches_layout)
        config_patches_scroll.setWidget(config_patches_widget)
        config_patches_outer_layout.addWidget(config_patches_scroll)
        config_patches_group.setLayout(config_patches_outer_layout)
        
        # Add both sections to horizontal layout (side by side)
        patches_container.addWidget(code_patches_group)
        patches_container.addWidget(config_patches_group)
        
        layout.addLayout(patches_container, 0)  # Fixed height - no stretch
        
        # Patch control buttons - Compact
        controls_layout = QHBoxLayout()
        controls_layout.setSpacing(8)
        
        # Select All button
        select_all_btn = QPushButton("Select All")
        select_all_btn.setMaximumWidth(100)
        select_all_btn.clicked.connect(self.select_all_patches)
        select_all_btn.setToolTip("Select all available patches")
        select_all_btn.setStyleSheet("padding: 5px 10px;")
        
        # Unselect All button  
        unselect_all_btn = QPushButton("Unselect All")
        unselect_all_btn.setMaximumWidth(100)
        unselect_all_btn.clicked.connect(self.unselect_all_patches)
        unselect_all_btn.setToolTip("Unselect all patches")
        unselect_all_btn.setStyleSheet("padding: 5px 10px;")
        
        # Patch Fixes button
        patch_fixes_btn = QPushButton("Apply Patch Fixes")
        patch_fixes_btn.clicked.connect(self.apply_selected_patches)
        patch_fixes_btn.setToolTip("Apply the selected patches using unified patch applicator")
        patch_fixes_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                font-weight: bold;
                padding: 6px 20px;
                border: none;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
            QPushButton:pressed {
                background-color: #005A9E;
            }
        """)
        
        # Run Terminal Commands button
        terminal_commands_btn = QPushButton("Run Investigation Commands")
        terminal_commands_btn.clicked.connect(self.run_code_terminal_commands)
        terminal_commands_btn.setToolTip("Execute the generated terminal commands for troubleshooting")
        terminal_commands_btn.setStyleSheet("""
            QPushButton {
                background-color: #107C10;
                color: white;
                font-weight: bold;
                padding: 6px 20px;
                border: none;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #0F5F0F;
            }
            QPushButton:pressed {
                background-color: #0A4A0A;
            }
        """)
        
        controls_layout.addWidget(select_all_btn)
        controls_layout.addWidget(unselect_all_btn)
        controls_layout.addStretch()
        controls_layout.addWidget(patch_fixes_btn)
        controls_layout.addWidget(terminal_commands_btn)
        
        # Progress bar
        self.code_analysis_progress = QProgressBar()
        self.code_analysis_progress.hide()
        
        layout.addLayout(controls_layout)
        layout.addWidget(self.code_analysis_progress)
        
        dialog.setLayout(layout)
        dialog.exec_()

    def load_bug_history(self):
        """Load bug analysis history from files."""
        self.bug_history_combo.clear()
        self.bug_history_combo.addItem("Select a bug analysis...")
        
        # Use absolute path to ensure consistency
        workspace_root = os.path.dirname(os.path.abspath(__file__))
        history_dir = os.path.join(workspace_root, "backend", "resources", "bug_history")
        
        if not os.path.exists(history_dir):
            print(f"DEBUG: Bug history directory not found: {history_dir}")
            return
        
        # Get all history files and sort by timestamp (newest first)
        history_files = [f for f in os.listdir(history_dir) if f.endswith('.json')]
        history_files.sort(reverse=True)  # Sort by filename (timestamp) - newest first
        
        # Store mapping of display text to file path
        self.bug_history_map = {}
        
        for filename in history_files:
            file_path = os.path.join(history_dir, filename)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    error_message = data.get('error_message', 'Unknown error')
                    timestamp = data.get('timestamp', '')
                    
                    # Handle None error message
                    if error_message is None:
                        error_message = 'Unknown error'
                    
                    # Format display text: "Error message (truncated)"
                    display_text = error_message[:100] + "..." if len(error_message) > 100 else error_message
                    
                    # Add timestamp to display
                    if timestamp:
                        try:
                            dt = datetime.fromisoformat(timestamp)
                            time_str = dt.strftime("%Y-%m-%d %H:%M")
                            display_text = f"[{time_str}] {display_text}"
                        except:
                            pass
                    
                    # Debug: Show patch counts in this history file
                    results = data.get('results', {})
                    phase3_fixes = results.get('phase3_fixes', {})
                    fix_suggestion = phase3_fixes.get('fix_suggestion', {})
                    code_count = len(fix_suggestion.get('code_patches', []))
                    config_count = len(fix_suggestion.get('config_patches', []))
                    
                    # Add crash indicator if this is crash analysis
                    is_crash = results.get('crash_analysis', False)
                    if is_crash:
                        display_text = f"🔬 {display_text} [CRASH] [Code:{code_count}]"
                    else:
                        display_text += f" [Code:{code_count}, Config:{config_count}]"
                    
                    self.bug_history_combo.addItem(display_text)
                    self.bug_history_map[display_text] = file_path
            except Exception as e:
                print(f"Error loading history file {filename}: {e}")
                continue

    def load_selected_bug_analysis(self):
        """Load the selected bug analysis and display patches."""
        selected_text = self.bug_history_combo.currentText()
        
        if selected_text == "Select a bug analysis...":
            QMessageBox.warning(self, "No Selection", "Please select a bug analysis from the dropdown.")
            return
        
        if selected_text not in self.bug_history_map:
            QMessageBox.warning(self, "Error", "Could not find the selected analysis file.")
            return
        
        file_path = self.bug_history_map[selected_text]
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Store loaded analysis
            self.loaded_bug_analysis = data
            
            # Display error details
            error_message = data.get('error_message', 'Unknown error')
            log_file = data.get('log_file', 'N/A')
            timestamp = data.get('timestamp', 'N/A')
            
            error_display = f"Error Message:\n{error_message}\n\n"
            error_display += f"Log File: {log_file}\n"
            error_display += f"Timestamp: {timestamp}\n"
            
            self.code_error_display.setText(error_display)
            
            # Extract and display patches
            results = data.get('results', {})
            phase3_fixes = results.get('phase3_fixes', {})
            fix_suggestion = phase3_fixes.get('fix_suggestion', {})
            
            code_patches = fix_suggestion.get('code_patches', [])
            config_patches = fix_suggestion.get('config_patches', [])
            
            # Debug output
            print(f"\nDEBUG [Code Assistant Load]: Loading from history file")
            print(f"DEBUG: Code Patches: {len(code_patches)}")
            print(f"DEBUG: Config Patches: {len(config_patches)}")
            if code_patches:
                print(f"DEBUG: Code patch names: {[p.get('function_name', 'N/A') for p in code_patches]}")
            if config_patches:
                print(f"DEBUG: Config patch names: {[p.get('config_name', p.get('parameter_name', 'N/A')) for p in config_patches]}")
            
            # Clear existing patches
            self.clear_patch_checkboxes()
            
            # Display code patches
            if code_patches:
                # Remove placeholder if exists
                if hasattr(self, 'code_patches_placeholder'):
                    self.code_patches_placeholder.hide()
                
                for patch in code_patches:
                    function_name = patch.get('function_name', 'Unknown')
                    file_path = patch.get('file_path', 'Unknown')
                    file_name = os.path.basename(file_path)
                    checkbox_text = f"{function_name} ({file_name})"
                    print(f"DEBUG: Creating code patch checkbox: {checkbox_text}")
                    checkbox = QCheckBox(checkbox_text)
                    checkbox.setToolTip(f"File: {file_path}\nDescription: {patch.get('description', 'N/A')}")
                    checkbox.setChecked(True)
                    checkbox.stateChanged.connect(self.update_patch_preview)
                    self.code_patches_layout.addWidget(checkbox)
                    self.code_patches_checkboxes.append(checkbox)
                print(f"DEBUG: Total code patch checkboxes created: {len(self.code_patches_checkboxes)}")
            
            # Display config patches
            if config_patches:
                # Remove placeholder if exists
                if hasattr(self, 'config_patches_placeholder'):
                    self.config_patches_placeholder.hide()
                
                for patch in config_patches:
                    # Try both parameter_name and config_name for compatibility
                    param_name = patch.get('config_name', patch.get('parameter_name', 'Unknown'))
                    file_path = patch.get('file_path', 'Unknown')
                    file_name = os.path.basename(file_path)
                    checkbox_text = f"{param_name} ({file_name})"
                    print(f"DEBUG: Creating config patch checkbox: {checkbox_text}")
                    checkbox = QCheckBox(checkbox_text)
                    checkbox.setToolTip(f"File: {file_path}\nDescription: {patch.get('description', 'N/A')}")
                    checkbox.setChecked(True)
                    checkbox.stateChanged.connect(self.update_patch_preview)
                    self.config_patches_layout.addWidget(checkbox)
                    self.config_patches_checkboxes.append(checkbox)
                print(f"DEBUG: Total config patch checkboxes created: {len(self.config_patches_checkboxes)}")
            
            # Save fix_suggestions.json for apply_selected_patches to use
            self.save_fix_suggestions_from_loaded_analysis(data)
            
            # Update the preview to show all selected patches
            self.update_patch_preview()
            
            QMessageBox.information(self, "Analysis Loaded", 
                                   f"Loaded {len(code_patches)} code patches and {len(config_patches)} config patches.")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load analysis:\n{str(e)}")
            print(f"Error loading bug analysis: {e}")
            import traceback
            traceback.print_exc()

    def clear_patch_checkboxes(self):
        """Clear all patch checkboxes."""
        # Clear code patches
        for checkbox in self.code_patches_checkboxes:
            self.code_patches_layout.removeWidget(checkbox)
            checkbox.deleteLater()
        self.code_patches_checkboxes = []
        
        # Clear config patches
        for checkbox in self.config_patches_checkboxes:
            self.config_patches_layout.removeWidget(checkbox)
            checkbox.deleteLater()
        self.config_patches_checkboxes = []
    
    def update_patch_preview(self):
        """Update the patch preview display based on selected checkboxes with editable fields."""
        try:
            # Clear existing preview content
            while self.preview_content_layout.count():
                child = self.preview_content_layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()
            
            # Clear editor references
            self.patch_value_editors.clear()
            
            # Get loaded patch data
            if not hasattr(self, 'loaded_bug_analysis'):
                placeholder = QLabel("No patches loaded. Please load a bug analysis first.")
                placeholder.setStyleSheet("color: #888; font-style: italic; padding: 10px;")
                self.preview_content_layout.addWidget(placeholder)
                return
            
            results = self.loaded_bug_analysis.get('results', {})
            phase3_fixes = results.get('phase3_fixes', {})
            fix_suggestion = phase3_fixes.get('fix_suggestion', {})
            
            code_patches = fix_suggestion.get('code_patches', [])
            config_patches = fix_suggestion.get('config_patches', [])
            
            # Count selected patches
            selected_code_count = 0
            selected_config_count = 0
            
            if hasattr(self, 'code_patches_checkboxes'):
                selected_code_count = sum(1 for i, cb in enumerate(self.code_patches_checkboxes) 
                                         if cb.isChecked() and i < len(code_patches))
            
            if hasattr(self, 'config_patches_checkboxes'):
                selected_config_count = sum(1 for i, cb in enumerate(self.config_patches_checkboxes) 
                                           if cb.isChecked() and i < len(config_patches))
            
            # Check for selected code patches
            if selected_code_count > 0:
                # Code patches header
                header = QLabel("💻 CODE PATCHES:")
                header.setStyleSheet("font-weight: bold; font-size: 11pt; color: #0078D4; padding: 5px 0px;")
                self.preview_content_layout.addWidget(header)
                
                patch_num = 1
                for i, checkbox in enumerate(self.code_patches_checkboxes):
                    if checkbox.isChecked() and i < len(code_patches):
                        patch = code_patches[i]
                        function_name = patch.get('function_name', 'Unknown')
                        file_path = patch.get('file_path', 'Unknown')
                        description = patch.get('description', 'No description')
                        patch_type = patch.get('patch_type', 'modification')
                        line_info = patch.get('line_number') or patch.get('line_numbers', '')
                        
                        # Patch info
                        info_text = f"{patch_num}. Function: {function_name}\n"
                        info_text += f"   File: {file_path}\n"
                        if patch_type:
                            info_text += f"   Type: {patch_type}\n"
                        if line_info:
                            info_text += f"   Lines: {line_info}\n"
                        info_text += f"   Description: {description}"
                        
                        info_label = QLabel(info_text)
                        info_label.setStyleSheet("font-family: 'Consolas', monospace; font-size: 9pt; padding: 3px; background-color: white;")
                        info_label.setWordWrap(True)
                        self.preview_content_layout.addWidget(info_label)
                        
                        # Display the patch content
                        if 'original_code' in patch:
                            orig_text = "🔴 Original Code:\n"
                            for line in patch['original_code'].split('\n'):
                                orig_text += f"  - {line}\n"
                            
                            orig_label = QLabel(orig_text)
                            orig_label.setStyleSheet("font-family: 'Consolas', monospace; font-size: 8pt; padding: 3px; background-color: #ffe6e6; border-left: 3px solid #ff4444;")
                            orig_label.setWordWrap(True)
                            self.preview_content_layout.addWidget(orig_label)
                        
                        # Check for patched code (try multiple field names)
                        patched_code = patch.get('suggested_code') or patch.get('patched_code') or patch.get('new_code')
                        if patched_code:
                            # Editable patched code section
                            patched_code_label = QLabel("🟢 Patched Code (Editable):")
                            patched_code_label.setStyleSheet("font-family: 'Consolas', monospace; font-size: 8pt; font-weight: bold; color: #107C10; padding: 3px;")
                            self.preview_content_layout.addWidget(patched_code_label)
                            
                            # Create editable text area for patched code
                            patched_code_edit = QTextEdit()
                            patched_code_edit.setPlainText(patched_code)
                            patched_code_edit.setStyleSheet("""
                                QTextEdit {
                                    font-family: 'Consolas', monospace;
                                    font-size: 8pt;
                                    padding: 4px;
                                    background-color: #e6ffe6;
                                    border: 2px solid #44ff44;
                                    border-radius: 3px;
                                    min-height: 60px;
                                    max-height: 150px;
                                }
                                QTextEdit:focus {
                                    border: 2px solid #107C10;
                                    background-color: #ffffff;
                                }
                            """)
                            patched_code_edit.setPlaceholderText("Edit patched code here...")
                            patched_code_edit.setAcceptRichText(False)
                            self.preview_content_layout.addWidget(patched_code_edit)
                            
                            # Store the editor reference for code patches
                            self.patch_value_editors[f'code_{i}'] = {
                                'type': 'code',
                                'editor': patched_code_edit,
                                'patch_index': i,
                                'original_value': patched_code
                            }
                        
                        # Separator
                        separator = QLabel("─" * 50)
                        separator.setStyleSheet("color: #d1d1d1; font-size: 8pt;")
                        self.preview_content_layout.addWidget(separator)
                        
                        patch_num += 1
            
            # Check for selected config patches
            if selected_config_count > 0:
                # Config patches header
                header = QLabel("⚙️  CONFIG PATCHES:")
                header.setStyleSheet("font-weight: bold; font-size: 11pt; color: #107C10; padding: 5px 0px;")
                self.preview_content_layout.addWidget(header)
                
                patch_num = 1
                for i, checkbox in enumerate(self.config_patches_checkboxes):
                    if checkbox.isChecked() and i < len(config_patches):
                        patch = config_patches[i]
                        config_name = patch.get('config_name', patch.get('parameter_name', 'Unknown'))
                        file_path = patch.get('file_path', 'Unknown')
                        description = patch.get('description', 'No description')
                        line_info = patch.get('line_number') or patch.get('line_numbers', '')
                        
                        # Patch info
                        info_text = f"{patch_num}. Config: {config_name}\n"
                        info_text += f"   File: {file_path}\n"
                        if line_info:
                            info_text += f"   Line: {line_info}\n"
                        info_text += f"   Description: {description}"
                        
                        info_label = QLabel(info_text)
                        info_label.setStyleSheet("font-family: 'Consolas', monospace; font-size: 9pt; padding: 3px; background-color: white;")
                        info_label.setWordWrap(True)
                        self.preview_content_layout.addWidget(info_label)
                        
                        # Display config changes
                        if 'current_value' in patch:
                            curr_label = QLabel(f"🔴 Current Value: {patch['current_value']}")
                            curr_label.setStyleSheet("font-family: 'Consolas', monospace; font-size: 8pt; padding: 3px; background-color: #ffe6e6; border-left: 3px solid #ff4444;")
                            curr_label.setWordWrap(True)
                            self.preview_content_layout.addWidget(curr_label)
                        
                        # Check for new value (try multiple field names)
                        new_value = patch.get('suggested_value') or patch.get('new_value') or patch.get('recommended_value')
                        if new_value:
                            # Editable new value section
                            new_value_layout = QHBoxLayout()
                            new_value_label = QLabel("🟢 New Value:")
                            new_value_label.setStyleSheet("font-family: 'Consolas', monospace; font-size: 8pt; font-weight: bold; color: #107C10;")
                            
                            new_value_edit = QLineEdit(str(new_value))
                            new_value_edit.setStyleSheet("""
                                QLineEdit {
                                    font-family: 'Consolas', monospace;
                                    font-size: 8pt;
                                    padding: 4px;
                                    background-color: #e6ffe6;
                                    border: 2px solid #44ff44;
                                    border-radius: 3px;
                                }
                                QLineEdit:focus {
                                    border: 2px solid #107C10;
                                    background-color: #ffffff;
                                }
                            """)
                            new_value_edit.setPlaceholderText("Edit new value here...")
                            
                            new_value_layout.addWidget(new_value_label)
                            new_value_layout.addWidget(new_value_edit, 1)
                            
                            self.preview_content_layout.addLayout(new_value_layout)
                            
                            # Store the editor reference
                            self.patch_value_editors[f'config_{i}'] = {
                                'type': 'config',
                                'editor': new_value_edit,
                                'patch_index': i,
                                'original_value': str(new_value)
                            }
                        
                        if 'reasoning' in patch:
                            reason_label = QLabel(f"💡 Reasoning: {patch['reasoning']}")
                            reason_label.setStyleSheet("font-family: 'Consolas', monospace; font-size: 8pt; padding: 3px; font-style: italic; color: #605E5C;")
                            reason_label.setWordWrap(True)
                            self.preview_content_layout.addWidget(reason_label)
                        
                        # Separator
                        separator = QLabel("─" * 50)
                        separator.setStyleSheet("color: #d1d1d1; font-size: 8pt;")
                        self.preview_content_layout.addWidget(separator)
                        
                        patch_num += 1
            
            # Show placeholder if no patches selected
            if selected_code_count == 0 and selected_config_count == 0:
                placeholder = QLabel("No patches selected. Check the checkboxes to preview patches.")
                placeholder.setStyleSheet("color: #888; font-style: italic; padding: 10px;")
                self.preview_content_layout.addWidget(placeholder)
                
        except Exception as e:
            error_label = QLabel(f"Error loading patch preview: {str(e)}")
            error_label.setStyleSheet("color: red; padding: 10px;")
            self.preview_content_layout.addWidget(error_label)
            print(f"Error in update_patch_preview: {e}")
            import traceback
            traceback.print_exc()

    def run_code_terminal_commands(self):
        """Run investigation commands for the loaded analysis."""
        try:
            if not hasattr(self, 'loaded_bug_analysis'):
                QMessageBox.warning(self, "No Analysis", "Please load a bug analysis first.")
                return
            
            results = self.loaded_bug_analysis.get('results', {})
            
            # Get terminal commands from phase4_commands (this is where investigation COMMANDS are stored)
            terminal_commands = []
            phase4_commands = results.get('phase4_commands', {})
            if phase4_commands:
                terminal_commands_data = phase4_commands.get('terminal_commands', [])
                if terminal_commands_data:
                    # Already in command format
                    for cmd in terminal_commands_data:
                        if isinstance(cmd, dict):
                            terminal_commands.append(cmd)
                        else:
                            # Convert string to command format
                            terminal_commands.append({'command': str(cmd), 'explanation': 'Investigation command'})
            
            # If no terminal commands found, show a message (don't fallback to investigation_steps)
            # Investigation steps are text descriptions, not executable commands
            if not terminal_commands:
                QMessageBox.information(
                    self, 
                    "No Commands", 
                    "No investigation commands available for this analysis.\n\n"
                    "Note: Investigation commands are generated during Phase 4 of the RCA analysis. "
                    "If this is an older analysis, it may not have investigation commands."
                )
                return
            
            # Show confirmation dialog with commands to be executed
            commands_text = "\n".join([f"{i+1}. {cmd['command']}" for i, cmd in enumerate(terminal_commands)])
            confirm_msg = f"Execute {len(terminal_commands)} terminal commands?\n\nCommands to run:\n{commands_text}\n\nNote: Some commands may require administrator privileges."
            
            reply = QMessageBox.question(
                self,
                "Confirm Command Execution",
                confirm_msg,
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            
            if reply != QMessageBox.Yes:
                return
            
            # Create a dialog to show command execution results
            results_dialog = QDialog(self)
            results_dialog.setWindowTitle("Terminal Commands Execution")
            results_dialog.setMinimumSize(800, 600)
            
            layout = QVBoxLayout()
            
            # Header
            header = QLabel("Terminal Commands Execution Results")
            header.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 10px;")
            layout.addWidget(header)
            
            # Results text area
            results_text = QTextEdit()
            results_text.setReadOnly(True)
            results_text.setStyleSheet("""
                QTextEdit {
                    font-family: 'Consolas', 'Monaco', monospace;
                    font-size: 11px;
                    background-color: #1E1E1E;
                    color: #FFFFFF;
                    border: 1px solid #333;
                }
            """)
            layout.addWidget(results_text)
            
            # Close button
            close_btn = QPushButton("Close")
            close_btn.clicked.connect(results_dialog.close)
            layout.addWidget(close_btn)
            
            results_dialog.setLayout(layout)
            
            # Show dialog
            results_dialog.setWindowFlags(Qt.Window | Qt.WindowStaysOnTopHint)
            results_dialog.show()
            results_dialog.raise_()
            results_dialog.activateWindow()
            
            # Execute commands one by one
            import subprocess
            
            for i, cmd_info in enumerate(terminal_commands, 1):
                command = cmd_info['command']
                explanation = cmd_info.get('explanation', 'No explanation provided')
                
                results_text.append(f"\n{'='*60}")
                results_text.append(f"COMMAND {i}/{len(terminal_commands)}: {command}")
                results_text.append(f"EXPLANATION: {explanation}")
                results_text.append(f"{'='*60}")
                
                QApplication.processEvents()  # Keep UI responsive
                
                try:
                    # Execute the command
                    result = subprocess.run(
                        command,
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    
                    # Display output
                    if result.stdout:
                        results_text.append(f"\n✅ OUTPUT:\n{result.stdout}")
                    if result.stderr:
                        results_text.append(f"\n⚠️ STDERR:\n{result.stderr}")
                    
                    if result.returncode == 0:
                        results_text.append(f"\n✅ Command completed successfully (exit code: 0)")
                    else:
                        results_text.append(f"\n❌ Command failed with exit code: {result.returncode}")
                        
                except subprocess.TimeoutExpired:
                    results_text.append(f"\n⏱️ Command timed out after 30 seconds")
                except Exception as e:
                    results_text.append(f"\n❌ Error executing command: {str(e)}")
                
                QApplication.processEvents()
            
            results_text.append(f"\n{'='*60}")
            results_text.append(f"✅ All commands executed ({len(terminal_commands)} total)")
            results_text.append(f"{'='*60}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to run terminal commands:\n{str(e)}")
            print(f"Error in run_code_terminal_commands: {e}")
            import traceback
            traceback.print_exc()

    def save_fix_suggestions_from_loaded_analysis(self, analysis_data):
        """Save fix_suggestions.json from loaded analysis for patch application."""
        try:
            resources_dir = "backend/resources"
            os.makedirs(resources_dir, exist_ok=True)
            
            # Create timestamp for unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            fix_suggestions_file = os.path.join(resources_dir, f"fix_suggestions_{timestamp}.json")
            
            # Extract fix_suggestion from results
            results = analysis_data.get('results', {})
            phase3_fixes = results.get('phase3_fixes', {})
            fix_suggestion = phase3_fixes.get('fix_suggestion', {})
            
            # Add code directory to fix_suggestion for apply_selected_patches to use
            if 'code_dir' in analysis_data:
                fix_suggestion['code_dir'] = analysis_data['code_dir']
            
            # Save to file
            with open(fix_suggestions_file, 'w', encoding='utf-8') as f:
                json.dump(fix_suggestion, f, indent=2)
            
            print(f"Saved fix_suggestions.json to: {fix_suggestions_file}")
            
        except Exception as e:
            print(f"Error saving fix_suggestions.json: {e}")

    def select_code_directory(self):
        """Select the OAI gNodeB source code directory."""
        folder_path = QFileDialog.getExistingDirectory(self, "Select OAI gNodeB Source Directory")
        if folder_path:
            # Validate if it's an OAI gNodeB directory
            if self.validate_oai_directory(folder_path):
                self.code_dir_edit.setText(folder_path)
            else:
                QMessageBox.warning(
                    self,
                    "Invalid Directory",
                    "The selected directory does not appear to be an OAI gNodeB source directory. Please select the correct directory."
                )

    def validate_oai_directory(self, directory):
        """Validate if the directory is an OAI gNodeB source directory."""
        if not directory or not os.path.exists(directory):
            QMessageBox.warning(
                self,
                "Invalid Directory",
                "Please select a valid directory."
            )
            return False

        if not os.path.isdir(directory):
            QMessageBox.warning(
                self,
                "Invalid Path",
                "The selected path is not a directory."
            )
            return False

        if not os.listdir(directory):
            QMessageBox.warning(
                self,
                "Empty Directory",
                "The selected directory is empty."
            )
            return False

        # Check if it's the openairinterface5g-develop folder (or similar)
        directory_name = os.path.basename(directory).lower()
        
        # Accept openairinterface5g-develop folder or similar variants
        if 'openairinterface5g' in directory_name or 'openair' in directory_name:
            return True
        
        # For all other folders, show not supported message
        QMessageBox.information(
                self,
                "Directory Not Supported",
                f"The selected directory '{os.path.basename(directory)}' is not currently supported.\n\nPlease select the 'openairinterface5g-develop' directory or a similar OpenAirInterface source directory."
                )
        return False

    def start_code_analysis(self):
        """Start the code analysis process."""
        if not self.code_dir_edit.text():
            QMessageBox.warning(self, "Missing Directory", "Please select the OAI gNodeB source directory.")
            return

        # Check if RCA analysis has been run (patches are available)
        has_patches = (hasattr(self, 'code_patches_checkboxes') and len(self.code_patches_checkboxes) > 0) or \
                     (hasattr(self, 'config_patches_checkboxes') and len(self.config_patches_checkboxes) > 0)
        
        if not has_patches:
            QMessageBox.warning(self, "No Patches Available", "Please run RCA Analysis first to generate patches.")
            return

        try:
            # Show and reset progress bar
            self.code_analysis_progress.setValue(0)
            self.code_analysis_progress.show()
            
            # Create and start worker thread
            self.code_worker_thread = QThread()
            
            # Get selected patches information
            selected_patches = {
                'code_patches': [self.extract_name_from_checkbox_text(cb.text()) for cb in getattr(self, 'code_patches_checkboxes', []) if cb.isChecked()],
                'config_patches': [self.extract_name_from_checkbox_text(cb.text()) for cb in getattr(self, 'config_patches_checkboxes', []) if cb.isChecked()]
            }
            
            self.code_worker = CodeAnalysisWorker(
                self.code_dir_edit.text(),
                self.code_log_file_combo.currentText() if self.code_log_file_combo.currentText() else "No log file selected",
                selected_patches
            )
            self.code_worker.moveToThread(self.code_worker_thread)
            
            # Connect signals
            self.code_worker_thread.started.connect(self.code_worker.run)
            self.code_worker.finished.connect(self.code_worker_thread.quit)
            self.code_worker.finished.connect(self.code_worker.deleteLater)
            self.code_worker_thread.finished.connect(self.code_worker_thread.deleteLater)
            self.code_worker.progress.connect(self.code_analysis_progress.setValue)
            self.code_worker.result.connect(self.handle_code_analysis_result)
            
            # Disable UI elements
            self.code_dir_edit.setEnabled(False)
            self.module_combo.setEnabled(False)
            self.code_structure_check.setEnabled(False)
            self.function_analysis_check.setEnabled(False)
            self.dependency_check.setEnabled(False)
            self.api_check.setEnabled(False)
            self.memory_check.setEnabled(False)
            self.threading_check.setEnabled(False)
            
            # Start analysis
            self.code_worker_thread.start()

        except Exception as e:
            self.code_analysis_progress.hide()
            QMessageBox.critical(self, "Error", f"Failed to start analysis: {str(e)}")
            self.enable_code_analysis_ui()

    def handle_code_analysis_result(self, result):
        """Handle the code analysis results."""
        try:
            self.code_analysis_progress.hide()
            self.enable_code_analysis_ui()
            
            if not result:
                self.code_results_text.setText("Analysis failed or no results found.")
                return
                
            # Format the results
            formatted_result = "Code Analysis Results\n"
            formatted_result += "===================\n\n"
            
            try:
                result_dict = json.loads(result)
                for category, data in result_dict.items():
                    formatted_result += f"{category}:\n"
                    formatted_result += "-" * len(category) + "\n"
                    
                    if isinstance(data, dict):
                        for key, value in data.items():
                            formatted_result += f"\n{key}:\n"
                            if isinstance(value, list):
                                for item in value:
                                    formatted_result += f"  - {item}\n"
                            else:
                                formatted_result += f"  {value}\n"
                    elif isinstance(data, list):
                        for item in data:
                            formatted_result += f"- {item}\n"
                    else:
                        formatted_result += f"{data}\n"
                    formatted_result += "\n"
                
            except json.JSONDecodeError:
                formatted_result += result
                
            self.code_results_text.setText(formatted_result)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error displaying analysis results: {str(e)}")

    def enable_code_analysis_ui(self):
        """Re-enable the code analysis UI elements."""
        self.code_dir_edit.setEnabled(True)
        self.module_combo.setEnabled(True)
        self.code_structure_check.setEnabled(True)
        self.function_analysis_check.setEnabled(True)
        self.dependency_check.setEnabled(True)
        self.api_check.setEnabled(True)
        self.memory_check.setEnabled(True)
        self.threading_check.setEnabled(True)

    def save_code_analysis(self):
        """Save the code analysis results."""
        try:
            # Get file save location
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"oai_code_analysis_{timestamp}.txt"
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Analysis Results",
                os.path.join(os.path.expanduser("~"), default_filename),
                "Text Files (*.txt);;All Files (*.*)"
            )

            if not file_path:
                return

            # Save to file with metadata
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write("OAI gNodeB Code Analysis Results\n")
                f.write("=" * 40 + "\n\n")
                f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Directory: {self.code_dir_edit.text()}\n")
                f.write(f"Log File: {self.code_log_file_combo.currentText()}\n")
                f.write(f"Analysis Options:\n")
                for checkbox in self.analysis_checkboxes:
                    if checkbox.isChecked():
                        f.write(f"- {checkbox.text()}\n")
                f.write("\nResults:\n")
                f.write("=" * 40 + "\n\n")
                f.write(results_text)

            # Save to database
            analysis_data = {
                'action_type': 'CODE_ANALYSIS',
                'description': f"Code analysis for {self.code_log_file_combo.currentText()} log file",
                'details': json.dumps({
                    'timestamp': datetime.now().isoformat(),
                    'directory': self.code_dir_edit.text(),
                    'log_file': self.code_log_file_combo.currentText(),
                    'analysis_options': [cb.text() for cb in self.analysis_checkboxes if cb.isChecked()],
                    'results': results_text,
                    'file_path': file_path
                })
            }
            
            # Initialize database if needed
            initialize_database()
            
            # Save to database
            save_to_database('user_history', analysis_data)

            QMessageBox.information(
                self,
                "Success",
                f"Analysis results have been saved to:\n{file_path}\nand recorded in the history database."
            )

        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to save analysis results: {str(e)}"
            )

    def create_prompt_templates_page(self):
        """Create the prompt templates page with role-based categories and templates."""
        templates_widget = QWidget()
        templates_widget.setStyleSheet(f"background-color: {THEME['colors']['background']};")
        
        # Create a scroll area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet("""
            QScrollArea {
                border: none;
                background: transparent;
            }
            QScrollBar:vertical {
                border: none;
                background: #F0F0F0;
                width: 8px;
                border-radius: 4px;
            }
            QScrollBar::handle:vertical {
                background: #C0C0C0;
                border-radius: 4px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                border: none;
                background: none;
            }
        """)
        
        # Create a container widget for the scroll area
        scroll_content = QWidget()
        templates_layout = QVBoxLayout(scroll_content)
        templates_layout.setContentsMargins(24, 24, 24, 24)
        templates_layout.setSpacing(24)

        # Header section
        title = QLabel("Prompt Templates")
        title.setStyleSheet(THEME['typography']['title'])
        templates_layout.addWidget(title)

        subtitle = QLabel("Select a role and customize templates for your specific needs")
        subtitle.setStyleSheet(THEME['typography']['subtitle'])
        templates_layout.addWidget(subtitle)

        # Main content area
        content_frame = QFrame()
        content_frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 24px;
            }}
        """)
        content_layout = QVBoxLayout(content_frame)
        content_layout.setSpacing(16)

        # Role selection dropdown
        role_label = QLabel("Select Role")
        role_label.setStyleSheet("""
            QLabel {
                font-size: 16px;
                font-weight: 600;
                color: #323130;
                margin-top: 16px;
                margin-bottom: 12px;
                padding-left: 4px;
            }
        """)
        content_layout.addWidget(role_label)

        role_combo = QComboBox()
        role_combo.setStyleSheet("""
            QComboBox {
                background-color: #FFFFFF;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                padding: 8px 12px;
                min-width: 200px;
                font-size: 14px;
                margin-bottom: 16px;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
                padding-right: 8px;
            }
            QComboBox::down-arrow {
                image: url(down_arrow.png);
                width: 12px;
                height: 12px;
            }
            QComboBox:hover {
                border-color: #0078D4;
            }
            QComboBox:focus {
                border-color: #0078D4;
                border-width: 2px;
            }
        """)
        role_combo.addItems(["Select Role", "Tester", "Developer", "Analyst"])
        role_combo.setCurrentText("Select Role")
        content_layout.addWidget(role_combo)

        # Template categories section
        categories_label = QLabel("Template Categories")
        categories_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                font-weight: 600;
                color: #323130;
                margin-top: 16px;
                margin-bottom: 8px;
            }
        """)
        content_layout.addWidget(categories_label)

        # Template categories grid
        categories_grid = QGridLayout()
        categories_grid.setSpacing(8)
        categories_grid.setContentsMargins(0, 0, 0, 16)
        content_layout.addLayout(categories_grid)

        # Template content
        template_content_layout = QHBoxLayout()
        template_content_layout.setContentsMargins(0, 0, 0, 0)
        template_content_layout.setSpacing(16)
        
        # Left side - Template content
        template_container = QWidget()
        template_container_layout = QVBoxLayout(template_container)
        template_container_layout.setContentsMargins(0, 0, 0, 0)
        
        template_label = QLabel("Template Content")
        template_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                font-weight: 600;
                color: #323130;
                margin-bottom: 8px;
            }
        """)
        template_container_layout.addWidget(template_label)

        template_edit = QTextEdit()
        template_edit.setStyleSheet("""
            QTextEdit {
                background-color: #FFFFFF;
                border: 1px solid #E1E5EA;
                border-radius: 4px;
                font-family: 'Consolas', monospace;
                font-size: 13px;
                line-height: 1.6;
                padding: 16px;
                color: #323130;
            }
            QTextEdit:focus {
                border-color: #0078D4;
            }
        """)
        template_edit.setMinimumHeight(400)
        template_edit.setReadOnly(True)
        template_container_layout.addWidget(template_edit)
        
        # Store the template editor as an instance variable
        self.prompt_text_edit = template_edit
        
        template_content_layout.addWidget(template_container, stretch=5)

        # Right side - Action buttons
        action_container = QWidget()
        action_container.setFixedWidth(120)
        action_layout = QVBoxLayout(action_container)
        action_layout.setAlignment(Qt.AlignTop)
        action_layout.setSpacing(8)
        action_layout.setContentsMargins(0, 0, 0, 0)
        
        action_button_style = f"""
            QPushButton {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 4px;
                color: {THEME['colors']['text']['primary']};
                font-size: 13px;
                padding: 8px 16px;
                width: 100px;
                text-align: center;
            }}
            QPushButton:hover {{
                background-color: {THEME['colors']['hover']};
                border-color: {THEME['colors']['primary']};
            }}
        """
        
        modify_button = QPushButton("✏️ Modify")
        modify_button.setStyleSheet(action_button_style)
        action_layout.addWidget(modify_button)

        save_button = QPushButton("💾 Save")
        save_button.setStyleSheet(action_button_style + f"""
            QPushButton {{
                background-color: {THEME['colors']['primary']};
                color: white;
                border: none;
            }}
            QPushButton:hover {{
                background-color: {THEME['colors']['secondary']};
            }}
        """)
        action_layout.addWidget(save_button)
        action_layout.addStretch()
        
        # Store buttons as instance variables
        self.modify_button = modify_button
        self.save_button = save_button
        
        template_content_layout.addWidget(action_container)
        content_layout.addLayout(template_content_layout)

        templates_layout.addWidget(content_frame)
        templates_layout.addStretch()

        # Set up scroll area
        scroll_area.setWidget(scroll_content)
        
        # Main layout for the templates widget
        main_layout = QVBoxLayout(templates_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.addWidget(scroll_area)

        # Store template editor and grid for access in other methods
        self.prompt_editor = template_edit
        self.categories_grid = categories_grid

        # Connect signals
        role_combo.currentTextChanged.connect(lambda text: self.show_role_templates(text.lower()))
        modify_button.clicked.connect(self.handle_modify_button)
        save_button.clicked.connect(self.handle_save_button)

        return templates_widget

    def show_role_templates(self, role):
        """Show templates for the selected role."""
        # Clear existing items in the grid
        while self.categories_grid.count():
            item = self.categories_grid.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Skip if "select role" is chosen
        if role == "select role":
            return

        # Add new template categories for the selected role
        templates = self.role_templates.get(role, [])
        row = 0
        col = 0
        max_cols = 3  # Display 3 categories per row

        for title, desc, icon in templates:
            # Create a category card
            category_card = QFrame()
            category_card.setStyleSheet(f"""
                QFrame {{
                    background-color: {THEME['colors']['surface']};
                    border: 1px solid {THEME['colors']['border']};
                    border-radius: 4px;
                    padding: 12px;
                    min-width: 200px;
                }}
                QFrame:hover {{
                    background-color: {THEME['colors']['hover']};
                    border-color: {THEME['colors']['primary']};
                }}
            """)
            
            # Create vertical layout for the card
            card_layout = QVBoxLayout(category_card)
            card_layout.setSpacing(8)
            card_layout.setContentsMargins(12, 12, 12, 12)

            # Add icon and title in horizontal layout
            header_layout = QHBoxLayout()
            icon_label = QLabel(icon)
            icon_label.setStyleSheet("font-size: 20px;")
            header_layout.addWidget(icon_label)

            title_label = QLabel(title)
            title_label.setStyleSheet(f"""
                QLabel {{
                    color: {THEME['colors']['text']['primary']};
                    font-size: 14px;
                    font-weight: 600;
                }}
            """)
            header_layout.addWidget(title_label)
            header_layout.addStretch()
            card_layout.addLayout(header_layout)

            # Add description
            desc_label = QLabel(desc)
            desc_label.setStyleSheet(f"""
                QLabel {{
                    color: {THEME['colors']['text']['secondary']};
                    font-size: 12px;
                }}
            """)
            desc_label.setWordWrap(True)
            card_layout.addWidget(desc_label)

            # Add select button
            select_btn = QPushButton("Select Template")
            select_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {THEME['colors']['surface']};
                    border: 1px solid {THEME['colors']['border']};
                    border-radius: 4px;
                    color: {THEME['colors']['text']['primary']};
                    font-size: 12px;
                    padding: 6px 12px;
                    margin-top: 8px;
                }}
                QPushButton:hover {{
                    background-color: {THEME['colors']['hover']};
                    border-color: {THEME['colors']['primary']};
                }}
            """)
            select_btn.clicked.connect(lambda checked, t=title: self.handle_prompt(t))
            card_layout.addWidget(select_btn)

            # Add the card to the grid
            self.categories_grid.addWidget(category_card, row, col)
            
            # Update row and column position
            col += 1
            if col >= max_cols:
                col = 0
                row += 1

    def create_user_history_page(self):
        """Create the user history page with enterprise styling and consistent filter controls."""
        history_widget = QWidget()
        history_widget.setStyleSheet(f"background-color: {THEME['colors']['background']};")
        history_layout = QVBoxLayout()
        history_layout.setContentsMargins(24, 24, 24, 24)
        history_layout.setSpacing(24)

        # Title section with professional content
        title = QLabel("Activity History")
        title.setStyleSheet(THEME['typography']['title'])
        history_layout.addWidget(title)

        subtitle = QLabel("View and manage your framework activities and analysis results")
        subtitle.setStyleSheet(THEME['typography']['subtitle'])
        history_layout.addWidget(subtitle)

        # Filter section with consistent styling
        filter_frame = QFrame()
        filter_frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 16px;
            }}
        """)
        filter_layout = QHBoxLayout(filter_frame)
        filter_layout.setSpacing(16)

        # Date filter with consistent size
        date_container = QWidget()
        date_layout = QVBoxLayout(date_container)
        date_layout.setSpacing(4)
        date_layout.setContentsMargins(0, 0, 0, 0)

        date_label = QLabel("Time Period")
        date_label.setStyleSheet(THEME['typography']['label'] + "font-weight: 600;")
        date_layout.addWidget(date_label)

        self.date_filter = QComboBox()
        self.date_filter.setFixedWidth(180)  # Set consistent width
        self.date_filter.setFixedHeight(32)  # Set consistent height
        self.date_filter.addItems(["All Dates", "Today", "Last 7 Days", "Last 30 Days"])
        self.date_filter.setStyleSheet(INPUT_STYLE + """
            QComboBox {
                font-size: 13px;
                padding: 6px 12px;
            }
        """)
        self.date_filter.currentTextChanged.connect(self.refresh_history)
        date_layout.addWidget(self.date_filter)
        filter_layout.addWidget(date_container)

        # Action type filter with consistent size
        action_container = QWidget()
        action_layout = QVBoxLayout(action_container)
        action_layout.setSpacing(4)
        action_layout.setContentsMargins(0, 0, 0, 0)

        action_label = QLabel("Activity Type")
        action_label.setStyleSheet(THEME['typography']['label'] + "font-weight: 600;")
        action_layout.addWidget(action_label)

        self.action_filter = QComboBox()
        self.action_filter.setFixedWidth(180)  # Set consistent width
        self.action_filter.setFixedHeight(32)  # Set consistent height
        self.action_filter.addItems([
            "All Activities",
            "Test Script Generation",
            "Test Case Creation",
            "Bug Analysis",
            "Code Analysis",
            "Prompt Templates"
        ])
        self.action_filter.setStyleSheet(INPUT_STYLE + """
            QComboBox {
                font-size: 13px;
                padding: 6px 12px;
            }
        """)
        self.action_filter.currentTextChanged.connect(self.refresh_history)
        action_layout.addWidget(self.action_filter)
        filter_layout.addWidget(action_container)

        filter_layout.addStretch()

        # Action buttons with consistent styling
        button_container = QWidget()
        button_layout = QHBoxLayout(button_container)
        button_layout.setSpacing(8)
        button_layout.setContentsMargins(0, 0, 0, 0)

        clear_btn = QPushButton("Clear")
        clear_btn.setFixedHeight(32)  # Set consistent height
        clear_btn.setStyleSheet(BUTTON_STYLE + """
            QPushButton {
                min-width: 80px;
                font-size: 13px;
                text-align: center;
            }
        """)
        clear_btn.clicked.connect(self.clear_filters)
        button_layout.addWidget(clear_btn)

        refresh_btn = QPushButton("Refresh")
        refresh_btn.setFixedHeight(32)  # Set consistent height
        refresh_btn.setStyleSheet(BUTTON_STYLE + """
            QPushButton {
                min-width: 80px;
                font-size: 13px;
                text-align: center;
                background-color: #0078D4;
                color: white;
                border: none;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
            QPushButton:pressed {
                background-color: #005A9E;
            }
        """)
        refresh_btn.clicked.connect(self.refresh_history)
        button_layout.addWidget(refresh_btn)

        filter_layout.addWidget(button_container)
        history_layout.addWidget(filter_frame)

        # Main content area
        content_frame = QFrame()
        content_frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 24px;
            }}
        """)
        self.history_content_layout = QVBoxLayout(content_frame)
        self.history_content_layout.setSpacing(16)

        # Scroll area with custom styling
        scroll = QScrollArea()
        scroll.setWidget(content_frame)
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background: transparent;
            }
            QScrollBar:vertical {
                border: none;
                background: #F0F0F0;
                width: 8px;
                border-radius: 4px;
            }
            QScrollBar::handle:vertical {
                background: #C0C0C0;
                border-radius: 4px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                border: none;
                background: none;
            }
        """)
        
        history_layout.addWidget(scroll)
        history_widget.setLayout(history_layout)

        # Initial population of history entries
        self.refresh_history()

        return history_widget

    def clear_filters(self):
        """Reset all filters to their default values."""
        self.date_filter.setCurrentText("All Dates")
        self.action_filter.setCurrentText("All Activities")
        self.refresh_history()

    def refresh_history(self):
        """Refresh the history entries with current filter settings."""
        # Clear existing entries
        for i in reversed(range(self.history_content_layout.count())): 
            widget = self.history_content_layout.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()

        # Get filter values
        date_filter = self.date_filter.currentText()
        action_filter = self.action_filter.currentText()

        # Build filters dictionary
        filters = {}
        if date_filter != "All Dates":
            today = datetime.now().date()
            if date_filter == "Today":
                filters['date'] = today.isoformat()
            elif date_filter == "Last 7 Days":
                filters['date_range'] = (today - timedelta(days=7), today)
            elif date_filter == "Last 30 Days":
                filters['date_range'] = (today - timedelta(days=30), today)

        if action_filter != "All Activities":
            filters['action_type'] = action_filter

        # Get filtered entries
        history_entries = fetch_from_database('user_history', filters)

        if history_entries:
            for entry in history_entries:
                # Create a card for each history entry
                entry_card = QFrame()
                entry_card.setStyleSheet("""
                    QFrame {
                        background-color: white;
                        border: 1px solid #E1E5EA;
                        border-radius: 6px;
                        padding: 16px;
                        margin-bottom: 8px;
                    }
                    QFrame:hover {
                        border-color: #0078D4;
                        background-color: #FAFAFA;
            }
        """)
                card_layout = QVBoxLayout(entry_card)
                card_layout.setSpacing(12)
                card_layout.setContentsMargins(16, 16, 16, 16)

                # Header with timestamp, action type, and buttons
                header_layout = QHBoxLayout()
                
                # Action type with icon
                action_type = entry[2]  # action_type column
                icon = self.get_action_icon(action_type)
                action_label = QLabel(f"{icon} {action_type}")
                action_label.setStyleSheet("""
                    QLabel {
                        font-size: 14px;
                        font-weight: 600;
                        color: #323130;
                    }
                """)
                header_layout.addWidget(action_label)
                
                header_layout.addStretch()
                
                # Timestamp
                timestamp = datetime.strptime(entry[1], "%Y-%m-%d %H:%M:%S")
                time_label = QLabel(timestamp.strftime("%B %d, %Y %I:%M %p"))
                time_label.setStyleSheet("""
                    QLabel {
                        font-size: 12px;
                        color: #605E5C;
                    }
                """)
                header_layout.addWidget(time_label)

                # Action buttons
                button_layout = QHBoxLayout()
                button_layout.setSpacing(8)

                # Copy button
                copy_btn = QPushButton("📋")
                copy_btn.setStyleSheet("""
                    QPushButton {
                        background: transparent;
                        border: none;
                        padding: 4px 8px;
                        font-size: 14px;
                        border-radius: 4px;
                    }
                    QPushButton:hover {
                        background-color: #E3F2FD;
                    }
                """)
                copy_btn.clicked.connect(lambda checked, text=entry[4]: self.copy_to_clipboard(text))
                button_layout.addWidget(copy_btn)

                # Delete button
                delete_btn = QPushButton("🗑️")
                delete_btn.setStyleSheet("""
                    QPushButton {
                        background: transparent;
                        border: none;
                        padding: 4px 8px;
                        font-size: 14px;
                        border-radius: 4px;
                    }
                    QPushButton:hover {
                        background-color: #FEE2E2;
                    }
                """)
                delete_btn.clicked.connect(lambda checked, entry_id=entry[0]: self.delete_history_entry(entry_id))
                button_layout.addWidget(delete_btn)

                header_layout.addLayout(button_layout)
                card_layout.addLayout(header_layout)

                # Description
                if entry[3]:  # description column
                    desc_label = QLabel(entry[3])
                    desc_label.setStyleSheet("""
                        QLabel {
                            font-size: 13px;
                            color: #323130;
                            line-height: 1.4;
                        }
                    """)
                    desc_label.setWordWrap(True)
                    card_layout.addWidget(desc_label)

                # Details
                if entry[4]:  # details column
                    try:
                        details = json.loads(entry[4])
                        details_frame = QFrame()
                        details_frame.setStyleSheet("""
                            QFrame {
                                background-color: #F8F9FA;
                                border: 1px solid #E1E5EA;
                                border-radius: 4px;
                                padding: 12px;
                            }
                        """)
                        details_layout = QVBoxLayout(details_frame)
                        details_layout.setSpacing(8)

                        # Add metadata
                        if isinstance(details, dict):
                            metadata_layout = QGridLayout()
                            row = 0
                            for key, value in details.items():
                                if key != 'results':  # Skip results as they'll be shown separately
                                    if isinstance(value, dict):
                                        value = json.dumps(value, indent=2)
                                    elif isinstance(value, list):
                                        value = ', '.join(str(v) for v in value)
                                    else:
                                        value = str(value)
                                    
                                    key_label = QLabel(f"{key.replace('_', ' ').title()}:")
                                    key_label.setStyleSheet("font-weight: 600; color: #323130;")
                                    value_label = QLabel(value)
                                    value_label.setWordWrap(True)
                                    value_label.setStyleSheet("color: #323130;")
                                    
                                    metadata_layout.addWidget(key_label, row, 0)
                                    metadata_layout.addWidget(value_label, row, 1)
                                    row += 1
                            
                            details_layout.addLayout(metadata_layout)

                            # Add results if present
                            if 'results' in details:
                                results_label = QLabel("Results:")
                                results_label.setStyleSheet("font-weight: 600; color: #323130; margin-top: 8px;")
                                details_layout.addWidget(results_label)
                                
                                results_text = QTextEdit()
                                results_text.setPlainText(details['results'])
                                results_text.setReadOnly(True)
                                results_text.setStyleSheet("""
                                    QTextEdit {
                                        background-color: white;
                                        border: 1px solid #E1E5EA;
                                        border-radius: 4px;
                                        padding: 8px;
                                        font-family: 'Consolas', monospace;
                                        font-size: 13px;
                                    }
                                """)
                                results_text.setMaximumHeight(200)
                                details_layout.addWidget(results_text)

                        card_layout.addWidget(details_frame)
                    except json.JSONDecodeError:
                        # If details is not valid JSON, display as plain text
                        details_text = QTextEdit()
                        details_text.setPlainText(entry[4])
                        details_text.setReadOnly(True)
                        details_text.setStyleSheet("""
                            QTextEdit {
                                background-color: #F8F9FA;
                                border: 1px solid #E1E5EA;
                                border-radius: 4px;
                                padding: 8px;
                                font-family: 'Consolas', monospace;
                                font-size: 13px;
                            }
                        """)
                        details_text.setMaximumHeight(200)
                        card_layout.addWidget(details_text)

                self.history_content_layout.addWidget(entry_card)
        else:
            # Show a message when no history entries are found
            no_data_label = QLabel("No history entries found")
            no_data_label.setStyleSheet("""
                QLabel {
                    font-size: 14px;
                    color: #605E5C;
                    padding: 24px;
                }
            """)
            no_data_label.setAlignment(Qt.AlignCenter)
            self.history_content_layout.addWidget(no_data_label)

    def copy_to_clipboard(self, text):
        """Copy the given text to clipboard."""
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(self, "Success", "Content copied to clipboard!")

    def delete_history_entry(self, entry_id):
        """Delete a history entry from the database."""
        reply = QMessageBox.question(
            self, 
            "Confirm Delete",
            "Are you sure you want to delete this history entry?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            conn = sqlite3.connect('agentic_ran.db')
            cursor = conn.cursor()
            cursor.execute('DELETE FROM user_history WHERE id = ?', (entry_id,))
            conn.commit()
            conn.close()
            self.refresh_history()

    def get_action_icon(self, action_type):
        """Return an appropriate icon for the action type."""
        icons = {
            'SAVE_PROMPT': '💾',
            'GENERATE_TEST': '🔧',
            'ANALYZE_BUG': '🔍',
            'CODE_ANALYSIS': '📊',
            'TEST_SCRIPT': '📝',
            'TEST_CASE': '✅',
            'BUG_ANALYSIS': '🐛',
            'DEFAULT': '📝'
        }
        return icons.get(action_type, icons['DEFAULT'])

    def handle_prompt(self, prompt_key):
        """Handle prompt selection and update the UI."""
        try:
            if prompt_key == "Custom":
                self.prompt_text_edit.clear()
                self.prompt_text_edit.setReadOnly(False)
                self.prompt_text_edit.setPlaceholderText("Enter your custom prompt here...")
                # Hide variables container for Custom
                if hasattr(self, 'variables_container'):
                    self.variables_container.setVisible(False)
                # Only enable/disable buttons if they exist (Test Script Generator context)
                if hasattr(self, 'modify_btn'):
                    self.modify_btn.setEnabled(False)
                if hasattr(self, 'save_template_btn'):
                    self.save_template_btn.setEnabled(True)
                # Store the latest custom prompt for use in evaluation (will be updated on save or modify)
                # This will be updated again when the user actually enters or modifies the custom prompt
                self.latest_custom_prompt = self.prompt_text_edit.toPlainText()  #chandu_update_eval
            elif prompt_key in self.prompts:
                self.prompt_text_edit.setReadOnly(True)
                self.prompt_text_edit.setPlaceholderText("")
                
                # Handle dictionary structure for Test Case template
                if prompt_key == "Test Case" and isinstance(self.prompts[prompt_key], dict):
                    # Merge system and user prompts
                    system_prompt = self.prompts[prompt_key].get("System Prompt", "")
                    user_prompt = self.prompts[prompt_key].get("User Prompt", "")
                    merged_prompt = f"{system_prompt}\n\n{user_prompt}"
                    self.prompt_text_edit.setText(merged_prompt)
                    # Show test case variables container
                    if hasattr(self, 'test_case_variables_container'):
                        self.test_case_variables_container.setVisible(True)
                    # Hide test script variables container
                    if hasattr(self, 'test_script_variables_container'):
                        self.test_script_variables_container.setVisible(False)
                elif prompt_key == "Test Script":
                    # Handle Test Script template
                    self.prompt_text_edit.setText(self.prompts[prompt_key])
                    # Show test script variables container
                    if hasattr(self, 'test_script_variables_container'):
                        self.test_script_variables_container.setVisible(True)
                    # Hide test case variables container
                    if hasattr(self, 'test_case_variables_container'):
                        self.test_case_variables_container.setVisible(False)
                else:
                    # Handle regular string prompts
                    self.prompt_text_edit.setText(self.prompts[prompt_key])
                    # Hide both variables containers for other templates
                    if hasattr(self, 'test_case_variables_container'):
                        self.test_case_variables_container.setVisible(False)
                    if hasattr(self, 'test_script_variables_container'):
                        self.test_script_variables_container.setVisible(False)
                
                # Only enable buttons if they exist (Test Script Generator context)
                if hasattr(self, 'modify_btn'):
                    self.modify_btn.setEnabled(True)
                if hasattr(self, 'save_template_btn'):
                    self.save_template_btn.setEnabled(False)
            else:
                self.prompt_text_edit.clear()
                self.prompt_text_edit.setReadOnly(True)
                self.prompt_text_edit.setPlaceholderText("Select a template or choose 'Custom' to enter your own prompt...")
                # Hide variables container for default case
                if hasattr(self, 'variables_container'):
                    self.variables_container.setVisible(False)
                # Only disable buttons if they exist (Test Script Generator context)
                if hasattr(self, 'modify_btn'):
                    self.modify_btn.setEnabled(False)
                if hasattr(self, 'save_template_btn'):
                    self.save_template_btn.setEnabled(False)
            
            self.current_prompt_key = prompt_key
            
        except Exception as e:
            self.show_error(f"Error handling prompt selection: {str(e)}")

    def handle_modify_button(self):
        """Handle Modify button click."""
        try:
            # Store current content before modification
            if not hasattr(self, 'original_content'):
                self.original_content = self.prompt_text_edit.toPlainText()
                
            # Enable text editing
            self.prompt_text_edit.setReadOnly(False)
            
            # Enable Save button if it exists
            if hasattr(self, 'save_button'):
                self.save_button.setEnabled(True)
            
            # Disable Modify button if it exists
            if hasattr(self, 'modify_button'):
                self.modify_button.setEnabled(False)
            
            # Set focus to the text edit
            self.prompt_text_edit.setFocus()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error in modify operation: {str(e)}")

    def handle_save_button(self):
        """Handle saving a new or modified template."""
        try:
            current_content = self.prompt_text_edit.toPlainText().strip()
            if not current_content:
                self.show_warning("Cannot save empty template.")
                return

            if self.current_prompt_key == "Custom":
                # Save as new template
                name, ok = QInputDialog.getText(self, "Save Template", "Enter template name:")
                if ok and name:
                    if name in self.prompts:
                        self.show_warning(f"Template '{name}' already exists. Choose a different name.")
                        return
                    self.prompts[name] = current_content
                    self.prompt_combo.addItem(name)
                    self.prompt_combo.setCurrentText(name)
                    self.show_info(f"Template '{name}' saved successfully.")
            else:
                # Update existing template
                reply = QMessageBox.question(self, "Update Template",
                    f"Do you want to update the '{self.current_prompt_key}' template?",
                    QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
                
                if reply == QMessageBox.Yes:
                    self.prompts[self.current_prompt_key] = current_content
                    self.show_info(f"Template '{self.current_prompt_key}' updated successfully.")

        except Exception as e:
            self.show_error(f"Error saving template: {str(e)}")

    def identify_root_cause(self, error_chain):
        """Identify the root cause in an error chain."""
        try:
            if not error_chain:
                return None
                
            # Start with the first error in the chain
            root = error_chain[0]
            
            # Look for more fundamental errors
            for error in error_chain[1:]:
                if self.is_more_fundamental(error, root):
                    root = error
            
            return {
                'error': root,
                'confidence': self.calculate_root_cause_confidence(root, error_chain),
                'related_errors': [e for e in error_chain if e != root and 
                                 self.are_errors_related(e, root)]
            }
            
        except Exception as e:
            logging.error(f"Root cause identification failed: {str(e)}")
            return None

    def is_more_fundamental(self, error1, error2):
        """Determine if error1 is more fundamental than error2."""
        try:
            # Check if error1 occurs in error2's stack trace
            if ('stack_trace' in error2 and 
                error1.get('location', '') in [frame.get('location', '') for frame in error2['stack_trace']]):
                return True
            
            # Check temporal relationship
            if ('timestamp' in error1 and 'timestamp' in error2 and 
                error1['timestamp'] < error2['timestamp']):
                return True
            
            # Check error categories
            fundamental_categories = {'memory', 'system', 'initialization'}
            if (error1.get('category', '') in fundamental_categories and 
                error2.get('category', '') not in fundamental_categories):
                return True
            
            return False
            
        except Exception as e:
            logging.error(f"Fundamental error check failed: {str(e)}")
            return False

    def assess_chain_impact(self, error_chain):
        """Assess the impact of an error chain."""
        try:
            return {
                'severity': self.calculate_chain_severity(error_chain),
                'scope': self.determine_chain_scope(error_chain),
                'affected_components': self.identify_affected_components(error_chain),
                'resolution_priority': self.calculate_resolution_priority(error_chain)
            }
        except Exception as e:
            logging.error(f"Chain impact assessment failed: {str(e)}")
            return None

    def analyze_file(self, file_path):
        """Analyze a log file for errors using configured patterns."""
        try:
            results = {
                'file': file_path,
                'errors': [],
                'warnings': [],
                'performance_issues': [],
                'timing_issues': [],
                'patterns': None  # Will store pattern mining results
            }

            # ... existing error detection code ...

            # Perform pattern mining if enabled
            if self.analysis_options.get('pattern_mining', False):
                results['patterns'] = self.mine_patterns(results['errors'])

            return results

        except Exception as e:
            logging.error(f"File analysis failed: {str(e)}")
            return {'file': file_path, 'error': str(e)}

    def generate_report(self, results):
        """Generate a comprehensive analysis report."""
        try:
            report = []
            report.append("Bug Analysis Report\n")
            report.append("=================\n\n")

            # ... existing report generation code ...

            # Add pattern mining results if available
            for result in results:
                if result.get('patterns'):
                    report.append("\nPattern Analysis:\n")
                    report.append("-----------------\n")
                    
                    # Temporal Patterns
                    if result['patterns'].get('temporal_patterns'):
                        report.append("\nTemporal Patterns:\n")
                        for pattern in result['patterns']['temporal_patterns']:
                            report.append(f"  • Sequence of {pattern['count']} related errors within {pattern['window']}s\n")
                    
                    # Error Chains
                    if result['patterns'].get('error_chains'):
                        report.append("\nError Chains:\n")
                        for chain in result['patterns']['error_chains']:
                            root = chain['root_cause']
                            report.append(f"  • Chain starting with: {root['error'].get('message', '')}\n")
                            report.append(f"    Confidence: {root['confidence']:.2f}\n")
                    
                    # Frequency Patterns
                    if result['patterns'].get('frequency_patterns'):
                        report.append("\nFrequent Patterns:\n")
                        for error, count in result['patterns']['frequency_patterns']['frequent_errors'].items():
                            report.append(f"  • {error}: occurred {count} times\n")
                    
                    # Location Patterns
                    if result['patterns'].get('location_patterns'):
                        report.append("\nHot Spots:\n")
                        for location, info in result['patterns']['location_patterns'].items():
                            report.append(f"  • {location}: {info['frequency']} errors\n")
                            report.append(f"    Types: {', '.join(info['error_types'])}\n")

            return "".join(report)

        except Exception as e:
            return f"Report generation failed: {str(e)}"

    def create_explore_page(self):
        """Create the explore page with modern card-based layout."""
        explore_widget = QWidget()
        explore_widget.setStyleSheet(f"background-color: {THEME['colors']['background']};")
        explore_layout = QVBoxLayout()
        explore_layout.setContentsMargins(24, 24, 24, 24)
        explore_layout.setSpacing(24)

        # Title section
        title = QLabel("Explore Test Scripts")
        title.setStyleSheet(THEME['typography']['title'])
        explore_layout.addWidget(title)

        subtitle = QLabel("Browse and analyze existing test scripts")
        subtitle.setStyleSheet(THEME['typography']['subtitle'])
        explore_layout.addWidget(subtitle)

        # Main content area with cards
        content_frame = QFrame()
        content_frame.setStyleSheet(f"""
            QFrame {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 8px;
                padding: 24px;
            }}
        """)
        content_layout = QVBoxLayout(content_frame)
        content_layout.setSpacing(16)

        # Test Scripts Grid
        scripts_grid = QGridLayout()
        scripts_grid.setSpacing(16)

        # Load test scripts from the database
        try:
            self.cursor.execute("SELECT * FROM test_scripts ORDER BY created_at DESC")
            scripts = self.cursor.fetchall()

            for i, script in enumerate(scripts):
                card = QFrame()
                card.setStyleSheet(f"""
                    QFrame {{
                        background-color: {THEME['colors']['surface']};
                        border: 1px solid {THEME['colors']['border']};
                        border-radius: 8px;
                        padding: 16px;
                        cursor: pointer;
                    }}
                    QFrame:hover {{
                        border-color: {THEME['colors']['primary']};
                        background-color: {THEME['colors']['hover']};
                    }}
                """)
                card_layout = QVBoxLayout(card)
                card_layout.setSpacing(8)

                # Script name and type
                name_label = QLabel(script[1])  # name
                name_label.setStyleSheet(THEME['typography']['body'] + "font-weight: 600;")
                card_layout.addWidget(name_label)

                type_label = QLabel(script[2])  # type
                type_label.setStyleSheet(THEME['typography']['caption'])
                card_layout.addWidget(type_label)

                # Creation date
                date_label = QLabel(f"Created: {script[4]}")  # created_at
                date_label.setStyleSheet(THEME['typography']['caption'])
                card_layout.addWidget(date_label)

                scripts_grid.addWidget(card, i // 3, i % 3)

        except Exception as e:
            error_label = QLabel(f"Error loading test scripts: {str(e)}")
            error_label.setStyleSheet("color: red;")
            content_layout.addWidget(error_label)

        content_layout.addLayout(scripts_grid)
        explore_layout.addWidget(content_frame)
        explore_widget.setLayout(explore_layout)
        return explore_widget

    def button_style(self):
        """Return the standard button style."""
        return f"""
            QPushButton {{
                background-color: {THEME['colors']['surface']};
                border: 1px solid {THEME['colors']['border']};
                border-radius: 4px;
                padding: 8px 16px;
                font-size: 13px;
                color: {THEME['colors']['text']};
            }}
            QPushButton:hover {{
                background-color: {THEME['colors']['hover']};
                border-color: {THEME['colors']['primary']};
            }}
            QPushButton:pressed {{
                background-color: {THEME['colors']['primary']};
                color: white;
            }}
        """

    def active_button_style(self):
        """Return the style for active/primary buttons."""
        return f"""
            QPushButton {{
                background-color: {THEME['colors']['primary']};
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-size: 13px;
                color: white;
            }}
            QPushButton:hover {{
                background-color: {THEME['colors']['secondary']};
            }}
            QPushButton:pressed {{
                background-color: {THEME['colors']['primary']};
            }}
        """

    def combo_style(self):
        """Return the style for combo boxes."""
        return """
            QComboBox {
                background-color: white;
                border: 1px solid #E0E0E0;
                border-radius: 4px;
                padding: 6px 8px;
                font-size: 12px;
                min-width: 120px;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                width: 12px;
                height: 12px;
            }
            QComboBox:hover {
                border-color: #0078D4;
            }
            QComboBox:focus {
                border-color: #0078D4;
                border-width: 2px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #E0E0E0;
                selection-background-color: #E3F2FD;
                selection-color: #1976D2;
                background: white;
            }
        """

    def compact_combo_style(self):
        """Return the style for compact combo boxes."""
        return """
            QComboBox {
                background-color: white;
                border: 1px solid #E0E0E0;
                border-radius: 4px;
                padding: 4px 6px;
                font-size: 11px;
                min-width: 120px;
                max-width: 140px;
            }
            QComboBox::drop-down {
                border: none;
                width: 16px;
            }
            QComboBox::down-arrow {
                width: 10px;
                height: 10px;
            }
            QComboBox:hover {
                border-color: #0078D4;
            }
            QComboBox:focus {
                border-color: #0078D4;
                border-width: 2px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #E0E0E0;
                selection-background-color: #E3F2FD;
                selection-color: #1976D2;
                background: white;
                max-height: 150px;
            }
        """

    def update_prompt_variables(self):
        """Update the prompt with selected variable values."""
        try:
            if not hasattr(self, 'prompt_text_edit'):
                return
                
            # Determine which template is currently selected
            current_prompt_key = None
            if hasattr(self, 'prompt_combo'):
                current_prompt_key = self.prompt_combo.currentText()
            
            # Get the original prompt template based on current selection
            if current_prompt_key == "Test Case" and hasattr(self, 'prompts') and 'Test Case' in self.prompts:
                test_case_prompt = self.prompts['Test Case']
                if isinstance(test_case_prompt, dict):
                    # Merge system and user prompts
                    system_prompt = test_case_prompt.get('System Prompt', '')
                    user_prompt = test_case_prompt.get('User Prompt', '')
                    original_prompt = f"{system_prompt}\n\n{user_prompt}"
                else:
                    original_prompt = test_case_prompt
            elif current_prompt_key == "Test Script" and hasattr(self, 'prompts') and 'Test Script' in self.prompts:
                original_prompt = self.prompts['Test Script']
            else:
                return
            
            # Create a copy to work with
            updated_prompt = original_prompt
            
            # Replace all variables with their current selections
            if hasattr(self, 'domain_combo') and not self.domain_combo.currentText().startswith("Select"):
                updated_prompt = updated_prompt.replace('{DOMAIN}', self.domain_combo.currentText())
            if hasattr(self, 'system_type_combo') and not self.system_type_combo.currentText().startswith("Select"):
                updated_prompt = updated_prompt.replace('{SYSTEM_TYPE}', self.system_type_combo.currentText())
            if hasattr(self, 'primary_feature_combo') and not self.primary_feature_combo.currentText().startswith("Select"):
                updated_prompt = updated_prompt.replace('{PRIMARY_FEATURE}', self.primary_feature_combo.currentText())
            if hasattr(self, 'connection_method_combo') and not self.connection_method_combo.currentText().startswith("Select"):
                updated_prompt = updated_prompt.replace('{CONNECTION_METHOD}', self.connection_method_combo.currentText())
            if hasattr(self, 'login_credentials_combo') and not self.login_credentials_combo.currentText().startswith("Select"):
                updated_prompt = updated_prompt.replace('{LOGIN_CREDENTIALS}', self.login_credentials_combo.currentText())
            if hasattr(self, 'access_mode_combo') and not self.access_mode_combo.currentText().startswith("Select"):
                updated_prompt = updated_prompt.replace('{ACCESS_MODE}', self.access_mode_combo.currentText())
            if hasattr(self, 'language_combo') and not self.language_combo.currentText().startswith("Select"):
                language_text = self.language_combo.currentText()
                # Remove "(Coming Soon)" text if present
                if " (Coming Soon)" in language_text:
                    language_text = language_text.replace(" (Coming Soon)", "")
                updated_prompt = updated_prompt.replace('{LANGUAGE}', language_text)
            
            # For Test Script, also replace the reference code placeholder
            if current_prompt_key == "Test Script" and hasattr(self, 'ue_attach_utils'):
                updated_prompt = updated_prompt.replace('{self.ue_attach_utils}', self.ue_attach_utils)
            
            # Update the prompt text
            self.prompt_text_edit.setPlainText(updated_prompt)
                
        except Exception as e:
            print(f"Error updating prompt variables: {str(e)}")

    def upload_reference_code(self):
        """Handle file upload for reference code."""
        try:
            from PyQt5.QtWidgets import QFileDialog
            
            # Get current language to determine file extensions
            language = self.language_combo.currentText() if hasattr(self, 'language_combo') else "Python"
            
            # Set file filter based on language
            if language == "Python":
                file_filter = "Python files (*.py *.txt);;All files (*.*)"
            elif language == "C":
                file_filter = "C files (*.c);;All files (*.*)"
            elif language == "C++":
                file_filter = "C++ files (*.cpp);;All files (*.*)"
            else:
                file_filter = "All files (*.*)"
            
            # Open file dialog
            file_path, _ = QFileDialog.getOpenFileName(
                self, 
                "Select Reference Code File", 
                "", 
                file_filter
            )
            
            if file_path:
                # Read the file content
                with open(file_path, 'r', encoding='utf-8') as file:
                    file_content = file.read()
                
                # Update the ue_attach_utils with the uploaded content
                self.ue_attach_utils = file_content
                
                # Update the file name label
                import os
                file_name = os.path.basename(file_path)
                self.uploaded_file_label.setText(f"Uploaded: {file_name}")
                self.uploaded_file_label.setStyleSheet("color: #28a745; font-weight: bold;")
                
                # Update the prompt if it's currently displayed
                self.update_prompt_variables()
                
        except Exception as e:
            print(f"Error uploading file: {str(e)}")
            self.uploaded_file_label.setText("Upload failed")
            self.uploaded_file_label.setStyleSheet("color: #dc3545; font-weight: bold;")

    def handle_language_selection(self):
        """Handle language selection with coming soon message for non-Python languages."""
        try:
            selected_language = self.language_combo.currentText()
            
            if "Coming Soon" in selected_language:
                # Show coming soon message
                from PyQt5.QtWidgets import QMessageBox
                msg = QMessageBox()
                msg.setWindowTitle("Coming Soon")
                msg.setText(f"{selected_language.split(' (')[0]} support is coming soon!")
                msg.setInformativeText("Currently only Python is supported for test script generation.")
                msg.setIcon(QMessageBox.Information)
                msg.exec_()
                
                # Reset to Python
                self.language_combo.setCurrentText("Python")
                return
            
            # Update prompt variables if Python is selected
            if selected_language == "Python":
                self.update_prompt_variables()
                
        except Exception as e:
            print(f"Error handling language selection: {str(e)}")

    def get_current_prompt(self):
        """Get the current prompt from either template or custom input."""
        try:
            if not hasattr(self, 'prompt_text_edit'):
                return None
                
            prompt_text = self.prompt_text_edit.toPlainText().strip()
            if not prompt_text:
                return None
                
            return prompt_text
        except Exception as e:
            print(f"Error getting current prompt: {str(e)}")
            return None

    def save_response(self):
        """Save the generated test script to a file."""
        try:
            if not hasattr(self, 'response_text') or not self.response_text.toPlainText():
                self.show_warning("No test script to save.")
                return

            response = self.response_text.toPlainText()
            if self.testcases_list == "Test Case":
                self.testcases_name = response
            # Get output directory
            output_dir = self.get_output_directory()
            if not output_dir:
                output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
                os.makedirs(output_dir, exist_ok=True)

            # Save the file
            success, result = self.save_response_to_file(response, output_dir)
            if success:
                self.show_info(f"Test script saved to: {result}")
            else:
                self.show_error(f"Failed to save test script: {result}")

        except Exception as e:
            self.show_error(f"Error saving test script: {str(e)}")

    def show_info(self, message):
        """Show an information message to the user."""
        QMessageBox.information(self, "Information", message)

    def show_warning(self, message):
        """Show a warning message to the user."""
        QMessageBox.warning(self, "Warning", message)

    def show_error(self, message):
        """Show an error message to the user."""
        QMessageBox.critical(self, "Error", message)

    def show_api_key_dialog(self, preset_key=None):
        """Show a dialog for entering OpenAI API key."""
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("OpenAI API Key Required")
            dialog.setWindowModality(Qt.ApplicationModal)
            dialog.setFixedSize(400, 150)
            
            layout = QVBoxLayout()
            layout.setContentsMargins(20, 20, 20, 20)
            layout.setSpacing(10)

            # Add explanation label
            label = QLabel("Please enter your OpenAI API key.\nYou can find this at https://platform.openai.com/api-keys")
            label.setWordWrap(True)
            layout.addWidget(label)

            # Add input field
            key_input = QLineEdit()
            key_input.setEchoMode(QLineEdit.Password)
            key_input.setPlaceholderText("Enter your API key here")
            if preset_key:
                key_input.setText(preset_key)
            layout.addWidget(key_input)

            # Add buttons
            button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
            button_box.accepted.connect(dialog.accept)
            button_box.rejected.connect(dialog.reject)
            layout.addWidget(button_box)

            dialog.setLayout(layout)

            # Set focus to input field
            key_input.setFocus()

            # Handle dialog result
            result = dialog.exec_()
            if result == QDialog.Accepted:
                api_key = key_input.text().strip()
                if api_key:
                    # Validate API key format
                    if api_key.startswith('sk-') and len(api_key) > 20:
                        # Store the API key in settings
                        self.settings.setValue('openai_api_key', api_key)
                        os.environ['OPENAI_API_KEY'] = api_key
                        return api_key
                    else:
                        QMessageBox.warning(self, "Invalid API Key", 
                                          "The API key format appears to be invalid.\nIt should start with 'sk-' and be longer than 20 characters.")
                        return self.show_api_key_dialog()  # Recursive call for retry
                else:
                    QMessageBox.warning(self, "Missing API Key", 
                                      "Please enter an API key.")
                    return self.show_api_key_dialog()  # Recursive call for retry
            return None
            
        except Exception as e:
            print(f"Error in API key dialog: {str(e)}")
            QMessageBox.critical(self, "Error", 
                               f"An error occurred while showing the API key dialog: {str(e)}")
            return None

    def create_eval_page(self):
        eval_widget = QWidget()
        layout = QVBoxLayout(eval_widget)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(16)

        title = QLabel("Response Evaluation")
        title.setStyleSheet(THEME['typography']['title'])
        layout.addWidget(title)

        file_layout = QGridLayout()
        # Remove prompt UI
        # self.eval_prompt_path = QLineEdit()
        self.eval_dataset_path = QLineEdit()
        self.eval_response_path = QLineEdit()
        for le in [self.eval_dataset_path, self.eval_response_path]:
            le.setReadOnly(True)
        dataset_btn = QPushButton("Load Dataset")
        response_btn = QPushButton("Load Response")
        use_latest_dataset_btn = QPushButton("Use Latest Dataset")
        use_latest_response_btn = QPushButton("Use Latest Response")
        dataset_btn.clicked.connect(lambda: self.select_eval_file(self.eval_dataset_path))
        response_btn.clicked.connect(lambda: self.select_eval_file(self.eval_response_path))
        use_latest_dataset_btn.clicked.connect(lambda: self.eval_dataset_path.setText("[LATEST_DATASET]"))
        use_latest_response_btn.clicked.connect(lambda: self.eval_response_path.setText("[LATEST_RESPONSE]"))
        file_layout.addWidget(QLabel("Dataset:"), 0, 0)
        file_layout.addWidget(self.eval_dataset_path, 0, 1)
        file_layout.addWidget(dataset_btn, 0, 2)
        file_layout.addWidget(use_latest_dataset_btn, 0, 3)
        file_layout.addWidget(QLabel("Response:"), 1, 0)
        file_layout.addWidget(self.eval_response_path, 1, 1)
        file_layout.addWidget(response_btn, 1, 2)
        file_layout.addWidget(use_latest_response_btn, 1, 3)
        layout.addLayout(file_layout)

        eval_btn = QPushButton("Response Evaluation")
        eval_btn.setStyleSheet(self.active_button_style())
        eval_btn.clicked.connect(self.run_response_evaluation)
        layout.addWidget(eval_btn)

        self.eval_results_box = QTextEdit()
        self.eval_results_box.setReadOnly(True)
        self.eval_results_box.setMinimumHeight(40)
        self.eval_results_box.setMaximumHeight(80)
        layout.addWidget(self.eval_results_box)

        self.ies_details_box = QTextEdit()
        self.ies_details_box.setReadOnly(True)
        self.ies_details_box.setMinimumHeight(120)
        self.ies_details_box.setMaximumHeight(250)
        self.ies_details_box.setStyleSheet("font-family: Consolas, monospace; font-size: 12px; background: #fafafa; border: 1px solid #E1E5EA; border-radius: 4px; padding: 8px;")
        layout.addWidget(self.ies_details_box)

        return eval_widget

    def run_response_evaluation(self):
        import os
        import tempfile
        import sys
        import importlib.util
        dataset = self.eval_dataset_path.text()
        response = self.eval_response_path.text()
        if not (dataset and response):
            self.show_warning("Please select both dataset and response files.")
            return
        temp_dataset_file = None
        temp_response_file = None
        # Handle latest dataset
        if dataset == "[LATEST_DATASET]" and self.latest_dataset_content:
            tmp = tempfile.NamedTemporaryFile(delete=False, mode='w', encoding='utf-8', suffix='.txt')
            tmp.write(self.latest_dataset_content)
            tmp.close()
            dataset_file = tmp.name
            temp_dataset_file = dataset_file
        else:
            dataset_file = dataset
        # Handle latest response
        if response == "[LATEST_RESPONSE]" and self.latest_generated_response:
            tmp = tempfile.NamedTemporaryFile(delete=False, mode='w', encoding='utf-8', suffix='.txt')
            tmp.write(self.latest_generated_response)
            tmp.close()
            response_file = tmp.name
            temp_response_file = response_file
        else:
            response_file = response
        try:
            # Dynamically import cross_encoder_validation
            spec = importlib.util.spec_from_file_location("cross_encoder_validation", os.path.join(os.getcwd(), "cross_encoder_validation.py"))
            cev = importlib.util.module_from_spec(spec)
            sys.modules["cross_encoder_validation"] = cev
            spec.loader.exec_module(cev)
            # Run evaluation (assume dataset is DOCX, response is JSON)
            result = cev.evaluate_coverage(dataset_file, response_file)
            percent = result.get("percentage", 0)
            covered = result.get("covered", 0)
            total = result.get("total", 0)
            details = result.get("details", "")
            self.eval_results_box.setHtml(f"<span style='font-size:18px; font-weight:bold; color:#0078D4;'>Coverage: {covered} / {total} ({percent}%)</span>")
            self.ies_details_box.setPlainText(details)
        except Exception as e:
            self.show_error(f"Evaluation failed: {str(e)}")
        finally:
            if temp_dataset_file:
                try:
                    os.remove(temp_dataset_file)
                except Exception:
                    pass
            if temp_response_file:
                try:
                    os.remove(temp_response_file)
                except Exception:
                    pass

    def select_eval_file(self, line_edit):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select File", "", "All Files (*.*)")
        if file_path:
            line_edit.setText(file_path)

    def generate_with_new_prompt(self):
        """Generate a new response using the new prompt, previous response, and dataset."""
        try:
            text_content = self.text_content
            if not text_content:
                self.show_warning("Please load a dataset first.")
                return
            new_prompt = self.new_prompt_edit.toPlainText().strip()
            if not new_prompt:
                self.show_warning("Please enter a new prompt.")
                return
            previous_script = self.previous_response or ""
            dataset_summary = text_content[:1000]  # Use first 1000 chars as summary (or improve as needed)
            # Build iterative refinement prompt
            iterative_prompt = f"""
Using the dataset provided and the current user prompt, generate or update a detailed, well-structured test script. If a previous script is provided, use it as a foundation to refine, expand, or optimize the script based on the new prompt.\n\nYour goals:\n1. Improve the test script with each new prompt.\n2. Retain relevant parts of the previous script.\n3. Integrate new instructions or validations as requested.\n4. Ensure the final script is complete, readable, and executable.\n\nInputs:\n- Current Prompt: {new_prompt}\n- Previous Test Script (previous response): {previous_script}\n- Dataset Summary: {dataset_summary}\n\nRespond only with the refined script. Format it using appropriate comments and structure.\n 
Note:Make Sure you update the script or testcase which is placed in {new_prompt} and don't modify any other thing which is not placed in {new_prompt}
"""
            # Show progress dialog
            self.progress_dialog = QProgressDialog("Generating refined response...", "Cancel", 0, 100, self)
            self.progress_dialog.setWindowModality(Qt.WindowModal)
            self.progress_dialog.setAutoClose(True)
            self.progress_dialog.setAutoReset(True)
            self.progress_dialog.show()
            # Use a thread for generation (reuse ExtractionThread for consistency)
            self.extraction_thread = ExtractionThread(iterative_prompt, self)
            self.progress_connection = self.extraction_thread.progress_signal.connect(self.update_progress)
            self.finished_connection = self.extraction_thread.finished.connect(self.handle_response_complete)
            self.error_connection = self.extraction_thread.error_signal.connect(self.show_error)
            self.extraction_thread.text_content = text_content
            self.extraction_thread.selected_prompt = iterative_prompt
            self.extraction_thread.start()
        except Exception as e:
            self.show_error(f"Error generating with new prompt: {str(e)}")
            if hasattr(self, 'progress_dialog'):
                self.progress_dialog.close()

    def update_previous_response(self, response):
        self.previous_response = response

    def update_latest_custom_prompt(self):
        # Stub method to avoid AttributeError. Implement logic if needed.
        pass







    def select_log_directory(self, line_edit):
        """Open a dialog to select a log directory and set it in the given QLineEdit."""
        dir_path = QFileDialog.getExistingDirectory(self, "Select Log Directory", "")
        if dir_path:
            line_edit.setText(dir_path)

    def cleanup_code_assistant_widgets(self):
        """Clean up code assistant widget references when dialog is closed."""
        if hasattr(self, 'start_code_rca_btn'):
            self.start_code_rca_btn = None
        if hasattr(self, 'code_log_file_combo'):
            self.code_log_file_combo = None
        if hasattr(self, 'code_rca_results_text'):
            self.code_rca_results_text = None
        if hasattr(self, 'fix_present_btn'):
            self.fix_present_btn = None
        if hasattr(self, 'existing_fix_result'):
            self.existing_fix_result = None

    def check_for_existing_fix(self):
        """Check if a fix already exists in embeddings for the selected log file"""
        try:
            if not hasattr(self, 'fix_present_btn'):
                return
            
            # Don't trigger if change was programmatic (folder upload)
            if hasattr(self, 'log_combo_programmatic_change') and self.log_combo_programmatic_change:
                return
            
            # Hide button by default
            self.fix_present_btn.setVisible(False)
            self.existing_fix_result = None
            
            # Get selected log file
            log_file = self.code_log_file_combo.currentText()
            log_dir = self.code_log_dir_edit.text()
            
            # Skip if placeholder is selected or no file/dir
            if not log_file or not log_dir or log_file == "Click to select the log file":
                return
            
            log_file_path = os.path.join(log_dir, log_file)
            
            if not os.path.exists(log_file_path):
                return
            
            # Create progress dialog
            progress = QProgressDialog("Checking for existing fixes...", "Cancel", 0, 100, self)
            progress.setWindowTitle("Analyzing Log File")
            progress.setWindowModality(Qt.WindowModal)
            progress.setMinimumDuration(0)  # Show immediately
            progress.setValue(0)
            QApplication.processEvents()
            
            print(f"\n🔍 Checking for existing fixes for: {log_file}")
            
            # Step 1: Extract error from log file
            progress.setLabelText("Extracting error message from log file...")
            progress.setValue(20)
            QApplication.processEvents()
            
            if progress.wasCanceled():
                return
            
            from Error_fixing_pipelin.parse_log_context import LogContextParser
            
            # Get the source code directory and extract folder name
            code_dir = self.code_dir_edit.text() if hasattr(self, 'code_dir_edit') else None
            if code_dir:
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            else:
                openair_codebase_file_name = "openairinterface5g-develop"  # Default fallback
            
            log_parser = LogContextParser(openair_codebase_file_name=openair_codebase_file_name)
            error_message = log_parser.extract_error_message(log_file_path)
            
            if not error_message:
                print("   No clear error extracted from log")
                progress.close()
                return
            
            print(f"   Error extracted: {error_message[:100]}...")
            
            # Step 2: Search embeddings
            progress.setLabelText("Searching 30,000+ commits for similar fixes...")
            progress.setValue(50)
            QApplication.processEvents()
            
            if progress.wasCanceled():
                return
            
            from Error_fixing_pipelin.smart_commit_selector import CommitSearcher, SmartSelector
            
            # Get the source code directory and extract folder name
            code_dir = self.code_dir_edit.text() if hasattr(self, 'code_dir_edit') else None
            if code_dir:
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            else:
                openair_codebase_file_name = "openairinterface5g-develop"  # Default fallback
            
            searcher = CommitSearcher(embeddings_dir='Error_fixing_pipelin/resources/embeddings', 
                                     validate_commits=False,  # Disabled to show commits across different repo versions
                                     openair_codebase_file_name=openair_codebase_file_name)
            search_results = searcher.search(error_message, top_k=10)
            
            if not search_results:
                print("   No similar commits found")
                progress.close()
                return
            
            # Step 3: Intelligent selection
            progress.setLabelText("Analyzing matches with smart selector...")
            progress.setValue(75)
            QApplication.processEvents()
            
            if progress.wasCanceled():
                return
            
            selector = SmartSelector(use_llm=False)
            selection_result = selector.select_best_fix(error_message, search_results)
            
            # Step 4: Check results
            progress.setLabelText("Finalizing results...")
            progress.setValue(90)
            QApplication.processEvents()
            
            # Check if a good fix was found
            if selection_result['status'] in ['auto_selected', 'suggested', 'llm_verified']:
                if selection_result['commit'] and selection_result['commit']['is_rca_commit']:
                    # RCA fix found!
                    confidence = selection_result['confidence']
                    score = selection_result['commit'].get('boosted_score', 
                                                           selection_result['commit']['similarity'])
                    
                    print(f"   ✅ Found existing RCA fix! Confidence: {confidence}, Score: {score:.2%}")
                    
                    # Store the result
                    self.existing_fix_result = selection_result
                    
                    # Show the button
                    self.fix_present_btn.setVisible(True)
                    self.fix_present_btn.setText(f"✅ Fix Already Present ({confidence}) - Click to View")
                else:
                    print("   Found similar commits but no RCA fixes")
            else:
                print("   No suitable fix found")
            
            # Complete
            progress.setValue(100)
            progress.close()
                
        except Exception as e:
            if 'progress' in locals():
                progress.close()
            print(f"⚠️ Error checking for existing fix: {e}")
            import traceback
            traceback.print_exc()
    
    def clear_previous_patches(self):
        """Clear previous RCA analysis patches from UI and backend."""
        try:
            # Remove UI checkboxes completely
            if hasattr(self, 'code_patches_checkboxes'):
                for cb in self.code_patches_checkboxes:
                    cb.setParent(None)  # Remove from parent widget
                    cb.deleteLater()   # Schedule for deletion
                self.code_patches_checkboxes.clear()
            
            if hasattr(self, 'config_patches_checkboxes'):
                for cb in self.config_patches_checkboxes:
                    cb.setParent(None)  # Remove from parent widget
                    cb.deleteLater()   # Schedule for deletion
                self.config_patches_checkboxes.clear()
            
            # Clear the patches layout if it exists
            if hasattr(self, 'patches_layout'):
                # Remove all widgets from the layout
                while self.patches_layout.count():
                    child = self.patches_layout.takeAt(0)
                    if child.widget():
                        child.widget().deleteLater()
            
            # Clear any patch-related labels
            if hasattr(self, 'code_patches_label'):
                self.code_patches_label.setParent(None)
                self.code_patches_label.deleteLater()
                delattr(self, 'code_patches_label')
            
            if hasattr(self, 'config_patches_label'):
                self.config_patches_label.setParent(None)
                self.config_patches_label.deleteLater()
                delattr(self, 'config_patches_label')
            
            # Clear backend suggestions files
            import os
            import glob
            
            # Remove old fix suggestions files
            resources_dir = "backend/resources"
            if os.path.exists(resources_dir):
                old_files = glob.glob(os.path.join(resources_dir, "fix_suggestions_*.json"))
                for file_path in old_files:
                    try:
                        os.remove(file_path)
                        print(f"Removed old suggestions file: {file_path}")
                    except Exception as e:
                        print(f"Could not remove {file_path}: {e}")
            
            # Clear git history suggestions file reference and commit message
            if hasattr(self, 'git_history_suggestions_file'):
                delattr(self, 'git_history_suggestions_file')
            if hasattr(self, 'git_history_commit_message'):
                delattr(self, 'git_history_commit_message')
            if hasattr(self, 'git_history_commit_info'):
                delattr(self, 'git_history_commit_info')
            
            print("✅ Cleared previous patches from UI and backend")
            
        except Exception as e:
            print(f"Error clearing previous patches: {e}")
    
    def show_existing_fix_details(self):
        """Show details of the existing fix found in embeddings"""
        # Clear previous RCA analysis patches when showing git history
        self.clear_previous_patches()
        
        if not hasattr(self, 'existing_fix_result') or not self.existing_fix_result:
            QMessageBox.information(self, "No Fix", "No existing fix data available.")
            return
        
        result = self.existing_fix_result
        commit = result.get('commit')
        
        if not commit:
            QMessageBox.information(self, "No Fix", "No fix commit available.")
            return
        
        try:
            # Get the source code directory and extract folder name
            code_dir = self.code_dir_edit.text() if hasattr(self, 'code_dir_edit') else None
            if code_dir:
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            else:
                openair_codebase_file_name = "openairinterface5g-develop"  # Default fallback
            
            # Get full commit details from git
            commit_hash = commit.get('commit_hash', commit.get('commit_hash_short', ''))
            if not commit_hash:
                QMessageBox.warning(self, "Error", "No commit hash available.")
                return
            
            # Get full commit details using git show with actual code changes
            import subprocess
            git_show_cmd = ['git', 'show', '--stat', '--format=fuller', '--unified=3', commit_hash]
            
            # Change to the git repository directory
            # Try multiple possible paths - including BOTH repositories for cross-repo fixes
            possible_paths = [
                # Try selected repository first
                os.path.join(os.getcwd(), openair_codebase_file_name),
                os.path.join(os.getcwd(), '..', openair_codebase_file_name),
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', openair_codebase_file_name),
                os.path.join(os.getcwd(), '..', 'Error_fixing_pipelin', openair_codebase_file_name),
                # Fallback: Try other known repositories
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', 'openairinterface5g-develop'),
                os.path.join(os.getcwd(), 'Error_fixing_pipelin', 'openairinterface5g-test'),
            ]
            
            git_dir = None
            commit_found_in = None
            
            # First try to find a repo where the commit exists
            for path in possible_paths:
                if os.path.exists(path) and os.path.exists(os.path.join(path, '.git')):
                    # Check if commit exists in this repo
                    try:
                        import subprocess
                        check_result = subprocess.run(
                            ['git', 'cat-file', '-e', commit_hash],
                            cwd=path,
                            capture_output=True,
                            timeout=2
                        )
                        if check_result.returncode == 0:
                            git_dir = path
                            commit_found_in = os.path.basename(path)
                            print(f"✅ Found commit in repository: {commit_found_in}")
                            break
                    except:
                        continue
            
            # If commit not found in any repo, just use the first available git repo
            if not git_dir:
                for path in possible_paths:
                    if os.path.exists(path) and os.path.exists(os.path.join(path, '.git')):
                        git_dir = path
                        print(f"⚠️ Commit not found in any repo, using: {os.path.basename(path)}")
                        break
            
            if not git_dir:
                print(f"Git repository not found. Tried paths: {possible_paths}")
                self._display_basic_fix_details(commit, result)
                return
            
            print(f"Running git show in directory: {git_dir}")
            print(f"Git command: {' '.join(git_show_cmd)}")
            
            # Show cross-repo info if applicable
            if commit_found_in and commit_found_in != openair_codebase_file_name:
                print(f"📌 Cross-repository fix: Found in '{commit_found_in}', will apply to '{openair_codebase_file_name}'")
            
            git_result = subprocess.run(
                git_show_cmd, 
                cwd=git_dir,
                capture_output=True, 
                text=True, 
                timeout=30
            )
            
            print(f"Git show return code: {git_result.returncode}")
            if git_result.stderr:
                print(f"Git show stderr: {git_result.stderr}")
            
            if git_result.returncode != 0:
                print("Git show failed, falling back to basic display")
                # Fallback to basic info if git show fails
                self._display_basic_fix_details(commit, result)
                return
            
            # Parse git show output
            git_output = git_result.stdout
            lines = git_output.split('\n')
            
            # Extract commit info
            commit_info = {}
            in_message = False
            message_lines = []
            
            for line in lines:
                if line.startswith('commit '):
                    commit_info['full_hash'] = line.split()[1]
                elif line.startswith('Author:'):
                    commit_info['author'] = line.replace('Author:', '').strip()
                elif line.startswith('AuthorDate:'):
                    commit_info['date'] = line.replace('AuthorDate:', '').strip()
                elif line.startswith('Commit:'):
                    commit_info['committer'] = line.replace('Commit:', '').strip()
                elif line.startswith('CommitDate:'):
                    commit_info['commit_date'] = line.replace('CommitDate:', '').strip()
                elif line.startswith('    ') and not in_message:
                    in_message = True
                    message_lines.append(line.strip())
                elif in_message and line.startswith('    '):
                    message_lines.append(line.strip())
                elif in_message and not line.startswith('    '):
                    break
            
            # Store the original commit message for later use
            self.git_history_commit_message = '\n'.join(message_lines)
            self.git_history_commit_info = commit_info
            
            # Format the detailed display
            details = self._format_fix_details(commit, result, commit_info, message_lines, git_output)
            
            # Display in the results text box
            if hasattr(self, 'code_rca_results_text'):
                self.code_rca_results_text.setPlainText(details)
            else:
                QMessageBox.information(self, "Existing Fix Details", details)
                
        except Exception as e:
            print(f"Error getting commit details: {e}")
            # Fallback to basic display
            self._display_basic_fix_details(commit, result)
    
    def _get_current_branch(self):
        """Get the current branch of the target repository"""
        try:
            # Get the source code directory - check both Code RCA and Bug Discovery fields
            code_dir = None
            if hasattr(self, 'code_dir_edit') and self.code_dir_edit.text():
                code_dir = self.code_dir_edit.text()
            elif hasattr(self, 'bug_code_dir_edit') and self.bug_code_dir_edit.text():
                code_dir = self.bug_code_dir_edit.text()
            
            if not code_dir or not os.path.exists(code_dir):
                return "N/A"
            
            # Get the git directory path
            git_dir = code_dir
            if not os.path.exists(os.path.join(git_dir, '.git')):
                return "N/A"
            
            # Run git command to get current branch
            import subprocess
            result = subprocess.run(
                ['git', 'rev-parse', '--abbrev-ref', 'HEAD'],
                cwd=git_dir,
                capture_output=True,
                text=True,
                timeout=5
            )
            
            if result.returncode == 0:
                branch = result.stdout.strip()
                return branch if branch else "N/A"
            else:
                return "N/A"
                
        except Exception as e:
            print(f"Error getting current branch: {e}")
            return "N/A"
    
    def _display_basic_fix_details(self, commit, result):
        """Display basic fix details when git show fails"""
        msg = f"🔍 EXISTING FIX FOUND\n"
        msg += f"{'='*60}\n\n"
        
        # Basic commit info
        msg += f"📋 COMMIT DETAILS\n"
        msg += f"{'-'*30}\n"
        msg += f"Hash: {commit.get('commit_hash_short', 'N/A')}\n"
        msg += f"Author: {commit.get('author_name', 'N/A')}\n"
        msg += f"Date: {commit.get('date_iso', 'N/A')}\n"
        msg += f"Subject: {commit.get('subject', 'N/A')}\n\n"
        
        # Branch information
        msg += f"🌿 BRANCH INFORMATION\n"
        msg += f"{'-'*30}\n"
        msg += f"Fix Source: origin/develop\n"
        
        # Get current branch of target repository
        current_branch = self._get_current_branch()
        msg += f"Target Branch: {current_branch}\n\n"
        
        # Match details
        msg += f"🎯 MATCH DETAILS\n"
        msg += f"{'-'*30}\n"
        msg += f"Confidence: {result.get('confidence', 'N/A')}\n"
        msg += f"Similarity Score: {commit.get('boosted_score', commit.get('similarity', 0)):.2%}\n"
        msg += f"Reasoning: {result.get('reasoning', 'N/A')}\n\n"
        
        # Commit message from embeddings
        if commit.get('body'):
            msg += f"📝 COMMIT MESSAGE\n"
            msg += f"{'-'*30}\n"
            msg += f"{commit.get('body', 'N/A')}\n\n"
        
        # Patch information
        msg += f"🔧 PATCH INFORMATION\n"
        msg += f"{'-'*30}\n"
        
        if commit.get('code_patches'):
            msg += f"Code Patches: {commit.get('code_patch_count', 0)}\n"
            for i, patch in enumerate(commit.get('code_patches', [])[:5], 1):  # Show first 5
                if isinstance(patch, dict):
                    msg += f"  {i}. {patch.get('function', 'N/A')} ({patch.get('file', 'N/A')})\n"
                else:
                    msg += f"  {i}. {patch}\n"
            if len(commit.get('code_patches', [])) > 5:
                msg += f"  ... and {len(commit.get('code_patches', [])) - 5} more\n"
        
        if commit.get('config_patches'):
            msg += f"Config Patches: {commit.get('config_patch_count', 0)}\n"
            for i, patch in enumerate(commit.get('config_patches', [])[:5], 1):  # Show first 5
                if isinstance(patch, dict):
                    msg += f"  {i}. {patch.get('parameter', 'N/A')} ({patch.get('file', 'N/A')})\n"
                else:
                    msg += f"  {i}. {patch}\n"
            if len(commit.get('config_patches', [])) > 5:
                msg += f"  ... and {len(commit.get('config_patches', [])) - 5} more\n"
        
        # Files changed
        if commit.get('files_changed'):
            msg += f"\n📁 FILES CHANGED\n"
            msg += f"{'-'*30}\n"
            for file_info in commit.get('files_changed', [])[:10]:  # Show first 10
                if isinstance(file_info, dict):
                    msg += f"  • {file_info.get('file', 'N/A')} ({file_info.get('changes', 'N/A')})\n"
                else:
                    msg += f"  • {file_info}\n"
            if len(commit.get('files_changed', [])) > 10:
                msg += f"  ... and {len(commit.get('files_changed', [])) - 10} more files\n"
        
        # Recommendation
        msg += f"\n💡 RECOMMENDATION\n"
        msg += f"{'-'*30}\n"
        msg += f"This fix was previously applied to resolve a similar issue.\n"
        msg += f"Review the patches above and apply similar changes to your codebase.\n"
        msg += f"\nNote: Full git details not available. The commit may have been\n"
        msg += f"deleted or the repository path is incorrect.\n"
        
        if hasattr(self, 'code_rca_results_text'):
            self.code_rca_results_text.setPlainText(msg)
        else:
            QMessageBox.information(self, "Existing Fix Details", msg)
    
    def _format_fix_details(self, commit, result, commit_info, message_lines, git_output):
        """Format comprehensive fix details"""
        details = f"🔍 EXISTING FIX FOUND\n"
        details += f"{'='*60}\n\n"
        
        # Basic commit info
        details += f"📋 COMMIT DETAILS\n"
        details += f"{'-'*30}\n"
        details += f"Hash: {commit_info.get('full_hash', commit.get('commit_hash_short', 'N/A'))}\n"
        details += f"Author: {commit_info.get('author', commit.get('author_name', 'N/A'))}\n"
        details += f"Date: {commit_info.get('date', commit.get('date_iso', 'N/A'))}\n"
        details += f"Committer: {commit_info.get('committer', 'N/A')}\n"
        details += f"Commit Date: {commit_info.get('commit_date', 'N/A')}\n\n"
        
        # Branch information
        details += f"🌿 BRANCH INFORMATION\n"
        details += f"{'-'*30}\n"
        details += f"Fix Source: origin/develop\n"
        
        # Get current branch of target repository
        current_branch = self._get_current_branch()
        details += f"Target Branch: {current_branch}\n\n"
        
        # Match details
        details += f"🎯 MATCH DETAILS\n"
        details += f"{'-'*30}\n"
        details += f"Confidence: {result.get('confidence', 'N/A')}\n"
        details += f"Similarity Score: {commit.get('boosted_score', commit.get('similarity', 0)):.2%}\n"
        details += f"Reasoning: {result.get('reasoning', 'N/A')}\n\n"
        
        # Commit message
        if message_lines:
            details += f"📝 COMMIT MESSAGE\n"
            details += f"{'-'*30}\n"
            details += '\n'.join(message_lines) + "\n\n"
        
        # Patch information
        details += f"🔧 PATCH INFORMATION\n"
        details += f"{'-'*30}\n"
        if commit.get('code_patches'):
            details += f"Code Patches: {commit.get('code_patch_count', 0)}\n"
            for i, patch in enumerate(commit.get('code_patches', [])[:3], 1):  # Show first 3
                if isinstance(patch, dict):
                    details += f"  {i}. {patch.get('function', 'N/A')} ({patch.get('file', 'N/A')})\n"
                else:
                    details += f"  {i}. {patch}\n"
            if len(commit.get('code_patches', [])) > 3:
                details += f"  ... and {len(commit.get('code_patches', [])) - 3} more\n"
        
        if commit.get('config_patches'):
            details += f"Config Patches: {commit.get('config_patch_count', 0)}\n"
            for i, patch in enumerate(commit.get('config_patches', [])[:3], 1):  # Show first 3
                if isinstance(patch, dict):
                    details += f"  {i}. {patch.get('parameter', 'N/A')} ({patch.get('file', 'N/A')})\n"
                else:
                    details += f"  {i}. {patch}\n"
            if len(commit.get('config_patches', [])) > 3:
                details += f"  ... and {len(commit.get('config_patches', [])) - 3} more\n"
        
        # Files changed (from git show)
        if 'files changed' in git_output.lower():
            details += f"\n📁 FILES CHANGED\n"
            details += f"{'-'*30}\n"
            # Extract file changes from git show output
            file_section = False
            for line in git_output.split('\n'):
                if 'files changed' in line.lower():
                    file_section = True
                    details += f"{line}\n"
                elif file_section and line.strip() and not line.startswith(' '):
                    break
                elif file_section and line.strip():
                    details += f"{line}\n"
        
        # Show actual code changes (diff)
        details += f"\n🔍 CODE CHANGES (DIFF)\n"
        details += f"{'-'*30}\n"
        
        # Parse and format the diff output
        diff_lines = []
        in_diff = False
        current_file = None
        
        for line in git_output.split('\n'):
            if line.startswith('diff --git'):
                in_diff = True
                current_file = line.split()[-1] if len(line.split()) > 3 else "Unknown file"
                diff_lines.append(f"\n📄 File: {current_file}")
                diff_lines.append("─" * 50)
            elif in_diff and line.startswith('@@'):
                # Skip the hunk header for cleaner display
                continue
            elif in_diff and (line.startswith('+') or line.startswith('-') or line.startswith(' ')):
                # Show the actual changes
                if line.startswith('+'):
                    diff_lines.append(f"➕ {line[1:]}")  # Green for additions
                elif line.startswith('-'):
                    diff_lines.append(f"➖ {line[1:]}")  # Red for deletions
                else:
                    diff_lines.append(f"  {line[1:]}")  # Context lines
            elif in_diff and line.strip() == '':
                # Empty line in diff
                diff_lines.append("")
            elif in_diff and not line.startswith(('diff', 'index', '+++', '---', '@@', '+', '-', ' ')):
                # End of diff section
                in_diff = False
        
        if diff_lines:
            details += '\n'.join(diff_lines[:100])  # Limit to first 100 lines to avoid overwhelming output
            if len(diff_lines) > 100:
                details += f"\n\n... and {len(diff_lines) - 100} more lines of changes"
        else:
            details += "No code changes visible in this commit."
        
        # Show actual code changes if available
        details += f"\n💡 RECOMMENDATION\n"
        details += f"{'-'*30}\n"
        details += f"This fix was previously applied to resolve a similar issue.\n"
        details += f"You can review the changes above and apply similar patches\n"
        details += f"to your current codebase.\n"
        
        return details
    
    def update_code_rca_button_state(self):
        """Update the state of the Start RCA Analysis button based on log file and code directory selection."""
        try:
            if (hasattr(self, 'start_code_rca_btn') and 
                hasattr(self, 'code_log_file_combo') and 
                hasattr(self, 'code_dir_edit') and
                self.start_code_rca_btn is not None and 
                self.code_log_file_combo is not None and
                self.code_dir_edit is not None):
                # Enable button only if both code directory and log file are selected
                has_code_dir = bool(self.code_dir_edit.text().strip())
                has_log_file = bool(self.code_log_file_combo.currentText())
                self.start_code_rca_btn.setEnabled(has_code_dir and has_log_file)
        except RuntimeError:
            # Widget has been deleted, ignore the error
            pass

    def load_patches_from_json(self, json_file_path):
        """Load patches from fix_suggestions.json and populate the checkboxes."""
        try:
            if not os.path.exists(json_file_path):
                return
                
            with open(json_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Store data for preview access
            if not hasattr(self, 'loaded_bug_analysis'):
                self.loaded_bug_analysis = {}
            if 'results' not in self.loaded_bug_analysis:
                self.loaded_bug_analysis['results'] = {}
            if 'phase3_fixes' not in self.loaded_bug_analysis['results']:
                self.loaded_bug_analysis['results']['phase3_fixes'] = {}
            self.loaded_bug_analysis['results']['phase3_fixes']['fix_suggestion'] = data.get('fix_suggestion', {})
            
            # Clear existing checkboxes
            self.clear_patches_checkboxes()
            
            # Load code patches
            if 'fix_suggestion' in data:
                fix_suggestion = data['fix_suggestion']
                
                # Load code patches
                code_patches = fix_suggestion.get('code_patches', [])
                if code_patches:
                    # Hide placeholder
                    self.code_patches_placeholder.hide()
                    
                    for i, patch in enumerate(code_patches):
                        function_name = patch.get('function_name', f'Unknown Function {i+1}')
                        file_path = patch.get('file_path', 'Unknown file')
                        
                        # Create checkbox with function name and file info
                        file_name = os.path.basename(file_path)
                        checkbox_text = f"{function_name} ({file_name})"
                        checkbox = QCheckBox(checkbox_text)
                        checkbox.setToolTip(f"File: {file_path}\nDescription: {patch.get('description', 'No description')}")
                        checkbox.setChecked(True)  # Default to checked
                        checkbox.stateChanged.connect(self.update_patch_preview)
                        
                        self.code_patches_checkboxes.append(checkbox)
                        self.code_patches_layout.addWidget(checkbox)
                else:
                    # Show placeholder if no code patches
                    self.code_patches_placeholder.show()
                
                # Load config patches
                config_patches = fix_suggestion.get('config_patches', [])
                if config_patches:
                    # Hide placeholder
                    self.config_patches_placeholder.hide()
                    
                    for i, patch in enumerate(config_patches):
                        config_name = patch.get('config_name', f'Unknown Config {i+1}')
                        file_path = patch.get('file_path', 'Unknown file')
                        
                        # Create checkbox with config name and file info
                        file_name = os.path.basename(file_path)
                        checkbox_text = f"{config_name} ({file_name})"
                        checkbox = QCheckBox(checkbox_text)
                        checkbox.setToolTip(f"File: {file_path}\nDescription: {patch.get('description', 'No description')}")
                        checkbox.setChecked(True)  # Default to checked
                        checkbox.stateChanged.connect(self.update_patch_preview)
                        
                        self.config_patches_checkboxes.append(checkbox)
                        self.config_patches_layout.addWidget(checkbox)
                else:
                    # Show placeholder if no config patches
                    self.config_patches_placeholder.show()
            
            # Update the preview to show all selected patches
            self.update_patch_preview()
                    
        except Exception as e:
            print(f"Error loading patches from JSON: {e}")

    def extract_name_from_checkbox_text(self, checkbox_text):
        """Extract the original name from checkbox text that includes file name in brackets."""
        # Format: "name (filename)" -> extract "name"
        if ' (' in checkbox_text and checkbox_text.endswith(')'):
            return checkbox_text.split(' (')[0]
        return checkbox_text

    def clear_patches_checkboxes(self):
        """Clear all existing patch checkboxes."""
        # Clear code patches checkboxes
        for checkbox in self.code_patches_checkboxes:
            checkbox.setParent(None)
            checkbox.deleteLater()
        self.code_patches_checkboxes.clear()
        
        # Clear config patches checkboxes
        for checkbox in self.config_patches_checkboxes:
            checkbox.setParent(None)
            checkbox.deleteLater()
        self.config_patches_checkboxes.clear()
        
        # Show placeholders
        self.code_patches_placeholder.show()
        self.config_patches_placeholder.show()

    def select_all_patches(self):
        """Select all available patches."""
        # Select all code patches
        if hasattr(self, 'code_patches_checkboxes'):
            for checkbox in self.code_patches_checkboxes:
                checkbox.setChecked(True)
        
        # Select all config patches
        if hasattr(self, 'config_patches_checkboxes'):
            for checkbox in self.config_patches_checkboxes:
                checkbox.setChecked(True)
        
        # Update preview
        self.update_patch_preview()

    def unselect_all_patches(self):
        """Unselect all patches."""
        # Unselect all code patches
        if hasattr(self, 'code_patches_checkboxes'):
            for checkbox in self.code_patches_checkboxes:
                checkbox.setChecked(False)
        
        # Unselect all config patches
        if hasattr(self, 'config_patches_checkboxes'):
            for checkbox in self.config_patches_checkboxes:
                checkbox.setChecked(False)
        
        # Update preview
        self.update_patch_preview()

    def get_edited_patch_values(self):
        """Get the edited values from input fields for both code and config patches."""
        edited_values = {
            'code_patches': {},
            'config_patches': {}
        }
        
        if hasattr(self, 'patch_value_editors'):
            for key, editor_info in self.patch_value_editors.items():
                if 'editor' in editor_info:
                    patch_index = editor_info['patch_index']
                    original_value = editor_info['original_value']
                    
                    if editor_info['type'] == 'config':
                        # Get the current value from the input field (QLineEdit)
                        current_value = editor_info['editor'].text()
                        
                        # Store if value was edited
                        if current_value != original_value:
                            print(f"DEBUG: Config patch {patch_index} edited: '{original_value}' -> '{current_value}'")
                        
                        edited_values['config_patches'][patch_index] = current_value
                    
                    elif editor_info['type'] == 'code':
                        # Get the current value from the text edit field (QTextEdit)
                        current_value = editor_info['editor'].toPlainText()
                        
                        # Store if value was edited
                        if current_value != original_value:
                            print(f"DEBUG: Code patch {patch_index} edited")
                            print(f"  Original length: {len(original_value)} chars")
                            print(f"  New length: {len(current_value)} chars")
                        
                        edited_values['code_patches'][patch_index] = current_value
        
        return edited_values
    
    def apply_selected_patches(self):
        """Apply the selected patches using unified patch applicator."""
        try:
            # Get the source code directory from loaded analysis
            code_dir = None
            if hasattr(self, 'loaded_bug_analysis'):
                code_dir = self.loaded_bug_analysis.get('code_dir', None)
            elif hasattr(self, 'code_dir_edit'):
                code_dir = self.code_dir_edit.text()
            
            if not code_dir:
                QMessageBox.warning(self, "Missing Code Directory", "Please load a bug analysis first.")
                return
            
            openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            
            # Get selected patches
            selected_code_patches = []
            selected_config_patches = []
            
            if hasattr(self, 'code_patches_checkboxes'):
                print(f"DEBUG: Total code patch checkboxes: {len(self.code_patches_checkboxes)}")
                for cb in self.code_patches_checkboxes:
                    print(f"DEBUG: Code checkbox '{cb.text()}' - Checked: {cb.isChecked()}")
                    if cb.isChecked():
                        selected_code_patches.append(cb.text())
            else:
                print("DEBUG: code_patches_checkboxes attribute not found")
            
            if hasattr(self, 'config_patches_checkboxes'):
                print(f"DEBUG: Total config patch checkboxes: {len(self.config_patches_checkboxes)}")
                for cb in self.config_patches_checkboxes:
                    print(f"DEBUG: Config checkbox '{cb.text()}' - Checked: {cb.isChecked()}")
                    if cb.isChecked():
                        selected_config_patches.append(cb.text())
            else:
                print("DEBUG: config_patches_checkboxes attribute not found")
            
            print(f"DEBUG: Final selected_code_patches: {selected_code_patches}")
            print(f"DEBUG: Final selected_config_patches: {selected_config_patches}")
            
            # Check if we're in git history mode (no checkboxes but text box has git commit details)
            if not selected_code_patches and not selected_config_patches:
                # Try to parse git commit details from the text box
                if hasattr(self, 'code_rca_results_text'):
                    text_content = self.code_rca_results_text.toPlainText()
                    if "EXISTING FIX FOUND" in text_content and "🔧 PATCH INFORMATION" in text_content:
                        print("DEBUG: Git history mode detected, parsing patches...")
                        # Parse git commit details and create patches
                        parsed_patches = self.parse_git_commit_for_patches(text_content, openair_codebase_file_name)
                        if parsed_patches:
                            selected_code_patches = parsed_patches.get('code_patches', [])
                            selected_config_patches = parsed_patches.get('config_patches', [])
                            print(f"DEBUG: Parsed patches - Code: {len(selected_code_patches)}, Config: {len(selected_config_patches)}")
                            print(f"DEBUG: Selected code patches: {selected_code_patches}")
                            print(f"DEBUG: Selected config patches: {selected_config_patches}")
                        else:
                            print("DEBUG: No patches parsed from git history")
            
            if not selected_code_patches and not selected_config_patches:
                QMessageBox.warning(self, "No Patches Selected", "Please select at least one patch to apply.")
                return
            
            # Show confirmation dialog
            patch_count = len(selected_code_patches) + len(selected_config_patches)
            confirm_msg = f"Apply {patch_count} selected patches?\n\n"
            confirm_msg += f"Code patches: {len(selected_code_patches)}\n"
            confirm_msg += f"Config patches: {len(selected_config_patches)}\n\n"
            confirm_msg += "Backup files will be created automatically."
            
            reply = QMessageBox.question(
                self, 
                "Confirm Patch Application", 
                confirm_msg,
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            
            if reply != QMessageBox.Yes:
                return
            
            # Find the most recent fix_suggestions.json file
            resources_dir = "backend/resources"
            fix_suggestions_file = None
            
            # Check if we have a git history suggestions file
            if hasattr(self, 'git_history_suggestions_file') and os.path.exists(self.git_history_suggestions_file):
                fix_suggestions_file = self.git_history_suggestions_file
                print(f"Using git history suggestions file: {fix_suggestions_file}")
            else:
                # Use regular RCA analysis file
                if os.path.exists(resources_dir):
                    json_files = [f for f in os.listdir(resources_dir) if f.startswith('fix_suggestions_') and f.endswith('.json')]
                    if json_files:
                        # Get the most recent file
                        json_files.sort(reverse=True)
                        fix_suggestions_file = os.path.join(resources_dir, json_files[0])
            
            if not fix_suggestions_file or not os.path.exists(fix_suggestions_file):
                QMessageBox.warning(self, "Fix Suggestions Not Found", "Could not find fix_suggestions.json file. Please run RCA analysis first.")
                return
            
            # If code_dir is not available from loaded_bug_analysis, try to get it from fix_suggestions file
            if not code_dir:
                try:
                    with open(fix_suggestions_file, 'r', encoding='utf-8') as f:
                        suggestions_data = json.load(f)
                    code_dir = suggestions_data.get('code_dir', None)
                    if code_dir:
                        openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
                        print(f"Found code directory in fix_suggestions file: {code_dir}")
                except Exception as e:
                    print(f"Error reading code directory from fix_suggestions file: {e}")
            
            if not code_dir:
                QMessageBox.warning(self, "Missing Code Directory", "Could not determine the source code directory. Please ensure the bug analysis includes the code directory.")
                return
            
            # Show progress dialog
            progress_dialog = QProgressDialog("Applying patches...", "Cancel", 0, 0, self)
            progress_dialog.setWindowModality(Qt.WindowModal)
            progress_dialog.show()
            QApplication.processEvents()
            
            # Import and use unified patch applicator
            try:
                import sys
                sys.path.append('Error_fixing_pipelin')
                from unified_patch_applicator import UnifiedPatchApplicator
                
                # Update patch values with edited values from input fields before applying
                edited_patch_values = self.get_edited_patch_values()
                
                # Create a filtered fix_suggestions.json with only selected patches and corrected paths
                filtered_suggestions_file = self.create_filtered_suggestions_file(
                    fix_suggestions_file, selected_code_patches, selected_config_patches, openair_codebase_file_name, edited_patch_values
                )
                
                print(f"\nDEBUG: Applying patches from: {filtered_suggestions_file}")
                print(f"DEBUG: Current working directory: {os.getcwd()}")
                
                # Create applicator with filtered file
                applicator = UnifiedPatchApplicator(filtered_suggestions_file)
                
                # Apply patches with backup
                print(f"DEBUG: Starting patch application...")
                result = applicator.apply_all_patches(dry_run=False, backup=True)
                print(f"DEBUG: Patch application completed")
                
                # Debug: Show what files were modified
                print(f"\n{'='*60}")
                print(f"DEBUG [Patch Application Results]:")
                print(f"  Total applied: {result.get('total_applied', 0)}")
                print(f"  Total failed: {result.get('total_failed', 0)}")
                if 'applied_code_patches' in result:
                    print(f"  Applied code patches: {len(result['applied_code_patches'])}")
                    for patch in result['applied_code_patches']:
                        if isinstance(patch, dict):
                            print(f"    - {patch.get('function_name', 'N/A')} in {patch.get('file_path', 'N/A')}")
                if 'applied_config_patches' in result:
                    print(f"  Applied config patches: {len(result['applied_config_patches'])}")
                    for patch in result['applied_config_patches']:
                        if isinstance(patch, dict):
                            print(f"    - {patch.get('config_name', 'N/A')} in {patch.get('file_path', 'N/A')}")
                print(f"{'='*60}\n")
                
                # Close progress dialog
                progress_dialog.close()
                
                # Show results
                if result["success"]:
                    # Show success dialog with Git commit option
                    self.show_patch_success_dialog(result, selected_code_patches, selected_config_patches)
                else:
                    QMessageBox.warning(
                        self, 
                        "Patch Application Issues", 
                        f"Some patches had issues:\n\nApplied: {result['total_applied']}\nFailed: {result['total_failed']}\n\nCheck console output for details."
                    )
                    
            except ImportError as e:
                progress_dialog.close()
                QMessageBox.critical(
                    self, 
                    "Module Import Error", 
                    f"Could not import unified patch applicator:\n{str(e)}\n\nPlease ensure the Error_fixing_pipelin module is available."
                )
            except Exception as e:
                progress_dialog.close()
                QMessageBox.critical(
                    self, 
                    "Patch Application Error", 
                    f"An error occurred while applying patches:\n{str(e)}"
                )
                
        except Exception as e:
            QMessageBox.critical(
                self, 
                "Error", 
                f"An unexpected error occurred:\n{str(e)}"
            )
            
    def get_meaningful_source_name(self, source):
        """Convert technical source names to meaningful display names"""
        source_mapping = {
            'pattern_hint_direct': 'Direct Retrieval',
            'context_aware_retrieval': 'Semantic Search',
            'error_handling_pattern': 'Pattern Matching',
            'call_chain_upstream': 'Call Chain (Upstream)',
            'call_chain_downstream': 'Call Chain (Downstream)',
            'call_chain_seed': 'Call Chain (Seed)',
            'symbolic_keyword': 'Keyword Search',
            'direct_function_extraction': 'Function Extraction',
            'semantic_search': 'Semantic Search',
            'pattern_matching': 'Pattern Matching',
            'keyword_search': 'Keyword Search',
            'call_chain_analysis': 'Call Chain Analysis'
        }
        return source_mapping.get(source, source)

    def show_error_artifacts(self):
        """Show error artifacts including error name, suspected functions, sources, and context"""
        try:
            # Check if analysis results exist
            artifacts_file = "Error_fixing_pipelin/output/complete_error_analysis.json"
            phase2_file = "Error_fixing_pipelin/output/phase2_results.json"
            fix_suggestions_file = "Error_fixing_pipelin/output/fix_suggestions.json"
            
            # Try to load the most complete file first
            artifacts_data = None
            if os.path.exists(artifacts_file):
                with open(artifacts_file, 'r', encoding='utf-8') as f:
                    artifacts_data = json.load(f)
            elif os.path.exists(phase2_file):
                with open(phase2_file, 'r', encoding='utf-8') as f:
                    artifacts_data = json.load(f)
            elif os.path.exists(fix_suggestions_file):
                with open(fix_suggestions_file, 'r', encoding='utf-8') as f:
                    artifacts_data = json.load(f)
            
            if not artifacts_data:
                QMessageBox.warning(self, "No Data", "No error analysis artifacts found. Please run RCA analysis first.")
                return
            
            # Create dialog for artifacts display
            dialog = QDialog(self)
            dialog.setWindowTitle("Error Analysis Artifacts")
            dialog.setMinimumSize(1200, 750)  # Optimized for better screen fit
            
            # Apply professional blue theme styling to the dialog
            dialog.setStyleSheet("""
                QDialog {
                    background-color: #f0f4f8;
                    color: #1e3a5f;
                }
                QTabWidget::pane {
                    border: 1px solid #b3d9ff;
                    background-color: #ffffff;
                    border-radius: 4px;
                }
                QTabBar::tab {
                    background-color: #e6f2ff;
                    color: #1e3a5f;
                    padding: 8px 16px;
                    margin-right: 2px;
                    border-top-left-radius: 4px;
                    border-top-right-radius: 4px;
                }
                QTabBar::tab:selected {
                    background-color: #ffffff;
                    color: #1e3a5f;
                    border-bottom: 2px solid #0066cc;
                }
                QTabBar::tab:hover {
                    background-color: #cce7ff;
                }
            """)
            
            layout = QVBoxLayout()
            
            # Create tab widget for different artifact types
            tab_widget = QTabWidget()
            
            # Error Information Tab
            error_tab = self.create_error_info_tab(artifacts_data)
            tab_widget.addTab(error_tab, "🔍 Error Information")
            
            # Suspected Functions Tab
            functions_tab = self.create_functions_artifacts_tab(artifacts_data)
            tab_widget.addTab(functions_tab, "⚙️ Suspected Functions")
            
            # Suspected Configs Tab
            configs_tab = self.create_configs_artifacts_tab(artifacts_data)
            tab_widget.addTab(configs_tab, "⚙️ Suspected Configs")
            
            # Context Information Tab
            context_tab = self.create_context_artifacts_tab(artifacts_data)
            tab_widget.addTab(context_tab, "📋 Context Information")
            
            # Analysis Summary Tab
            summary_tab = self.create_analysis_summary_tab(artifacts_data)
            tab_widget.addTab(summary_tab, "📊 Analysis Summary")
            
            # 3GPP Spec Reference Tab
            spec_tab = self.create_3gpp_spec_tab(artifacts_data)
            tab_widget.addTab(spec_tab, "📚 3GPP Spec Reference")
            
            layout.addWidget(tab_widget)
            
            # Export button
            export_btn = QPushButton("Export Artifacts")
            export_btn.clicked.connect(lambda: self.export_artifacts(artifacts_data))
            layout.addWidget(export_btn)
            
            dialog.setLayout(layout)
            dialog.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to show error artifacts: {str(e)}")
            print(f"Error in show_error_artifacts: {e}")

    def create_crash_backtrace_tab(self, data):
        """Create tab showing crash backtrace"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        backtrace_group = QGroupBox("Call Stack (Backtrace)")
        backtrace_layout = QVBoxLayout()
        
        backtrace = data.get('backtrace', [])
        if backtrace:
            backtrace_text = QTextEdit()
            backtrace_text.setReadOnly(True)
            backtrace_text.setStyleSheet("font-family: 'Consolas', monospace; font-size: 11px;")
            
            text = ""
            for frame in backtrace:
                text += f"Frame #{frame.get('frame_number', '?')}: {frame.get('function', 'Unknown')}\n"
                text += f"   File: {frame.get('file', 'Unknown')}\n"
                text += f"   Line: {frame.get('line', '?')}\n\n"
            
            backtrace_text.setPlainText(text)
            backtrace_layout.addWidget(backtrace_text)
        else:
            backtrace_layout.addWidget(QLabel("No backtrace data available"))
        
        backtrace_group.setLayout(backtrace_layout)
        layout.addWidget(backtrace_group)
        tab.setLayout(layout)
        return tab
    
    def create_crash_scenario_tab(self, data):
        """Create tab showing scenario flow before crash"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        scenario_group = QGroupBox("Scenario Flow (Steps Before Crash)")
        scenario_layout = QVBoxLayout()
        
        scenario_flow = data.get('scenario_flow', [])
        if scenario_flow:
            scenario_text = QTextEdit()
            scenario_text.setReadOnly(True)
            scenario_text.setStyleSheet("font-family: 'Consolas', monospace; font-size: 11px;")
            
            text = ""
            for i, step in enumerate(scenario_flow, 1):
                text += f"{i}. {step}\n"
            
            scenario_text.setPlainText(text)
            scenario_layout.addWidget(scenario_text)
        else:
            scenario_layout.addWidget(QLabel("No scenario flow data available"))
        
        scenario_group.setLayout(scenario_layout)
        layout.addWidget(scenario_group)
        
        # Add Root Cause Analysis Section
        root_cause_group = QGroupBox("🔬 DETAILED ROOT CAUSE:")
        root_cause_layout = QVBoxLayout()
        
        # Extract root_cause_analysis from phase3_fixes
        phase3_fixes = data.get('phase3_fixes', {})
        fix_suggestion = phase3_fixes.get('fix_suggestion', {})
        root_cause_analysis = fix_suggestion.get('root_cause_analysis', '')
        
        if root_cause_analysis:
            root_cause_text = QTextEdit()
            root_cause_text.setReadOnly(True)
            root_cause_text.setStyleSheet("font-family: 'Consolas', monospace; font-size: 11px; background-color: #F8F9FA; padding: 10px;")
            root_cause_text.setPlainText(root_cause_analysis)
            root_cause_layout.addWidget(root_cause_text)
        else:
            root_cause_layout.addWidget(QLabel("Root cause analysis not available"))
        
        root_cause_group.setLayout(root_cause_layout)
        layout.addWidget(root_cause_group)
        
        tab.setLayout(layout)
        return tab
    
    def create_error_info_tab(self, data):
        """Create tab showing error information"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Error Message Section
        error_group = QGroupBox("Error Details")
        error_layout = QVBoxLayout()
        
        # Error message
        error_msg = data.get('error_message', 'Unknown error')
        error_label = QLabel(f"<b>Error Message:</b><br>{error_msg}")
        error_label.setWordWrap(True)
        error_label.setStyleSheet("padding: 10px; background-color: #FFF3CD; border: 1px solid #FFEAA7; border-radius: 4px;")
        error_layout.addWidget(error_label)
        
        # Timestamp
        timestamp = data.get('timestamp', 'Unknown')
        timestamp_label = QLabel(f"<b>Analysis Time:</b> {timestamp}")
        error_layout.addWidget(timestamp_label)
        
        # Log file
        log_file = data.get('log_file', 'None')
        log_label = QLabel(f"<b>Log File:</b> {log_file}")
        log_label.setWordWrap(True)
        error_layout.addWidget(log_label)
        
        error_group.setLayout(error_layout)
        layout.addWidget(error_group)
        
        layout.addStretch()
        tab.setLayout(layout)
        return tab

    def create_functions_artifacts_tab(self, data):
        """Create tab showing suspected functions with sources and context"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Get Phase 3 suspected function names from fix_suggestions
        phase3_suspected_function_names = []
        phase3_data = data.get('phase3_fixes', {})
        fix_suggestion = phase3_data.get('fix_suggestion', {})
        if fix_suggestion.get('suspected_functions'):
            phase3_suspected_function_names = fix_suggestion['suspected_functions']
        
        # Get all suspected functions from Phase 2 data
        all_suspected_functions = []
        phase2_data = data.get('phase2_analysis', {})
        if phase2_data.get('suspected_functions'):
            all_suspected_functions = phase2_data['suspected_functions']
        elif data.get('suspected_functions'):
            all_suspected_functions = data['suspected_functions']
        
        # Create a map of function names to their detailed data
        function_details_map = {}
        for func in all_suspected_functions:
            func_name = func.get('function_name', '') if isinstance(func, dict) else str(func)
            if func_name:
                function_details_map[func_name] = func
        
        # Build list of all Phase 3 suspected functions, with details if available
        suspected_functions = []
        if phase3_suspected_function_names:
            for func_name in phase3_suspected_function_names:
                if func_name in function_details_map:
                    # Use detailed data from Phase 2
                    suspected_functions.append(function_details_map[func_name])
                else:
                    # Create basic entry for functions without Phase 2 details
                    suspected_functions.append({
                        'function_name': func_name,
                        'file_path': 'Unknown (details not available)',
                        'relevance_score': 0.0,
                        'source': 'Phase 3 Analysis',
                        'context': 'No detailed context available',
                        'reason': 'Identified as suspected function but detailed analysis not available'
                    })
        
        if not suspected_functions:
            if phase3_suspected_function_names:
                no_data_label = QLabel(f"No detailed information found for Phase 3 suspected functions: {', '.join(phase3_suspected_function_names)}")
            else:
                no_data_label = QLabel("No suspected functions found in the analysis.")
            no_data_label.setAlignment(Qt.AlignCenter)
            no_data_label.setStyleSheet("color: #888; font-style: italic; padding: 20px;")
            layout.addWidget(no_data_label)
        else:
            # Create table for suspected functions
            functions_table = QTableWidget()
            functions_table.setColumnCount(6)
            functions_table.setHorizontalHeaderLabels([
                "Function Name", "File Path", "Relevance", "Source", "Context Preview", "Reason"
            ])
            
            # Populate table
            for i, func in enumerate(suspected_functions):
                functions_table.insertRow(i)
                
                # Function name
                func_name = func.get('function_name', 'Unknown')
                functions_table.setItem(i, 0, QTableWidgetItem(func_name))
                
                # File path
                file_path = func.get('file_path', 'Unknown')
                functions_table.setItem(i, 1, QTableWidgetItem(file_path))
                
                # Relevance score
                score = func.get('relevance_score', 0)
                functions_table.setItem(i, 2, QTableWidgetItem(f"{score:.2f}"))
                
                # Source
                source = func.get('source', 'Unknown')
                meaningful_source = self.get_meaningful_source_name(source)
                functions_table.setItem(i, 3, QTableWidgetItem(meaningful_source))
                
                # Context preview (code snippet)
                code_snippet = func.get('code_snippet', '')
                # Show more content but truncate very long snippets
                context_preview = code_snippet[:300] + "..." if len(code_snippet) > 300 else code_snippet
                functions_table.setItem(i, 4, QTableWidgetItem(context_preview))
                
                # Reason
                reason = func.get('reason', 'No reason provided')
                functions_table.setItem(i, 5, QTableWidgetItem(reason))
            
            # Adjust column widths and enable text wrapping
            functions_table.resizeColumnsToContents()
            functions_table.setAlternatingRowColors(True)
            functions_table.setSelectionBehavior(QTableWidget.SelectRows)
            
            # Enable text wrapping for all columns
            functions_table.setWordWrap(True)
            functions_table.setTextElideMode(Qt.ElideNone)
            
            # Set specific column widths for better display (reason column gets most space)
            functions_table.setColumnWidth(0, 170)  # Function Name
            functions_table.setColumnWidth(1, 250)  # File Path
            functions_table.setColumnWidth(2, 85)   # Relevance
            functions_table.setColumnWidth(3, 100)  # Source
            functions_table.setColumnWidth(4, 320)  # Context Preview (code)
            functions_table.setColumnWidth(5, 350)  # Reason (widest column for most content)
            
            # Set row height to accommodate wrapped text
            functions_table.verticalHeader().setDefaultSectionSize(80)
            
            # Apply professional blue theme styling to the functions table
            functions_table.setStyleSheet("""
                QTableWidget {
                    background-color: #ffffff;
                    alternate-background-color: #f0f4f8;
                    gridline-color: #b3d9ff;
                    border: 1px solid #b3d9ff;
                    border-radius: 4px;
                }
                QTableWidget::item {
                    padding: 8px;
                    border: none;
                    color: #1e3a5f;
                }
                QTableWidget::item:selected {
                    background-color: #0066cc;
                    color: #ffffff;
                }
                QTableWidget::item:hover {
                    background-color: #e6f2ff;
                }
                QHeaderView::section {
                    background-color: #1e3a5f;
                    color: #ffffff;
                    padding: 10px;
                    border: none;
                    font-weight: bold;
                    text-align: center;
                }
                QHeaderView::section:first {
                    border-top-left-radius: 4px;
                }
                QHeaderView::section:last {
                    border-top-right-radius: 4px;
                }
            """)
            
            # Add double-click handler to show full context
            functions_table.itemDoubleClicked.connect(lambda item: self.show_function_full_context(item, suspected_functions))
            
            layout.addWidget(functions_table)
        
        tab.setLayout(layout)
        return tab

    def create_configs_artifacts_tab(self, data):
        """Create tab showing suspected configurations with sources and context"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Get Phase 3 suspected config names from fix_suggestions
        phase3_suspected_config_names = []
        phase3_data = data.get('phase3_fixes', {})
        fix_suggestion = phase3_data.get('fix_suggestion', {})
        if fix_suggestion.get('suspected_configs'):
            phase3_suspected_config_names = fix_suggestion['suspected_configs']
        
        # Get all suspected configs from Phase 2 data
        all_suspected_configs = []
        phase2_data = data.get('phase2_analysis', {})
        if phase2_data.get('suspected_configs'):
            all_suspected_configs = phase2_data['suspected_configs']
        elif data.get('suspected_configs'):
            all_suspected_configs = data['suspected_configs']
        
        # Filter configs to show only those mentioned in Phase 3
        suspected_configs = []
        if phase3_suspected_config_names:
            for config in all_suspected_configs:
                config_name = config.get('param_name', '') if isinstance(config, dict) else str(config)
                if config_name in phase3_suspected_config_names:
                    suspected_configs.append(config)
        
        if not suspected_configs:
            if phase3_suspected_config_names:
                no_data_label = QLabel(f"No detailed information found for Phase 3 suspected configs: {', '.join(phase3_suspected_config_names)}")
            else:
                no_data_label = QLabel("No suspected configurations found in the analysis.")
            no_data_label.setAlignment(Qt.AlignCenter)
            no_data_label.setStyleSheet("color: #888; font-style: italic; padding: 20px;")
            layout.addWidget(no_data_label)
        else:
            # Create table for suspected configs
            configs_table = QTableWidget()
            configs_table.setColumnCount(7)
            configs_table.setHorizontalHeaderLabels([
                "Parameter Name", "File Path", "Current Value", "Relevance", "Source", "Line Number", "Reason"
            ])
            
            # Populate table
            for i, config in enumerate(suspected_configs):
                configs_table.insertRow(i)
                
                # Parameter name
                param_name = config.get('param_name', 'Unknown')
                configs_table.setItem(i, 0, QTableWidgetItem(param_name))
                
                # File path
                file_path = config.get('file_path', 'Unknown')
                configs_table.setItem(i, 1, QTableWidgetItem(file_path))
                
                # Current value (show more content since we have text wrapping)
                param_value = config.get('param_value', 'Unknown')
                # Show more content but truncate very long values
                if len(str(param_value)) > 100:
                    param_value = str(param_value)[:97] + "..."
                configs_table.setItem(i, 2, QTableWidgetItem(str(param_value)))
                
                # Relevance score
                score = config.get('relevance_score', 0)
                configs_table.setItem(i, 3, QTableWidgetItem(f"{score:.2f}"))
                
                # Source
                source = config.get('source', 'Unknown')
                meaningful_source = self.get_meaningful_source_name(source)
                configs_table.setItem(i, 4, QTableWidgetItem(meaningful_source))
                
                # Line number
                line_number = config.get('line_number', 'Unknown')
                configs_table.setItem(i, 5, QTableWidgetItem(str(line_number)))
                
                # Reason
                reason = config.get('reason', 'No reason provided')
                configs_table.setItem(i, 6, QTableWidgetItem(reason))
            
            # Adjust column widths and enable text wrapping
            configs_table.resizeColumnsToContents()
            configs_table.setAlternatingRowColors(True)
            configs_table.setSelectionBehavior(QTableWidget.SelectRows)
            
            # Enable text wrapping for all columns
            configs_table.setWordWrap(True)
            configs_table.setTextElideMode(Qt.ElideNone)
            
            # Set specific column widths for better display (reason column gets most space)
            configs_table.setColumnWidth(0, 200)  # Parameter Name
            configs_table.setColumnWidth(1, 250)  # File Path
            configs_table.setColumnWidth(2, 160)  # Current Value
            configs_table.setColumnWidth(3, 85)   # Relevance
            configs_table.setColumnWidth(4, 100)  # Source
            configs_table.setColumnWidth(5, 70)   # Line Number
            configs_table.setColumnWidth(6, 350)  # Reason (widest column for most content)
            
            # Set row height to accommodate wrapped text
            configs_table.verticalHeader().setDefaultSectionSize(80)
            
            # Apply professional blue theme styling to the configs table
            configs_table.setStyleSheet("""
                QTableWidget {
                    background-color: #ffffff;
                    alternate-background-color: #f0f4f8;
                    gridline-color: #b3d9ff;
                    border: 1px solid #b3d9ff;
                    border-radius: 4px;
                }
                QTableWidget::item {
                    padding: 8px;
                    border: none;
                    color: #1e3a5f;
                }
                QTableWidget::item:selected {
                    background-color: #004080;
                    color: #ffffff;
                }
                QTableWidget::item:hover {
                    background-color: #e6f2ff;
                }
                QHeaderView::section {
                    background-color: #1e3a5f;
                    color: #ffffff;
                    padding: 10px;
                    border: none;
                    font-weight: bold;
                    text-align: center;
                }
                QHeaderView::section:first {
                    border-top-left-radius: 4px;
                }
                QHeaderView::section:last {
                    border-top-right-radius: 4px;
                }
            """)
            
            # Add double-click handler to show full context
            configs_table.itemDoubleClicked.connect(lambda item: self.show_config_full_context(item, suspected_configs))
            
            layout.addWidget(configs_table)
        
        tab.setLayout(layout)
        return tab

    def create_context_artifacts_tab(self, data):
        """Create tab showing context information"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Deployment Context Section (same for both crash and normal analysis)
        deployment_context = data.get('deployment_context')
        if deployment_context:
            context_group = QGroupBox("Deployment Context")
            context_layout = QVBoxLayout()
            
            # Role
            role = deployment_context.get('role', 'Unknown')
            role_label = QLabel(f"<b>Role:</b> {role}")
            context_layout.addWidget(role_label)
            
            # Network parameters - Load from error_patterns_structured.json
            network_config = {}
            try:
                with open('Error_fixing_pipelin/database/error_patterns_structured.json', 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    deployment_context_json = data.get('deployment_context', {})
                    network_config = {
                        'CU IP Address': deployment_context_json.get('cu_ip_address', 'Unknown'),
                        'DU IP Address': deployment_context_json.get('du_ip_address', 'Unknown'),
                        'gNB IP Address': deployment_context_json.get('gnb_ip_address', 'Unknown'),
                        'GNB IP Address for NG AMF': deployment_context_json.get('gnb_ip_address', 'Unknown'),
                        'AMF IP Address': deployment_context_json.get('amf_ip_address', 'Unknown'),
                        'Core Network Machine IP': deployment_context_json.get('core_network_machine_ip', 'Unknown'),
                        'Local SCTP Port C': deployment_context_json.get('local_s_portc', 'Unknown'),
                        'Local SCTP Port D': deployment_context_json.get('local_s_portd', 'Unknown'),
                        'Remote SCTP Port C': deployment_context_json.get('remote_s_portc', 'Unknown'),
                        'Remote SCTP Port D': deployment_context_json.get('remote_s_portd', 'Unknown'),
                        'NSSAI SST': deployment_context_json.get('nssai_sst', 'Unknown'),
                        'NSSAI SD': deployment_context_json.get('nssai_sd', 'Unknown'),
                        'NMC Size': deployment_context_json.get('nmc_size', 'Unknown'),
                        'DNN': deployment_context_json.get('dnn', 'Unknown')
                    }
            except Exception as e:
                # Fallback to original network_params if JSON file not available
                network_params = deployment_context.get('network_params', {})
                if network_params:
                    network_config = {key: value for key, value in network_params.items()}
            
            if network_config:
                network_label = QLabel("<b>Network Parameters:</b>")
                context_layout.addWidget(network_label)
                
                for key, value in network_config.items():
                    param_label = QLabel(f"  • {key}: {value}")
                    context_layout.addWidget(param_label)
            
            # Deployment Commands Section
            try:
                with open('Error_fixing_pipelin/database/error_patterns_structured.json', 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    deployment_context_json = data.get('deployment_context', {})
                    
                    cu_command = deployment_context_json.get('Deploy_command_cu_gnb_conf', '')
                    du_command = deployment_context_json.get('Deploy_command_du_gnb_conf', '')
                    
                    if cu_command or du_command:
                        commands_label = QLabel("<b>Deployment Commands:</b>")
                        context_layout.addWidget(commands_label)
                        
                        if cu_command:
                            cu_label = QLabel(f"  🚀 <b>CU Deployment:</b> {cu_command}")
                            cu_label.setWordWrap(True)
                            cu_label.setStyleSheet("color: #0066cc; font-weight: bold;")
                            context_layout.addWidget(cu_label)
                        
                        if du_command:
                            du_label = QLabel(f"  🚀 <b>DU Deployment:</b> {du_command}")
                            du_label.setWordWrap(True)
                            du_label.setStyleSheet("color: #0066cc; font-weight: bold;")
                            context_layout.addWidget(du_label)
                        
                        # Add guidance about config files
                        guidance_label = QLabel("<b>📁 Config File Guidance:</b>")
                        context_layout.addWidget(guidance_label)
                        
                        guidance_text = QLabel("  • CU errors → cu_gnb.conf only<br/>  • DU errors → du_gnb.conf only<br/>  • Only these two config files are supported for configuration changes")
                        guidance_text.setWordWrap(True)
                        guidance_text.setStyleSheet("color: #1e3a5f; font-style: italic;")
                        context_layout.addWidget(guidance_text)
                        
            except Exception as e:
                print(f"Could not load deployment commands: {e}")
            
            # Active configs
            active_configs = deployment_context.get('active_configs', [])
            if active_configs:
                configs_label = QLabel(f"<b>Active Configurations:</b> {len(active_configs)} found")
                context_layout.addWidget(configs_label)
                
                for i, config in enumerate(active_configs[:5], 1):  # Show first 5
                    if isinstance(config, dict):
                        config_info = f"  {i}. {config.get('used', 'Unknown')}"
                        config_label = QLabel(config_info)
                        config_label.setWordWrap(True)
                        context_layout.addWidget(config_label)
            
            context_group.setLayout(context_layout)
            layout.addWidget(context_group)
        
        # Investigation Steps Section
        phase3_data = data.get('phase3_fixes', {})
        fix_suggestion = phase3_data.get('fix_suggestion', {})
        investigation_steps = fix_suggestion.get('investigation_steps', [])
        
        if investigation_steps:
            steps_group = QGroupBox("Investigation Steps")
            steps_layout = QVBoxLayout()
            
            for i, step in enumerate(investigation_steps, 1):
                step_label = QLabel(f"{i}. {step}")
                step_label.setWordWrap(True)
                steps_layout.addWidget(step_label)
            
            steps_group.setLayout(steps_layout)
            layout.addWidget(steps_group)
        
        layout.addStretch()
        tab.setLayout(layout)
        return tab

    def create_analysis_summary_tab(self, data):
        """Create tab showing analysis summary"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # Summary statistics
        summary_group = QGroupBox("Analysis Summary")
        summary_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: #1e3a5f;
                border: 2px solid #b3d9ff;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
                background-color: #ffffff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 10px 0 10px;
                background-color: #f0f4f8;
                color: #1e3a5f;
            }
        """)
        summary_layout = QVBoxLayout()
        
        # Phase 2 summary (filtered by Phase 3)
        phase2_data = data.get('phase2_analysis', {})
        all_suspected_functions = phase2_data.get('suspected_functions', [])
        all_suspected_configs = phase2_data.get('suspected_configs', [])
        
        # Get Phase 3 suspected names
        phase3_data = data.get('phase3_fixes', {})
        fix_suggestion = phase3_data.get('fix_suggestion', {})
        phase3_suspected_function_names = fix_suggestion.get('suspected_functions', [])
        phase3_suspected_config_names = fix_suggestion.get('suspected_configs', [])
        
        # Filter and count functions
        filtered_functions = []
        if phase3_suspected_function_names:
            for func in all_suspected_functions:
                func_name = func.get('function_name', '') if isinstance(func, dict) else str(func)
                if func_name in phase3_suspected_function_names:
                    filtered_functions.append(func)
        
        # Filter and count configs
        filtered_configs = []
        if phase3_suspected_config_names:
            for config in all_suspected_configs:
                config_name = config.get('param_name', '') if isinstance(config, dict) else str(config)
                if config_name in phase3_suspected_config_names:
                    filtered_configs.append(config)
        
        functions_count = len(filtered_functions)
        configs_count = len(filtered_configs)
        suspected_functions = filtered_functions
        suspected_configs = filtered_configs
        retrieval_method = phase2_data.get('retrieval_method', 'standard')
        
        phase2_label = QLabel(f"<b>Phase 2 - Error Analysis (Phase 3 Filtered):</b>")
        summary_layout.addWidget(phase2_label)
        summary_layout.addWidget(QLabel(f"  • Suspected Functions (from Phase 3): {functions_count}"))
        summary_layout.addWidget(QLabel(f"  • Suspected Configurations (from Phase 3): {configs_count}"))
        summary_layout.addWidget(QLabel(f"  • Retrieval Method: {retrieval_method}"))
        
        # Show top configs if available
        if configs_count > 0:
            top_configs_label = QLabel(f"<b>Top Suspected Configurations:</b>")
            summary_layout.addWidget(top_configs_label)
            for i, config in enumerate(suspected_configs[:3], 1):
                param_name = config.get('param_name', 'Unknown')
                score = config.get('relevance_score', 0)
                source = config.get('source', 'unknown')
                meaningful_source = self.get_meaningful_source_name(source)
                summary_layout.addWidget(QLabel(f"    {i}. {param_name} (score: {score:.2f}, source: {meaningful_source})"))
        
        # Phase 3 summary
        phase3_data = data.get('phase3_fixes', {})
        fix_suggestion = phase3_data.get('fix_suggestion', {})
        code_patches = fix_suggestion.get('code_patches', [])
        config_patches = fix_suggestion.get('config_patches', [])
        
        phase3_label = QLabel(f"<b>Phase 3 - Fix Suggestions:</b>")
        summary_layout.addWidget(phase3_label)
        summary_layout.addWidget(QLabel(f"  • Code Patches: {len(code_patches)}"))
        summary_layout.addWidget(QLabel(f"  • Config Patches: {len(config_patches)}"))
        summary_layout.addWidget(QLabel(f"  • Root Cause Analysis: {'Available' if fix_suggestion.get('reason') else 'Not Available'}"))
        
        # Terminal commands summary
        terminal_commands = data.get('phase4_commands', {}).get('terminal_commands', [])
        if terminal_commands:
            commands_label = QLabel(f"<b>Phase 4 - Terminal Commands:</b>")
            summary_layout.addWidget(commands_label)
            summary_layout.addWidget(QLabel(f"  • Commands Generated: {len(terminal_commands)}"))
        
        summary_group.setLayout(summary_layout)
        layout.addWidget(summary_group)
        
        # Overall summary
        overall_group = QGroupBox("Overall Analysis Status")
        overall_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: #1e3a5f;
                border: 2px solid #b3d9ff;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
                background-color: #ffffff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 10px 0 10px;
                background-color: #f0f4f8;
                color: #1e3a5f;
            }
        """)
        overall_layout = QVBoxLayout()
        
        context_aware = data.get('summary', {}).get('context_aware', False)
        fix_available = bool(fix_suggestion.get('config_fix') or fix_suggestion.get('code_patches'))
        deployment_context = data.get('deployment_context')
        
        status_label = QLabel(f"<b>Analysis Status:</b>")
        overall_layout.addWidget(status_label)
        overall_layout.addWidget(QLabel(f"  • Context-Aware Analysis: {'Yes' if context_aware else 'No'}"))
        overall_layout.addWidget(QLabel(f"  • Fix Suggestions Available: {'Yes' if fix_available else 'No'}"))
        overall_layout.addWidget(QLabel(f"  • Deployment Context: {'Available' if deployment_context else 'Not Available'}"))
        
        overall_group.setLayout(overall_layout)
        layout.addWidget(overall_group)
        
        layout.addStretch()
        tab.setLayout(layout)
        return tab

    def create_3gpp_spec_tab(self, data):
        """Create tab showing 3GPP specification context with modern formatting"""
        tab = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Get specification context from different possible locations
        spec_context = None
        
        # Try to get from phase3_fixes first
        phase3_data = data.get('phase3_fixes', {})
        fix_suggestion = phase3_data.get('fix_suggestion', {})
        if fix_suggestion.get('specification_context'):
            spec_context = fix_suggestion['specification_context']
        
        # Try to get from the complete data
        if not spec_context and data.get('specification_context'):
            spec_context = data['specification_context']
        
        # Try to get from phase2_analysis
        if not spec_context:
            phase2_data = data.get('phase2_analysis', {})
            if phase2_data.get('specification_context'):
                spec_context = phase2_data['specification_context']
        
        if not spec_context or spec_context.strip() == "- 3GPP specification context not available":
            no_data_label = QLabel("No 3GPP specification context found in the analysis.")
            no_data_label.setAlignment(Qt.AlignCenter)
            no_data_label.setStyleSheet("color: #888; font-style: italic; padding: 40px; font-size: 14px;")
            layout.addWidget(no_data_label)
        else:
            # Create scroll area for content
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setMinimumHeight(600)
            scroll.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            scroll.setStyleSheet("""
                QScrollArea {
                    border: none;
                    background-color: #f8f9fa;
                }
            """)
            
            content_widget = QWidget()
            content_layout = QVBoxLayout(content_widget)
            content_layout.setSpacing(20)
            content_layout.setContentsMargins(20, 20, 20, 20)
            
            # Parse and create modern formatted sections
            spec_sections = self.parse_3gpp_spec_context(spec_context)
            
            if spec_sections:
                # Header
                header_label = QLabel("📚 3GPP Specification Reference - 38.331")
                header_label.setStyleSheet("""
                    QLabel {
                        font-size: 18px;
                        font-weight: bold;
                        color: #1e3a5f;
                        padding: 15px 0;
                        border-bottom: 2px solid #0066cc;
                        margin-bottom: 20px;
                    }
                """)
                content_layout.addWidget(header_label)
                
                # Add each section
                for section in spec_sections:
                    section_widget = self.create_spec_section_widget(section)
                    content_layout.addWidget(section_widget)
                
                # Summary
                summary_widget = self.create_spec_summary_widget(len(spec_sections))
                content_layout.addWidget(summary_widget)
            
            content_layout.addStretch()
            scroll.setWidget(content_widget)
            layout.addWidget(scroll, 1)
        
        tab.setLayout(layout)
        return tab

    def format_3gpp_spec_context(self, spec_context):
        """Format the 3GPP specification context for better readability"""
        if not spec_context:
            return "No specification context available"
        
        # Handle the actual format with "=== Extraction Section No. : X ===" headers
        sections = spec_context.strip().split("=== Extraction Section No.")
        formatted_text = "3GPP SPECIFICATION REFERENCE\n"
        formatted_text += "=" * 80 + "\n\n"
        
        section_count = 0
        
        for section in sections:
            if not section.strip():
                continue
                
            section_count += 1
            lines = section.strip().split('\n')
            
            # Extract section number from the first line
            first_line = lines[0].strip() if lines else ""
            if first_line.startswith(": "):
                section_num = first_line.replace(": ", "").replace(" ===", "").strip()
                formatted_text += f"EXTRACTION SECTION {section_num}\n"
                formatted_text += "=" * 60 + "\n\n"
            else:
                formatted_text += f"EXTRACTION SECTION {section_count}\n"
                formatted_text += "=" * 60 + "\n\n"
            
            # Process the content
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # Skip the section header line
                if line.startswith(": ") and "===" in line:
                    continue
                
                # Check for page information
                if line.startswith("Page:"):
                    page_info = line
                    formatted_text += f"Page: {page_info.replace('Page:', '').strip()}\n\n"
                    
                elif line.startswith("## "):
                    # Main section header (## Section Title)
                    section_title = line.replace("## ", "").strip()
                    formatted_text += f"SECTION TITLE: {section_title}\n"
                    formatted_text += "-" * 50 + "\n"
                    
                elif line.startswith("### "):
                    # Subsection header (### Section Number)
                    section_number = line.replace("### ", "").strip()
                    formatted_text += f"SUBSECTION: {section_number}\n\n"
                    
                else:
                    # Regular content - add proper spacing and formatting
                    if line.startswith("--") or line.startswith("ETSI"):
                        # ASN.1 or technical content - keep as is with indentation
                        formatted_text += f"   {line}\n"
                    else:
                        # Regular text content
                        formatted_text += f"{line}\n"
            
            # Add spacing between sections
            formatted_text += "\n" + "-" * 80 + "\n\n"
        
        # Add summary at the end
        if section_count > 0:
            formatted_text += "SUMMARY\n"
            formatted_text += "=" * 40 + "\n"
            formatted_text += f"Total specification sections found: {section_count}\n"
            formatted_text += "These sections contain relevant 3GPP standards information\n"
            formatted_text += "related to the error analysis and potential fixes.\n"
        
        return formatted_text

    def extract_readable_text_from_pdf_page(self, page_number, section_number=None):
        """Extract clean, readable text from the PDF by searching for section number
        
        Returns:
            tuple: (text, actual_page_number) where actual_page_number is the real PDF page found
        """
        try:
            pdf_paths = [
                os.path.join(os.path.dirname(__file__), 'Error_fixing_pipelin', 'resources', 'ts_138331v160301p.pdf'),
                os.path.join(os.path.dirname(__file__), 'backend', 'resources', 'ts_138331v160301p.pdf')
            ]
            
            pdf_path = None
            for path in pdf_paths:
                if os.path.exists(path):
                    pdf_path = path
                    break
            
            if not pdf_path:
                return ("", page_number)
            
            pdf_doc = fitz.open(pdf_path)
            found_page_idx = None
            found_text = ""
            
            # Search for section number in PDF
            if section_number:
                search_patterns = [
                    f"{section_number} ",
                    f"{section_number}\n",
                    f"\n{section_number} ",
                ]
                
                start_page = max(0, int(page_number) - 20) if page_number else 0
                end_page = min(len(pdf_doc), start_page + 40)
                
                for page_idx in range(start_page, end_page):
                    page = pdf_doc[page_idx]
                    page_text = page.get_text()
                    
                    for pattern in search_patterns:
                        if pattern in page_text:
                            found_page_idx = page_idx
                            found_text = page_text
                            break
                    
                    if found_page_idx is not None:
                        break
            
            # Fallback to original page number if section not found
            if not found_text:
                page_idx = int(page_number) - 1
                if 0 <= page_idx < len(pdf_doc):
                    page = pdf_doc[page_idx]
                    found_text = page.get_text()
                    found_page_idx = page_idx
            
            pdf_doc.close()
            
            if not found_text:
                return ("", page_number)
            
            # Clean up the text
            # Remove ETSI headers/footers
            lines = found_text.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                # Skip header/footer lines
                if 'ETSI' in line and 'TS 138 331' in line:
                    continue
                if line.startswith('3GPP TS 38.331 version'):
                    continue
                if line and not line.isdigit():  # Skip page numbers
                    cleaned_lines.append(line)
            
            text = ' '.join(cleaned_lines)
            
            # Remove excessive whitespace
            text = ' '.join(text.split())
            
            # Limit to reasonable length (first 1500 chars or so)
            if len(text) > 1500:
                text = text[:1500] + "..."
            
            # Return text and the actual page number found (1-indexed for display)
            actual_page_display = found_page_idx + 1 if found_page_idx is not None else page_number
            return (text, actual_page_display)
            
        except Exception as e:
            return ("", page_number)
    
    def parse_3gpp_original_chunks(self, chunks, filtered_spec_context=None):
        """Parse original 3GPP chunks, selecting top 3 by similarity or LLM-filtered sections"""
        if not chunks:
            return []
        
        # Check if LLM filtered specific sections
        selected_sections = None
        if filtered_spec_context and "=== Extraction Section No." in filtered_spec_context:
            import re
            matches = re.findall(r'=== Extraction Section No\. : (\d+) ===', filtered_spec_context)
            selected_sections = set(int(m) for m in matches)
        else:
            # Use top 3 by similarity score
            chunks_sorted = sorted(chunks, key=lambda x: x.get('similarity', 0), reverse=True)
            chunks = chunks_sorted[:3]
        
        sections = []
        
        # Process chunks (only selected ones if filtering is enabled)
        for idx, chunk in enumerate(chunks, start=1):
            # Skip if this section wasn't selected by LLM
            if selected_sections is not None and idx not in selected_sections:
                continue
            
            raw_section_title = chunk.get('section_title', '')
            section_number = chunk.get('section_number', '')
            chunk_text_content = chunk.get('text', '').strip()
            page_number = chunk.get('page_number')
            
            # Extract readable text directly from PDF by searching for section number
            text_content, actual_page_number = self.extract_readable_text_from_pdf_page(page_number, section_number)
            
            # Fallback to chunk text if PDF extraction fails
            if not text_content or len(text_content) < 50:
                text_content = chunk_text_content
                actual_page_number = page_number
            
            # Extract clean title from section_title
            if raw_section_title:
                title_lines = raw_section_title.split('\n')
                first_line = title_lines[0].strip() if title_lines else ''
                
                if len(first_line) > 100:
                    sentence_end = first_line.find('. ')
                    if sentence_end > 0 and sentence_end < 100:
                        first_line = first_line[:sentence_end]
                    else:
                        first_line = first_line[:100]
                
                section_title_clean = first_line
            else:
                section_title_clean = ''
            
            # Create a cleaner title
            if section_number and section_title_clean:
                clean_title = f"{section_number} - {section_title_clean}"
            elif section_title_clean:
                clean_title = section_title_clean
            elif section_number:
                clean_title = f"Section {section_number}"
            else:
                clean_title = f"Section {idx}"
            
            # Truncate if too long
            if len(clean_title) > 80:
                clean_title = clean_title[:77] + "..."
            
            section = {
                'section_number': str(idx),
                'items': [{
                    'title': clean_title,
                    'page': str(actual_page_number),
                    'section_title': section_title_clean,
                    'content': text_content
                }]
            }
            sections.append(section)
        
        return sections
    
    def parse_3gpp_spec_context(self, spec_context):
        """Parse 3GPP specification context into structured sections"""
        if not spec_context:
            return []
        
        sections = []
        
        # Remove header if present
        if spec_context.startswith("3GPP SPECIFICATION REFERENCE"):
            lines = spec_context.split('\n')
            start_idx = 0
            for i, line in enumerate(lines):
                if line.strip().startswith('='):
                    start_idx = i + 1
                    break
            spec_context = '\n'.join(lines[start_idx:])
        
        # Split by extraction sections
        if "=== Extraction Section No." in spec_context:
            spec_parts = spec_context.strip().split("=== Extraction Section No.")
        else:
            spec_parts = [spec_context]
        
        for idx, part in enumerate(spec_parts, start=1):
            if not part.strip():
                continue
            
            current_section = {
                'section_number': str(idx),
                'items': []
            }
            
            lines = part.strip().split('\n')
            
            # Extract section number if present
            first_line = lines[0].strip() if lines else ""
            if first_line.startswith(": "):
                section_num = first_line.replace(": ", "").replace(" ===", "").strip()
                current_section['section_number'] = section_num
                lines = lines[1:]  # Remove the section header line
            
            # Parse content - handle two different formats
            current_item = None
            i = 0
            
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    i += 1
                    continue
                
                # Check for Format 1 title: **Text** followed by **Page:** or **Section Title:**
                is_format1_title = False
                if (line.startswith("**") and 
                    line.endswith("**") and 
                    not line.startswith("**Page:") and
                    not line.startswith("**Section Title:") and
                    not line.startswith("**Content:")):
                    # Lookahead to confirm it's a title (next line should be **Page:** or **Section Title:**)
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if (next_line.startswith("**Page:") or 
                            next_line.startswith("**Section Title:")):
                            is_format1_title = True
                    else:
                        # Last line, treat as title
                        is_format1_title = True
                
                # Format 1: **Title** followed by **Page:** **Section Title:** **Content:**
                if is_format1_title:
                    # New item header (Format 1)
                    if current_item and (current_item['page'] or current_item['section_title'] or current_item['content']):
                        current_section['items'].append(current_item)
                    current_item = {
                        'title': line.replace("**", "").strip(),
                        'page': '',
                        'section_title': '',
                        'content': ''
                    }
                    i += 1
                    
                elif line.startswith("**Page:**"):
                    if current_item:
                        current_item['page'] = line.replace("**Page:**", "").strip()
                    i += 1
                    
                elif line.startswith("**Section Title:**"):
                    if current_item:
                        current_item['section_title'] = line.replace("**Section Title:**", "").strip()
                    i += 1
                    
                elif line.startswith("**Content:**"):
                    if current_item:
                        # Get content on same line if any
                        content_on_line = line.replace("**Content:**", "").strip()
                        if content_on_line:
                            current_item['content'] = content_on_line
                        # Mark that we're now in content mode
                        i += 1
                        # Continue collecting content lines (including those starting with **)
                        content_lines_collected = 0
                        while i < len(lines):
                            next_line = lines[i].strip()
                            if not next_line:
                                i += 1
                                continue
                            
                            # Stop if we hit a new item (Format 1: **Title**)
                            # A new item is: **Something** followed by **Page:** or **Section Title:**
                            is_new_item_title = False
                            if (next_line.startswith("**") and 
                                next_line.endswith("**") and 
                                not next_line.startswith("**Page:") and
                                not next_line.startswith("**Section Title:") and
                                not next_line.startswith("**Content:")):
                                # Check if next line is **Page:** or **Section Title:**
                                # This confirms it's a new item, not just content with **
                                if i + 1 < len(lines):
                                    next_next_line = lines[i + 1].strip()
                                    if (next_next_line.startswith("**Page:") or 
                                        next_next_line.startswith("**Section Title:")):
                                        is_new_item_title = True
                                else:
                                    # If it's the last line and looks like a title, treat as new item
                                    is_new_item_title = True
                            
                            # Stop if we hit a new section marker
                            if is_new_item_title:
                                break
                            if next_line.startswith("**Section Title:") and next_line.endswith("**"):
                                break
                            if next_line.startswith("**Page:") and next_line.endswith("**"):
                                break
                            if next_line.startswith("=== Extraction Section No."):
                                break
                            
                            # It's content - add it
                            if current_item['content']:
                                current_item['content'] += " " + next_line
                            else:
                                current_item['content'] = next_line
                            i += 1
                        continue  # Don't increment i again
                    i += 1
                    
                # Format 2: **Section Title: value** followed by **Page: value** then content
                elif line.startswith("**Section Title:") and line.endswith("**"):
                    # New item header (Format 2)
                    if current_item and (current_item['page'] or current_item['section_title'] or current_item['content']):
                        current_section['items'].append(current_item)
                    
                    section_title_text = line.replace("**Section Title:", "").replace("**", "").strip()
                    current_item = {
                        'title': section_title_text,
                        'page': '',
                        'section_title': section_title_text,
                        'content': ''
                    }
                    i += 1
                    
                elif line.startswith("**Page:") and line.endswith("**"):
                    if current_item:
                        current_item['page'] = line.replace("**Page:", "").replace("**", "").strip()
                    i += 1
                    
                else:
                    # Content text (for Format 2 - content without **Content:** marker)
                    if current_item:
                        # Check if this is the start of a new item/section using same logic as is_format1_title
                        is_new_header = (
                            (line.startswith("**") and line.endswith("**") and 
                             not line.startswith("**Page:") and 
                             not line.startswith("**Section Title:") and 
                             not line.startswith("**Content:")) or  # Format 1 title
                            (line.startswith("**Section Title:") and line.endswith("**")) or  # Format 2 title
                            line.startswith("=== Extraction Section No.")  # New section
                        )
                        
                        if not is_new_header:
                            # It's content
                            if current_item['content']:
                                current_item['content'] += " " + line
                            else:
                                current_item['content'] = line
                    i += 1
            
            # Add last item
            if current_item and (current_item['page'] or current_item['section_title'] or current_item['content']):
                current_section['items'].append(current_item)
            
            if current_section['items']:
                sections.append(current_section)
        
        return sections

    def create_spec_section_widget(self, section):
        """Create a modern widget for a specification section"""
        section_widget = QGroupBox(f"📄 Extraction Section {section['section_number']}")
        section_widget.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: #1e3a5f;
                border: 1px solid #b3d9ff;
                border-radius: 4px;
                margin-top: 10px;
                padding: 15px;
                background-color: #ffffff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                padding: 4px 8px;
                color: #1e3a5f;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setSpacing(15)
        
        for item in section['items']:
            item_widget = self.create_spec_item_widget(item)
            layout.addWidget(item_widget)
        
        section_widget.setLayout(layout)
        return section_widget

    def create_spec_item_widget(self, item):
        """Create a modern widget for a specification item"""
        item_widget = QWidget()
        item_widget.setStyleSheet("""
            QWidget {
                background-color: #f8f9fa;
                border: 1px solid #e1e5ea;
                border-radius: 4px;
                padding: 15px;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setSpacing(8)
        
        # Title
        title_label = QLabel(f"🔹 {item['title']}")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                font-weight: bold;
                color: #0066cc;
                margin-bottom: 5px;
            }
        """)
        layout.addWidget(title_label)
        
        # Metadata row with View in Doc button
        metadata_layout = QHBoxLayout()
        metadata_layout.setSpacing(20)
        
        if item['page']:
            page_label = QLabel(f"📄 Page: {item['page']}")
            page_label.setStyleSheet("""
                QLabel {
                    font-size: 12px;
                    color: #666;
                    background-color: #e3f2fd;
                    padding: 4px 8px;
                    border-radius: 3px;
                }
            """)
            metadata_layout.addWidget(page_label)
        
        if item['section_title']:
            section_label = QLabel(f"📋 Section: {item['section_title']}")
            section_label.setStyleSheet("""
                QLabel {
                    font-size: 12px;
                    color: #666;
                    background-color: #e8f5e8;
                    padding: 4px 8px;
                    border-radius: 3px;
                }
            """)
            metadata_layout.addWidget(section_label)
        
        metadata_layout.addStretch()
        
        # Add "View in Doc" button if page number is available
        if item['page']:
            view_doc_btn = QPushButton("📖 View in Doc")
            view_doc_btn.setStyleSheet("""
                QPushButton {
                    background-color: #0066cc;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px 12px;
                    font-size: 12px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #0052a3;
                }
                QPushButton:pressed {
                    background-color: #003d7a;
                }
            """)
            view_doc_btn.clicked.connect(lambda: self.open_pdf_at_page(item['page']))
            metadata_layout.addWidget(view_doc_btn)
        
        layout.addLayout(metadata_layout)
        
        # Content
        if item['content']:
            content_label = QLabel(item['content'])
            content_label.setWordWrap(True)
            content_label.setStyleSheet("""
                QLabel {
                    font-size: 13px;
                    color: #1e3a5f;
                    line-height: 1.6;
                    padding: 10px;
                    background-color: #ffffff;
                    border: 1px solid #e1e5ea;
                    border-radius: 4px;
                }
            """)
            layout.addWidget(content_label)
        
        item_widget.setLayout(layout)
        return item_widget

    def create_spec_summary_widget(self, section_count):
        """Create a summary widget for the specification"""
        summary_widget = QGroupBox("📊 Summary")
        summary_widget.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: #1e3a5f;
                border: 1px solid #b3d9ff;
                border-radius: 4px;
                margin-top: 10px;
                padding: 15px;
                background-color: #ffffff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                padding: 4px 8px;
                color: #1e3a5f;
            }
        """)
        
        layout = QVBoxLayout()
        
        summary_text = f"""
        <div style='font-size: 13px; line-height: 1.6; color: #1e3a5f;'>
            <b>Total specification sections found:</b> {section_count}<br><br>
            These sections contain relevant 3GPP standards information related to the error analysis and potential fixes.
        </div>
        """
        
        summary_label = QLabel(summary_text)
        summary_label.setWordWrap(True)
        layout.addWidget(summary_label)
        
        summary_widget.setLayout(layout)
        return summary_widget

    def open_pdf_at_page(self, page_number):
        """Open the 3GPP specification PDF at a specific page in built-in viewer"""
        try:
            # Path to the PDF - try both locations
            pdf_paths = [
                os.path.join(os.path.dirname(__file__), 'Error_fixing_pipelin', 'resources', 'ts_138331v160301p.pdf'),
                os.path.join(os.path.dirname(__file__), 'backend', 'resources', 'ts_138331v160301p.pdf')
            ]
            
            pdf_path = None
            for path in pdf_paths:
                if os.path.exists(path):
                    pdf_path = path
                    break
            
            if not pdf_path:
                QMessageBox.warning(
                    self, 
                    "PDF Not Found", 
                    f"Could not find ts_138331v160301p.pdf in:\n{chr(10).join(pdf_paths)}"
                )
                return
            
            # Extract numeric page number
            page_num = str(page_number).strip()
            # Try to extract just the number if it contains extra text
            import re
            match = re.search(r'\d+', page_num)
            if match:
                page_num = match.group()
            
            page_num = int(page_num)
            
            # Open PDF to check page numbers
            pdf_doc = fitz.open(pdf_path)
            
            actual_page_index = None
            # Try to find the page with the matching page label (printed page number)
            for i, page in enumerate(pdf_doc):
                page_label = page.get_label()
                if page_label == str(page_num):
                    actual_page_index = i
                    break
            
            pdf_doc.close()
            
            # If not found, try common offsets for 3GPP specs
            if actual_page_index is None:
                # 3GPP specs typically have front matter, so real content starts around page 10-20
                actual_page_index = page_num + 10  # Common offset for TS 38.331
            
            # Open the built-in PDF viewer
            self.show_pdf_viewer(pdf_path, actual_page_index)
                
        except ValueError as e:
            QMessageBox.warning(
                self,
                "Invalid Page Number",
                f"Could not parse page number '{page_number}'.\n\nError: {str(e)}"
            )
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error Opening PDF",
                f"An error occurred while trying to open the PDF:\n\n{str(e)}"
            )
            import traceback
            traceback.print_exc()
    
    def show_pdf_viewer(self, pdf_path, page_index):
        """Show a built-in PDF viewer dialog with the specified page"""
        try:
            # Create dialog
            dialog = QDialog(self)
            dialog.setWindowTitle(f"3GPP Specification - TS 38.331")
            dialog.setMinimumSize(1000, 800)
            
            # Main layout
            layout = QVBoxLayout()
            layout.setContentsMargins(0, 0, 0, 0)
            layout.setSpacing(0)
            
            # Open PDF
            pdf_doc = fitz.open(pdf_path)
            total_pages = len(pdf_doc)
            
            # Validate page index
            if page_index < 0:
                page_index = 0
            elif page_index >= total_pages:
                page_index = total_pages - 1
            
            # Toolbar
            toolbar = QWidget()
            toolbar.setStyleSheet("""
                QWidget {
                    background-color: #f0f0f0;
                    border-bottom: 1px solid #ccc;
                    padding: 8px;
                }
            """)
            toolbar_layout = QHBoxLayout(toolbar)
            toolbar_layout.setContentsMargins(10, 5, 10, 5)
            
            # Page navigation controls
            page_label = QLabel(f"Page:")
            page_label.setStyleSheet("font-weight: bold; color: #333;")
            toolbar_layout.addWidget(page_label)
            
            page_input = QLineEdit()
            page_input.setMaximumWidth(60)
            page_input.setText(str(page_index + 1))
            page_input.setStyleSheet("""
                QLineEdit {
                    padding: 4px 8px;
                    border: 1px solid #ccc;
                    border-radius: 3px;
                    background-color: white;
                }
            """)
            toolbar_layout.addWidget(page_input)
            
            total_label = QLabel(f"of {total_pages}")
            total_label.setStyleSheet("color: #666;")
            toolbar_layout.addWidget(total_label)
            
            # Navigation buttons
            prev_btn = QPushButton("◀ Previous")
            prev_btn.setStyleSheet("""
                QPushButton {
                    background-color: #0066cc;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px 12px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #0052a3;
                }
                QPushButton:disabled {
                    background-color: #ccc;
                }
            """)
            toolbar_layout.addWidget(prev_btn)
            
            next_btn = QPushButton("Next ▶")
            next_btn.setStyleSheet("""
                QPushButton {
                    background-color: #0066cc;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px 12px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #0052a3;
                }
                QPushButton:disabled {
                    background-color: #ccc;
                }
            """)
            toolbar_layout.addWidget(next_btn)
            
            toolbar_layout.addStretch()
            
            # Zoom controls
            zoom_label = QLabel("Zoom:")
            zoom_label.setStyleSheet("font-weight: bold; color: #333;")
            toolbar_layout.addWidget(zoom_label)
            
            zoom_combo = QComboBox()
            zoom_combo.addItems(["50%", "75%", "100%", "125%", "150%", "200%"])
            zoom_combo.setCurrentText("100%")
            zoom_combo.setStyleSheet("""
                QComboBox {
                    padding: 4px 8px;
                    border: 1px solid #ccc;
                    border-radius: 3px;
                    background-color: white;
                }
            """)
            toolbar_layout.addWidget(zoom_combo)
            
            # Open in external viewer button
            external_btn = QPushButton("🔗 Open in External Viewer")
            external_btn.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px 12px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
            """)
            toolbar_layout.addWidget(external_btn)
            
            # Keyboard shortcuts hint
            hint_label = QLabel("💡 Tip: Use ← → or PgUp/PgDown to navigate, Ctrl+/- to zoom")
            hint_label.setStyleSheet("color: #666; font-size: 11px; font-style: italic; margin-left: 15px;")
            toolbar_layout.addWidget(hint_label)
            
            layout.addWidget(toolbar)
            
            # Scroll area for PDF pages
            scroll = QScrollArea()
            scroll.setWidgetResizable(False)  # Don't resize widget - allow scrolling
            scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            scroll.setAlignment(Qt.AlignCenter)
            scroll.setStyleSheet("""
                QScrollArea {
                    border: none;
                    background-color: #525252;
                }
            """)
            
            # Label to display PDF page directly (no container needed)
            pdf_label = QLabel()
            pdf_label.setStyleSheet("""
                QLabel {
                    background-color: white;
                    border: 1px solid #333;
                    padding: 10px;
                    margin: 20px;
                }
            """)
            pdf_label.setScaledContents(False)
            
            scroll.setWidget(pdf_label)
            layout.addWidget(scroll, 1)
            
            dialog.setLayout(layout)
            
            # Function to render and display a page
            current_zoom = [1.0]  # Use list to modify in nested function
            
            def render_page(page_idx):
                if 0 <= page_idx < total_pages:
                    page = pdf_doc[page_idx]
                    zoom = current_zoom[0]
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to QImage
                    img = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format_RGB888)
                    pixmap = QPixmap.fromImage(img)
                    
                    pdf_label.setPixmap(pixmap)
                    pdf_label.adjustSize()  # Ensure label resizes to pixmap size
                    page_input.setText(str(page_idx + 1))
                    
                    # Update button states
                    prev_btn.setEnabled(page_idx > 0)
                    next_btn.setEnabled(page_idx < total_pages - 1)
            
            # Connect navigation
            current_page = [page_index]  # Use list to modify in nested function
            
            def go_to_page(idx):
                if 0 <= idx < total_pages:
                    current_page[0] = idx
                    render_page(idx)
            
            def prev_page():
                go_to_page(current_page[0] - 1)
            
            def next_page():
                go_to_page(current_page[0] + 1)
            
            def page_input_changed():
                try:
                    page_num = int(page_input.text())
                    if 1 <= page_num <= total_pages:
                        go_to_page(page_num - 1)
                except ValueError:
                    pass
            
            def zoom_changed(text):
                try:
                    zoom_percent = int(text.replace("%", ""))
                    current_zoom[0] = zoom_percent / 100.0
                    render_page(current_page[0])
                except ValueError:
                    pass
            
            def open_external():
                if sys.platform.startswith('win'):
                    os.startfile(pdf_path)
                elif sys.platform.startswith('darwin'):
                    subprocess.Popen(['open', pdf_path])
                else:
                    subprocess.Popen(['xdg-open', pdf_path])
            
            prev_btn.clicked.connect(prev_page)
            next_btn.clicked.connect(next_page)
            page_input.returnPressed.connect(page_input_changed)
            zoom_combo.currentTextChanged.connect(zoom_changed)
            external_btn.clicked.connect(open_external)
            
            # Add keyboard shortcuts
            # Navigation shortcuts
            QShortcut(QKeySequence("Left"), dialog).activated.connect(prev_page)
            QShortcut(QKeySequence("Right"), dialog).activated.connect(next_page)
            QShortcut(QKeySequence("PgUp"), dialog).activated.connect(prev_page)
            QShortcut(QKeySequence("PgDown"), dialog).activated.connect(next_page)
            
            # Zoom shortcuts
            def zoom_in():
                current_index = zoom_combo.currentIndex()
                if current_index < zoom_combo.count() - 1:
                    zoom_combo.setCurrentIndex(current_index + 1)
            
            def zoom_out():
                current_index = zoom_combo.currentIndex()
                if current_index > 0:
                    zoom_combo.setCurrentIndex(current_index - 1)
            
            QShortcut(QKeySequence("Ctrl++"), dialog).activated.connect(zoom_in)
            QShortcut(QKeySequence("Ctrl+="), dialog).activated.connect(zoom_in)
            QShortcut(QKeySequence("Ctrl+-"), dialog).activated.connect(zoom_out)
            
            # Render initial page
            render_page(page_index)
            
            # Show dialog
            dialog.exec_()
            
            # Cleanup
            pdf_doc.close()
            
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error Opening PDF",
                f"Failed to open PDF viewer:\n\n{str(e)}"
            )
            import traceback
            traceback.print_exc()

    def show_function_full_context(self, item, suspected_functions):
        """Show full context for a selected function"""
        try:
            row = item.row()
            if row < len(suspected_functions):
                func = suspected_functions[row]
                
                dialog = QDialog(self)
                dialog.setWindowTitle(f"Function Context: {func.get('function_name', 'Unknown')}")
                dialog.setMinimumSize(800, 600)
                
                # Apply professional blue theme styling to the context dialog
                dialog.setStyleSheet("""
                    QDialog {
                        background-color: #f0f4f8;
                        color: #1e3a5f;
                    }
                """)
                
                layout = QVBoxLayout()
                
                # Function details
                source = func.get('source', 'Unknown')
                meaningful_source = self.get_meaningful_source_name(source)
                details_text = f"""
Function Name: {func.get('function_name', 'Unknown')}
File Path: {func.get('file_path', 'Unknown')}
Relevance Score: {func.get('relevance_score', 0):.2f}
Source: {meaningful_source}
Reason: {func.get('reason', 'No reason provided')}

Code Context:
{func.get('code_snippet', 'No code snippet available')}
                """
                
                text_edit = QTextEdit()
                text_edit.setPlainText(details_text)
                text_edit.setReadOnly(True)
                text_edit.setStyleSheet("""
                    QTextEdit {
                        font-family: 'Consolas', 'Monaco', monospace;
                        font-size: 11px;
                        background-color: #ffffff;
                        border: 2px solid #b3d9ff;
                        border-radius: 6px;
                        padding: 15px;
                        color: #1e3a5f;
                        line-height: 1.4;
                    }
                    QTextEdit:focus {
                        border-color: #0066cc;
                    }
                """)
                
                layout.addWidget(text_edit)
                
                # Close button
                close_btn = QPushButton("Close")
                close_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #1e3a5f;
                        color: #ffffff;
                        border: none;
                        padding: 10px 20px;
                        border-radius: 4px;
                        font-weight: bold;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: #2c4a6b;
                    }
                    QPushButton:pressed {
                        background-color: #153157;
                    }
                """)
                close_btn.clicked.connect(dialog.accept)
                layout.addWidget(close_btn)
                
                dialog.setLayout(layout)
                dialog.exec_()
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to show function context: {str(e)}")

    def show_config_full_context(self, item, suspected_configs):
        """Show full context for a selected configuration"""
        try:
            row = item.row()
            if row < len(suspected_configs):
                config = suspected_configs[row]
                
                dialog = QDialog(self)
                dialog.setWindowTitle(f"Configuration Context: {config.get('param_name', 'Unknown')}")
                dialog.setMinimumSize(800, 600)
                
                # Apply professional blue theme styling to the context dialog
                dialog.setStyleSheet("""
                    QDialog {
                        background-color: #f0f4f8;
                        color: #1e3a5f;
                    }
                """)
                
                layout = QVBoxLayout()
                
                # Configuration details
                source = config.get('source', 'Unknown')
                meaningful_source = self.get_meaningful_source_name(source)
                details_text = f"""
Parameter Name: {config.get('param_name', 'Unknown')}
File Path: {config.get('file_path', 'Unknown')}
Current Value: {config.get('param_value', 'Unknown')}
Relevance Score: {config.get('relevance_score', 0):.2f}
Source: {meaningful_source}
Line Number: {config.get('line_number', 'Unknown')}
Reason: {config.get('reason', 'No reason provided')}

Configuration Context:
{config.get('config_context', 'No configuration context available')}
                """
                
                text_edit = QTextEdit()
                text_edit.setPlainText(details_text)
                text_edit.setReadOnly(True)
                text_edit.setStyleSheet("""
                    QTextEdit {
                        font-family: 'Consolas', 'Monaco', monospace;
                        font-size: 11px;
                        background-color: #ffffff;
                        border: 2px solid #b3d9ff;
                        border-radius: 6px;
                        padding: 15px;
                        color: #1e3a5f;
                        line-height: 1.4;
                    }
                    QTextEdit:focus {
                        border-color: #0066cc;
                    }
                """)
                
                layout.addWidget(text_edit)
                
                # Close button
                close_btn = QPushButton("Close")
                close_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #1e3a5f;
                        color: #ffffff;
                        border: none;
                        padding: 10px 20px;
                        border-radius: 4px;
                        font-weight: bold;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: #2c4a6b;
                    }
                    QPushButton:pressed {
                        background-color: #153157;
                    }
                """)
                close_btn.clicked.connect(dialog.accept)
                layout.addWidget(close_btn)
                
                dialog.setLayout(layout)
                dialog.exec_()
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to show configuration context: {str(e)}")

    def export_artifacts(self, data):
        """Export artifacts to a file"""
        try:
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Export Error Artifacts", 
                "error_artifacts.json", 
                "JSON Files (*.json)"
            )
            
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
                QMessageBox.information(self, "Export Successful", f"Artifacts exported to: {file_path}")
                
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export artifacts: {str(e)}")

    def run_terminal_commands(self):
        """Run the generated terminal commands from the RCA analysis."""
        try:
            # Find the most recent fix_suggestions.json file
            resources_dir = "backend/resources"
            fix_suggestions_file = None
            
            if os.path.exists(resources_dir):
                json_files = [f for f in os.listdir(resources_dir) if f.startswith('fix_suggestions_') and f.endswith('.json')]
                if json_files:
                    # Get the most recent file
                    json_files.sort(reverse=True)
                    fix_suggestions_file = os.path.join(resources_dir, json_files[0])
            
            if not fix_suggestions_file or not os.path.exists(fix_suggestions_file):
                QMessageBox.warning(self, "Fix Suggestions Not Found", "Could not find fix_suggestions.json file. Please run RCA analysis first.")
                return
            
            # Load terminal commands from the JSON file
            with open(fix_suggestions_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            terminal_commands = data.get('terminal_commands', {}).get('terminal_commands', [])
            
            if not terminal_commands:
                QMessageBox.information(
                    self,
                    "No Commands Found",
                    "No terminal commands found in the RCA analysis results.\n\nTerminal commands are generated during the RCA analysis process."
                )
                return
            
            # Show confirmation dialog with commands to be executed
            commands_text = "\n".join([f"{i+1}. {cmd['command']}" for i, cmd in enumerate(terminal_commands)])
            confirm_msg = f"Execute {len(terminal_commands)} terminal commands?\n\nCommands to run:\n{commands_text}\n\nNote: Some commands may require administrator privileges."
            
            reply = QMessageBox.question(
                self,
                "Confirm Command Execution",
                confirm_msg,
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            
            if reply != QMessageBox.Yes:
                return
            
            # Create a dialog to show command execution results
            results_dialog = QDialog(self)
            results_dialog.setWindowTitle("Terminal Commands Execution")
            results_dialog.setMinimumSize(800, 600)
            
            layout = QVBoxLayout()
            
            # Header
            header = QLabel("Terminal Commands Execution Results")
            header.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 10px;")
            layout.addWidget(header)
            
            # Results text area
            results_text = QTextEdit()
            results_text.setReadOnly(True)
            results_text.setStyleSheet("""
                QTextEdit {
                    font-family: 'Consolas', 'Monaco', monospace;
                    font-size: 11px;
                    background-color: #1E1E1E;
                    color: #FFFFFF;
                    border: 1px solid #333;
                }
            """)
            layout.addWidget(results_text)
            
            # Close button
            close_btn = QPushButton("Close")
            close_btn.clicked.connect(results_dialog.close)
            layout.addWidget(close_btn)
            
            results_dialog.setLayout(layout)
            
            # Bring dialog to front and make it modal
            results_dialog.setWindowFlags(Qt.Window | Qt.WindowStaysOnTopHint)
            results_dialog.show()
            results_dialog.raise_()
            results_dialog.activateWindow()
            
            # Force focus and bring to front (works on Windows, Linux, macOS)
            results_dialog.setFocus()
            results_dialog.setWindowState(results_dialog.windowState() & ~Qt.WindowMinimized | Qt.WindowActive)
            
            # Execute commands one by one
            for i, cmd_info in enumerate(terminal_commands, 1):
                command = cmd_info['command']
                explanation = cmd_info.get('explanation', 'No explanation provided')
                
                results_text.append(f"\n{'='*60}")
                results_text.append(f"COMMAND {i}/{len(terminal_commands)}: {command}")
                results_text.append(f"EXPLANATION: {explanation}")
                results_text.append(f"{'='*60}")
                
                QApplication.processEvents()  # Keep UI responsive
                
                try:
                    # Execute the command
                    import subprocess
                    result = subprocess.run(
                        command,
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=30  # 30 second timeout
                    )
                    
                    # Display results
                    if result.returncode == 0:
                        results_text.append(f"✅ SUCCESS (Exit Code: {result.returncode})")
                        if result.stdout:
                            results_text.append(f"OUTPUT:\n{result.stdout}")
                    else:
                        results_text.append(f"❌ FAILED (Exit Code: {result.returncode})")
                        if result.stderr:
                            results_text.append(f"ERROR:\n{result.stderr}")
                        if result.stdout:
                            results_text.append(f"OUTPUT:\n{result.stdout}")
                    
                except subprocess.TimeoutExpired:
                    results_text.append(f"⏰ TIMEOUT: Command took longer than 30 seconds")
                except Exception as cmd_error:
                    results_text.append(f"❌ EXECUTION ERROR: {str(cmd_error)}")
                
                # Scroll to bottom
                results_text.verticalScrollBar().setValue(results_text.verticalScrollBar().maximum())
                QApplication.processEvents()
            
            results_text.append(f"\n{'='*60}")
            results_text.append(f"EXECUTION COMPLETED")
            results_text.append(f"Total commands: {len(terminal_commands)}")
            results_text.append(f"{'='*60}")
            
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"An error occurred while running terminal commands:\n{str(e)}"
            )

    def show_patch_success_dialog(self, result, selected_code_patches, selected_config_patches):
        """Show success dialog with Git commit option."""
        # Store patches for embedding update later
        self.last_applied_code_patches = selected_code_patches
        self.last_applied_config_patches = selected_config_patches
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Patches Applied Successfully")
        dialog.setMinimumSize(500, 400)
        
        layout = QVBoxLayout()
        
        # Success message
        success_label = QLabel("🎉 Patches Applied Successfully!")
        success_label.setStyleSheet("font-size: 16px; font-weight: bold; color: #107C10; margin-bottom: 10px;")
        layout.addWidget(success_label)
        
        # Patch summary
        summary_text = f"Applied: {result['total_applied']} patches\n"
        summary_text += f"Code patches: {len(selected_code_patches)}\n"
        summary_text += f"Config patches: {len(selected_config_patches)}\n\n"
        summary_text += "Backup files created automatically.\n\n"
        summary_text += "Selected patches:\n"
        
        for patch in selected_code_patches:
            summary_text += f"• Code: {patch}\n"
        for patch in selected_config_patches:
            summary_text += f"• Config: {patch}\n"
        
        summary_label = QLabel(summary_text)
        summary_label.setStyleSheet("font-size: 12px; margin-bottom: 20px;")
        layout.addWidget(summary_label)
        
        # Git options section
        git_group = QGroupBox("Git Operations")
        git_layout = QVBoxLayout()
        
        # Git commit checkbox
        self.git_commit_check = QCheckBox("Commit changes to Git")
        self.git_commit_check.setChecked(True)  # Default to checked
        self.git_commit_check.setToolTip("Create a Git commit with the applied patches")
        git_layout.addWidget(self.git_commit_check)
        
        # Git push checkbox
        self.git_push_check = QCheckBox("Push to remote repository")
        self.git_push_check.setToolTip("Push the commit to the remote Git repository")
        git_layout.addWidget(self.git_push_check)
        
        # Commit message input
        commit_msg_label = QLabel("Commit Message:")
        git_layout.addWidget(commit_msg_label)
        
        # Generate dynamic commit message based on selected patches and RCA analysis
        default_msg = self.generate_dynamic_commit_message(selected_code_patches, selected_config_patches)
        
        self.commit_msg_edit = QTextEdit()
        self.commit_msg_edit.setPlainText(default_msg)
        self.commit_msg_edit.setMaximumHeight(200)  # Increased height for better visibility
        git_layout.addWidget(self.commit_msg_edit)
        
        git_group.setLayout(git_layout)
        layout.addWidget(git_group)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        # Just Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        
        # Git Commit & Push button
        git_btn = QPushButton("Commit & Push")
        git_btn.clicked.connect(lambda: self.handle_git_operations(dialog))
        git_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D4;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #106EBE;
            }
        """)
        
        button_layout.addWidget(close_btn)
        button_layout.addStretch()
        button_layout.addWidget(git_btn)
        
        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        dialog.exec_()

    def generate_dynamic_commit_message(self, selected_code_patches, selected_config_patches):
        """Generate dynamic commit message based on selected patches and RCA analysis."""
        try:
            # Check if this is from git history - use original commit message
            if hasattr(self, 'git_history_commit_message') and self.git_history_commit_message:
                print("Using original commit message from git history")
                return self.git_history_commit_message
            
            # First, try to use loaded bug analysis data (from Code Assistant)
            rca_data = None
            error_text = None
            
            if hasattr(self, 'loaded_bug_analysis') and self.loaded_bug_analysis:
                print("Using loaded bug analysis data for commit message")
                # Get error message from loaded analysis
                error_text = self.loaded_bug_analysis.get('error_message', None)
                
                # Get RCA data from loaded analysis
                results = self.loaded_bug_analysis.get('results', {})
                phase3_fixes = results.get('phase3_fixes', {})
                rca_data = phase3_fixes
            
            # If no loaded analysis, try to find the most recent fix_suggestions.json file
            if not rca_data:
                print("No loaded bug analysis, searching for fix_suggestions file")
                resources_dir = "backend/resources"
                fix_suggestions_file = None
                
                if os.path.exists(resources_dir):
                    json_files = [f for f in os.listdir(resources_dir) if f.startswith('fix_suggestions_') and f.endswith('.json')]
                    if json_files:
                        # Get the most recent file
                        json_files.sort(reverse=True)
                        fix_suggestions_file = os.path.join(resources_dir, json_files[0])
                
                if not fix_suggestions_file or not os.path.exists(fix_suggestions_file):
                    # Fallback to hardcoded message if no RCA data available
                    return self.generate_fallback_commit_message(selected_code_patches, selected_config_patches)
                
                # Load the RCA analysis data from file
                with open(fix_suggestions_file, 'r', encoding='utf-8') as f:
                    rca_data = json.load(f)
                
                # Try to get error_text from file (might be error_text or error_message)
                error_text = rca_data.get('error_text') or rca_data.get('error_message', 'Unknown error')
            
            # Extract fix suggestion data
            fix_suggestion = rca_data.get('fix_suggestion', {})
            
            # Use error_text if we got it, otherwise default
            if not error_text:
                error_text = 'Unknown error'
            
            print(f"DEBUG: Commit message using error: {error_text}")
            
            # Generate commit subject based on error type
            commit_subject = self.generate_commit_subject(error_text, fix_suggestion)
            
            # Build commit message
            commit_msg = f"{commit_subject}\n\n"
            
            # Add error message
            commit_msg += f"Error Message: \"{error_text}\"\n\n"
            
            # Add patch details
            commit_msg += f"Applied {len(selected_code_patches)} code patches and {len(selected_config_patches)} config patches.\n\n"
            
            # Add selected patch descriptions
            if selected_code_patches or selected_config_patches:
                commit_msg += "Selected Patch Details:\n"
                
                # Get code patches from RCA data
                code_patches_data = fix_suggestion.get('code_patches', [])
                for patch_name in selected_code_patches:
                    # Find the patch data
                    for patch_data in code_patches_data:
                        if patch_data.get('function_name') == patch_name:
                            description = patch_data.get('description', 'No description available')
                            file_path = patch_data.get('file_path', 'Unknown file')
                            commit_msg += f"- Code: {patch_name} ({file_path.split('/')[-1]})\n"
                            commit_msg += f"  Description: {description}\n"
                            break
                    else:
                        commit_msg += f"- Code: {patch_name}\n"
                
                # Get config patches from RCA data
                config_patches_data = fix_suggestion.get('config_patches', [])
                for patch_name in selected_config_patches:
                    # Find the patch data
                    for patch_data in config_patches_data:
                        if patch_data.get('config_name') == patch_name:
                            description = patch_data.get('description', 'No description available')
                            file_path = patch_data.get('file_path', 'Unknown file')
                            commit_msg += f"- Config: {patch_name} ({file_path.split('/')[-1]})\n"
                            commit_msg += f"  Description: {description}\n"
                            break
                    else:
                        commit_msg += f"- Config: {patch_name}\n"
                
                commit_msg += "\n"
            
            # Add root cause analysis from RCA data
            root_cause = fix_suggestion.get('root_cause_analysis', '')
            if root_cause:
                commit_msg += f"Root Cause Analysis:\n{root_cause}\n\n"
            
            # Add reason if available (but don't duplicate root cause analysis)
            reason = fix_suggestion.get('reason', '')
            if reason and reason != root_cause:
                commit_msg += f"Analysis Summary:\n{reason}\n\n"
            
            return commit_msg
            
        except Exception as e:
            print(f"Error generating dynamic commit message: {e}")
            # Fallback to hardcoded message
            return self.generate_fallback_commit_message(selected_code_patches, selected_config_patches)
    
    def generate_commit_subject(self, error_text, fix_suggestion):
        """Generate commit subject based on error type and analysis."""
        error_lower = error_text.lower()
        
        # Extract error type and generate appropriate subject
        if 'contention' in error_lower and 'timer' in error_lower:
            return "Fix contention resolution timer configuration"
        elif 'amf' in error_lower and 'association' in error_lower:
            return "Fix AMF association issue"
        elif 'rrc' in error_lower and 'setup' in error_lower:
            return "Fix RRC setup procedure"
        elif 'timeout' in error_lower:
            return "Fix timeout configuration issue"
        elif 'random access' in error_lower or 'ra procedure' in error_lower:
            return "Fix Random Access procedure"
        elif 'ngap' in error_lower:
            return "Fix NGAP protocol issue"
        elif 'f1ap' in error_lower:
            return "Fix F1AP protocol issue"
        else:
            # Try to extract meaningful keywords
            keywords = []
            if 'gnb' in error_lower:
                keywords.append('gNB')
            if 'ue' in error_lower:
                keywords.append('UE')
            if 'network' in error_lower:
                keywords.append('network')
            
            if keywords:
                return f"Fix {'/'.join(keywords)} configuration issue"
            else:
                return "Fix network configuration issue"
    
    def generate_fallback_commit_message(self, selected_code_patches, selected_config_patches):
        """Generate fallback commit message when RCA data is not available."""
        default_msg = f"Apply RCA patches for network configuration issue\n\n"
        default_msg += f"- Applied {len(selected_code_patches)} code patches\n"
        default_msg += f"- Applied {len(selected_config_patches)} config patches\n\n"
        
        if selected_code_patches:
            default_msg += "Code patches:\n"
            for patch in selected_code_patches:
                default_msg += f"  - {patch}\n"
            default_msg += "\n"
        
        if selected_config_patches:
            default_msg += "Config patches:\n"
            for patch in selected_config_patches:
                default_msg += f"  - {patch}\n"
            default_msg += "\n"
        
        default_msg += "Generated by AgenticRAN RCA Analysis"
        return default_msg

    def handle_git_operations(self, dialog):
        """Handle Git commit and push operations."""
        try:
            if not self.git_commit_check.isChecked():
                dialog.close()
                return
            
            commit_message = self.commit_msg_edit.toPlainText().strip()
            if not commit_message:
                QMessageBox.warning(dialog, "Missing Commit Message", "Please enter a commit message.")
                return
            
            # Show progress dialog
            progress_dialog = QProgressDialog("Performing Git operations...", "Cancel", 0, 0, dialog)
            progress_dialog.setWindowModality(Qt.WindowModal)
            progress_dialog.show()
            QApplication.processEvents()
            
            # Perform Git operations
            git_result = self.perform_git_commit_and_push(commit_message, self.git_push_check.isChecked())
            
            progress_dialog.close()
            
            if git_result["success"]:
                success_msg = "Git operations completed successfully!\n\n"
                if git_result["committed"]:
                    success_msg += f"✅ Committed: {git_result['commit_hash']}\n"
                if git_result["pushed"]:
                    success_msg += f"✅ Pushed to remote repository\n"
                
                QMessageBox.information(dialog, "Git Operations Successful", success_msg)
                dialog.close()
            else:
                QMessageBox.warning(
                    dialog, 
                    "Git Operations Failed", 
                    f"Git operations failed:\n\n{git_result['error']}\n\nPlease check your Git configuration and try again."
                )
                
        except Exception as e:
            QMessageBox.critical(
                dialog,
                "Git Error",
                f"An error occurred during Git operations:\n{str(e)}"
            )

    def perform_git_commit_and_push(self, commit_message, should_push=False):
        """Perform Git commit and optionally push to remote."""
        try:
            import subprocess
            
            result = {
                "success": False,
                "committed": False,
                "pushed": False,
                "commit_hash": None,
                "error": None
            }
            
            # Get the source code directory and extract folder name
            code_dir = None
            
            # Try multiple sources for code directory
            if hasattr(self, 'loaded_bug_analysis'):
                code_dir = self.loaded_bug_analysis.get('code_dir', None)
            elif hasattr(self, 'code_dir_edit'):
                code_dir = self.code_dir_edit.text()
            
            # If still not found, try from the most recent fix_suggestions.json
            if not code_dir:
                try:
                    resources_dir = "backend/resources"
                    if os.path.exists(resources_dir):
                        json_files = [f for f in os.listdir(resources_dir) if f.startswith('fix_suggestions_') and f.endswith('.json')]
                        if json_files:
                            json_files.sort(reverse=True)
                            fix_suggestions_file = os.path.join(resources_dir, json_files[0])
                            with open(fix_suggestions_file, 'r', encoding='utf-8') as f:
                                suggestions_data = json.load(f)
                            code_dir = suggestions_data.get('code_dir', None)
                except Exception as e:
                    print(f"Error reading code directory from fix_suggestions: {e}")
            
            if code_dir:
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            else:
                openair_codebase_file_name = "openairinterface5g-develop"  # Default fallback
            
            # Get the OpenAirInterface5G directory path
            oai_dir = os.path.join('Error_fixing_pipelin', openair_codebase_file_name)
            
            print(f"\n{'='*60}")
            print(f"DEBUG [Git Operations]:")
            print(f"  Code directory: {code_dir}")
            print(f"  Codebase folder name: {openair_codebase_file_name}")
            print(f"  Git repository path: {oai_dir}")
            print(f"  Current working directory: {os.getcwd()}")
            print(f"{'='*60}\n")
            
            if not os.path.exists(oai_dir):
                result["error"] = f"OpenAirInterface5G directory not found at: {oai_dir}\n\nPlease clone the repository first."
                return result
            
            # Check if we're in a Git repository
            git_check = subprocess.run(
                ["git", "status"],
                capture_output=True,
                text=True,
                cwd=oai_dir
            )
            
            print(f"Git status check:")
            print(f"  Return code: {git_check.returncode}")
            if git_check.stdout:
                print(f"  Status:\n{git_check.stdout}")
            
            if git_check.returncode != 0:
                result["error"] = "Not a Git repository or Git not available"
                return result
            
            # Add all changes
            add_result = subprocess.run(
                ["git", "add", "."],
                capture_output=True,
                text=True,
                cwd=oai_dir
            )
            
            if add_result.returncode != 0:
                result["error"] = f"Failed to add changes: {add_result.stderr}"
                return result
            
            # Check if there are changes to commit
            diff_check = subprocess.run(
                ["git", "diff", "--cached", "--quiet"],
                capture_output=True,
                text=True,
                cwd=oai_dir
            )
            
            # Show git status for debugging
            status_result = subprocess.run(
                ["git", "status", "--short"],
                capture_output=True,
                text=True,
                cwd=oai_dir
            )
            
            print(f"Git status (short):")
            if status_result.stdout:
                print(f"{status_result.stdout}")
            else:
                print(f"  No changes detected by Git")
            
            # Show what was staged
            staged_result = subprocess.run(
                ["git", "diff", "--cached", "--name-only"],
                capture_output=True,
                text=True,
                cwd=oai_dir
            )
            
            print(f"Staged files:")
            if staged_result.stdout:
                print(f"{staged_result.stdout}")
            else:
                print(f"  No files staged for commit")
            print()
            
            if diff_check.returncode == 0:
                result["error"] = "No changes to commit - patches may have been applied outside the Git repository or files are unchanged"
                return result
            
            # Commit changes
            commit_result = subprocess.run(
                ["git", "commit", "-m", commit_message],
                capture_output=True,
                text=True,
                cwd=oai_dir
            )
            
            if commit_result.returncode != 0:
                result["error"] = f"Failed to commit: {commit_result.stderr}"
                return result
            
            result["committed"] = True
            
            # Get commit hash
            hash_result = subprocess.run(
                ["git", "rev-parse", "HEAD"],
                capture_output=True,
                text=True,
                cwd=oai_dir
            )
            
            if hash_result.returncode == 0:
                full_hash = hash_result.stdout.strip()
                result["commit_hash"] = full_hash[:8]  # Short for display
                result["full_commit_hash"] = full_hash  # Full for embedding update
            
            # Git-commit embeddings: updated once when "Start RCA Analysis" runs (commit-gated), not after each commit.
            result["embedding_updated"] = False
            # ========================================
            
            # Push if requested
            if should_push:
                push_result = subprocess.run(
                    ["git", "push"],
                    capture_output=True,
                    text=True,
                    cwd=oai_dir
                )
                
                if push_result.returncode != 0:
                    result["error"] = f"Commit successful but push failed: {push_result.stderr}"
                    result["success"] = True  # Commit was successful
                    return result
                
                result["pushed"] = True
            
            result["success"] = True
            return result
            
        except Exception as e:
            return {
                "success": False,
                "committed": False,
                "pushed": False,
                "commit_hash": None,
                "error": str(e)
            }

    def create_filtered_suggestions_file(self, original_file, selected_code_patches, selected_config_patches, openair_codebase_file_name="openairinterface5g-develop", edited_patch_values=None):
        """Create a filtered fix_suggestions.json with only selected patches, corrected paths, and edited values."""
        try:
            if edited_patch_values is None:
                edited_patch_values = {
                    'code_patches': {},
                    'config_patches': {}
                }
            
            print(f"\nDEBUG: === create_filtered_suggestions_file ===")
            print(f"DEBUG: Original file: {original_file}")
            print(f"DEBUG: Selected code patches: {selected_code_patches}")
            print(f"DEBUG: Selected config patches: {selected_config_patches}")
            print(f"DEBUG: Codebase folder name: {openair_codebase_file_name}")
            
            # Load original suggestions
            with open(original_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Create filtered data
            filtered_data = data.copy()
            
            # Initialize filtered patches lists
            filtered_code_patches = []
            filtered_config_patches = []
            
            # Handle both nested and direct fix_suggestion formats
            if 'fix_suggestion' in filtered_data:
                fix_suggestion = filtered_data['fix_suggestion']
            elif 'code_patches' in filtered_data or 'config_patches' in filtered_data:
                # Direct format - the data itself is the fix_suggestion
                fix_suggestion = filtered_data
            else:
                # No valid structure found
                print("Warning: No valid fix_suggestion structure found in file")
                fix_suggestion = {}
            
            # Filter and fix code patches
            original_code_patches = fix_suggestion.get('code_patches', [])
            
            for idx, patch in enumerate(original_code_patches):
                function_name = patch.get('function_name', '')
                file_path = patch.get('file_path', '')
                # Create unique identifier: function_name (filename) - same format as checkbox text
                file_name = os.path.basename(file_path)
                unique_identifier = f"{function_name} ({file_name})"
                
                print(f"DEBUG: Checking code patch: {unique_identifier}")
                print(f"DEBUG: Original path: {file_path}")
                
                if unique_identifier in selected_code_patches:
                    print(f"DEBUG: Match found for code patch: {unique_identifier}")
                    
                    # Update with edited code if available
                    if idx in edited_patch_values.get('code_patches', {}):
                        edited_code = edited_patch_values['code_patches'][idx]
                        # Update all possible code field names
                        if 'patched_code' in patch:
                            patch['patched_code'] = edited_code
                        if 'suggested_code' in patch:
                            patch['suggested_code'] = edited_code
                        if 'new_code' in patch:
                            patch['new_code'] = edited_code
                        print(f"DEBUG: Updated code patch {idx} with edited code ({len(edited_code)} chars)")
                    
                    # Fix the file path
                    original_path = patch.get('file_path', '')
                    
                    # Check if path needs correction
                    if not original_path.startswith('Error_fixing_pipelin'):
                        if original_path.startswith(openair_codebase_file_name + '/'):
                            # Path has codebase prefix, prepend Error_fixing_pipelin
                            corrected_path = 'Error_fixing_pipelin/' + original_path
                        else:
                            # Path doesn't have codebase prefix, add both
                            corrected_path = f'Error_fixing_pipelin/{openair_codebase_file_name}/{original_path}'
                        patch['file_path'] = corrected_path
                        print(f"Corrected code patch path: {original_path} -> {corrected_path}")
                    else:
                        print(f"Path already correct: {original_path}")
                    
                    filtered_code_patches.append(patch)
            
            # Filter and fix config patches
            original_config_patches = fix_suggestion.get('config_patches', [])
            
            for idx, patch in enumerate(original_config_patches):
                # Try both config_name and parameter_name for compatibility
                config_name = patch.get('config_name', patch.get('parameter_name', ''))
                file_path = patch.get('file_path', '')
                # Create unique identifier: config_name (filename) - same format as checkbox text
                file_name = os.path.basename(file_path)
                unique_identifier = f"{config_name} ({file_name})"
                
                print(f"DEBUG: Checking config patch: {unique_identifier}")
                print(f"DEBUG: Original path: {file_path}")
                
                if unique_identifier in selected_config_patches:
                    print(f"DEBUG: Match found for config patch: {unique_identifier}")
                    
                    # Update with edited value if available
                    if idx in edited_patch_values.get('config_patches', {}):
                        edited_value = edited_patch_values['config_patches'][idx]
                        # Update all possible value field names
                        if 'new_value' in patch:
                            patch['new_value'] = edited_value
                        if 'suggested_value' in patch:
                            patch['suggested_value'] = edited_value
                        if 'recommended_value' in patch:
                            patch['recommended_value'] = edited_value
                        print(f"DEBUG: Updated config patch {idx} with edited value: {edited_value}")
                    # Fix the file path
                    original_path = patch.get('file_path', '')
                    
                    # Check if path needs correction
                    if not original_path.startswith('Error_fixing_pipelin'):
                        if original_path.startswith(openair_codebase_file_name + '/'):
                            # Path has codebase prefix, prepend Error_fixing_pipelin
                            corrected_path = 'Error_fixing_pipelin/' + original_path
                        elif original_path.startswith('targets/'):
                            # Handle relative paths that should be under the OAI codebase folder
                            corrected_path = f'Error_fixing_pipelin/{openair_codebase_file_name}/{original_path}'
                        else:
                            # Generic relative path - add full prefix
                            corrected_path = f'Error_fixing_pipelin/{openair_codebase_file_name}/{original_path}'
                        patch['file_path'] = corrected_path
                        print(f"Corrected config patch path: {original_path} -> {corrected_path}")
                    else:
                        print(f"Path already correct: {original_path}")
                    
                    filtered_config_patches.append(patch)
                
            # Update the filtered data - Always ensure 'fix_suggestion' wrapper for UnifiedPatchApplicator
            if 'fix_suggestion' in data:
                # Nested format - update the nested structure
                fix_suggestion['code_patches'] = filtered_code_patches
                fix_suggestion['config_patches'] = filtered_config_patches
                filtered_data['fix_suggestion'] = fix_suggestion
            else:
                # Direct format - wrap patches in 'fix_suggestion' structure
                filtered_data = {
                    'fix_suggestion': {
                        'code_patches': filtered_code_patches,
                        'config_patches': filtered_config_patches
                    }
                }
            
            print(f"DEBUG: Final structure - Has 'fix_suggestion': {'fix_suggestion' in filtered_data}")
            print(f"DEBUG: Patches in final structure - Code: {len(filtered_data.get('fix_suggestion', {}).get('code_patches', []))}, Config: {len(filtered_data.get('fix_suggestion', {}).get('config_patches', []))}")
            
            # Save filtered suggestions to a temporary file
            filtered_file = os.path.join("backend/resources", "filtered_fix_suggestions.json")
            os.makedirs("backend/resources", exist_ok=True)
            
            with open(filtered_file, 'w', encoding='utf-8') as f:
                json.dump(filtered_data, f, indent=2, ensure_ascii=False)
            
            print(f"Created filtered suggestions file: {filtered_file}")
            print(f"Selected code patches: {len(filtered_code_patches)}")
            print(f"Selected config patches: {len(filtered_config_patches)}")
            
            return filtered_file
            
        except Exception as e:
            print(f"Error creating filtered suggestions file: {e}")
            return original_file  # Fallback to original file

    def parse_fix_suggestions_json(self, json_file_path):
        """Parse and format fix_suggestions.json file for comprehensive display."""
        try:
            if not os.path.exists(json_file_path):
                return "Error: Fix suggestions file not found."
            
            with open(json_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Format the results with enhanced styling
            formatted_text = "🔧 RCA ANALYSIS RESULTS\n"
            formatted_text += "=" * 80 + "\n\n"
            
            # Error Information
            if 'error_text' in data:
                formatted_text += f"🔥 ERROR DETECTED:\n"
                formatted_text += "-" * 40 + "\n"
                formatted_text += f"{data['error_text']}\n\n"
            
            # Phase 3 Analysis Results
            if 'fix_suggestion' in data:
                fix_suggestion = data['fix_suggestion']
                
                # Root Cause Analysis
                if fix_suggestion.get('reason'):
                    formatted_text += "🎯 ROOT CAUSE ANALYSIS:\n"
                    formatted_text += "-" * 40 + "\n"
                    formatted_text += f"{fix_suggestion['reason']}\n\n"
                
                # Suspected Functions
                suspected_functions = fix_suggestion.get('suspected_functions', [])
                if suspected_functions:
                    formatted_text += "🔍 SUSPECTED FUNCTIONS:\n"
                    formatted_text += "-" * 40 + "\n"
                    for i, func in enumerate(suspected_functions, 1):
                        formatted_text += f"  {i}. {func}\n"
                    formatted_text += "\n"
                
                # Suspected Configs
                suspected_configs = fix_suggestion.get('suspected_configs', [])
                if suspected_configs:
                    formatted_text += "⚙️ SUSPECTED CONFIGURATIONS:\n"
                    formatted_text += "-" * 40 + "\n"
                    for i, config in enumerate(suspected_configs, 1):
                        formatted_text += f"  {i}. {config}\n"
                    formatted_text += "\n"
                
                # Configuration Fixes (Enhanced parsing)
                if fix_suggestion.get('config_fix'):
                    formatted_text += "🔧 CONFIGURATION FIXES:\n"
                    formatted_text += "-" * 40 + "\n"
                    try:
                        # Try to parse as JSON if it's a string representation
                        config_fix_data = fix_suggestion['config_fix']
                        if isinstance(config_fix_data, str):
                            # Handle string representation of list/dict
                            import ast
                            config_fix_data = ast.literal_eval(config_fix_data)
                        
                        if isinstance(config_fix_data, list):
                            for i, config in enumerate(config_fix_data, 1):
                                formatted_text += f"  {i}. Configuration: {config.get('config_name', 'Unknown')}\n"
                                formatted_text += f"     File: {config.get('file_path', 'Unknown')}\n"
                                formatted_text += f"     Line: {config.get('line_number', 'Unknown')}\n"
                                formatted_text += f"     Current: {config.get('current_value', 'Unknown')}\n"
                                formatted_text += f"     New: {config.get('new_value', 'Unknown')}\n"
                                formatted_text += f"     Reason: {config.get('description', 'No description')}\n\n"
                        else:
                            formatted_text += f"{config_fix_data}\n\n"
                    except:
                        formatted_text += f"{fix_suggestion['config_fix']}\n\n"
                
                # Code Patches (NEW - Enhanced)
                code_patches = fix_suggestion.get('code_patches', [])
                if code_patches:
                    formatted_text += "\n💻 CODE PATCHES:\n"
                    formatted_text += "─" * 50 + "\n"
                    for i, patch in enumerate(code_patches, 1):
                        formatted_text += f"  {i}. Function: {patch.get('function_name', 'Unknown')}\n"
                        formatted_text += f"     File: {patch.get('file_path', 'Unknown')}\n"
                        formatted_text += f"     Type: {patch.get('patch_type', 'Unknown')}\n"
                        formatted_text += f"     Lines: {patch.get('line_numbers', 'Unknown')}\n"
                        formatted_text += f"     Description: {patch.get('description', 'No description')}\n"
                        
                        if patch.get('original_code') and patch.get('patched_code'):
                            # Format as git-style diff
                            formatted_text += f"     📝 Code Changes:\n"
                            formatted_text += f"     {'─' * 50}\n"
                            
                            # Parse original and patched code for diff-style display
                            original_lines = patch['original_code'].strip().split('\n')
                            patched_lines = patch['patched_code'].strip().split('\n')
                            
                            # Show original code in red (deletion style)
                            formatted_text += f"     🔴 Original Code:\n"
                            for line in original_lines:
                                if line.strip():
                                    formatted_text += f"     - {line}\n"
                            
                            # Show patched code in green (addition style)
                            formatted_text += f"\n     🟢 Patched Code:\n"
                            for line in patched_lines:
                                if line.strip():
                                    formatted_text += f"     + {line}\n"
                            
                            formatted_text += f"     {'─' * 50}\n"
                        elif patch.get('original_code'):
                            formatted_text += f"     Original Code:\n"
                            formatted_text += f"     {patch['original_code']}\n"
                        elif patch.get('patched_code'):
                            formatted_text += f"     Patched Code:\n"
                            formatted_text += f"     {patch['patched_code']}\n"
                        formatted_text += "\n"
                
                # Config Patches (NEW)
                config_patches = fix_suggestion.get('config_patches', [])
                if config_patches:
                    formatted_text += "\n⚙️ CONFIG PATCHES:\n"
                    formatted_text += "─" * 50 + "\n"
                    for i, patch in enumerate(config_patches, 1):
                        formatted_text += f"  {i}. Config: {patch.get('config_name', 'Unknown')}\n"
                        formatted_text += f"     File: {patch.get('file_path', 'Unknown')}\n"
                        formatted_text += f"     Line: {patch.get('line_number', 'Unknown')}\n"
                        formatted_text += f"     Description: {patch.get('description', 'No description')}\n"
                        
                        # Format config changes as git-style diff
                        if patch.get('current_value') and patch.get('new_value'):
                            formatted_text += f"     📝 Configuration Changes:\n"
                            formatted_text += f"     {'─' * 50}\n"
                            formatted_text += f"     🔴 Current Value: {patch.get('current_value', 'Unknown')}\n"
                            formatted_text += f"     🟢 New Value:     {patch.get('new_value', 'Unknown')}\n"
                            formatted_text += f"     {'─' * 50}\n"
                        else:
                            formatted_text += f"     Current: {patch.get('current_value', 'Unknown')}\n"
                            formatted_text += f"     New: {patch.get('new_value', 'Unknown')}\n"
                        
                        formatted_text += "\n"
                
                # Investigation Steps
                investigation_steps = fix_suggestion.get('investigation_steps', [])
                if investigation_steps:
                    formatted_text += "\n📋 INVESTIGATION STEPS:\n"
                    formatted_text += "─" * 50 + "\n"
                    for i, step in enumerate(investigation_steps, 1):
                        formatted_text += f"  {i}. {step}\n"
                    formatted_text += "\n"
                
                # Root Cause Analysis (if separate from reason)
                if fix_suggestion.get('root_cause_analysis'):
                    formatted_text += "\n🔬 DETAILED ROOT CAUSE:\n"
                    formatted_text += "─" * 50 + "\n"
                    formatted_text += f"{fix_suggestion['root_cause_analysis']}\n\n"
            
            # Terminal Commands (Phase 4)
            if 'terminal_commands' in data:
                terminal_commands = data['terminal_commands']
                if terminal_commands.get('terminal_commands'):
                    formatted_text += "\n💻 INVESTIGATION COMMANDS:\n"
                    formatted_text += "─" * 50 + "\n"
                    for i, cmd in enumerate(terminal_commands['terminal_commands'], 1):
                        formatted_text += f"  {i}. {cmd['command']}\n"
                        formatted_text += f"     💡 {cmd['explanation']}\n\n"
            
            # Context Summary (if available) - support both old and new field names
            context = data.get('context_summary') or data.get('summary')
            if context:
                formatted_text += "📊 ANALYSIS CONTEXT:\n"
                formatted_text += "-" * 40 + "\n"
                # Support both old field names (candidate_functions_count) and new ones (total_functions_analyzed)
                formatted_text += f"  Functions analyzed: {context.get('total_functions_analyzed', context.get('candidate_functions_count', 0))}\n"
                formatted_text += f"  Configs analyzed: {context.get('total_configs_analyzed', context.get('candidate_configs_count', 0))}\n"
                formatted_text += f"  Call graph entries: {context.get('call_graph_entries', 0)}\n"
                formatted_text += f"  Pattern matched: {context.get('pattern_matched', False)}\n\n"
            
            # Summary Information
            formatted_text += "📊 ANALYSIS SUMMARY:\n"
            formatted_text += "-" * 40 + "\n"
            formatted_text += f"Analysis completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            formatted_text += f"Results file: {os.path.basename(json_file_path)}\n"
            formatted_text += f"Total sections: {len([k for k in data.keys() if k != 'terminal_commands'])}\n"
            
            return formatted_text
            
        except Exception as e:
            return f"Error parsing fix suggestions: {str(e)}\n\nPlease check the file format and try again."

    def extract_error_from_log(self, log_file_path):
        """Extract error message from log file using comprehensive pattern matching from RCA analysis."""
        try:
            if not os.path.exists(log_file_path):
                return None
                
            with open(log_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Comprehensive error patterns based on the RCA error patterns database
            error_patterns = [
                # ========== HIGH PRIORITY: ROOT CAUSE ERRORS ==========
                
                # # PDU Session patterns (HIGH PRIORITY - often root cause)
                # r'Received PDU Session Establishment Reject',
                # r'PDU Session Establishment Reject',
                # r'PDU session setup failure',
                # r'PDU session establishment failed',
                # r'PDU session.*failed',
                # r'unknow message type \d+',  # Typo in OAI code, but important pattern
                # r'unknown message type \d+',
                # r'Received unexpected message in DLinformationTransfer',


                # AMF Association patterns (from error_patterns_enhanced.json)
                r'No AMF is associated to the gNB',
                r'No AMF is associated with gNB',
                r'AMF association failed',
                r'AMF not connected',
                r'No AMF available',
                r'AMF selection failed',
                r'AMF.*not.*associated',
                r'AMF.*connection.*failed',
                
                # # RRC patterns
                # r'RRC setup failed',
                # r'RRC setup timeout',
                # r'Invalid UE capability',
                # r'Missing security capabilities',
                # r'RRCConnectionSetup failed, returning to IDLE state',
                # r'RRC connection establishment failed',
                # r'\[RRC\].*not defined',
                # r'RRC.*warning',
                # r'RRC.*error',
                
                # # PDCP patterns
                # r'PDCP integrity check failed',
                # r'PDCP ciphering failed',
                # r'PDCP data forwarding failed',
                # r'PDCP error',
                # r'PDCP security verification failed',
                # r'ERROR: PDCP integrity check failed .* during handover',
                # r'PDCP data request failed',
                # r'PDCP.*failed',
                # r'\[PDCP\]\[E\].*failed',
                
                # # RLC patterns
                # r'Data sending request over RLC failed with \'Bad Parameter\' reason!',
                # r'Data sending request over RLC failed with \'Internal Error\' reason!',
                # r'RLC_OP_STATUS_INTERNAL_ERROR',
                # r'Data sending request over RLC failed with \'Out of Resources\' reason!',
                # r'RLC error',
                
                # # MAC patterns
                # r'Module id .* Contention resolution timer expired, RA failed',
                # r'\[UE .*\] Contention resolution failed',
                # r'MAC error',
                # r'Random access failed',
                # r'Contention resolution timer expired',
                # r'RA contention resolution failed',
                # r'\[MAC\]\[E\].*contention.*failed',
                # r'nr_ra_contention_resolution_failed',
                
                # # Handover patterns
                # r'\[Hh]andover.*failed',
                # r'Aborting handover',
                # r'Handover.*failure',
                
                # # Authentication patterns
                # r'Authentication failed',
                # r'Auth.*failed',
                # r'Security.*failed',
                
                # # Registration patterns
                # r'Registration failed',
                # r'UE registration.*failed',
                # r'Registration.*failure',
                
                # # PDU Session patterns
                # r'PDU session setup failure',
                # r'PDU session.*failed',
                # r'Session.*failure',
                
                # # General error patterns
                # r'ERROR[:\s]+(.+?)(?:\n|$)',
                # r'error[:\s]+(.+?)(?:\n|$)',
                # r'Error[:\s]+(.+?)(?:\n|$)',
                # r'FAILED[:\s]+(.+?)(?:\n|$)',
                # r'failed[:\s]+(.+?)(?:\n|$)',
                # r'Failed[:\s]+(.+?)(?:\n|$)',
                # r'Exception[:\s]+(.+?)(?:\n|$)',
                # r'FATAL[:\s]+(.+?)(?:\n|$)',
                # r'fatal[:\s]+(.+?)(?:\n|$)',
                # r'CRITICAL[:\s]+(.+?)(?:\n|$)',
                # r'critical[:\s]+(.+?)(?:\n|$)',
                # r'Connection establishment failed',
                # r'Connection.*failed',
                # r'Setup.*failed',
                # r'Initialization.*failed',
                # r'Configuration.*failed',
                # r'Timeout.*expired',
                # r'Timer.*expired'
            ]
            
            # Look for error patterns
            for pattern in error_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE | re.MULTILINE)
                if matches:
                    # Return the first significant error found
                    for match in matches:
                        if isinstance(match, tuple):
                            match = match[0] if match[0] else match[1]
                        if match and len(match.strip()) > 10:  # Ensure it's a meaningful error
                            return match.strip()
            
            # If no specific error pattern found, look for lines containing common error keywords
            # lines = content.split('\n')
            # for line in lines:
            #     line_lower = line.lower()
            #     if any(keyword in line_lower for keyword in [
            #         'error', 'failed', 'exception', 'fatal', 'critical', 'timeout', 'expired',
            #         'amf', 'associated', 'gnb', 'rrc', 'pdcp', 'handover', 'authentication',
            #         'registration', 'pdu', 'session', 'setup', 'connection', 'association'
            #     ]):
            #         if len(line.strip()) > 20:  # Ensure it's a meaningful line
            #             return line.strip()
            
            return None
            
        except Exception as e:
            print(f"Error extracting error from log: {e}")
            return None

    def generate_terminal_commands(self, error_message: str, investigation_steps: list, deployment_context: dict = None, troubleshooting_hints: list = None, openair_codebase_file_name: str = "openairinterface5g-develop") -> list:
        """
    Phase 4: Generate exact terminal commands for verification and implementation
        
        Args:
            error_message: The original error
            investigation_steps: List of investigation steps from Phase 3
            deployment_context: Deployment context with IP addresses
            troubleshooting_hints: Troubleshooting hints from error patterns
            openair_codebase_file_name: Name of the OpenAirInterface5G codebase folder
            
        Returns:
            List of exact terminal commands
        """
        # logger.info("🔧 PHASE 4 - GENERATING TERMINAL COMMANDS")
        # logger.info("=" * 60)
        
        try:
            if deployment_context and deployment_context.get("dependency_advice_mode"):
                cmds = []
                for step in investigation_steps or []:
                    if not isinstance(step, str):
                        continue
                    st = step.strip()
                    if st.startswith("$"):
                        cmds.append(st[1:].strip())
                    elif st.startswith(">"):
                        cmds.append(st[1:].strip())
                if cmds:
                    return cmds

            # Prepare context for command generation
            context_parts = []
            
            # Error information
            context_parts.append(f"Error: {error_message}")
            
            # Deployment context (IP addresses, network info)
            # Always try to get from JSON file first (has real IP addresses), then fallback to deployment context
            network_config = {}
            try:
                with open('database/error_patterns_structured.json', 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    deployment_context_json = data.get('deployment_context', {})
                    network_config = {
                        'gNB IP': deployment_context_json.get('cu_ip_address', 'Unknown'),
                        'DU IP': deployment_context_json.get('du_ip_address', 'Unknown'),
                        'AMF IP': deployment_context_json.get('amf_ip_address', 'Unknown'),
                        'Core Network IP': deployment_context_json.get('core_network_machine_ip', 'Unknown'),
                        'Local SCTP Port C': deployment_context_json.get('local_s_portc', 'Unknown'),
                        'Local SCTP Port D': deployment_context_json.get('local_s_portd', 'Unknown'),
                        'Remote SCTP Port C': deployment_context_json.get('remote_s_portc', 'Unknown'),
                        'Remote SCTP Port D': deployment_context_json.get('remote_s_portd', 'Unknown'),
                        'NSSAI SST': deployment_context_json.get('nssai_sst', 'Unknown'),
                        'NSSAI SD': deployment_context_json.get('nssai_sd', 'Unknown'),
                        'NMC Size': deployment_context_json.get('nmc_size', 'Unknown'),
                        'DNN': deployment_context_json.get('dnn', 'Unknown')
                    }
            except Exception as e:
                # logger.warning(f"Could not load deployment context from JSON: {e}")
                # Fallback to deployment context from pipeline
                if deployment_context:
                    network_params = deployment_context.get('network_params', {})
                    network_config = {
                        'gNB IP': network_params.get('gnb_ipv4', 'Unknown'),
                        'AMF IP': network_params.get('amf_ipv4', 'Unknown'),
                        'Core Network IP': network_params.get('core_network_ip', 'Unknown'),
                        'Local SCTP Port': network_params.get('local_s_portc', 'Unknown'),
                        'Remote SCTP Port': network_params.get('remote_s_portc', 'Unknown')
                    }
            
            if network_config:
                context_parts.append(f"Network Configuration:")
                for key, value in network_config.items():
                    context_parts.append(f"- {key}: {value}")
            
            # Investigation steps
            context_parts.append(f"Investigation Steps to Convert:")
            for i, step in enumerate(investigation_steps, 1):
                context_parts.append(f"{i}. {step}")
            
            # Troubleshooting hints
            if troubleshooting_hints:
                context_parts.append(f"Troubleshooting Hints:")
                for hint in troubleshooting_hints:
                    context_parts.append(f"- {hint}")
            
            context = "\n".join(context_parts)
            
            # Error-specific prompt for command generation
            prompt = f"""You are a 5G/LTE network troubleshooting expert. Given the error and context below, generate ONLY the ESSENTIAL terminal commands needed to verify and fix the issue.

    {context}

    IMPORTANT: Generate ONLY 2-3 essential commands, not a comprehensive list. 

    For RA (Random Access) timer errors, focus on:
    1. ONE command to check/verify the configuration file
    2. ONE command to test network connectivity (if applicable)

    For AMF association errors, focus on:
    1. ONE connectivity test command (ping to AMF)
    2. ONE routing command (if applicable)

    Requirements:
    - Provide EXACT commands that can be copy-pasted
    - Use the specific IP addresses from the context (if available)
    - Keep it minimal - only the most critical commands
    - Format each command on a new line starting with "COMMAND: "
    - Add brief explanation after each command
    - If IP addresses are not available, use placeholder format like <AMF_IP> or <GNB_IP>

    Example format for RA timer error:
    COMMAND: grep -n "ra_ContentionResolutionTimer" du_gnb.conf
    EXPLANATION: Check current value of contention resolution timer in gNB config

    COMMAND: ping <GNB_IP>
    EXPLANATION: Test connectivity to gNB (replace <GNB_IP> with actual IP)

    Example format for AMF error:
    COMMAND: ping <AMF_IP>
    EXPLANATION: Test basic connectivity to AMF

    COMMAND: ip route add <AMF_IP> via <CORE_NETWORK_IP> dev eth0
    EXPLANATION: Add static route for AMF through core network machine
    """

            # Use the same Azure client from phase3_pipeline
            from fix_suggestion_pipeline import FixSuggestionPipeline
            phase3_pipeline = FixSuggestionPipeline()
            
            response = phase3_pipeline.azure_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a network troubleshooting expert who generates exact terminal commands."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.1,
                seed=33333  # Terminal command generation seed
            )
            
            commands_text = response.choices[0].message.content.strip()
            
            # Parse commands from response
            commands = []
            current_command = None
            current_explanation = None
            
            for line in commands_text.split('\n'):
                line = line.strip()
                if line.startswith('COMMAND:'):
                    # Save previous command if exists
                    if current_command:
                        commands.append({
                            "command": current_command,
                            "explanation": current_explanation or "No explanation provided"
                        })
                    current_command = line.replace('COMMAND:', '').strip()
                    current_explanation = None
                elif line.startswith('EXPLANATION:'):
                    current_explanation = line.replace('EXPLANATION:', '').strip()
                elif line and not line.startswith('COMMAND:') and not line.startswith('EXPLANATION:'):
                    # Only treat as command if it looks like a command (starts with common command words)
                    if (current_command is None and line and 
                        any(line.startswith(cmd) for cmd in ['grep', 'ping', 'ip route', 'systemctl', 'cat', 'ls', 'cd', 'sudo', 'ssh', 'telnet', 'traceroute', 'netstat', 'ifconfig', 'iptables'])):
                        current_command = line
                    elif current_explanation is None and line and not any(line.startswith(cmd) for cmd in ['grep', 'ping', 'ip route', 'systemctl', 'cat', 'ls', 'cd', 'sudo', 'ssh', 'telnet', 'traceroute', 'netstat', 'ifconfig', 'iptables']):
                        current_explanation = line
            
            # Add the last command
            if current_command:
                commands.append({
                    "command": current_command,
                    "explanation": current_explanation or "No explanation provided"
                })
            
            # logger.info(f"✅ Generated {len(commands)} terminal commands")
            return commands
            
        except Exception as e:
            # logger.error(f"❌ Failed to generate terminal commands: {e}")
            return []

    def handle_start_code_rca_analysis(self):
        """Handle the Start RCA Analysis button click in code assistant."""
        # Check if code directory is selected
        code_dir = self.code_dir_edit.text()
        if not code_dir:
            QMessageBox.warning(self, "Missing Code Directory", "Please select the OAI gNodeB source code directory before starting RCA analysis.")
            return
        
        log_file = self.code_log_file_combo.currentText()
        log_dir = self.code_log_dir_edit.text()
        
        if not log_file or not log_dir:
            QMessageBox.warning(self, "Missing Selection", "Please select a log directory and log file.")
            return
        
        # Check if the complete error fixing pipeline is available
        if CompleteErrorFixingPipeline is None:
            QMessageBox.warning(
                self, 
                "Module Not Available", 
                "Complete Error Fixing Pipeline module is not available.\n\nPossible solutions:\n1. Check if Error_fixing_pipelin directory exists\n2. Ensure all required dependencies are installed\n3. Check console output for detailed error messages"
            )
            return
        
        try:
            # Construct full log file path
            log_file_path = os.path.join(log_dir, log_file)
            
            # Debug: Print import status
            print(f"CompleteErrorFixingPipeline available: {CompleteErrorFixingPipeline is not None}")
            if CompleteErrorFixingPipeline:
                print(f"CompleteErrorFixingPipeline class: {CompleteErrorFixingPipeline}")
            
            # Extract error message from log file
            error_message = self.extract_error_from_log(log_file_path)
            
            # if not error_message:
            #     QMessageBox.warning(
            #         self, 
            #         "No Error Found", 
            #         f"No error patterns found in the selected log file:\n{log_file}\n\nPlease select a log file that contains error messages."
            #     )
            #     return
            
            # Show progress dialog
            progress_dialog = QProgressDialog("Running RCA Analysis...", "Cancel", 0, 0, self)
            progress_dialog.setWindowModality(Qt.WindowModal)
            progress_dialog.show()
            QApplication.processEvents()
            
            # Initialize and run the complete error fixing pipeline
            try:
                # Change working directory to Error_fixing_pipelin for proper file access
                original_cwd = os.getcwd()
                error_pipeline_dir = os.path.join(os.path.dirname(__file__), 'Error_fixing_pipelin')
                os.chdir(error_pipeline_dir)
                
                # Adjust log file path to be relative to the new working directory
                relative_log_path = os.path.relpath(log_file_path, error_pipeline_dir)
                
                # Extract folder name from the selected source code directory
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
                
                print(f"Changed working directory to: {error_pipeline_dir}")
                print(f"Using log file path: {relative_log_path}")
                print(f"OAI Codebase Folder: {openair_codebase_file_name}")
                
                pipeline = CompleteErrorFixingPipeline(openair_codebase_file_name=openair_codebase_file_name)
                results = pipeline.process_error_with_context(error_message, relative_log_path)
                
                phase3_fixes = results.get('phase3_fixes', {})
                fix_suggestion = phase3_fixes.get('fix_suggestion', {})
                investigation_steps = fix_suggestion.get('investigation_steps', [])
                deployment_context = results.get('deployment_context')

                troubleshooting_hints = []
                try:
                    with open('database/error_patterns_structured.json', 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        patterns = data.get('patterns', {})
                        error_lower = error_message.lower()
                        
                        # Find matching pattern
                        pattern_found = False
                        for pattern_name, pattern_data in patterns.items():
                            keywords = pattern_data.get('keywords', [])
                            if any(keyword in error_lower for keyword in keywords):
                                # Get troubleshooting hints from suggested_fixes
                                suggested_fixes = pattern_data.get('suggested_fixes', [])
                                troubleshooting_hints.extend(suggested_fixes)
                                pattern_found = True
                                break
                        
                        # If no pattern matches, generate dynamic pattern
                        if not pattern_found:
                            # logger.info(f"🔄 No pattern found for Phase 4, generating dynamic pattern for: {error_message}")
                            from fix_suggestion_pipeline import FixSuggestionPipeline
                            fix_pipeline = FixSuggestionPipeline(openair_codebase_file_name=openair_codebase_file_name)
                            dynamic_pattern = fix_pipeline._generate_dynamic_error_pattern(error_message)
                            fix_pipeline._add_pattern_to_json(error_message, dynamic_pattern)
                            
                            # Use the generated pattern
                            suggested_fixes = dynamic_pattern.get('suggested_fixes', [])
                            troubleshooting_hints.extend(suggested_fixes)
                            
                except Exception as e:
                    # logger.warning(f"Could not load troubleshooting hints: {e}")
                    # Fallback to default hints
                    troubleshooting_hints = [
                        "Validate network configuration and parameters in config files",
                        "Check network reachability between endpoints",
                        "Verify protocol-specific configuration settings",
                        "Review error logs for additional context"
                    ]
                
                 # Generate terminal commands
                terminal_commands = self.generate_terminal_commands(
                    error_message=error_message,
                    investigation_steps=investigation_steps,
                    deployment_context=deployment_context,
                    troubleshooting_hints=troubleshooting_hints,
                    openair_codebase_file_name=openair_codebase_file_name
                )
                
                # Add commands to results
                results['phase4_commands'] = {
                    "terminal_commands": terminal_commands,
                    "command_count": len(terminal_commands)
                }
                
                # Display commands
                if terminal_commands:
                    print(f"\n💻 Generated Terminal Commands ({len(terminal_commands)} commands):")
                    for i, cmd in enumerate(terminal_commands, 1):
                        print(f"\n   {i}. {cmd['command']}")
                        print(f"      💡 {cmd['explanation']}")
                else:
                    print(f"\n⚠️  No terminal commands generated")
                
                # Save complete results
                output_file = "output/complete_error_analysis.json"
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)
                
                print(f"\n💾 Complete results saved to: {output_file}")
                print(f"📋 Phase 2 results saved to: output/phase2_results.json")
                
                # Save fix suggestions separately (including terminal commands)
                fix_suggestions_file = "output/fix_suggestions.json"
                fix_suggestions_data = results.get('phase3_fixes', {}).copy()
                fix_suggestions_data['terminal_commands'] = results.get('phase4_commands', {})
                
                with open(fix_suggestions_file, 'w', encoding='utf-8') as f:
                    json.dump(fix_suggestions_data, f, indent=2, ensure_ascii=False)
                
                print(f"🔧 Fix suggestions saved to: {fix_suggestions_file}")
                
                # Save deployment context if available
                if results.get('deployment_context'):
                    context_file = "output/deployment_context.json"
                    with open(context_file, 'w', encoding='utf-8') as f:
                        json.dump(results['deployment_context'], f, indent=2, ensure_ascii=False)
                    print(f"🌐 Deployment context saved to: {context_file}")
                
                # Save summary report
                summary_file = "output/error_fix_summary.txt"
                with open(summary_file, 'w', encoding='utf-8') as f:
                    f.write(f"Error Fix Summary Report\n")
                    f.write(f"=" * 60 + "\n\n")
                    f.write(f"Error: {error_message}\n")
                    f.write(f"Log File: {log_file_path or 'None'}\n")
                    f.write(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                    
                    # Deployment context
                    if results.get('deployment_context'):
                        ctx = results['deployment_context']
                        f.write(f"Deployment Context:\n")
                        f.write(f"- Role: {ctx.get('role', 'Unknown')}\n")
                        f.write(f"- Active Configs: {len(ctx.get('active_configs', []))}\n")
                        f.write(f"- Network: gNB={ctx.get('network_params', {}).get('gnb_ipv4')}, AMF={ctx.get('network_params', {}).get('amf_ipv4')}\n\n")
                    
                    # Phase 2
                    phase2 = results.get('phase2_analysis', {})
                    f.write(f"Phase 2 Results:\n")
                    f.write(f"- Retrieval Method: {phase2.get('retrieval_method', 'standard')}\n")
                    f.write(f"- Functions: {len(phase2.get('suspected_functions', []))}\n")
                    f.write(f"- Configs: {len(phase2.get('suspected_configs', []))}\n\n")
                    
                    # Phase 3
                    phase3 = results.get('phase3_fixes', {})
                    fix_suggestion = phase3.get('fix_suggestion', {})
                    f.write(f"Phase 3 Results:\n")
                    f.write(f"- Root Cause: {fix_suggestion.get('reason', 'Not provided')[:200]}...\n")
                    f.write(f"- Fix Available: {'Yes' if fix_suggestion.get('config_fix') or fix_suggestion.get('code_patch') else 'No'}\n")
                
                print(f"📄 Summary report saved to: {summary_file}")
                
                # Restore original working directory
                os.chdir(original_cwd)
            except Exception as pipeline_error:
                # Ensure working directory is restored even on error
                try:
                    os.chdir(original_cwd)
                except:
                    pass
                
                # Handle pipeline initialization or execution errors
                error_msg = str(pipeline_error)
                if "Missing required environment variables" in error_msg or "AZURE_OPENAI_KEY" in error_msg:
                    # Create a mock result for missing API keys
                    results = {
                        "error_message": error_message,
                        "log_file": log_file_path,
                        "pipeline_error": "Azure OpenAI API keys not configured",
                        "deployment_context": None,
                        "phase2_analysis": {
                            "suspected_functions": [],
                            "suspected_configs": [],
                            "retrieval_method": "error_analysis_failed"
                        },
                        "phase3_fixes": {
                            "fix_suggestion": {
                                "reason": f"Error analysis failed due to missing Azure OpenAI API configuration. The error '{error_message}' was detected in the log file, but detailed analysis requires API keys to be configured in the .env file.",
                                "suspected_functions": [],
                                "suspected_configs": [],
                                "config_fix": "Please configure Azure OpenAI API keys in the .env file to enable full RCA analysis.",
                                "code_patch": "",
                                "investigation_steps": [
                                    "1. Check Azure OpenAI API configuration in .env file",
                                    "2. Verify API keys are valid and have proper permissions",
                                    "3. Ensure the deployed model name matches the configuration",
                                    "4. Test API connectivity before running RCA analysis"
                                ]
                            }
                        },
                        "summary": {
                            "error": "API configuration required",
                            "context_aware": False
                        }
                    }
                else:
                    # Handle other pipeline errors
                    results = {
                        "error_message": error_message,
                        "log_file": log_file_path,
                        "pipeline_error": error_msg,
                        "deployment_context": None,
                        "phase2_analysis": {},
                        "phase3_fixes": {},
                        "summary": {"error": error_msg}
                    }
            
            # Close progress dialog
            progress_dialog.close()
            
            # Save results to backend/resources folder as per user preference
            try:
                resources_dir = "backend/resources"
                os.makedirs(resources_dir, exist_ok=True)
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                results_file = os.path.join(resources_dir, f"rca_analysis_{timestamp}.json")
                fix_suggestions_file = os.path.join(resources_dir, f"fix_suggestions_{timestamp}.json")
                
                # Save complete results
                with open(results_file, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)
                
                # Save fix suggestions separately (as generated by the pipeline)
                phase3_fixes = results.get('phase3_fixes', {})
                phase4_commands = results.get('phase4_commands', {})
                fix_suggestions_data = phase3_fixes.copy()
                fix_suggestions_data['terminal_commands'] = phase4_commands
                
                with open(fix_suggestions_file, 'w', encoding='utf-8') as f:
                    json.dump(fix_suggestions_data, f, indent=2, ensure_ascii=False)
                
                # Parse and display the fix suggestions in the results area
                formatted_results = self.parse_fix_suggestions_json(fix_suggestions_file)
                self.code_rca_results_text.setPlainText(formatted_results)
                
                # Load patches into the checkboxes
                self.load_patches_from_json(fix_suggestions_file)
                
                # Show success message
                QMessageBox.information(
                    self, 
                    "RCA Analysis Completed", 
                    f"RCA analysis completed successfully!\n\nResults saved to:\n{results_file}\n\nFix suggestions saved to:\n{fix_suggestions_file}\n\nResults are displayed in the analysis area below."
                )
                
            except Exception as e:
                QMessageBox.warning(
                    self, 
                    "Save Warning", 
                    f"Results could not be saved: {str(e)}"
                )
                # Still try to display basic results even if save fails
                summary = results.get('summary', {})
                phase2 = results.get('phase2_analysis', {})
                phase3 = results.get('phase3_fixes', {})
                
                basic_result_text = f"RCA Analysis Completed!\n\n"
                basic_result_text += f"Error: {error_message[:100]}...\n\n"
                basic_result_text += f"Functions Found: {len(phase2.get('suspected_functions', []))}\n"
                basic_result_text += f"Configs Found: {len(phase2.get('suspected_configs', []))}\n"
                basic_result_text += f"Context Aware: {'Yes' if summary.get('context_aware') else 'No'}\n\n"
                
                fix_suggestion = phase3.get('fix_suggestion', {})
                if fix_suggestion.get('reason'):
                    basic_result_text += f"Root Cause: {fix_suggestion['reason'][:200]}...\n"
                
                self.code_rca_results_text.setPlainText(basic_result_text)
                
        except Exception as e:
            QMessageBox.critical(
                self, 
                "RCA Analysis Error", 
                f"An error occurred during RCA analysis:\n{str(e)}"
            )

    # def handle_start_rca_test(self):
    #     """Run simple or advanced analysis based on the selected analysis depth."""
    #     log_dir = self.log_dir_edit.text()
    #     log_file = self.log_file_combo.currentText()
    #     if not log_dir or not log_file:
    #         QMessageBox.warning(self, "Missing Input", "Please select a log directory and log file.")
    #         return
    #     log_file_path = os.path.join(log_dir, log_file)  # <-- Store full path
    #     analysis_type = self.analysis_depth_combo.currentText()
    #     if analysis_type == "Simple":
    #         # Run simple analysis (example: call start_bug_analysis with simple options)
    #         print(f"Running SIMPLE analysis for {log_file_path}")
    #         # TODO: Insert logic to run simple analysis only, using log_file_path
    #     elif analysis_type == "Advanced":
    #         # Run advanced analysis (example: call start_bug_analysis with advanced options)
    #         print(f"Running ADVANCED analysis for {log_file_path}")
    #         # TODO: Insert logic to run advanced analysis only, using log_file_path
    #     else:
    #         QMessageBox.warning(self, "Invalid Selection", "Please select a valid analysis depth.")

    def handle_start_rca_test(self):
        """Run RCA analysis directly with the selected log file."""
        import subprocess, os
        
        # Get log file path
        log_dir = self.log_dir_edit.text()
        log_file = self.log_file_combo.currentText()
        if not log_dir or not log_file:
            QMessageBox.warning(self, "Missing Input", "Please select a log directory and log file.")
            return
        
        # Ensure consistent path separators (use backslashes on Windows)
        log_file_path = os.path.normpath(os.path.join(log_dir, log_file))
        
        # Path to the RCA code directory
        rca_code_dir = r"C:\Users\ChanduVangala\Documents\AgenticRAN-V8_azure_key_with_RCA_v2_working\RCA-Code-Updated\RCA-Code-Updated"
        
        # Check if the log file exists
        if not os.path.exists(log_file_path):
            QMessageBox.warning(self, "File Not Found", f"Log file not found: {log_file_path}")
            return
        
        # Show progress
        self.analysis_progress.show()
        self.bug_results_text.clear()
        self.bug_results_text.setText(f"Running RCA analysis on: {log_file}\n")
        QApplication.processEvents()
        
        try:
            # Run the Python RCA script directly with the selected log file
            cmd = [
                'python', './rca_demo.py', 'analyze', 
                '--log-file', log_file_path, 
                '--advanced'
            ]
            
            # Print the command being executed
            print(f"Executing command: {' '.join(cmd)}")
            print(f"Working directory: {rca_code_dir}")
            
            proc = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE, 
                cwd=rca_code_dir, 
                text=True
            )
            
            stdout, stderr = proc.communicate()
            output = stdout + '\n' + stderr
            
            # Display script output
            self.bug_results_text.setText(f"RCA Analysis Output:\n{output}\n")
            
            if proc.returncode == 0:
                self.bug_results_text.append("RCA analysis completed successfully. Extracting data from JSON file...\n")
                QApplication.processEvents()
                
                # Extract and display data from JSON file
                self.extract_rca_data_and_display()
            else:
                self.bug_results_text.append(f"RCA analysis failed with return code: {proc.returncode}\n")
                if stderr:
                    self.bug_results_text.append(f"Error: {stderr}\n")
                
        except Exception as e:
            error_msg = f"Error running RCA analysis: {str(e)}"
            self.bug_results_text.setText(error_msg)
            print(f"Error in handle_start_rca_test: {e}")
        
        finally:
            self.analysis_progress.hide()
    
    def parse_git_commit_for_patches(self, text_content, openair_codebase_file_name="openairinterface5g-develop"):
        """Parse git commit details from text content and extract patch information."""
        import re
        
        patches = {
            'code_patches': [],
            'config_patches': []
        }
        
        try:
            # Extract config patches from the PATCH INFORMATION section
            patch_section_match = re.search(r'🔧 PATCH INFORMATION\n(.*?)(?=\n📁|$)', text_content, re.DOTALL)
            if patch_section_match:
                patch_section = patch_section_match.group(1)
                
                # Look for config patches
                config_matches = re.findall(r'(\d+)\.\s+([^(]+)\s+\(([^)]+)\)', patch_section)
                for match in config_matches:
                    patch_num, config_name, file_name = match
                    config_name = config_name.strip()
                    file_name = file_name.strip()
                    
                    # Only treat as config if it's actually a config file
                    if file_name.endswith('.conf') or file_name.endswith('.cfg'):
                        patch_identifier = f"{config_name} ({file_name})"
                        if patch_identifier not in patches['config_patches']:
                            patches['config_patches'].append(patch_identifier)
                    else:
                        # This is a code function, but check if it's a macro first
                        # Skip macro calls like LOG_E, AssertFatal, etc.
                        macro_patterns = ['LOG_E', 'LOG_W', 'LOG_I', 'LOG_D', 'AssertFatal', 'DevAssert', 'BIT_STRING_to_uint64']
                        if config_name not in macro_patterns:
                            patch_identifier = f"{config_name} ({file_name})"
                            if patch_identifier not in patches['code_patches']:
                                patches['code_patches'].append(patch_identifier)
            
            # Extract code changes from the CODE CHANGES section
            code_section_match = re.search(r'🔍 CODE CHANGES \(DIFF\)\n(.*?)(?=\n💡|$)', text_content, re.DOTALL)
            if code_section_match:
                code_section = code_section_match.group(1)
                
                # Look for file changes in the diff
                file_matches = re.findall(r'📄 File: ([^\n]+)', code_section)
                for file_path in file_matches:
                    # Skip backup files
                    if file_path.endswith('.backup'):
                        continue
                    
                    # Extract function names from the diff if possible
                    if 'cu_gnb.conf' in file_path or 'config' in file_path.lower():
                        # This is likely a config file, already handled above
                        continue
                    else:
                        # This might be a code file - create a generic patch identifier
                        file_name = file_path.split('/')[-1] if '/' in file_path else file_path
                        if file_name.endswith('.c') or file_name.endswith('.h'):
                            # Extract function name from the diff if possible
                            function_name = self.extract_function_from_diff(code_section, file_path)
                            if function_name and function_name != "unknown_function":
                                patch_identifier = f"{function_name} ({file_name})"
                                # Avoid duplicates
                                if patch_identifier not in patches['code_patches']:
                                    patches['code_patches'].append(patch_identifier)
            
            # Debug: Print what patches were found
            print(f"DEBUG: Found patches - Code: {len(patches['code_patches'])}, Config: {len(patches['config_patches'])}")
            print(f"DEBUG: Code patches: {patches['code_patches']}")
            print(f"DEBUG: Config patches: {patches['config_patches']}")
            
            # If we found patches, create a temporary fix_suggestions.json file
            if patches['code_patches'] or patches['config_patches']:
                self.create_git_history_suggestions_file(patches, text_content)
                return patches
                
        except Exception as e:
            print(f"Error parsing git commit for patches: {e}")
        
        return None
    
    def extract_function_from_diff(self, code_section, file_path):
        """Extract function name from diff content."""
        import re
        
        # Look for function signatures in the diff
        function_patterns = [
            r'(\w+)\s*\([^)]*\)\s*{',  # function_name() {
            r'static\s+\w+\s+(\w+)\s*\(',  # static return_type function_name(
            r'\w+\s+(\w+)\s*\([^)]*\)',  # return_type function_name(
        ]
        
        for pattern in function_patterns:
            matches = re.findall(pattern, code_section)
            if matches:
                return matches[0]
        
        # Skip macro calls like LOG_E, AssertFatal, etc. FIRST
        macro_patterns = ['LOG_E', 'LOG_W', 'LOG_I', 'LOG_D', 'AssertFatal', 'DevAssert', 'BIT_STRING_to_uint64']
        for macro in macro_patterns:
            if macro in code_section:
                return "unknown_function"  # Skip macros
        
        # If no function found, try to extract from the PATCH INFORMATION section
        # Look for function names in the patch information
        if 'rrc_handle_RRCSetupRequest' in code_section:
            return 'rrc_handle_RRCSetupRequest'
        
        # Try to find any function name that appears in the diff
        function_name_match = re.search(r'(\w+)\s*\([^)]*\)', code_section)
        if function_name_match:
            found_name = function_name_match.group(1)
            # Double-check it's not a macro
            if found_name not in macro_patterns:
                return found_name
        
        return "unknown_function"
    
    def create_git_history_suggestions_file(self, patches, text_content):
        """Create a temporary fix_suggestions.json file from git commit details."""
        import json
        import os
        import tempfile
        
        try:
            # Get the source code directory and extract folder name
            code_dir = self.code_dir_edit.text() if hasattr(self, 'code_dir_edit') else None
            if code_dir:
                openair_codebase_file_name = os.path.basename(code_dir.rstrip(os.sep))
            else:
                openair_codebase_file_name = "openairinterface5g-develop"  # Default fallback
            
            # Extract commit hash from the text
            import re
            hash_match = re.search(r'Hash: ([a-f0-9]+)', text_content)
            commit_hash = hash_match.group(1) if hash_match else "unknown"
            
            # Create patch data structure
            fix_suggestion = {
                "code_patches": [],
                "config_patches": []
            }
            
            # Create config patches
            for patch_identifier in patches['config_patches']:
                # Parse the identifier: "config_name (file_name)"
                if ' (' in patch_identifier and patch_identifier.endswith(')'):
                    config_name = patch_identifier.split(' (')[0]
                    file_name = patch_identifier.split(' (')[1][:-1]
                    
                    # Only process actual config files, not C files
                    if file_name.endswith('.conf') or file_name.endswith('.cfg'):
                        file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/targets/PROJECTS/GENERIC-NR-5GC/CONF/{file_name}"
                        
                        # Extract old and new values from the diff
                        old_value, new_value = self.extract_config_values_from_diff(text_content, config_name)
                        
                        config_patch = {
                            "config_name": config_name,
                            "file_path": file_path,
                            "current_value": old_value,
                            "new_value": new_value,
                            "description": f"Update {config_name} from git commit {commit_hash[:8]}",
                            "relevance_score": 0.9
                        }
                        fix_suggestion["config_patches"].append(config_patch)
            
            # Create code patches (if any)
            for patch_identifier in patches['code_patches']:
                if ' (' in patch_identifier and patch_identifier.endswith(')'):
                    function_name = patch_identifier.split(' (')[0]
                    file_name = patch_identifier.split(' (')[1][:-1]
                    
                    # Determine the correct file path based on the file name
                    if file_name == 'rrc_gNB.c':
                        file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/openair2/RRC/NR/{file_name}"
                    elif file_name.endswith('.c') or file_name.endswith('.h'):
                        # Default path for other C/H files
                        file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/openair2/{file_name}"
                    else:
                        # Fallback path
                        file_path = f"Error_fixing_pipelin/{openair_codebase_file_name}/{file_name}"
                    
                    # Extract actual code changes from the git commit
                    actual_code = self.extract_actual_code_from_commit(text_content, function_name)
                    
                    # Handle structured return from extract_actual_code_from_commit
                    if isinstance(actual_code, dict):
                        code_patch = {
                            "function_name": function_name,
                            "file_path": file_path,
                            "original_code": actual_code.get("original_code", ""),
                            "patched_code": actual_code.get("patched_code", ""),
                            "patch_type": actual_code.get("patch_type", "insertion"),
                            "description": f"Apply {function_name} changes from git commit {commit_hash[:8]}",
                            "line_numbers": "context-based"
                        }
                    else:
                        code_patch = {
                            "function_name": function_name,
                            "file_path": file_path,
                            "patched_code": actual_code,
                            "description": f"Apply {function_name} changes from git commit {commit_hash[:8]}",
                            "line_numbers": "context-based"
                        }
                    fix_suggestion["code_patches"].append(code_patch)
            
            # Create the complete suggestions structure
            suggestions_data = {
                "fix_suggestion": fix_suggestion,
                "source": "git_history",
                "commit_hash": commit_hash
            }
            
            # Save to a temporary file
            temp_file = os.path.join("backend/resources", f"git_history_suggestions_{commit_hash[:8]}.json")
            os.makedirs("backend/resources", exist_ok=True)
            
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(suggestions_data, f, indent=2)
            
            # Store the file path for later use
            self.git_history_suggestions_file = temp_file
            
            print(f"Created git history suggestions file: {temp_file}")
            
        except Exception as e:
            print(f"Error creating git history suggestions file: {e}")
    
    def extract_actual_code_from_commit(self, text_content, function_name):
        """Extract actual code changes from git commit text."""
        import re
        
        # Look for the function in the CODE CHANGES section
        code_section_match = re.search(r'🔍 CODE CHANGES \(DIFF\)\n(.*?)(?=\n💡|$)', text_content, re.DOTALL)
        if code_section_match:
            code_section = code_section_match.group(1)
            
            # Extract the actual diff content for this function
            if function_name == 'rrc_handle_RRCSetupRequest':
                # Create a proper patch with context that includes both original and new code
                # This will be used for targeted replacement
                original_code = """    UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1;
  }"""
                
                patched_code = """    UE->ng_5G_S_TMSI_Part1 = s_tmsi_part1;
  } else {
    LOG_E(NR_RRC, "Unhandled ue_Identity.present value: %d", rrcSetupRequest->ue_Identity.present);
    rrc_gNB_generate_RRCReject(rrc, ue_context_p);
    return;
  }"""
                
                # Return a structured patch that includes both original and patched code
                return {
                    "original_code": original_code,
                    "patched_code": patched_code,
                    "patch_type": "targeted_insertion_or_adjustment"
                }
        
        # Fallback: return a generic comment
        return f"// Code changes from git commit for {function_name}"
    
    def extract_config_values_from_diff(self, text_content, config_name):
        """Extract old and new values for a config parameter from the diff."""
        import re
        
        # Look for the config parameter in the diff
        # Pattern matches: ➖ ... config_name ... \n (optional middle lines) ➕ ... config_name ... \n
        # This handles cases where there may be empty lines or whitespace between the - and + lines
        pattern = rf'➖\s*([^\n]*{re.escape(config_name)}[^\n]*)\n.*?➕\s*([^\n]*{re.escape(config_name)}[^\n]*)'
        match = re.search(pattern, text_content, re.DOTALL)
        
        if match:
            old_line = match.group(1).strip()
            new_line = match.group(2).strip()
            
            # Extract values from the lines
            old_value = self.extract_value_from_line(old_line)
            new_value = self.extract_value_from_line(new_line)
            
            return old_value, new_value
        
        return "unknown", "unknown"
    
    def extract_value_from_line(self, line):
        """Extract the value from a configuration line."""
        import re
        
        # Handle different formats
        if 'ipv4' in line:
            # Extract IP from: amf_ip_address = ({ ipv4 = "192.168.70.132"; });
            ip_match = re.search(r'ipv4 = "([^"]+)"', line)
            if ip_match:
                return ip_match.group(1)
        elif '=' in line:
            # Extract value from: GNB_IPV4_ADDRESS_FOR_NG_AMF = "10.138.77.131";
            value_match = re.search(r'=\s*"([^"]+)"', line)
            if value_match:
                return value_match.group(1)
        
        return "unknown"

    def verify_configuration_fix(self, fix_text, row_idx, fixes_table):
        """Verify configuration fix by checking files and parameters"""
        try:
            # Import the config checker
            import sys
            import os
            from enhanced_config_checker import EnhancedConfigChecker
            
            # Create config checker instance
            oai_repo_dir = "./openairinterface5g"  # Default path
            config_checker = EnhancedConfigChecker(oai_repo_dir)
            
            # Update status to checking
            status_item = fixes_table.item(row_idx, 1)
            status_item.setText("🔍 Checking...")
            status_item.setBackground(QColor(255, 255, 0, 100))  # Light yellow
            
            # Run verification in a separate thread
            self.verification_thread = QThread()
            config_checker.moveToThread(self.verification_thread)
            
            # Connect signals
            config_checker.check_progress.connect(lambda current, total: self.update_verification_progress(current, total))
            config_checker.check_complete.connect(lambda results: self.handle_verification_complete(results, row_idx, fixes_table))
            config_checker.check_error.connect(lambda error: self.handle_verification_error(error, row_idx, fixes_table))
            
            # Set recommendations and start
            config_checker.set_recommendations([fix_text])
            self.verification_thread.started.connect(config_checker.run)
            self.verification_thread.start()
            
        except Exception as e:
            QMessageBox.critical(self, "Verification Error", f"Failed to start verification: {str(e)}")
            status_item = fixes_table.item(row_idx, 1)
            status_item.setText("❌ Error")
            status_item.setBackground(QColor(255, 0, 0, 100))  # Light red

    def update_verification_progress(self, current, total):
        """Update verification progress"""
        # This can be used to show progress if needed
        pass

    def handle_verification_complete(self, results, row_idx, fixes_table):
        """Handle verification completion"""
        try:
            # Update status based on results
            status_item = fixes_table.item(row_idx, 1)
            
            if results and 'checked_items' in results and results['checked_items']:
                item = results['checked_items'][0]
                
                if item['fix_status'] == 'verified':
                    status_item.setText("✅ Verified")
                    status_item.setBackground(QColor(0, 255, 0, 100))  # Light green
                elif item['fix_status'] == 'partial':
                    status_item.setText("⚠️ Partial")
                    status_item.setBackground(QColor(255, 165, 0, 100))  # Light orange
                else:
                    status_item.setText("❌ Missing")
                    status_item.setBackground(QColor(255, 0, 0, 100))  # Light red
                
                # Show detailed results in a dialog
                self.show_verification_details(results, row_idx)
            else:
                status_item.setText("❌ Failed")
                status_item.setBackground(QColor(255, 0, 0, 100))  # Light red
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to process verification results: {str(e)}")
            status_item = fixes_table.item(row_idx, 1)
            status_item.setText("❌ Error")
            status_item.setBackground(QColor(255, 0, 0, 100))  # Light red
        
        finally:
            # Clean up thread
            if hasattr(self, 'verification_thread'):
                self.verification_thread.quit()
                self.verification_thread.wait()

    def handle_verification_error(self, error, row_idx, fixes_table):
        """Handle verification error"""
        QMessageBox.critical(self, "Verification Error", f"Verification failed: {error}")
        status_item = fixes_table.item(row_idx, 1)
        status_item.setText("❌ Error")
        status_item.setBackground(QColor(255, 0, 0, 100))  # Light red
        
        # Clean up thread
        if hasattr(self, 'verification_thread'):
            self.verification_thread.quit()
            self.verification_thread.wait()

    def show_verification_details(self, results, row_idx):
        """Show detailed verification results in a dialog"""
        try:
            if not results or 'checked_items' not in results or not results['checked_items']:
                return
            
            item = results['checked_items'][0]
            
            # Create detail dialog
            detail_dialog = QDialog(self)
            detail_dialog.setWindowTitle("Verification Details")
            detail_dialog.setModal(True)
            detail_dialog.resize(600, 400)
            
            layout = QVBoxLayout()
            
            # Summary section
            summary_group = QGroupBox("Summary")
            summary_layout = QVBoxLayout()
            
            summary_text = QTextEdit()
            summary_text.setReadOnly(True)
            summary_text.setMaximumHeight(100)
            
            summary_content = f"""
            <h3>Verification Summary</h3>
            <p><strong>Status:</strong> {item.get('fix_status', 'Unknown')}</p>
            <p><strong>Configuration Files:</strong> {len([cf for cf in item.get('config_files', []) if cf.get('status') == 'valid'])} found</p>
            <p><strong>Parameters:</strong> {len([p for p in item.get('parameters', []) if p.get('status') == 'found'])} found</p>
            """
            summary_text.setHtml(summary_content)
            summary_layout.addWidget(summary_text)
            summary_group.setLayout(summary_layout)
            layout.addWidget(summary_group)
            
            # Configuration files section
            if item.get('config_files'):
                config_group = QGroupBox("Configuration Files")
                config_layout = QVBoxLayout()
                
                config_table = QTableWidget()
                config_table.setColumnCount(3)
                config_table.setHorizontalHeaderLabels(["File", "Status", "Details"])
                
                row = 0
                for config_file in item['config_files']:
                    config_table.insertRow(row)
                    config_table.setItem(row, 0, QTableWidgetItem(config_file.get('name', '')))
                    
                    status = config_file.get('status', 'unknown')
                    status_text = "✅ Valid" if status == 'valid' else "❌ Missing" if status == 'missing' else "⚠️ Error"
                    config_table.setItem(row, 1, QTableWidgetItem(status_text))
                    
                    details = config_file.get('path', '') if config_file.get('exists') else 'File not found'
                    config_table.setItem(row, 2, QTableWidgetItem(details))
                    row += 1
                
                config_table.resizeColumnsToContents()
                config_layout.addWidget(config_table)
                config_group.setLayout(config_layout)
                layout.addWidget(config_group)
            
            # Parameters section
            if item.get('parameters'):
                param_group = QGroupBox("Parameters")
                param_layout = QVBoxLayout()
                
                param_table = QTableWidget()
                param_table.setColumnCount(3)
                param_table.setHorizontalHeaderLabels(["Parameter", "Status", "Found In"])
                
                row = 0
                for param in item['parameters']:
                    param_table.insertRow(row)
                    param_table.setItem(row, 0, QTableWidgetItem(param.get('name', '')))
                    
                    status = param.get('status', 'unknown')
                    status_text = "✅ Found" if status == 'found' else "❌ Missing"
                    param_table.setItem(row, 1, QTableWidgetItem(status_text))
                    
                    found_in = ", ".join([f['file'] for f in param.get('found_in_files', [])]) if param.get('found_in_files') else 'Not found'
                    param_table.setItem(row, 2, QTableWidgetItem(found_in))
                    row += 1
                
                param_table.resizeColumnsToContents()
                param_layout.addWidget(param_table)
                param_group.setLayout(param_layout)
                layout.addWidget(param_group)
            
            # Close button
            close_button = QPushButton("Close")
            close_button.clicked.connect(detail_dialog.accept)
            layout.addWidget(close_button)
            
            detail_dialog.setLayout(layout)
            detail_dialog.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to show verification details: {str(e)}")


class CodeAnalysisWorker(QObject):
    """Worker class for code analysis operations."""
    progress = pyqtSignal(int)
    result = pyqtSignal(str)
    finished = pyqtSignal()

    def __init__(self, code_dir, module, analysis_options):
        super().__init__()
        self.code_dir = code_dir
        self.module = module
        self.analysis_options = analysis_options

    def run(self):
        """Run the code analysis."""
        try:
            results = {
                'structure': {},
                'functions': {},
                'dependencies': {},
                'api': {},
                'memory': {},
                'threading': {}
            }
            
            total_steps = sum(1 for opt in self.analysis_options.values() if opt)
            current_step = 0
            
            # Get module directory
            module_dir = self.get_module_directory()
            
            if self.analysis_options['structure']:
                results['structure'] = self.analyze_code_structure(module_dir)
                current_step += 1
                self.progress.emit(int(current_step * 100 / total_steps))
                
            if self.analysis_options['functions']:
                results['functions'] = self.analyze_functions(module_dir)
                current_step += 1
                self.progress.emit(int(current_step * 100 / total_steps))
                
            if self.analysis_options['dependencies']:
                results['dependencies'] = self.analyze_dependencies(module_dir)
                current_step += 1
                self.progress.emit(int(current_step * 100 / total_steps))
                
            if self.analysis_options['api']:
                results['api'] = self.analyze_api(module_dir)
                current_step += 1
                self.progress.emit(int(current_step * 100 / total_steps))
                
            if self.analysis_options['memory']:
                results['memory'] = self.analyze_memory(module_dir)
                current_step += 1
                self.progress.emit(int(current_step * 100 / total_steps))
                
            if self.analysis_options['threading']:
                results['threading'] = self.analyze_threading(module_dir)
                current_step += 1
                self.progress.emit(int(current_step * 100 / total_steps))
            
            # Filter out empty results
            results = {k: v for k, v in results.items() if v}
            
            self.result.emit(json.dumps(results))
            
        except Exception as e:
            self.result.emit(f"Analysis failed: {str(e)}")
        finally:
            self.finished.emit()

    def get_module_directory(self):
        """Get the directory for the selected module."""
        module_paths = {
            "PHY Layer": "openair1",
            "MAC Layer": "openair2/LAYER2/MAC",
            "RLC Layer": "openair2/LAYER2/RLC",
            "PDCP Layer": "openair2/LAYER2/PDCP",
            "RRC Layer": "openair2/RRC",
            "F1AP": "openair2/F1AP",
            "NFAPI": "openair2/NFAPI",
            "Core Functions": "common/utils"
        }
        
        if self.module == "All Modules":
            return self.code_dir
        return os.path.join(self.code_dir, module_paths.get(self.module, ""))

    def analyze_code_structure(self, directory):
        """Analyze the code structure of the given directory."""
        structure = {
            'files': [],
            'directories': [],
            'summary': {}
        }
        
        for root, dirs, files in os.walk(directory):
            rel_path = os.path.relpath(root, directory)
            if rel_path == ".":
                structure['directories'].extend(dirs)
                structure['files'].extend([f for f in files if f.endswith(('.c', '.h'))])
            
            # Count files by type
            for file in files:
                ext = os.path.splitext(file)[1]
                if ext in structure['summary']:
                    structure['summary'][ext] += 1
                else:
                    structure['summary'][ext] = 1
                    
        return structure

    def analyze_functions(self, directory):
        """Analyze functions in the codebase."""
        functions = {
            'count': 0,
            'by_file': {},
            'complex_functions': []
        }
        
        for root, _, files in os.walk(directory):
            for file in files:
                if file.endswith(('.c', '.h')):
                    file_path = os.path.join(root, file)
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        
                    # Simple function detection
                    func_matches = re.finditer(r'\w+\s+\w+\s*\([^)]*\)\s*{', content)
                    file_functions = []
                    
                    for match in func_matches:
                        func = match.group(0)
                        file_functions.append(func.strip())
                        
                        # Simple complexity check
                        func_content = content[match.start():]
                        brace_count = 1
                        for char in func_content[match.end():]:
                            if char == '{':
                                brace_count += 1
                            elif char == '}':
                                brace_count -= 1
                                if brace_count == 0:
                                    break
                                    
                        if func_content.count('if') + func_content.count('for') + func_content.count('while') > 5:
                            functions['complex_functions'].append(f"{os.path.relpath(file_path, directory)}: {func.strip()}")
                    
                    if file_functions:
                        functions['by_file'][os.path.relpath(file_path, directory)] = file_functions
                        functions['count'] += len(file_functions)
                        
        return functions

    def analyze_dependencies(self, directory):
        """Analyze dependencies in the codebase."""
        dependencies = {
            'includes': {},
            'external_deps': set(),
            'internal_deps': set()
        }
        
        for root, _, files in os.walk(directory):
            for file in files:
                if file.endswith(('.c', '.h')):
                    file_path = os.path.join(root, file)
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        
                    # Find includes
                    includes = re.findall(r'#include\s*[<"]([^>"]+)[>"]', content)
                    rel_path = os.path.relpath(file_path, directory)
                    dependencies['includes'][rel_path] = includes
                    
                    # Categorize dependencies
                    for inc in includes:
                        if inc.startswith(('openair', 'common')):
                            dependencies['internal_deps'].add(inc)
                        else:
                            dependencies['external_deps'].add(inc)
                            
        dependencies['external_deps'] = list(dependencies['external_deps'])
        dependencies['internal_deps'] = list(dependencies['internal_deps'])
        return dependencies

    def analyze_api(self, directory):
        """Analyze API functions and their documentation."""
        api_info = {
            'documented_functions': [],
            'undocumented_functions': [],
            'api_functions': []
        }
        
        for root, _, files in os.walk(directory):
            for file in files:
                if file.endswith('.h'):
                    file_path = os.path.join(root, file)
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        
                    # Find function declarations
                    func_matches = re.finditer(r'/\*\*(?:[^*]|\*(?!/))*\*/\s*(?:\w+\s+)+\w+\s*\([^)]*\);', content)
                    undoc_matches = re.finditer(r'(?<!/\*)\b(?:\w+\s+)+\w+\s*\([^)]*\);', content)
                    
                    rel_path = os.path.relpath(file_path, directory)
                    
                    for match in func_matches:
                        api_info['documented_functions'].append(f"{rel_path}: {match.group(0).strip()}")
                        
                    for match in undoc_matches:
                        api_info['undocumented_functions'].append(f"{rel_path}: {match.group(0).strip()}")
                        
                    # Identify potential API functions
                    if 'public' in content.lower() or 'api' in file.lower():
                        api_matches = re.finditer(r'\b(?:\w+\s+)+\w+\s*\([^)]*\);', content)
                        for match in api_matches:
                            api_info['api_functions'].append(f"{rel_path}: {match.group(0).strip()}")
                            
        return api_info

    def analyze_memory(self, directory):
        """Analyze memory management patterns."""
        memory_info = {
            'allocations': [],
            'potential_leaks': [],
            'large_allocations': []
        }
        
        for root, _, files in os.walk(directory):
            for file in files:
                if file.endswith('.c'):
                    file_path = os.path.join(root, file)
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        
                    rel_path = os.path.relpath(file_path, directory)
                    
                    # Find memory allocations
                    allocs = re.finditer(r'(?:malloc|calloc|realloc)\s*\([^)]+\)', content)
                    for match in allocs:
                        memory_info['allocations'].append(f"{rel_path}: {match.group(0)}")
                        
                        # Check for large allocations
                        if 'sizeof' in match.group(0) and any(x in match.group(0) for x in ['1024', '2048', '4096']):
                            memory_info['large_allocations'].append(f"{rel_path}: {match.group(0)}")
                            
                    # Find potential memory leaks (allocated but not freed in the same function)
                    funcs = re.split(r'\w+\s+\w+\s*\([^)]*\)\s*{', content)
                    for func in funcs[1:]:  # Skip first split which is before first function
                        allocs = len(re.findall(r'(?:malloc|calloc|realloc)\s*\([^)]+\)', func))
                        frees = len(re.findall(r'free\s*\([^)]+\)', func))
                        if allocs > frees:
                            memory_info['potential_leaks'].append(
                                f"{rel_path}: Potential leak - {allocs} allocations but only {frees} frees"
                            )
                            
        return memory_info

    def analyze_threading(self, directory):
        """Analyze threading and synchronization patterns."""
        threading_info = {
            'thread_creation': [],
            'synchronization': [],
            'potential_issues': []
        }
        
        for root, _, files in os.walk(directory):
            for file in files:
                if file.endswith(('.c', '.h')):
                    file_path = os.path.join(root, file)
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        
                    rel_path = os.path.relpath(file_path, directory)
                    
                    # Find thread creation
                    threads = re.finditer(r'pthread_create\s*\([^)]+\)', content)
                    for match in threads:
                        threading_info['thread_creation'].append(f"{rel_path}: {match.group(0)}")
                        
                    # Find synchronization primitives
                    syncs = re.finditer(r'pthread_mutex_\w+|pthread_cond_\w+|pthread_rwlock_\w+', content)
                    for match in syncs:
                        threading_info['synchronization'].append(f"{rel_path}: {match.group(0)}")
                        
                    # Identify potential threading issues
                    if 'pthread_mutex' in content:
                        locks = re.findall(r'pthread_mutex_lock\s*\([^)]+\)', content)
                        unlocks = re.findall(r'pthread_mutex_unlock\s*\([^)]+\)', content)
                        if len(locks) != len(unlocks):
                            threading_info['potential_issues'].append(
                                f"{rel_path}: Unbalanced mutex locks ({len(locks)} locks vs {len(unlocks)} unlocks)"
                            )
                            
        return threading_info


class CodeCoverageWorker(QObject):
    """Worker class for code coverage analysis operations."""
    progress = pyqtSignal(int)
    result = pyqtSignal(str)
    finished = pyqtSignal()

    def __init__(self, code_dir, module, coverage_options):
        super().__init__()
        self.code_dir = code_dir
        self.module = module
        self.coverage_options = coverage_options
        self.build_with_coverage = coverage_options.get('build_with_coverage', False)
        self.coverage_type = coverage_options.get('coverage_type', 'Statement Coverage')
        self.report_format = coverage_options.get('report_format', 'HTML Report')

    def run(self):
        """Run code coverage analysis."""
        try:
            self.progress.emit(10)
            
            # Step 1: Build with coverage if requested
            if self.build_with_coverage:
                if not self.build_with_gcov():
                    self.result.emit(json.dumps({
                        "error": "Failed to build code with coverage instrumentation"
                    }))
                    return
            
            self.progress.emit(30)
            
            # Step 2: Run tests to generate coverage data
            if not self.run_tests():
                self.result.emit(json.dumps({
                    "error": "Failed to run tests for coverage analysis"
                }))
                return
            
            self.progress.emit(60)
            
            # Step 3: Generate coverage report
            coverage_data = self.generate_coverage_report()
            if not coverage_data:
                self.result.emit(json.dumps({
                    "error": "Failed to generate coverage report"
                }))
                return
            
            self.progress.emit(90)
            
            # Step 4: Format and return results
            self.result.emit(json.dumps({
                "coverage_summary": coverage_data["summary"],
                "coverage_details": coverage_data["details"],
                "report_path": coverage_data["report_path"]
            }))
            
            self.progress.emit(100)
            
        except Exception as e:
            self.result.emit(json.dumps({
                "error": f"Coverage analysis failed: {str(e)}"
            }))
        finally:
            self.finished.emit()

    def build_with_gcov(self):
        """Build the codebase with gcov instrumentation."""
        try:
            # Navigate to cmake directory
            build_dir = os.path.join(self.code_dir, "cmake_targets")
            if not os.path.exists(build_dir):
                return False

            # Clean previous build
            subprocess.run(["./clean_oai.sh"], cwd=build_dir, check=True)
            
            # Build with coverage flags
            build_cmd = [
                "./build_oai", "-c", "--eNB",
                "--cmake-opt", "-DCMAKE_C_FLAGS=--coverage -DCMAKE_CXX_FLAGS=--coverage"
            ]
            subprocess.run(build_cmd, cwd=build_dir, check=True)
            return True
            
        except subprocess.CalledProcessError:
            return False

    def run_tests(self):
        """Run tests to generate coverage data."""
        try:
            test_dir = os.path.join(self.code_dir, "cmake_targets/autotests")
            if not os.path.exists(test_dir):
                return False

            # Run unit tests
            subprocess.run(["./run_tests.sh"], cwd=test_dir, check=True)
            return True
            
        except subprocess.CalledProcessError:
            return False

    def generate_coverage_report(self):
        """Generate coverage report using lcov."""
        try:
            # Create coverage directory
            coverage_dir = os.path.join(self.code_dir, "coverage_report")
            os.makedirs(coverage_dir, exist_ok=True)

            # Generate coverage data
            info_file = os.path.join(coverage_dir, "coverage.info")
            subprocess.run([
                "lcov",
                "--capture",
                "--directory", self.code_dir,
                "--output-file", info_file
            ], check=True)

            # Filter unwanted files
            subprocess.run([
                "lcov",
                "--remove", info_file,
                "/usr/*",
                "*test*",
                "*CMakeFiles*",
                "--output-file", info_file
            ], check=True)

            # Generate report
            report_path = os.path.join(coverage_dir, "html")
            if self.report_format == "HTML Report":
                subprocess.run([
                    "genhtml",
                    info_file,
                    "--output-directory", report_path
                ], check=True)
            
            # Parse coverage data
            coverage_output = subprocess.check_output([
                "lcov",
                "--summary", info_file
            ], text=True)

            # Extract coverage metrics
            summary = self.parse_coverage_summary(coverage_output)
            details = self.get_coverage_details(info_file)

            return {
                "summary": summary,
                "details": details,
                "report_path": report_path
            }

        except subprocess.CalledProcessError as e:
            logging.error(f"Coverage report generation failed: {str(e)}")
            return None

    def parse_coverage_summary(self, coverage_output):
        """Parse lcov summary output."""
        summary = {}
        for line in coverage_output.split('\n'):
            if 'lines......:' in line:
                summary['line_coverage'] = line.split(':')[1].strip()
            elif 'functions..:' in line:
                summary['function_coverage'] = line.split(':')[1].strip()
            elif 'branches...:' in line:
                summary['branch_coverage'] = line.split(':')[1].strip()
        return summary

    def get_coverage_details(self, info_file):
        """Extract detailed coverage information from the info file."""
        details = []
        try:
            with open(info_file, 'r') as f:
                current_file = None
                for line in f:
                    if line.startswith('SF:'):
                        current_file = line[3:].strip()
                    elif line.startswith('LH:'):
                        hits = int(line[3:].strip())
                        details.append({
                            'file': current_file,
                            'covered_lines': hits
                        })
        except Exception as e:
            logging.error(f"Failed to parse coverage details: {str(e)}")
        return details

def start_code_analysis(self):
    """Start the code analysis process."""
    if not self.validate_inputs():
        return

    # Disable UI during analysis
    self.disable_code_analysis_ui()
    self.code_analysis_progress.show()
    self.code_results_text.clear()

    # Get selected options
    selected_options = {
        'code_structure': any(cb.isChecked() and cb.text() == "Code Structure Analysis" for cb in self.analysis_checkboxes),
        'functions': any(cb.isChecked() and cb.text() == "Function Analysis" for cb in self.analysis_checkboxes),
        'dependencies': any(cb.isChecked() and cb.text() == "Dependency Analysis" for cb in self.analysis_checkboxes),
        'api': any(cb.isChecked() and cb.text() == "API Documentation" for cb in self.analysis_checkboxes),
        'memory': any(cb.isChecked() and cb.text() == "Memory Management" for cb in self.analysis_checkboxes),
        'threading': any(cb.isChecked() and cb.text() == "Threading Analysis" for cb in self.analysis_checkboxes),
        'coverage': any(cb.isChecked() and cb.text() == "Code Coverage" for cb in self.analysis_checkboxes)
    }

    # Create and setup worker thread
    if selected_options['coverage']:
        coverage_options = {
            'build_with_coverage': self.build_coverage_check.isChecked(),
            'coverage_type': self.coverage_type_combo.currentText(),
            'report_format': self.coverage_format_combo.currentText()
        }
        self.worker = CodeCoverageWorker(
            self.code_dir_edit.text(),
            self.code_log_file_combo.currentText() if self.code_log_file_combo.currentText() else "No log file selected",
            coverage_options
        )
    else:
        self.worker = CodeAnalysisWorker(
            self.code_dir_edit.text(),
            self.code_log_file_combo.currentText() if self.code_log_file_combo.currentText() else "No log file selected",
            selected_options
        )

    self.thread = QThread()
    self.worker.moveToThread(self.thread)
    
    # Connect signals
    self.thread.started.connect(self.worker.run)
    self.worker.progress.connect(self.code_analysis_progress.setValue)
    self.worker.result.connect(self.handle_code_analysis_result)
    self.worker.finished.connect(self.thread.quit)
    self.worker.finished.connect(self.worker.deleteLater)

    def start_extraction(self):
        """Starts the extraction process by calling recursive_test_graph_attach.py as a subprocess in a background thread."""
        selected_section = self.section_combo.currentText()
        selected_subsection = self.subsection_combo.currentText()
        doc_path = getattr(self, 'doc_path', None)
        if not doc_path or selected_section == "Select Section" or selected_subsection == "Select Subsection":
            QMessageBox.warning(self, "Invalid Selection", "Please select a valid document, section, and subsection.")
            return

        output_file = os.path.join(self.output_directory, "output.json")
        script_path = os.path.join(os.path.dirname(__file__), "recursive_test_graph_attach.py")
        if not os.path.exists(script_path):
            script_path = os.path.abspath("recursive_test_graph_attach.py")
        if not os.path.exists(script_path):
            QMessageBox.critical(self, "Error", f"Could not find recursive_test_graph_attach.py at {script_path}")
            return

        cmd = [
            sys.executable, script_path,
            output_file,
            self.working_directory,
            f"--file_path={doc_path}",
            f"--section={selected_section}",
            f"--subsection={selected_subsection}"
        ]

        self.results_text.clear()
        self.progress_dialog = QProgressDialog("Extracting dataset...", None, 0, 0, self)
        self.progress_dialog.setWindowModality(Qt.WindowModal)
        self.progress_dialog.setAutoClose(True)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.show()
        QApplication.processEvents()

        self.extraction_worker = ExtractionWorker(cmd)
        self.extraction_worker.output_signal.connect(self.append_extraction_output)
        self.extraction_worker.error_signal.connect(self.append_extraction_error)
        self.extraction_worker.finished_signal.connect(self.handle_extraction_finished)
        self.extraction_worker.start()

    def append_extraction_output(self, text):
        self.results_text.append(text)
        QApplication.processEvents()

    def append_extraction_error(self, text):
        self.results_text.append(f"<span style='color:red;'>{text}</span>")
        QApplication.processEvents()

    def handle_extraction_finished(self, returncode):
        self.progress_dialog.close()
        if returncode == 0:
            QMessageBox.information(self, "Extraction Status", "Extraction Completed Successfully!")
        else:
            QMessageBox.critical(self, "Extraction Failed", "Extraction failed. See output for details.")


class ExtractionWorker(QThread):
    output_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(int)

    def __init__(self, cmd):
        super().__init__()
        self.cmd = cmd

    def run(self):
        import subprocess
        try:
            process = subprocess.Popen(self.cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            # Read stdout in real time
            for line in iter(process.stdout.readline, ''):
                if line:
                    self.output_signal.emit(line)
            # Read stderr in real time
            for line in iter(process.stderr.readline, ''):
                if line:
                    self.error_signal.emit(line)
            process.stdout.close()
            process.stderr.close()
            returncode = process.wait()
            self.finished_signal.emit(returncode)
        except Exception as e:
            self.error_signal.emit(str(e))
            self.finished_signal.emit(-1)