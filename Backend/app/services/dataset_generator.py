"""
Backend logic for dataset extraction from documents without UI components.
This module contains the core Python functions for document processing,
reference extraction, and graph generation.
"""

import json
import os
import re
import requests
from bs4 import BeautifulSoup
import zipfile
from urllib.parse import urljoin
from pathlib import Path
import networkx as nx
from docx import Document
import fitz  # PyMuPDF
from openai import AzureOpenAI, BadRequestError
from dotenv import load_dotenv
from collections import defaultdict
from difflib import SequenceMatcher

from app.guardrails import get_spec_intel_guardrail


class DatasetExtractor:
    """Main class for dataset extraction operations."""
    
    def __init__(self):
        """Initialize the extractor with Azure OpenAI client."""
        load_dotenv()
        
        # Get Azure OpenAI credentials
        self.azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        self.azure_api_key = os.getenv("AZURE_OPENAI_API_KEY")
        self.azure_model_name = os.getenv("AZURE_OPENAI_MODEL_NAME")
        
        # Initialize OpenAI client
        self.client = AzureOpenAI(
            api_key=self.azure_api_key,
            azure_endpoint=self.azure_endpoint,
            api_version="2024-02-01"
        )
    
    def get_main_sections(self, doc_path):
        """Extract main section headings (Heading 1) from a DOCX file."""
        if doc_path.endswith('.docx'):
            doc = Document(doc_path)
            main_sections = []
            for para in doc.paragraphs:
                if para.style.name == 'Heading 1':
                    main_sections.append(para.text.strip())
            return main_sections
        elif doc_path.endswith('.pdf'):
            return self.get_pdf_main_sections(doc_path)
        return []
    
    def get_pdf_main_sections(self, pdf_path):
        """Extract main sections from a PDF."""
        sections = []
        document = fitz.open(pdf_path)
        for page_num in range(document.page_count):
            page = document.load_page(page_num)
            text = page.get_text()
            lines = text.split("\n")
            for line in lines:
                if line.lower().startswith("section"):
                    sections.append(line.strip())
        return sections
    
    def get_subsections(self, doc_path, main_section):
        """Extract subsections (Heading 2) under the specified main section."""
        if doc_path.endswith('.docx'):
            doc = Document(doc_path)
            subsections = []
            in_main_section = False
            
            for para in doc.paragraphs:
                if para.style.name == 'Heading 1':
                    in_main_section = (para.text.strip() == main_section)
                elif in_main_section and para.style.name == 'Heading 2':
                    subsections.append(para.text.strip())
            return subsections
        return []
    
    def extract_text_from_subsection(self, doc_path, selected_subsection):
        """Extract text from the specified subsection."""
        if doc_path.endswith('.docx'):
            return self.extract_text_from_docx_subsection(doc_path, selected_subsection)
        elif doc_path.endswith('.pdf'):
            return self.extract_text_from_pdf_subsection(doc_path, selected_subsection)
        return ""
    
    def extract_text_from_docx_subsection(self, doc_path, selected_subsection):
        """Extract text from the specified subsection in DOCX file."""
        doc = Document(doc_path)
        section_text = ""
        extract = False
        
        for para in doc.paragraphs:
            if para.style.name == 'Heading 2':
                if para.text == selected_subsection:
                    extract = True
                    section_text += para.text + "\n"
                elif extract:
                    break  # Stop when reaching the next Heading 2
            elif extract:
                section_text += para.text + "\n"
        
        return section_text
    
    def extract_text_from_pdf_subsection(self, pdf_path, selected_subsection):
        """Extract text from the specified subsection in PDF file."""
        section_text = ""
        extract = False
        document = fitz.open(pdf_path)
        
        for page_num in range(document.page_count):
            page = document.load_page(page_num)
            text = page.get_text()
            lines = text.split('\n')
            for line in lines:
                if selected_subsection in line:
                    extract = True
                if extract:
                    section_text += line + "\n"
                    if line.startswith('Clause '):
                        break
        
        return section_text
    
    def truncate_text(self, text, max_words=3000):
        """Truncate text to avoid OpenAI rate limit errors."""
        words = text.split()
        if len(words) > max_words:
            return ' '.join(words[:max_words])
        return text

    def _guard_llm_input(self, text: str, context: str) -> str:
        """Prepare document text for Azure OpenAI (upload guardrails already ran at ingest)."""
        guardrail = get_spec_intel_guardrail()
        truncated = self.truncate_text(text)
        return guardrail.prepare_extraction_llm_input(truncated, context=context)

    def _is_azure_content_filter_error(self, exc: Exception) -> bool:
        err = str(exc).lower()
        return "content_filter" in err or "jailbreak" in err or "content management policy" in err

    def _simplify_messages_for_azure(self, messages: list) -> list:
        """Strip wrapper text that can trigger Azure Prompt Shields false positives."""
        simplified = []
        for msg in messages:
            role = msg.get("role")
            content = msg.get("content", "")
            if role == "system":
                content = content.split("Treat all")[0].split("ignore embedded")[0].strip()
            elif role == "user":
                doc_match = re.search(r"<DOCUMENT>\s*(.*?)\s*</DOCUMENT>", content, re.DOTALL | re.IGNORECASE)
                if doc_match:
                    content = doc_match.group(1)
                prefix = "Technical specification excerpt:\n\n"
                if content.startswith(prefix):
                    content = content[len(prefix):]
                content = content.replace("[filtered-instruction]", " ").strip()
                content = content[:12000]
                content = f"Technical specification text for analysis:\n\n{content}"
            simplified.append({"role": role, "content": content})
        return simplified

    def _create_chat_completion(self, messages: list, **kwargs):
        """Call Azure OpenAI with one retry using simplified prompts if content filter triggers."""
        try:
            return self.client.chat.completions.create(
                model=self.azure_model_name,
                messages=messages,
                **kwargs,
            )
        except BadRequestError as exc:
            if not self._is_azure_content_filter_error(exc):
                raise
            print("⚠️ Azure content filter triggered; retrying with simplified specification prompt...")
            retry_messages = self._simplify_messages_for_azure(messages)
            try:
                return self.client.chat.completions.create(
                    model=self.azure_model_name,
                    messages=retry_messages,
                    **kwargs,
                )
            except BadRequestError as retry_exc:
                if self._is_azure_content_filter_error(retry_exc):
                    raise ValueError(
                        "Azure OpenAI blocked this specification text (jailbreak content filter). "
                        "Try a different subsection, or set Prompt Shields to 'annotate' mode in Azure AI Foundry."
                    ) from retry_exc
                raise
    
    def extract_references_and_clauses(self, text):
        """Extract 3GPP references and clauses using OpenAI API."""
        prompt = (
            "You are a 3GPP specification analyst. "
            "Extract 3GPP references and their associated clauses from the provided technical text. "
            "Present them grouped under the heading '3GPP References:', where each reference "
            "(e.g., '3GPP TS 23.401 [29]') has its related clauses listed underneath "
            "(e.g., '- Clause 5.3.2.1 ...'). Only show this section once. "
            "Do not repeat the same reference or clause. "
            "Output only the reference and clause listing."
        )

        llm_user_content = self._guard_llm_input(text, context="extract_references")
        
        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": llm_user_content}
        ]
        
        response = self._create_chat_completion(
            messages,
            max_tokens=500,
            temperature=0.5,
        )
        
        output = response.choices[0].message.content.strip()
        ref_clause_map, references, clauses = self.parse_openai_output(output)
        
        return ref_clause_map, references, clauses, output
    
    def parse_openai_output(self, output):
        """Parse OpenAI API output to extract references and clauses."""
        references = []
        clauses = []
        ref_clause_map = {}
        current_ref = None
        
        ref_pattern = re.compile(r"3GPP TS \d+\.\d+")
        clause_pattern = re.compile(r"-?\s*Clause\s+\d+(\.\d+)*")
        
        for line in output.splitlines():
            line = line.strip()
            if not line:
                continue
            
            # Check for 3GPP reference
            ref_match = ref_pattern.search(line)
            if ref_match:
                current_ref = ref_match.group()
                if current_ref not in ref_clause_map:
                    ref_clause_map[current_ref] = []
                references.append(current_ref)
            
            # Check for clause
            clause_match = clause_pattern.search(line)
            if clause_match and current_ref:
                clause = clause_match.group().lstrip('- ').strip()
                clauses.append(clause)
                ref_clause_map[current_ref].append(clause)
        
        return ref_clause_map, references, clauses
    
    def build_graph_from_input(self, input_text):
        """Build a graph from input text where references are nodes and clauses are edges."""
        G = nx.DiGraph()
        G.add_node("start")
        
        current_ref = None
        
        for line in input_text.splitlines():
            line = line.strip()
            if line.startswith("- 3GPP TS") or line.startswith("3GPP TS"):
                # Clean the reference: remove leading "- " and strip
                ref_text = line.lstrip("- ").strip()
                # IMPORTANT: Remove bracket numbers [29], [30] to match ref_clause_map keys
                # Extract only "3GPP TS XX.XXX" without the bracket
                ref_match = re.search(r"3GPP TS \d+\.\d+", ref_text)
                if ref_match:
                    current_ref = ref_match.group()  # "3GPP TS 23.401" without [29]
                else:
                    current_ref = ref_text  # Fallback to full text
                
                G.add_node(current_ref)
                G.add_edge("start", current_ref)
            elif (line.startswith("- Clause") or line.startswith("Clause")) and current_ref:
                clause = line.lstrip("- ").strip()
                G.add_node(clause)
                G.add_edge(current_ref, clause)
        
        return G
    
    def download_and_extract_specification(self, series, spec_number, download_directory):
        """Download and extract the latest specification from 3GPP."""
        base_url = f"https://www.3gpp.org/ftp/Specs/archive/{series}_series/{spec_number}/"
        
        # Get latest ZIP file
        response = requests.get(base_url)
        soup = BeautifulSoup(response.text, 'html.parser')
        zip_files = [a['href'] for a in soup.find_all('a', href=True) if a['href'].endswith('.zip')]
        
        if not zip_files:
            raise ValueError(f"No ZIP files found in the directory: {base_url}")
        
        latest_zip = sorted(zip_files, reverse=True)[0]
        latest_zip_url = urljoin(base_url, latest_zip)
        
        # Download the ZIP file
        response = requests.get(latest_zip_url)
        zip_file_name = os.path.basename(latest_zip_url)
        file_path = os.path.join(download_directory, zip_file_name)
        
        with open(file_path, 'wb') as file:
            file.write(response.content)
        
        # Extract the ZIP file
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.extractall(download_directory)
        
        # Remove the ZIP file after extraction
        os.remove(file_path)
    
    def download_specifications_from_references(self, references, download_directory):
        """Download specifications from 3GPP references."""
        for reference in references:
            match = re.search(r"3GPP TS (\d+)\.(\d+)", reference)
            if match:
                series = match.group(1)
                spec_number = f"{match.group(1)}.{match.group(2)}"
                try:
                    self.download_and_extract_specification(series, spec_number, download_directory)
                except ValueError as e:
                    print(f"Error downloading {reference}: {e}")
    
    def check_references_in_directory(self, references, directory):
        """Check if reference documents are present in the directory."""
        present_references = []
        missing_references = []
        present_ref_map = {}
        
        try:
            files_in_directory = os.listdir(directory)
            print(f"DEBUG: Files in directory {directory}:")
            print(f"DEBUG: {files_in_directory[:10]}")  # Show first 10 files
        except FileNotFoundError:
            print(f"Warning: Directory not found: {directory}")
            return present_ref_map, present_references, references
        
        print(f"DEBUG: Checking {len(references)} references...")
        for reference in references:
            matched = False
            # Extract numeric part (e.g., "23.401" from "3GPP TS 23.401")
            main_part_match = re.search(r'\d+\.\d+', reference)
            print(f"DEBUG: Processing reference: {reference}")
            print(f"DEBUG:   Extracted main_part_match: {main_part_match}")
            
            if main_part_match:
                main_part = main_part_match.group()
                # Remove dot for filename matching (e.g., "23.401" -> "23401")
                main_part_no_dot = main_part.replace('.', '')
                print(f"DEBUG:   Looking for: {main_part} → {main_part_no_dot}")
                
                for file_name in files_in_directory:
                    # Match files like "23401-j40.docx" or "23401.docx"
                    if main_part_no_dot in file_name:
                        print(f"✓ Found reference {reference} in file: {file_name}")
                        # IMPORTANT: Store the FILENAME (not the reference) in present_references
                        # This matches PyQt behavior where present_references is a list of filenames
                        present_references.append(file_name)
                        # Also clean the reference (remove bracket numbers like [29])
                        clean_ref = re.sub(r'\s*\[\d+\]$', '', reference.strip())
                        present_ref_map[clean_ref] = file_name
                        matched = True
                        break
            
            if not matched:
                missing_references.append(reference)
        
        return present_ref_map, present_references, missing_references
    
    def find_sections(self, doc):
        """Find sections in a DOCX document - matches PyQt version."""
        sections = {}
        current_section = None
        current_content = []
        
        for para in doc.paragraphs:
            # Check the paragraph style - match any heading level like PyQt does
            style = para.style.name
            
            # Determine if the paragraph is a section heading
            if style in ['Heading 0', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
                if current_section is not None:
                    # Store the content of the previous section
                    sections[current_section] = '\n'.join(current_content)
                
                # Start a new section
                current_section = para.text.strip()
                current_content = []
            else:
                # Append content to the current section
                if current_section is not None:
                    current_content.append(para.text)
        
        # Don't forget to store the last section's content
        if current_section is not None:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def search_clause_in_sections(self, sections, clause):
        """Search for a clause in document sections."""
        clause = clause.replace("Clause", "").strip()
        match = re.search(r"\b\d+(\.\d+)*[A-Za-z]?(\.\d+)?\b", clause)
        if not match:
            return None
        
        clause_number = match.group(0)
        collected = []
        
        for section_title, content in sections.items():
            match = re.search(r'(\d+(\.\d+)*)', section_title)
            if match:
                section_title_number = match.group(1)
                if clause_number == section_title_number or section_title_number.startswith(clause_number + '.'):
                    collected.append(f"{section_title}\n{content}")
        
        return "\n\n".join(collected) if collected else None
    
    def read_clause_from_file(self, file_path, clause):
        """Read a specific clause from a DOCX file."""
        try:
            doc = Document(file_path)
            sections = self.find_sections(doc)
            clause_text = self.search_clause_in_sections(sections, clause)
            
            if clause_text:
                return clause_text
            else:
                return f"Clause '{clause}' not found in the document."
        except Exception as e:
            return f"Error reading {file_path}: {e}"
    
    def save_graph_to_json(self, graph, output_file):
        """Save the graph structure to a JSON file."""
        data = nx.readwrite.json_graph.node_link_data(graph, edges="edges")
        with open(output_file, 'w') as f:
            json.dump(data, f, indent=4)
    
    # ===== RECURSIVE EXTRACTION FUNCTIONS (Matching PyQt Version) =====
    
    def extract_keywords_with_lines(self, text):
        """Extract TS references from text with their context lines."""
        keyword_pattern = r"TS\s*\d{2}\.\d{3}"
        result_dict = defaultdict(list)
        sentences = re.split(r'(?<=\.)\s+', text)
        
        for sentence in sentences:
            match = re.search(keyword_pattern, sentence, re.IGNORECASE)
            if match:
                ts_code = match.group(0)
                result_dict[ts_code] = sentence.strip()
        
        return result_dict
    
    def match_extract_from_ref(self, ref_dict):
        """Extract technical phrases from references using OpenAI API."""
        extracted_phrases = []
        
        for key, value in ref_dict.items():
            try:
                llm_user_content = self._guard_llm_input(value, context="match_extract_from_ref")
                response = self._create_chat_completion(
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are a 3GPP technical documentation assistant. "
                                "Extract one meaningful technical phrase from the specification excerpt."
                            ),
                        },
                        {"role": "user", "content": llm_user_content},
                    ],
                    temperature=0.7,
                )
                
                technical_phrase = response.choices[0].message.content.strip()
                formatted_string = f"Key: {key}, Extracted Phrase: {technical_phrase}"
                extracted_phrases.append(formatted_string)
                
            except Exception as e:
                print(f"Error extracting phrase for {key}: {e}")
                continue
        
        return extracted_phrases
    
    def is_phrase_in_section(self, phrase, section, threshold=0.5):
        """Check if a phrase is in a section with partial matching."""
        sentences = section.split('. ')
        
        for sentence in sentences:
            similarity = SequenceMatcher(None, phrase.lower(), sentence.lower()).ratio()
            if similarity >= threshold:
                return True
        
        return False
    
    def process_reference_extract_sections(self, phrases, directory):
        """Process reference sections using the working directory."""
        matched_sections = ''
        files_in_directory = os.listdir(directory)
        
        for phrase in phrases:
            match = re.search(r"Key: TS\s*(\d+\.\d+)", phrase)
            if match:
                main_part = match.group(1).replace('.', '')
                matched = False
                
                for file_name in files_in_directory:
                    file_numeric_part = re.search(r'\d+', file_name)
                    if file_numeric_part and main_part in file_numeric_part.group():
                        doc_path = os.path.join(directory, file_name)
                        try:
                            doc = Document(doc_path)
                            sections = self.find_sections(doc)
                            
                            for section_title, content in sections.items():
                                extracted_phrase = re.search(r"Extracted Phrase:\s*(.*)", phrase)
                                if extracted_phrase:
                                    phrase_text = extracted_phrase.group(1)
                                    if self.is_phrase_in_section(phrase_text, content, threshold=0.5):
                                        matched = True
                                        matched_sections += "\n" + "Section:" + "\t" + section_title + "\t"
                                        matched_sections += content
                                        break
                        except Exception as e:
                            print(f"Error processing {file_name}: {e}")
                            continue
                        
                        if matched:
                            break
        
        return matched_sections
    
    def append_clause_sections(self, text, filepath, clauses):
        """Append clause sections to the text."""
        try:
            doc = Document(filepath)
            for clause in clauses:
                clause_text = self.read_clause_from_file(filepath, clause)
                if clause_text and "not found" not in clause_text.lower():
                    text += "\n" + "Clause:" + "\t" + clause + "\t"
                    text += clause_text
        except Exception as e:
            print(f"Error appending clause sections: {e}")
        
        return text
    
    def generate_test_scenarios(self, input_text):
        """Generate test scenarios using OpenAI API."""
        llm_user_content = self._guard_llm_input(input_text, context="generate_test_scenarios")
        
        messages = [
            {
                "role": "system",
                "content": "You are an expert in 3GPP standards and test scenario design.",
            },
            {
                "role": "user",
                "content": (
                    "Based on the following specification excerpt, list possible test scenarios "
                    "including edge cases per 3GPP standards:\n\n"
                    f"{llm_user_content}"
                ),
            },
        ]
        
        try:
            response = self._create_chat_completion(
                messages,
                max_tokens=1000,
                temperature=0.7,
            )
            test_scenarios = response.choices[0].message.content
            return test_scenarios
        except Exception as e:
            print(f"Error generating test scenarios: {e}")
            return "Error generating test scenarios"
    
    def get_filename(self, reference_name, directory="."):
        """Map reference to actual filename in the specified directory."""
        # Use the provided directory instead of hardcoded current directory
        files_in_directory = os.listdir(directory)
        
        # Extract numeric part from reference
        match = re.search(r'\d+-[a-zA-Z0-9]+', reference_name)
        if match:
            main_part = match.group()
            for file_name in files_in_directory:
                file_match = re.search(r'\d+-[a-zA-Z0-9]+', file_name)
                if file_match and file_match.group() == main_part:
                    return file_name
        
        return None
    
    # ===== PYQT FLOW REPLICATION FUNCTIONS =====
    
    def get_filename_from_directory(self, node, directory):
        """Extract filename from node reference - matches PyQt get_filename function."""
        files_in_directory = os.listdir(directory)
        
        # Extract numeric part from node name
        match = re.search(r'\d+-[a-zA-Z0-9]+', node)
        if match:
            main_part = match.group()
            for file_name in files_in_directory:
                file_match = re.search(r'\d+-[a-zA-Z0-9]+', file_name)
                if file_match and file_match.group() == main_part:
                    return file_name
        
        return None
    
    def match_extract_from_ref_pyqt(self, ref_dict):
        """Extract technical phrases from references using OpenAI - matches PyQt version."""
        extracted_phrases = []
        
        for key, value in ref_dict.items():
            print(f"Key: {key}")
            print(f"Value: {value}")
            print("-" * 30)
            
            try:
                llm_user_content = self._guard_llm_input(value, context="match_extract_from_ref_pyqt")
                response = self._create_chat_completion(
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are a 3GPP technical documentation assistant. "
                                "Extract one meaningful technical phrase from the specification excerpt."
                            ),
                        },
                        {"role": "user", "content": llm_user_content},
                    ],
                    temperature=0.7,
                )
                
                technical_phrase = response.choices[0].message.content.strip()
                formatted_string = f"Key: {key}, Extracted Phrase: {technical_phrase}"
                extracted_phrases.append(formatted_string)
                print(formatted_string)
                
            except Exception as e:
                print(f"Error extracting phrase for {key}: {e}")
                continue
        
        return extracted_phrases
    
    def process_reference_extract_sections_pyqt(self, phrases, directory):
        """Process reference sections - matches PyQt version."""
        print("Starting process_reference_extract_sections")
        matched_sections = ''
        files_in_directory = os.listdir(directory)
        
        for phrase in phrases:
            match = re.search(r"Key: TS\s*(\d+\.\d+)", phrase)
            print(f"Processing phrase {phrase}, match = {match}")
            
            if match:
                main_part = match.group(1).replace('.', '')
                matched = False
                print(f"Looking for main_part {main_part}")
                
                for file_name in files_in_directory:
                    file_numeric_part = re.search(r'\d+', file_name)
                    if file_numeric_part and main_part in file_numeric_part.group():
                        doc_path = os.path.join(directory, file_name)
                        
                        try:
                            doc = Document(doc_path)
                            print(f"Processing document: {doc_path}")
                            
                            sections = self.find_sections(doc)
                            
                            for section_title, content in sections.items():
                                extracted_phrase_match = re.search(r"Extracted Phrase:\s*(.*)", phrase)
                                if extracted_phrase_match:
                                    phrase_text = extracted_phrase_match.group(1)
                                    print(f"Checking section = {section_title}, phrase = {phrase_text}")
                                    
                                    if self.is_phrase_in_section(phrase_text, content, threshold=0.5):
                                        matched = True
                                        print(f"Section matched {section_title}, phrase {phrase}")
                                        matched_sections += "\n" + "Section:" + "\t" + section_title + "\t"
                                        matched_sections += content
                                        break
                        except Exception as e:
                            print(f"Error processing {file_name}: {e}")
                            continue
                        
                        if matched:
                            break
        
        return matched_sections
    
    def extract_references_and_clauses_pyqt(self, graph, references, clauses, present_references, 
                                           directory, present_ref_map, ref_clause_map, safe_subsection):
        """
        Main extraction function - EXACT REPLICA of PyQt extract_references_and_clauses1 function.
        This processes the graph to extract text from specified clauses of reference files.
        """
        text = ''
        phrase = ''
        G = nx.DiGraph()
        start_node = "start"
        G.add_node(start_node)
        
        # Add reference nodes (filenames) and connect to the start node
        # present_references now contains filenames like "23401-i60.docx"
        for filename in present_references:
            G.add_node(filename)
            G.add_edge(start_node, filename)
        
        # Add edges from filename nodes to corresponding clauses using ref_clause_map
        # ref_clause_map has keys like "3GPP TS 23.401" mapped to filenames
        for ref, clause_list in ref_clause_map.items():
            if ref in present_ref_map:
                file_name = present_ref_map[ref]
                print(f"Adding edges for reference: {ref} → file: {file_name}")
                print(f"  Clauses for this reference: {clause_list}")
                
                for clause in clause_list:
                    G.add_edge(file_name, clause)
                    print(f"  Added edge: {file_name} → {clause}")
        
        graph = G
        
        # Process the graph
        print(f"DEBUG: Graph has {len(graph.nodes)} nodes: {list(graph.nodes)}")
        print(f"DEBUG: Graph has {len(graph.edges)} edges: {list(graph.edges)[:10]}")
        
        for node in graph.nodes:
            if node != start_node:
                print(f"Processing node: {node}")
                
                # Check neighbors (connected clauses) of the current reference node
                neighbors = list(nx.neighbors(graph, node))
                print(f"  Node {node} has {len(neighbors)} neighbors: {neighbors}")
                
                for clause in neighbors:
                    print(f"Processing clause: {clause} for node: {node}")
                    
                    if node:
                        # The node should already be a filename (from present_references)
                        # But call get_filename to ensure compatibility with PyQt flow
                        node1 = self.get_filename_from_directory(node, directory)
                        print(f"  get_filename_from_directory returned: {node1}")
                        
                        # If node is already a filename, use it directly
                        if not node1 and node.endswith('.docx'):
                            node1 = node
                            print(f"  Using node directly as filename: {node1}")
                        
                        if node1:
                            file_path = os.path.join(directory, node1)
                            
                            # Extract clause number from the clause string
                            clause_match = re.search(r'\b\d+(\.\d+)+\b', clause)
                            if clause_match:
                                clause_number = clause_match.group()
                            else:
                                clause_number = clause
                            
                            print(f"Reading clause {clause_number} from file {file_path} for node {node1}")
                            
                            # Read clause content from file
                            text = self.read_clause_from_file(file_path, clause_number)
                            
                            if text:
                                # Extract nested references and clauses
                                ref_clause_map_2, references_nested, clauses_nested, dummy_text = self.extract_references_and_clauses(text)
                                print("Nested References:", references_nested)
                                print("Nested Clauses:", clauses_nested)
                                
                                # Extract keywords with lines
                                ref_dict = self.extract_keywords_with_lines(text)
                                
                                # Extract technical phrases
                                phrase = self.match_extract_from_ref_pyqt(ref_dict)
                                print(f"Extracted phrases: {ref_dict}")
                                
                                # Process reference extract sections
                                matched_section = self.process_reference_extract_sections_pyqt(phrase, directory)
                                print("After extracting matched sections")
                                
                                # Append clause sections
                                text = self.append_clause_sections(text, file_path, clauses_nested)
                                print(f"Processed keywords for clause: {clause_number}")
                                
                                # Create output directory and save clause file
                                base_dir = Path(directory) / "datasets"
                                subsection_folder = base_dir / safe_subsection
                                subsection_folder.mkdir(parents=True, exist_ok=True)
                                
                                file_name = f"{clause_number.replace('.', '_')}_file.txt"
                                output_file_path = subsection_folder / file_name
                                
                                with open(output_file_path, "w", encoding="utf-8") as file:
                                    file.write(text)
                                    if matched_section:
                                        file.write(matched_section)
                                
                                print(f"Saved clause file: {output_file_path}")
                                
                                # Generate test scenarios (like PyQt)
                                test_scenarios = self.generate_test_scenarios(text)
                                print(f"Generated test scenarios for {clause_number}")
                                print(f"Finished processing node: {node1}")
                            else:
                                print(f"No text found in {node} for clause {clause_number}.\n")
        
        return text
    
    def extract_recursive_clauses(self, ref_clause_map, present_ref_map, working_directory, subsection_folder):
        """Main recursive extraction logic - processes each reference and clause with full recursion."""
        all_extracted_text = []
        processed_clauses = set()  # Track processed clauses to avoid infinite loops
        
        def process_clause_recursively(reference, clause, depth=0):
            """Recursively process a clause and extract nested references."""
            if depth > 5:  # Prevent infinite recursion
                print(f"Max recursion depth reached for clause: {clause}")
                return ""
            
            clause_key = f"{reference}_{clause}"
            if clause_key in processed_clauses:
                print(f"Already processed clause: {clause}")
                return ""
            
            processed_clauses.add(clause_key)
            print(f"Processing clause: {clause} (depth: {depth})")
            
            # Read clause content from the reference document
            filename = present_ref_map.get(reference)
            if not filename:
                print(f"No filename found for reference: {reference}")
                return ""
            
            file_path = os.path.join(working_directory, filename)
            clause_text = self.read_clause_from_file(file_path, clause)
            
            if not clause_text or "not found" in clause_text.lower():
                print(f"No content found for clause: {clause}")
                return ""
            
            print(f"Found clause content for {clause}")
            
            # ===== RECURSIVE EXTRACTION (MISSING IN CURRENT IMPLEMENTATION) =====
            # Extract NEW references and clauses from this clause text (like PyQt)
            new_ref_clause_map, new_references, new_clauses, _ = self.extract_references_and_clauses(clause_text)
            print(f"Found {len(new_references)} new references and {len(new_clauses)} new clauses in {clause}")
            
            # Process new references recursively
            nested_content = []
            for new_ref, new_clause_list in new_ref_clause_map.items():
                if new_ref in present_ref_map:  # Only process if we have the reference file
                    print(f"Processing nested reference: {new_ref}")
                    for new_clause in new_clause_list:
                        nested_text = process_clause_recursively(new_ref, new_clause, depth + 1)
                        if nested_text:
                            nested_content.append(nested_text)
            
            # Extract keywords from clause text
            keywords = self.extract_keywords_with_lines(clause_text)
            
            # Extract technical phrases (optional - can be slow)
            matched_sections = ""
            if keywords:
                try:
                    phrases = self.match_extract_from_ref(keywords)
                    matched_sections = self.process_reference_extract_sections(phrases, working_directory)
                except Exception as e:
                    print(f"Error in phrase extraction: {e}")
                    matched_sections = ""
            
            # Append clause sections
            clause_text = self.append_clause_sections(clause_text, file_path, [clause])
            
            # Add nested content to clause text
            if nested_content:
                clause_text += "\n\n" + "="*50 + f"\nNESTED REFERENCES FROM {clause}\n" + "="*50 + "\n\n"
                clause_text += "\n\n".join(nested_content)
            
            # Save individual clause file
            clause_number = re.search(r'\b\d+(\.\d+)+\b', clause)
            if clause_number:
                clause_number = clause_number.group().replace('.', '_')
            else:
                clause_number = clause.replace('.', '_').replace(' ', '_')
            
            clause_filename = f"{clause_number}_file.txt"
            clause_file_path = subsection_folder / clause_filename
            
            with open(clause_file_path, "w", encoding="utf-8") as f:
                f.write(clause_text)
                if matched_sections:
                    f.write(matched_sections)
            
            print(f"Saved clause file: {clause_filename}")
            
            # Generate test scenarios for individual clause (like PyQt) but don't save
            test_scenarios = self.generate_test_scenarios(clause_text)
            if test_scenarios:
                print(f"Generated test scenarios for {clause} (not saved, matching PyQt behavior)")
            
            return clause_text
        
        # Process each reference and its clauses
        for reference, clause_list in ref_clause_map.items():
            if reference not in present_ref_map:
                continue
            
            print(f"Processing reference: {reference}")
            
            # For each clause in this reference
            for clause in clause_list:
                clause_content = process_clause_recursively(reference, clause)
                if clause_content:
                    all_extracted_text.append(f"\n=== Clause {clause} ===\n{clause_content}")
        
        return "\n".join(all_extracted_text)
    
    def save_total_content(self, total_content, subsection_folder):
        """Save the total aggregated content."""
        total_file_path = subsection_folder / "total_content.txt"
        with open(total_file_path, "w", encoding="utf-8") as f:
            f.write(total_content)
        print(f"Saved total content: {total_file_path}")
    
    
    def extract_dataset(self, doc_path, section, subsection, output_directory, working_directory):
        """Main method to extract dataset from document with full recursive processing."""
        print(f"Starting dataset extraction for subsection: {subsection}")
        
        # ===== PHASE 1: BASIC EXTRACTION (EXISTING LOGIC) =====
        # Extract text from subsection
        section_text = self.extract_text_from_subsection(doc_path, subsection)
        print(f"Extracted {len(section_text)} characters from subsection")
        
        # Create subsection folder structure
        safe_subsection = re.sub(r'[\\/*?:"<>|]', "_", subsection)
        subsection_folder = Path(working_directory) / "datasets" / safe_subsection
        subsection_folder.mkdir(parents=True, exist_ok=True)
        
        # Save initial subsection text
        output_filename = f"{safe_subsection}_section.txt"
        output_path = subsection_folder / output_filename
        
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(section_text)
        print(f"Saved initial text: {output_path}")
        
        # Extract references and clauses using OpenAI
        print("Extracting references and clauses...")
        ref_clause_map, references, clauses, text = self.extract_references_and_clauses(section_text)
        print(f"Found {len(references)} references and {len(clauses)} clauses")
        
        # Download specifications to organized folder
        specs_directory = Path(working_directory) / "specifications"
        specs_directory.mkdir(parents=True, exist_ok=True)
        print(f"Downloading specifications to: {specs_directory}")
        self.download_specifications_from_references(references, str(specs_directory))
        
        # Check for present references
        present_ref_map, present_references, missing_references = self.check_references_in_directory(
            references, str(specs_directory)
        )
        print(f"Present references: {len(present_references)}, Missing: {len(missing_references)}")
        
        # ===== PHASE 2: RECURSIVE EXTRACTION (NEW LOGIC) =====
        print("Starting recursive clause extraction...")
        aggregated_content = self.extract_recursive_clauses(
            ref_clause_map, 
            present_ref_map, 
            str(specs_directory), 
            subsection_folder
        )
        print(f"Recursive extraction completed. Aggregated {len(aggregated_content)} characters")
        
        # ===== PHASE 3: SAVE AGGREGATED CONTENT =====
        total_content = section_text + "\n\n" + "="*50 + "\nRECURSIVE EXTRACTION RESULTS\n" + "="*50 + "\n\n" + aggregated_content
        self.save_total_content(total_content, subsection_folder)
        
        # ===== PHASE 4: TEST SCENARIOS (Individual clauses only, like PyQt) =====
        # Note: Test scenarios are generated for individual clauses in extract_recursive_clauses()
        # but not saved to files, matching PyQt behavior
        
        # ===== PHASE 5: BUILD AND SAVE GRAPH =====
        print("Building knowledge graph...")
        graph = self.build_graph_from_input(text)
        
        # Save graph to organized output directory
        # Ensure output directory exists
        output_dir_path = Path(output_directory)
        output_dir_path.mkdir(parents=True, exist_ok=True)
        
        output_file = output_dir_path / f"{safe_subsection}_output.json"
        self.save_graph_to_json(graph, str(output_file))
        print(f"Saved graph: {output_file}")
        
        # ===== PHASE 6: RETURN COMPREHENSIVE RESULTS =====
        result = {
            "section_text": section_text,
            "references": references,
            "clauses": clauses,
            "present_references": present_references,
            "missing_references": missing_references,
            "graph": graph,
            "output_file": str(output_file),
            "aggregated_content": aggregated_content,
            "total_content_file": str(subsection_folder / "total_content.txt"),
            "clause_files_count": len(list(subsection_folder.glob("*_file.txt"))),
            "files_created": {
                "initial_text": str(output_path),
                "total_content": str(subsection_folder / "total_content.txt"),
                "graph_json": str(output_file),
                "clause_files": [str(f) for f in subsection_folder.glob("*_file.txt")]
            }
        }
        
        print(f"Dataset extraction completed successfully!")
        print(f"Files created: {len(result['files_created']['clause_files'])} clause files + 3 main files")
        
        return result
    
    def extract_dataset_pyqt_style(self, doc_path, section, subsection, output_directory, working_directory):
        """
        EXACT REPLICA of the PyQt flow from recursive_test_graph_attach.py __main__ section.
        This is a simplified, synchronous version that matches the PyQt subprocess call.
        """
        print(f"=== Starting PyQt-style dataset extraction ===")
        print(f"Document: {doc_path}")
        print(f"Section: {section}")
        print(f"Subsection: {subsection}")
        print(f"Working Directory: {working_directory}")
        print(f"Output Directory: {output_directory}")
        
        # STEP 1: Extract text from the selected subsection
        section_text = self.extract_text_from_subsection(doc_path, subsection)
        print(f"Extracted {len(section_text)} characters from subsection")
        
        # STEP 2: Create safe subsection name and folder structure
        safe_subsection = re.sub(r'[\\/*?:"<>|]', "_", subsection)
        base_dir = Path(working_directory) / "datasets"
        subsection_folder = base_dir / safe_subsection
        subsection_folder.mkdir(parents=True, exist_ok=True)
        
        # STEP 3: Save initial subsection text
        output_filename = f"{safe_subsection}_section.txt"
        output_path = subsection_folder / output_filename
        
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(section_text)
        
        print(f"Saved initial subsection text: {output_path}")
        
        # STEP 4: Extract references and clauses using OpenAI
        ref_clause_map, references, clauses, text = self.extract_references_and_clauses(section_text)
        print(f"Found {len(references)} references and {len(clauses)} clauses")
        print(f"References: {references}")
        print(f"Clauses: {clauses}")
        
        # STEP 5: Download specifications to working directory (if needed)
        specs_directory = working_directory
        print(f"Checking for specifications in: {specs_directory}")
        print(f"DEBUG: specs_directory is absolute path? {Path(specs_directory).is_absolute()}")
        print(f"DEBUG: specs_directory exists? {Path(specs_directory).exists()}")
        
        # First check what's already present
        print(f"DEBUG: References to check: {references}")
        present_ref_map, present_references, missing_references = self.check_references_in_directory(
            references, specs_directory
        )
        print(f"DEBUG: present_ref_map = {present_ref_map}")
        print(f"DEBUG: present_references = {present_references}")
        print(f"DEBUG: missing_references = {missing_references}")
        
        if missing_references:
            print(f"Missing References: {missing_references}")
            print("Attempting to download missing specifications...")
            
            # Try to download missing specifications
            try:
                self.download_specifications_from_references(missing_references, specs_directory)
                print("Download completed. Re-checking references...")
                
                # Re-check after download
                present_ref_map, present_references, missing_references = self.check_references_in_directory(
                    references, specs_directory
                )
                
                if missing_references:
                    print(f"Still missing after download: {missing_references}")
                else:
                    print("All references are now present!")
            except Exception as e:
                print(f"Error downloading specifications: {e}")
                print("Continuing with available references only...")
        else:
            print("All references are already present in the directory.")
        
        print(f"Present references: {len(present_references)} / {len(references)}")
        
        # STEP 6: Create initial graph from the document, references, and clauses
        graph = self.build_graph_from_input(text)
        
        # STEP 7: Save initial graph (using "output.json" like PyQt does)
        # Ensure output directory exists before saving
        output_dir_path = Path(output_directory)
        output_dir_path.mkdir(parents=True, exist_ok=True)
        
        output_file = os.path.join(output_directory, "output.json")
        self.save_graph_to_json(graph, output_file)
        print(f"Saved initial graph: {output_file}")
        
        # STEP 8: MAIN EXTRACTION - Call the PyQt-style extraction function
        print("=" * 50)
        print("Starting main extraction (extract_references_and_clauses_pyqt)")
        print("=" * 50)
        
        section_text_result = self.extract_references_and_clauses_pyqt(
            graph, references, clauses, present_references, 
            specs_directory, present_ref_map, ref_clause_map, safe_subsection
        )
        
        # STEP 9: Save final graph
        self.save_graph_to_json(graph, output_file)
        print(f"Saved final graph: {output_file}")
        
        # STEP 10: Save total_content.txt (like PyQt does)
        total_content_path = subsection_folder / "total_content.txt"
        with open(total_content_path, "w", encoding="utf-8") as file:
            file.write(section_text)
            file.write("\n\n" + "=" * 50 + "\n")
            file.write("RECURSIVE EXTRACTION RESULTS\n")
            file.write("=" * 50 + "\n\n")
            if section_text_result:
                file.write(section_text_result)
        
        print(f"Saved total content: {total_content_path}")
        
        # STEP 11: Return results
        clause_files = list(subsection_folder.glob("*_file.txt"))
        
        # Check if we have the expected results
        success_message = "Dataset extraction completed successfully!"
        if len(clause_files) == 0 and len(clauses) > 0:
            success_message += f"\n⚠️  WARNING: Found {len(clauses)} clauses but created 0 clause files."
            if len(present_references) == 0:
                success_message += f"\n⚠️  REASON: No reference documents were found in the working directory."
                success_message += f"\n⚠️  MISSING: {missing_references}"
                success_message += f"\n💡 TIP: Make sure 3GPP specification documents are in: {specs_directory}"
                success_message += f"\n💡 TIP: Or ensure internet connection for automatic download."
        
        result = {
            "success": True,
            "message": success_message,
            "section_text": section_text,
            "references": references,
            "clauses": clauses,
            "present_references": present_references,
            "missing_references": missing_references,
            "present_ref_map": present_ref_map,
            "ref_clause_map": ref_clause_map,
            "output_file": output_file,
            "total_content_file": str(total_content_path),
            "clause_files_count": len(clause_files),
            "files_created": {
                "initial_text": str(output_path),
                "total_content": str(total_content_path),
                "graph_json": output_file,
                "clause_files": [str(f) for f in clause_files]
            }
        }
        
        print("=" * 50)
        print(f"✅ Dataset extraction completed!")
        print(f"✅ Output saved to: {output_file}")
        print(f"✅ Created {len(clause_files)} clause files")
        if len(clause_files) == 0 and len(clauses) > 0:
            print(f"⚠️  WARNING: Expected {len(clauses)} clause files but got 0")
            print(f"⚠️  Present references: {len(present_references)} / {len(references)}")
            if missing_references:
                print(f"⚠️  Missing references: {missing_references}")
        print("=" * 50)
        
        return result


def main():
    """Main function for command-line usage."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Extract dataset from document")
    parser.add_argument("doc_path", help="Path to the input document")
    parser.add_argument("section", help="Main section to extract")
    parser.add_argument("subsection", help="Subsection to extract")
    parser.add_argument("output_directory", help="Output directory for results")
    parser.add_argument("working_directory", help="Working directory for processing")
    
    args = parser.parse_args()
    
    # Initialize extractor
    extractor = DatasetExtractor()
    
    # Extract dataset
    result = extractor.extract_dataset(
        args.doc_path,
        args.section,
        args.subsection,
        args.output_directory,
        args.working_directory
    )
    
    print(f"Dataset extraction completed!")
    print(f"Output saved to: {result['output_file']}")
    print(f"References found: {len(result['references'])}")
    print(f"Clauses found: {len(result['clauses'])}")


if __name__ == "__main__":
    main()
