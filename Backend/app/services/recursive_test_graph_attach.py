import json
import os
import sys
from docx import Document
#import fitz  # PyMuPDF
import networkx as nx
import argparse
import openai
import re
import requests
from bs4 import BeautifulSoup
import zipfile
from urllib.parse import urljoin
from collections import defaultdict
import re
from difflib import SequenceMatcher
from pathlib import Path
import subprocess
from openai import AzureOpenAI
from dotenv import load_dotenv

# Check if .env file exists
if not os.path.exists('.env'):
    print("Warning: .env file not found in the current directory.")

# Load environment variables from .env file
load_dotenv()

# Get Azure OpenAI credentials from environment variables
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_MODEL_NAME = os.getenv("AZURE_OPENAI_MODEL_NAME")

missing_vars = []
if not AZURE_OPENAI_ENDPOINT:
    missing_vars.append("AZURE_OPENAI_ENDPOINT")
if not AZURE_OPENAI_API_KEY:
    missing_vars.append("AZURE_OPENAI_API_KEY")
if not AZURE_OPENAI_MODEL_NAME:
    missing_vars.append("AZURE_OPENAI_MODEL_NAME")

if missing_vars:
    print(f"Error: Missing environment variables: {', '.join(missing_vars)}. Please set them in your .env file.")
    sys.exit(1)

try:
    client = AzureOpenAI(
        api_key=AZURE_OPENAI_API_KEY,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_version="2024-02-01"  # Check Azure OpenAI documentation for the latest supported API version
    )
except Exception as e:
    print(f"Failed to initialize AzureOpenAI client: {e}")
    sys.exit(1)


# Removed hardcoded file_path - use command line arguments instead
# file_path = "C:\\Users\\ChanduVangala\\Documents\\GenAI\\Input_documents\\O-RAN.TIFG.E2E-Test.0-R003-v06.00.docx"
#file_path = "C:\\Users\\teamt\\Downloads\\genai\\23401-i6.docx"
def get_main_sections(doc_path):
    """Extracts main section headings (Heading 1) from a DOCX file."""
    doc = Document(doc_path)
    main_sections = []
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':  # Explicitly check for Heading 1 style
            main_sections.append(para.text.strip())
    return main_sections

def get_subsections(doc_path, main_section):
    """Extracts nested subsections (Heading 2) under the specified main section."""
    doc = Document(doc_path)
    subsections = []
    in_main_section = False

    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            in_main_section = (para.text.strip() == main_section)
        elif in_main_section and para.style.name == 'Heading 2':
            subsections.append(para.text.strip())

    return subsections

def extract_text_from_subsection(doc_path, selected_subsection):
    """Extracts text from the specified subsection and its nested subsections."""
    if doc_path.endswith('.docx'):
        return extract_text_from_docx_subsection(doc_path, selected_subsection)
    elif doc_path.endswith('.pdf'):
        return extract_text_from_pdf_subsection(doc_path, selected_subsection)
    return ""

def extract_text_from_docx_subsection(doc_path, selected_subsection):
    """Extracts text from the specified subsection and its nested subsections in a DOCX file."""
    doc = Document(doc_path)
    section_text = ""
    extract = False

    for para in doc.paragraphs:
        if para.style.name == 'Heading 2':
            if para.text == selected_subsection:
                extract = True
                section_text += para.text + "\n"
            elif extract:
                break  # Stop when reaching the next Heading 2
        elif extract:
            section_text += para.text + "\n"

    return section_text

def extract_text_from_pdf_subsection(pdf_path, selected_subsection):
    """Extracts text from the specified subsection in a PDF file."""
    section_text = ""
    extract = False
    document = fitz.open(pdf_path)

    for page_num in range(document.page_count):
        page = document.load_page(page_num)
        text = page.get_text()
        lines = text.split('\n')
        for line in lines:
            if selected_subsection in line:
                extract = True
            if extract:
                section_text += line + "\n"
                if line.startswith('Clause '):
                    break  # Assuming next clause starts with 'Clause '

    return section_text
def build_graph_from_input(input_text):
    """
    Function to build a graph from the given input text, where references are nodes and clauses are edges.
    
    Args:
    - input_text (str): The input text containing references and clauses.

    Returns:
    - G (networkx.DiGraph): The directed graph created from the input.
    """
    # Initialize a directed graph
    G = nx.DiGraph()

    # Add a start node
    G.add_node("start")

    current_ref = None
    
    # Process each line in the input text
    for line in input_text.splitlines():
        line = line.strip()
        if line.startswith("- 3GPP TS"):  # This identifies a reference node
            current_ref = line[2:].strip()
            G.add_node(current_ref)
            G.add_edge("start", current_ref)
        elif line.startswith("3GPP TS"):  #Arshit's Changes: This identifies a reference node
            current_ref = line.strip()
            G.add_node(current_ref)
            G.add_edge("start", current_ref)
        elif line.startswith("- Clause") and current_ref:  # This identifies a clause (edge)
            clause = line[2:].strip()
            G.add_node(clause)
            G.add_edge(current_ref, clause)
        elif line.startswith("Clause") and current_ref: #Arshit's Changes: This identifies a clause (edge)
            clause = line.strip()
            G.add_node(clause)
            G.add_edge(current_ref, clause)

    return G

def truncate_text(text, max_words=3000):
    """Truncate text to a maximum number of words to avoid OpenAI rate limit errors."""
    words = text.split()
    if len(words) > max_words:
        return ' '.join(words[:max_words])
    return text

#Arshit's Changes: Adding a function to parse the reference and clause structure and store it in a map with reference as key and clauses as value
def parse_reference_clause_structure(output):
    ref_to_clauses = {}
    current_ref = None

    ref_pattern = re.compile(r"3GPP TS \d+\.\d+")
    clause_pattern = re.compile(r"-?\s*Clause\s+\d+(\.\d+)*")

    for line in output.splitlines():
        line = line.strip()
        if not line:
            continue

        # Check for 3GPP reference
        ref_match = ref_pattern.search(line)
        if ref_match:
            current_ref = ref_match.group()
            if current_ref not in ref_to_clauses:
                ref_to_clauses[current_ref] = []

        # Check for clause (even if it's on the same line as reference)
        clause_match = clause_pattern.search(line)
        if clause_match and current_ref:
            ref_to_clauses[current_ref].append(clause_match.group().lstrip('- ').strip())

    return ref_to_clauses

def extract_references_and_clauses(text):
    """Extracts 3GPP references and clauses from the given text using OpenAI API."""
    # text = truncate_text(text)

    cont= ("Extract 3GPP references and their associated clauses from the given text. "
        "Present them grouped under the heading '3GPP References:', where each reference (e.g., '3GPP TS 23.401 [29]') "
        "has its related clauses listed underneath (e.g., '- Clause 5.3.2.1 ...'). "
        "Only show this section once. Do not repeat the same reference or clause. "
        "Avoid unnecessary formatting or cosmetic additions beyond what's described.")
    messages = [
        {"role": "system", "content": cont},
        {"role": "user", "content": text}
    ]
    #"Extract 3GPP references and clauses seperately from the following text,dont add any special characters as cosmetic to the output, keep the clauses under respective references"}
    #response = openai.chat.completions.create(
    response = client.chat.completions.create(
        #model="gpt-4o-mini",
        model=AZURE_OPENAI_MODEL_NAME,
        messages=messages,
        max_tokens=500,  # Increased token limit for potentially longer responses
        temperature=0.5
    )
    output = response.choices[0].message.content.strip()
    print("OpenAI Response for ref:", output)  # Debug print to see the response
    references, clauses = parse_openai_output(output)
    #Arshit's Changes: Adding a map to store respective clause under each reference
    ref_clause_map = parse_reference_clause_structure(output)
    print("before returning to 1")
    print("Ref Clause Map:", ref_clause_map) #Arshit's Changes: Printing the ref_clause_map
    return ref_clause_map,references, clauses, output

def get_openai_output(prompt):
    prompt = truncate_text(prompt)
    #response = openai.chat.completions.create(
    response = client.chat.completions.create(
        #model="gpt-4",  # or any other appropriate model
        model=AZURE_OPENAI_MODEL_NAME,
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content

def parse_openai_output(output):
    """Parses the OpenAI API output to extract references and clauses."""
    references = []
    clauses = []
    #print(f"jayasvi {output}")
    # Prompt to send to the OpenAI API
    prompt = f"""
    Extract the references and clauses from the following text. Separate the references and clauses clearly under headings "References" and "Clauses", dont assign any serial numbers:

    {output}
    """

    # Get the response from the OpenAI API
    api_output = get_openai_output(prompt)
    print("OpenAI Response for api_output:", api_output)
    # Split the output by lines and process
    lines = api_output.split('\n')
    current_section = None

    for line in lines:
        line = line.strip()
        if line.startswith("**References**") or line.startswith("References"):
            current_section = "references"
            continue
        elif line.startswith("**Clauses**:") or line.startswith("Clauses"):
            current_section = "clauses"
            continue

        if current_section == "references":
            if line:
                cleaned_line = re.sub(r'^(\d+\.\s*|[-•*]\s*)', '', line) #Arshit's Changes: Removing the serial numbers and special characters from the references
                references.append(cleaned_line)
        elif current_section == "clauses":
            if line:
                cleaned_line = re.sub(r'^(\d+\.\s*|[-•*]\s*)', '', line) #Arshit's Changes: Removing the serial numbers and special characters from the clauses
                clauses.append(cleaned_line)
    print("OpenAI Response for api_output after")
    return references, clauses

def create_graph_with_references_and_clauses(doc_name, references, clauses):
    """Creates a graph with the document as the main node, references as edges, and clauses as nodes."""
    G = nx.DiGraph()
    start_node = "start"
    G.add_node(start_node)
    G.add_edge(start_node, doc_name)
    
    for reference, clause in zip(references, clauses):
        G.add_edge(doc_name, reference)
        clause_node = f"Clause: {clause}"
        G.add_node(clause_node)
        G.add_edge(reference, clause_node)
    
    return G

def build_graph(text):
    # Create a directed graph
    graph = nx.DiGraph()

    # Add a starting node
    start_node = "start"
    graph.add_node(start_node)

    # Process the input text line by line
    for line in text.splitlines():
        if line.strip():
            # Split the line into reference (node) and clause (edge)
            parts = line.split(' Clause ')
            if len(parts) == 2:
                reference = parts[0].strip()
                clause = 'Clause ' + parts[1].strip()

                # Add the reference as a node
                if not graph.has_node(reference):
                    graph.add_node(reference)
                
                # Add the edge from the reference node to the clause
                graph.add_edge(reference, clause)

                # Also add an edge from the start node to the reference
                if not graph.has_edge(start_node, reference):
                    graph.add_edge(start_node, reference)

    return graph
def append_references_and_clauses_to_graph(G, doc_name, references, clauses):
    """Appends references and clauses to the existing graph."""
    for reference, clause in zip(references, clauses):
        if not G.has_node(reference):
            G.add_edge(doc_name, reference)
        clause_node = f"Clause: {clause}"
        if not G.has_node(clause_node):
            G.add_node(clause_node)
        G.add_edge(reference, clause_node)
    
    return G

def extend_graph_with_references_and_clauses(graph, directory):
    """Extends the graph by reading references in the specified directory and appending additional references and clauses."""
    new_references = set(graph.nodes)
    for reference in list(graph.nodes):
        if reference.startswith("Clause:"):
            continue
        for file_name in os.listdir(directory):
            if reference.replace('.', '').split()[-1] in file_name:
                file_path = os.path.join(directory, file_name)
                if file_path.endswith(('.docx', '.pdf')):
                    main_sections = get_main_sections(file_path)
                    for main_section in main_sections:
                        subsections = get_subsections(file_path, main_section)
                        for subsection in subsections:
                            section_text = extract_text_from_subsection(file_path, subsection)
                            references, clauses, text = extract_references_and_clauses(section_text)
                            for new_reference, clause in zip(references, clauses):
                                if new_reference not in new_references:
                                    graph.add_edge(reference, new_reference)
                                    clause_node = f"Clause: {clause}"
                                    graph.add_node(clause_node)
                                    graph.add_edge(new_reference, clause_node)
                                    new_references.add(new_reference)
    return graph

def save_graph_to_json(graph, output_file):
    """Saves the graph structure to a JSON file."""
    data = nx.readwrite.json_graph.node_link_data(graph, edges="edges")
    with open(output_file, 'w') as f:
        json.dump(data, f, indent=4)
# Function to find the latest ZIP file within the release
def get_latest_zip_file(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    zip_files = [a['href'] for a in soup.find_all('a', href=True) if a['href'].endswith('.zip')]

    if not zip_files:
        raise ValueError(f"No ZIP files found in the directory: {url}")
    
    latest_zip = sorted(zip_files, reverse=True)[0]
    return latest_zip

# Main function to download and extract the latest specification
def download_and_extract_specification(series, spec_number, download_directory):
    base_url = f"https://www.3gpp.org/ftp/Specs/archive/{series}_series/{spec_number}/"
    print(f"Base URL: {base_url}")
    
    latest_zip_file = get_latest_zip_file(base_url)
    
    # Use urljoin to correctly handle URL construction
    latest_zip_url = urljoin(base_url, latest_zip_file)
    
    print(f"Final URL with ZIP file: {latest_zip_url}")
    print(f"Downloading {latest_zip_file} from {latest_zip_url}")
    
    # Download the ZIP file
    response = requests.get(latest_zip_url)
    
    # Extract just the filename from the latest_zip_file
    zip_file_name = os.path.basename(latest_zip_url)
    
    # Save the file
    file_path = os.path.join(download_directory, zip_file_name)
    with open(file_path, 'wb') as file:
        file.write(response.content)
    
    print(f"Downloaded: {file_path}")
    
    # Unzip the downloaded file
    try:
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.extractall(download_directory)
            print(f"Extracted contents to: {download_directory}")
    except zipfile.BadZipFile:
        print(f"Error: The file downloaded is not a valid ZIP file: {file_path}")
    
    # Remove the ZIP file after extraction
    os.remove(file_path)
    print(f"Removed ZIP file: {file_path}")

# Function to parse the input references and download the corresponding specifications
def download_specifications_from_references(references, download_directory):
    for reference in references:
        match = re.search(r"3GPP TS (\d+)\.(\d+)", reference)
        if match:
            series = match.group(1)
            spec_number = f"{match.group(1)}.{match.group(2)}"
            print(f"Processing: Series {series}, Specification {spec_number}")
            try:
                download_and_extract_specification(series, spec_number, download_directory)
            except ValueError as e:
                print(f"Error: {e}")
        else:
            print(f"Invalid reference format: {reference}")
def check_references_in_directory(references, directory):
    """Checks if the reference documents are present in the specified directory using regex for pattern matching."""
    present_references = []
    missing_references = []
    present_ref_map={}   #Arshit's Changes: Adding a map to store the present references and their corresponding files
    files_in_directory = os.listdir(directory)
    
    for reference in references:
        matched = False
        # Extract the main part of the reference for comparison (e.g., "23.401" from "3GPP TS 23.401 [29]")
        main_part = re.search(r'\d+\.\d+', reference)
        if main_part:
            main_part = main_part.group().replace('.', '')
            for file_name in files_in_directory:
                # Extract numeric part of the file name for comparison (e.g., "23401" from "23401-i60.docx")
                file_numeric_part = re.search(r'\d+', file_name)
                if file_numeric_part and main_part in file_numeric_part.group():
                    present_references.append(file_name)  # Store the exact file name with extension
                    ref=re.sub(r'\s*\[\d+\]$', '', reference.strip())
                    present_ref_map[ref]=file_name #Arshit's Changes: Adding the reference and its corresponding file to the map
                    matched = True
                    break
        if not matched:
            missing_references.append(reference)
    
    print("Present References Files:", present_references)
    print("Present References Map:", present_ref_map) #Arshit's Changes: Printing the present references map
    return present_ref_map,present_references, missing_references
def append_clause_sections(text, filepath, clauses):
    doc = Document(filepath)
    print ("jayasvi in append clause")
    for clause in clauses:
        print ("jayasvi in append clause for:", clause )
        return_text = read_clause_from_file(filepath, clause)
        #print("jayasvi after return text")
        #print(return_text)
        text += "\n" + "Clause:" + "\t" + clause + "\t"
        text += return_text
    with open("total_content.txt", "w") as file:
        file.write(text)
    #print(text)
    return text
def get_filename(main_part):
    # Extract the main part of the reference for comparison (e.g., "23.401" from "3GPP TS 23.401 [29]")
    files_in_directory = os.listdir(directory)
    print("main_part0",main_part)
    #main_part = re.search(r'\d+\.\d+', main_part)
    main_part=re.search(r'\d+-[a-zA-Z0-9]+', main_part)
    print("main_part",main_part)
    if main_part:
        main_part = main_part.group().replace('.', '')
        print("main_part1",main_part)
        """for file_name in files_in_directory:
                # Extract numeric part of the file name for comparison (e.g., "23401" from "23401-i60.docx")
            file_numeric_part = re.search(r'\d+', file_name)
            print("file_numeric_part",file_numeric_part)
            if file_numeric_part and main_part in file_numeric_part.group():
                print("file_name",file_name)
                return file_name  # Store the exact file name with extension"""
        for file_name in files_in_directory:
          match = re.search(r'\d+-[a-zA-Z0-9]+', file_name)
          if match and match.group() == main_part:
            return file_name

def match_extract_from_ref(ref_dict):
    for key, value in ref_dict.items():
        print(f"Key: {key}")
        print(f"Value: {value}")
        print("-" * 30)
    technical_phrase='' 
    extracted_phrases = []
    for key, value in ref_dict.items():
        # Prepare the prompt for OpenAI API
        prompt = f"Extract a meaningful 3GPP technical phrase from the following text:\n\n{value}"
            
            # Make an API call to ChatGPT for completion
        #response = openai.chat.completions.create(
        response = client.chat.completions.create(
            #model="gpt-4",
            model=AZURE_OPENAI_MODEL_NAME,
            temperature=0.7,
            messages=[
                {"role": "system", "content": "You are an assistant specialized in 3GPP technical documentation."},
                {"role": "user", "content": prompt}
            ]
        )
            
            # Extract the content from the API response
        technical_phrase = response.choices[0].message.content.strip()
        formatted_string = f"Key: {key}, Extracted Phrase: {technical_phrase}"
        extracted_phrases.append(formatted_string)
        print(formatted_string)
            # Print the extracted phrase along with the key
        #print(f"Key: {key}, Extracted Phrase: {technical_phrase}")
    return extracted_phrases
def find_files_for_extracted_phrases(extracted_phrases, directory='.'):
    present_references = []
    missing_references = []
    files_in_directory = os.listdir(directory)
    
    for phrase in extracted_phrases:
        # Extract the number pattern (e.g., "23.122") from the phrase
        match = re.search(r"Key: TS\s*(\d+\.\d+)", phrase)
        if match:
            main_part = match.group(1).replace('.', '')  # e.g., "23122" from "23.122"
            matched = False
            
            # Search for files in the directory that contain this numeric part
            for file_name in files_in_directory:
                file_numeric_part = re.search(r'\d+', file_name)
                if file_numeric_part and main_part in file_numeric_part.group():
                    present_references.append(file_name)  # File is present
                    matched = True
                    break
            
            if not matched:
                missing_references.append(match.group(1))  # Add as missing reference if no match found
    print("jayasvi in find_files_for_extracted_phrases")
    print(present_references)
    return present_references, missing_references
def is_phrase_in_section(phrase, section, threshold=0.5):
    """
    Check if a phrase is in a section with partial matching.
    
    Args:
        phrase (str): The phrase to search for.
        section (str): The section of text in which to search.
        threshold (float): Similarity threshold for a match (0 to 1).
        
    Returns:
        bool: True if the phrase is found in the section with a similarity above the threshold, False otherwise.
    """
    # Tokenize the section into sentences for more granular matching
    sentences = section.split('. ')
    
    for sentence in sentences:
        similarity = SequenceMatcher(None, phrase.lower(), sentence.lower()).ratio()
        if similarity >= threshold:
            return True  # Found a similar enough match

    return False  # No sufficient match found
'''def process_reference_extract_sections(present, phrase):
    for i, file_name in enumerate(present):
        # Load the document
        try:
            doc_path = os.path.join('.', file_name)  # Adjust the path as needed
            doc = Document(doc_path)
            # Extract sections from the document
            sections = find_sections(doc)
            #for section in sections:
            print (f"jayasvi in sections phrase{phrase}")
            matched_sections = []
            # Check if the phrase is in any of the extracted sections
            for section in sections:
                print(f"jayasvi in matching sections {section}")
                #if phrase[i].lower() in section.lower():  # Match phrase within section text
                phrase_text=str(phrase)
                section_text=str(section)
                if is_phrase_in_section(phrase_text, section_text, threshold=0.7):
                    matched_sections.append(section)
                    print("jayasvi matched_section", matched_sections)
        except Exception as e:
            print(f"Failed to process {file_name} due to: {e}")
    #return matched_sections'''
def process_reference_extract_sections(phrases, directory):
    print("jaysvi in start")
    matched_sections = ''
    files_in_directory = os.listdir(directory)
    for phrase in phrases:
        match = re.search(r"Key: TS\s*(\d+\.\d+)", phrase)
        print(f"jayasvi in phrase {phrase}, match = {match}")
        if match:
            main_part = match.group(1).replace('.', '')  # e.g., "23122" from "23.122"
            matched = False
            print(f"jayasvi in main_part {main_part}")
            # Search for files in the directory that contain this numeric part
            for file_name in files_in_directory:
                file_numeric_part = re.search(r'\d+', file_name)
                if file_numeric_part and main_part in file_numeric_part.group():
                    doc_path = os.path.join(directory, file_name)  # Use the provided directory
                    doc = Document(doc_path)
                    print(f"docpath={doc_path}")
                    # Extract sections from the document
                    sections = find_sections(doc)
                    for section in sections:
                        extracted_phrase = re.search(r"Extracted Phrase:\s*(.*)", phrase).group(1)
                        print (f" section = {section}, phrase = {extracted_phrase}")
                        if is_phrase_in_section(extracted_phrase, section, threshold=0.5):
                            matched = True
                            print(f"section matched {section}, phrase {phrase}") 
                            matched_sections += "\n" + "Section:" + "\t" + section + "\t"
                            matched_sections += sections[section]
                            if matched:
                                break
    return matched_sections
def extract_references_and_clauses1(graph, references, clauses, present_references, directory,present_ref_map,ref_clause_map,safe_subsection):
    """Processes the graph to extract text from specified clauses of reference files.
    
    
    # Add nodes and edges to the graph
    for node in present_references:
        G.add_node(node)
        print("jayasvi node:", node)
    for i in range(len(clauses)):
        if i < len(present_references):
            G.add_edge(present_references[i], clauses[i])
            print("jayasvi node1:", present_references[i], clauses[i])    
    # Process the graph
    for node in G.nodes:
        for clause in nx.neighbors(G, node):
            file_path = os.path.join(directory, node)
"""
    text=''
    phrase=''
    G = nx.DiGraph()
    start_node = "start"
    G.add_node(start_node)  # Add the start node

# Add reference nodes and connect to the start node
    for node in present_references:
        G.add_node(node)  # Add each reference node
        G.add_edge(start_node, node)  # Connect start node to reference node

# Add edges from reference nodes to corresponding clauses
    # for i in range(len(clauses)):
    #     if i < len(present_references):
    #        print("jayasvi1")
    #        G.add_edge(present_references[i], clauses[i]) # Reference to clause
    # if len(present_references) == 1 or present_references[0] == present_references[1]:
    #    for i in range(len(clauses)):
    #      if i < len(clauses):
    #        print("jayasvi2", present_references[0], clauses[i], type(present_references[0]), type(clauses[i]))
    #        G.add_edge(present_references[0], clauses[i])
    #        #print("jayasvi3")
    #        print("Edge added:", present_references[0], "->", clauses[i])
    #        print("Current edges in graph:", list(G.edges()))

    #Arshit's Changes: Adding this code to add edges to the graph using the ref_clause_map
    for ref, clauses in ref_clause_map.items():
        if ref in present_ref_map:
            print("Hi in ref map")
            file_name = present_ref_map[ref]
            for clause in clauses:
                G.add_edge(file_name, clause)
    graph = G
# Process the graph
    print("All nodes in the graph:", graph.nodes)
    for node in graph.nodes:
        if node != start_node:
            print("jayasvi node ",  node)
    # Check neighbors (connected clauses) of the current reference node
            for clause in nx.neighbors(graph, node):
                print(f"jayasvi start node: {node}")
                if(node):
                    print(f"jayasvi start1 node: {node}")
                    node1 = get_filename(node)
                    print (f"jayasvi in extract2nd: {node1} , clause: {clause}")
                    if(node1):
                        file_path = os.path.join(directory, node1)
                        clause = re.search(r'\b\d+(\.\d+)+\b', clause).group()
                        print(f"jayasvi before calling read from clause {clause} from file {file_path} for node {node1}")
                        text = read_clause_from_file(file_path, clause)
                        if text:
                            #print(f"Extracted text from {node} for clause {clause}:", text)
                            ref_clause_map_2,references, clauses, dummy_text = extract_references_and_clauses(text)
                            print("Jayasvi References:", references)
                            print("Jayasvi Clauses:", clauses)
                            ref_dict = extract_keywords_with_lines(text)

                            phrase = match_extract_from_ref(ref_dict)
                            print(f"jayasvi phrases {ref_dict}")
                            #if (phrase):
                            #    references,dummy_clauses, dummy_text = extract_references_and_clauses(phrase)
                            #present, missing = find_files_for_extracted_phrases(phrase)
                            matched_section=process_reference_extract_sections(phrase, directory)
                            print("Jayasvi after level5 References:", references)
                            text=append_clause_sections(text,file_path,clauses)
                            print("Jayasvi before keywords:", clause)
                            # Use the directory parameter instead of hardcoded path
                            base_dir = Path(directory) / "datasets"
                            subsection_folder = base_dir / safe_subsection
                            subsection_folder.mkdir(parents=True, exist_ok=True)
                            file_name = f"{clause}_file.txt"
                            file_path = subsection_folder / file_name
                            with open(file_path, "w") as file:
                                file.write(text)
                                file.write(matched_section)
                            test_scenarios=generate_test_scenarios(text)
                            #print("Generated Test Scenarios:\n")
                            #print(test_scenarios)
                            print(f"jayasvi end node: {node1}")
                        else:
                            print(f"No text found in {node} for clause {clause}.\n")
    return text

'''def extract_keywords_with_lines(text):
    # Split text into lines
    #lines = text.splitlines()
    lines = text.split('\n')
    for line in lines[:3]:
        print(line)
    keyword_pattern = r"TS\s*\d{2}\.\d{3}"
    # Initialize an empty dictionary to store the results
    result_dict = {}
    sentences = re.split(r'(?<=\.)\s+',text)
    # Loop through each line in the text
    for sentence in sentences:
        # Search for the keyword pattern in the line
        match = re.search(keyword_pattern, sentence,re.IGNORECASE)
        #print("jayasvi keyword match",match, len(text))
        if match:
            # Use the matched keyword as the key and store the whole line as the value
            result_dict[match.group(0)] = sentence.strip()
            #print(f"Keyword: {match.group()}\nLine: {sentence.strip()}\n")
    return result_dict'''

def extract_keywords_with_lines(text):
 
    keyword_pattern = r"TS\s*\d{2}\.\d{3}"
    # Initialize an empty dictionary to store the results
    result_dict = defaultdict(list)
    sentences = re.split(r'(?<=\.)\s+',text)
    # Loop through each line in the text
    for sentence in sentences:
        # Search for the keyword pattern in the line
        match = re.search(keyword_pattern, sentence,re.IGNORECASE)
        #print("jayasvi keyword match",match, len(text))
        if match:
            ts_code = match.group(0)
            # Use the matched keyword as the key and store the whole line as the value
            #if ts_code in result_dict:
            #result_dict[ts_code].append(sentence.strip())
            #else:
            result_dict[ts_code] = sentence.strip()
            #print(f"Keyword: {match.group()}\nLine: {sentence.strip()}\n")
    return result_dict

def find_sections(doc):
    sections = {}
    current_section = None
    current_content = []

    for paragraph in doc.paragraphs:
        # Check the paragraph style
        style = paragraph.style.name

        # Determine if the paragraph is a section heading
        if style in ['Heading 0', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
            if current_section is not None:
                # Store the content of the previous section
                sections[current_section] = '\n'.join(current_content)
            
            # Start a new section
            current_section = paragraph.text
            current_content = []
            #print(f"Jayasvi in Section: {current_section}")
        else:
            # Append content to the current section
            if current_section is not None:
                current_content.append(paragraph.text)
    
    # Don't forget to store the last section's content
    if current_section is not None:
        sections[current_section] = '\n'.join(current_content)
        if current_section == '5.3.4B.6':
            print(f"Jayasvi in Section: {current_section}")
            print(f"Jayasvi in Section1: {sections[current_section]}")
    return sections

def query_openai_for_clause(section_text, clause):
    section_text = truncate_text(section_text)
    prompt = (
        f"Find the section or sub-section similar to '{clause}' in the following text:\n\n"
        f"{section_text}\n\n"
        f"Please extract the text around the clause till the next clause or subsection if found, or indicate if it's not present."
    )
    prompt = truncate_text(prompt)
    print(f"Querying OpenAI with prompt:\n{prompt[:500]}...")  # Debug: Print the first 500 characters of the prompt
    try:
        #response = openai.chat.completions.create(
        response = client.chat.completions.create(
            #model="gpt-4",
            model=AZURE_OPENAI_MODEL_NAME,
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500
        )
        result = response.choices[0].message.content.strip()
        print(f"Received response: {result[:500]}...")  # Debug: Print the first 500 characters of the response
        return result
    except Exception as e:
        error_message = f"Error querying OpenAI: {e}"
        print(error_message)  # Debug: Print the error message
        return error_message


def search_clause_in_sections(sections, clause):
    """Searches for the specified clause in the section titles and contents."""
    clause = clause.replace("Clause", "").strip()  # Adjust clause text for matching
    print(f"jayasvi {clause}")
    section_title_number=''
    #clause_number = clause.split(" ", 1)[0]  # Extract the number part of the clause
    #clause_number = clause.split("Clause")[1].strip().split(" ")[0]
    #match = re.search(r"\b\d+\.\d+\.\d+\.\d+\b", clause) or re.search(r"\b\d+(\.\d+)*[A-Za-z]?\b", clause)
    match = re.search(r"\b\d+(\.\d+)*[A-Za-z]?(\.\d+)?\b", clause) or re.search(r"\b\d+\.\d+\.\d+\.\d+\b", clause)
    clause_number = match.group(0)
    print(f"jayasvi clause number: {clause_number}")
    #print(f"jayasvi {sections}")
    
    # for section in sections:
    #     #print(f"jayasvi inside loop")
    #     #section_title_number = data['title'].split("\t")[0].strip()  # Extract the section number from title
    #     match = re.search(r'(\d+(\.\d+)*)', section)
    #     if (match):
    #         section_title_number = match.group(1)
    #         #print(f"jayasvi section_titile_number {section_title_number}")
    #     #print(f"Checking Section {section} with title: {data['title']}")  # Debug: Print section title
    #     #print(f"jayasvi: clause_number {clause_number} and matched in Section is {section_title_number} ")
    #     # Check if clause number matches the section title number
    #     if clause_number == section_title_number:
    #         print(f"Clause '{clause}' matched in Section {section}.")
    #         #result = query_openai_for_clause(data['content'], clause)
    #         result = sections[section]
    #         return result
    #     #else:
    #         #print(f"Clause '{clause}' not found in any section.")
    #         #return None

    #Arshit's Changes: the above function is not working as expected, so we are using the below function to search the clause in the sections
    collected=[]
    for section_title, content in sections.items():
        match = re.search(r'(\d+(\.\d+)*)', section_title)
        if match:
            section_title_number = match.group(1)

            if clause_number == section_title_number or section_title_number.startswith(clause_number+'.'):
                print(f"Clause '{clause}' matched in Section {section_title}.")
                collected.append(f"{section_title}\n{content}")
    if collected:
        return "\n\n".join(collected)
    return None

def read_clause_from_file(file_path, clause):
    """Reads the specified clause (section) from the given DOCX file by searching through sections/subsections."""
    try:
        print(f"Attempting to read DOCX file: {file_path}")  # Debug: file path

        # Load the DOCX file
        doc = Document(file_path)

        # Identify sections and subsections in the document
        sections = find_sections(doc)
        #print("jayasvi sections:", sections)
        print("Clauses:", clause)
        # Search for the clause within the identified sections
        clause_text = search_clause_in_sections(sections, clause)
        #print(f"jayasvi return {clause_text}")
        if clause_text:
            print(f"Extracted text from {file_path} for clause '{clause}':")
        else:
            clause_text = f"Clause '{clause}' not found in the document."

    except Exception as e:
        clause_text = f"Error reading {file_path}: {e}"
    
    return clause_text

def generate_test_scenarios(input_text):
    input_text = truncate_text(input_text)
    messages = [
        {"role": "system", "content": "You are an expert in 3GPP standards."},
        {"role": "user", "content": f"Based on the following description/text/input, generate all possible test scenarios python scripts for every section or case or scenario including edge cases as per 3GPP standards, take the redirected /referred clauses from the same text:\n\n{input_text}"}
    ]
    try:
        #response = openai.chat.completions.create(
        response = client.chat.completions.create(
            #model="gpt-4",
            model=AZURE_OPENAI_MODEL_NAME,
            messages=messages,
            max_tokens=1000,  # Adjust based on the desired length of output
            temperature=0.7,  # Controls randomness, 0.7 is good for balanced output
        )
        test_scenarios = response.choices[0].message.content
        return test_scenarios
    except openai.error.OpenAIError as e:
        print(f"OpenAI API error: {e}")
        return None


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract 3GPP references and clauses from a DOCX file.")
    parser.add_argument("output_file", help="Path to the output JSON file")
    parser.add_argument("directory", help="Directory to check for reference documents")
    parser.add_argument("--file_path", required=True, help="Path to the input DOCX/PDF file")
    parser.add_argument("--section", required=True, help="Main section to extract")
    parser.add_argument("--subsection", required=True, help="Subsection to extract")
    args = parser.parse_args()

    output_file = args.output_file
    directory = args.directory
    file_path = args.file_path
    selected_main_section = args.section
    selected_subsection = args.subsection

    # Validate main section and subsection
    main_sections = get_main_sections(file_path)
    if not main_sections or selected_main_section not in main_sections:
        print(f"Main section '{selected_main_section}' not found in the document.")
        exit(1)
    subsections = get_subsections(file_path, selected_main_section)
    if not subsections or selected_subsection not in subsections:
        print(f"Subsection '{selected_subsection}' not found in main section '{selected_main_section}'.")
        exit(1)

    # Extract text from the selected subsection
    section_text = extract_text_from_subsection(file_path, selected_subsection)
    # Extract references and clauses using OpenAI API

    # Changes: Adding this code to store the extracted text from initial section in a file
    # Use the directory parameter instead of hardcoded path
    base_dir = Path(directory) / "datasets"
    safe_subsection=re.sub(r'[\\/*?:"<>|]', "_", selected_subsection)
    #create a folder for the subsection_start
    subsection_folder = base_dir / safe_subsection
    subsection_folder.mkdir(parents=True, exist_ok=True)

    output_filename = f"{safe_subsection}_section.txt"
    #output_path = os.path.join(base_dir, output_filename)
    output_path = subsection_folder / output_filename

    # Save the text
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(section_text)

    ref_clause_map,references, clauses, text = extract_references_and_clauses(section_text)
    # Check for reference documents in the directory using OpenAI API
    present_ref_map,present_references, missing_references = check_references_in_directory(references, directory)
    if missing_references:
        print("Missing References:", missing_references)
    else:
        print("All references are present in the directory.")
    # Create a graph from the document, references, and clauses
    graph = build_graph_from_input(text)
    save_graph_to_json(graph, output_file)
    print("jayasvi before calling extract")
    section_text = extract_references_and_clauses1(graph, references, clauses, present_references, directory,present_ref_map,ref_clause_map,safe_subsection)  #Arshit's Changes: Adding present_ref_map and ref_clause_map to the function
    save_graph_to_json(graph, output_file)
    print(f"Graph saved to {output_file}")
